#!/usr/bin/env python3
"""
Test script to generate an actual double template output file.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.data.excel_processor import ExcelProcessor
from src.core.generation.template_processor import TemplateProcessor

def test_double_template_output():
    """Test double template output generation."""
    print("🔍 Testing Double Template Output Generation...")
    
    # Initialize processor
    processor = ExcelProcessor()
    
    # Try to load a default file
    from src.core.data.excel_processor import get_default_upload_file
    default_file = get_default_upload_file()
    
    if not default_file or not os.path.exists(default_file):
        print("❌ No default file found for testing")
        return False
    
    print(f"📁 Loading file: {default_file}")
    
    # Load the file
    if not processor.load_file(default_file):
        print("❌ Failed to load file")
        return False
    
    print("✅ File loaded successfully")
    
    # Get available tags
    tags = processor.get_available_tags()
    if not tags:
        print("❌ No tags found")
        return False
    
    print(f"📊 Found {len(tags)} tags")
    
    # Select first few tags for testing
    test_tags = tags[:8]  # Use 8 tags to fill the double template
    processor.select_tags([tag['Product Name*'] for tag in test_tags])
    
    print(f"📋 Selected {len(test_tags)} tags for testing")
    
    # Initialize template processor
    template_processor = TemplateProcessor('double', 'arial')
    
    # Get selected records
    selected_records = processor.get_selected_records('double')
    print(f"📋 Got {len(selected_records)} selected records")
    
    # Process the records
    print("🔄 Processing records...")
    result = template_processor.process_records(selected_records)
    
    if result:
        print("✅ Template generation successful!")
        
        # Save the output file
        output_file = "test_double_template_output.docx"
        result.save(output_file)
        print(f"💾 Saved output to: {output_file}")
        
        # Check the file size
        file_size = os.path.getsize(output_file)
        print(f"📏 File size: {file_size} bytes")
        
        # Check if the file contains the expected content
        print("🔍 Checking file content...")
        
        # Read the file and check for placeholders
        from docx import Document
        doc = Document(output_file)
        
        placeholder_count = 0
        content_count = 0
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if '{{Label' in text:
                            placeholder_count += 1
                            print(f"❌ Found placeholder: {text}")
                        elif text.strip():
                            content_count += 1
                            print(f"✅ Found content: {text[:50]}...")
        
        print(f"📊 Content analysis:")
        print(f"  - Placeholders found: {placeholder_count}")
        print(f"  - Content blocks found: {content_count}")
        
        if placeholder_count == 0:
            print("🎉 SUCCESS: No placeholders found in output!")
        else:
            print("⚠️  WARNING: Some placeholders still present")
        
        return True
    else:
        print("❌ Template generation failed")
        return False

if __name__ == "__main__":
    test_double_template_output() 