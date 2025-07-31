#!/usr/bin/env python3
"""
Test to verify ProductVendor marker preservation during processing
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docxtpl import DocxTemplate
from docx import Document
import tempfile
import shutil

def test_marker_preservation():
    """Test if ProductVendor markers are preserved through the entire process"""
    print("=== Testing ProductVendor Marker Preservation ===")
    
    # Create a simple template with ProductVendor placeholder
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "{{Label1.ProductVendor}}"
    
    # Save template
    template_path = "test_productvendor_template.docx"
    doc.save(template_path)
    
    # Create context with ProductVendor
    context = {
        'Label1': {
            'ProductVendor': 'PRODUCTVENDOR_STARTTest Vendor CompanyPRODUCTVENDOR_END'
        }
    }
    
    print(f"Context ProductVendor: '{context['Label1']['ProductVendor']}'")
    
    # Render template
    template = DocxTemplate(template_path)
    template.render(context)
    
    # Save rendered document
    output_path = "test_productvendor_rendered.docx"
    template.save(output_path)
    
    # Check rendered content
    rendered_doc = Document(output_path)
    rendered_text = rendered_doc.tables[0].cell(0, 0).text
    print(f"Rendered text: '{rendered_text}'")
    
    # Check if markers are present
    if 'PRODUCTVENDOR_START' in rendered_text and 'PRODUCTVENDOR_END' in rendered_text:
        print("✅ ProductVendor markers preserved after rendering")
    else:
        print("❌ ProductVendor markers lost after rendering")
    
    # Now test with TemplateProcessor
    print("\n=== Testing with TemplateProcessor ===")
    
    from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
    
    # Create processor
    font_scheme = get_font_scheme('horizontal')
    processor = TemplateProcessor('horizontal', font_scheme)
    
    # Create test record
    test_record = {
        'Vendor': 'Test Vendor Company',
        'ProductName': 'Test Product',
        'ProductStrain': 'Test Strain'
    }
    
    # Build context
    context = processor._build_label_context(test_record, doc)
    print(f"TemplateProcessor ProductVendor: '{context.get('ProductVendor', 'NOT_FOUND')}'")
    
    # Check if markers are present in context
    if 'ProductVendor' in context:
        vendor_value = context['ProductVendor']
        if 'PRODUCTVENDOR_START' in vendor_value and 'PRODUCTVENDOR_END' in vendor_value:
            print("✅ ProductVendor markers present in TemplateProcessor context")
        else:
            print("❌ ProductVendor markers missing in TemplateProcessor context")
    else:
        print("❌ ProductVendor not found in TemplateProcessor context")
    
    # Clean up
    os.remove(template_path)
    os.remove(output_path)
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    test_marker_preservation() 