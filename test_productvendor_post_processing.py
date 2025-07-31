#!/usr/bin/env python3
"""
Test to check if ProductVendor markers are stripped during post-processing
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docxtpl import DocxTemplate
from docx import Document
import pandas as pd

def test_post_processing_marker_preservation():
    """Test if ProductVendor markers are preserved during post-processing"""
    print("=== Testing ProductVendor Marker Preservation During Post-Processing ===")
    
    # Create a simple template with ProductVendor placeholder
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "{{Label1.ProductVendor}}"
    
    # Save template
    template_path = "test_productvendor_template.docx"
    doc.save(template_path)
    
    # Create context with ProductVendor
    context = {
        'Label1': {
            'ProductVendor': 'PRODUCTVENDOR_STARTTest Vendor CompanyPRODUCTVENDOR_END'
        }
    }
    
    print(f"Context ProductVendor: '{context['Label1']['ProductVendor']}'")
    
    # Render template
    template = DocxTemplate(template_path)
    template.render(context)
    
    # Save rendered document to check content
    rendered_path = "test_productvendor_rendered.docx"
    template.save(rendered_path)
    
    # Check content before post-processing
    rendered_doc = Document(rendered_path)
    rendered_text_before = rendered_doc.tables[0].cell(0, 0).text
    print(f"Before post-processing: '{rendered_text_before}'")
    
    if 'PRODUCTVENDOR_START' in rendered_text_before and 'PRODUCTVENDOR_END' in rendered_text_before:
        print("✅ ProductVendor markers present before post-processing")
    else:
        print("❌ ProductVendor markers missing before post-processing")
    
    # Now test with TemplateProcessor including post-processing
    print("\n=== Testing with TemplateProcessor Post-Processing ===")
    
    from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
    
    # Create processor
    font_scheme = get_font_scheme('horizontal')
    processor = TemplateProcessor('horizontal', font_scheme)
    
    # Create test record
    test_record = {
        'Vendor': 'Test Vendor Company',
        'ProductName': 'Test Product',
        'ProductStrain': 'Test Strain'
    }
    
    # Create a DataFrame with the test record
    records_df = pd.DataFrame([test_record])
    
    # Process records (this includes post-processing)
    try:
        output_path = processor.process_records(records_df)
        print(f"✅ Processing completed, output saved to: {output_path}")
        
        if output_path:
            # Check the output document
            output_doc = Document(output_path)
            
            # Check all tables for ProductVendor content
            found_vendor = False
            found_markers = False
            
            for table_idx, table in enumerate(output_doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text
                        if 'Test Vendor Company' in cell_text:
                            found_vendor = True
                            print(f"✅ Found vendor text in table {table_idx}, cell [{row_idx},{col_idx}]: '{cell_text}'")
                        
                        if 'PRODUCTVENDOR_START' in cell_text and 'PRODUCTVENDOR_END' in cell_text:
                            found_markers = True
                            print(f"✅ Found ProductVendor markers in table {table_idx}, cell [{row_idx},{col_idx}]: '{cell_text}'")
            
            if not found_vendor:
                print("❌ No vendor text found in output document")
            
            if not found_markers:
                print("❌ No ProductVendor markers found in output document")
                print("   This indicates markers are being stripped during post-processing")
            
            # Clean up
            os.remove(output_path)
        else:
            print("❌ No output path returned from processing")
        
    except Exception as e:
        print(f"❌ Error during processing: {e}")
        import traceback
        traceback.print_exc()
    
    # Clean up files
    if os.path.exists(template_path):
        os.remove(template_path)
    if os.path.exists(rendered_path):
        os.remove(rendered_path)
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    test_post_processing_marker_preservation() 