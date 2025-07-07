#!/usr/bin/env python3
"""
Debug script to investigate mini template lineage coloring issues.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.docx_formatting import apply_lineage_colors, COLORS
from docx import Document
from docx.oxml.ns import qn

def debug_mini_template_lineage():
    """Debug the mini template lineage coloring issue."""
    
    print("Debugging Mini Template Lineage Coloring")
    print("=" * 50)
    
    # Test data
    test_record = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'MIXED',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'Description': 'Test description text',
        'ProductStrain': 'Mixed',
        'ProductType': 'tincture'
    }
    
    try:
        # Create mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Process a single record
        result_doc = processor.process_records([test_record])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check each cell for lineage content and coloring
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    print(f"\nCell ({row_idx}, {col_idx}):")
                    print(f"  Cell text: '{cell.text}'")
                    
                    # Check if cell has background color
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = tcPr.find(qn('w:shd'))
                    
                    if shading is not None:
                        fill_color = shading.get(qn('w:fill'))
                        print(f"  Background color: {fill_color}")
                    else:
                        print(f"  No background color")
                    
                    # Check if cell contains lineage markers
                    if 'LINEAGE_START' in cell.text and 'LINEAGE_END' in cell.text:
                        print(f"  ✅ Contains lineage markers")
                        
                        # Extract lineage content
                        start_idx = cell.text.find('LINEAGE_START') + len('LINEAGE_START')
                        end_idx = cell.text.find('LINEAGE_END')
                        lineage_content = cell.text[start_idx:end_idx].strip()
                        print(f"  Lineage content: '{lineage_content}'")
                        
                        # Test if this should get colored
                        if 'MIXED' in lineage_content:
                            print(f"  Should be colored: {COLORS['MIXED']}")
                        elif 'CBD' in lineage_content:
                            print(f"  Should be colored: {COLORS['CBD']}")
                    else:
                        print(f"  ❌ No lineage markers found")
                    
                    # Check paragraph content
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        print(f"  Paragraph {para_idx}: '{paragraph.text}'")
                        for run_idx, run in enumerate(paragraph.runs):
                            print(f"    Run {run_idx}: '{run.text}'")
        else:
            print("❌ Failed to generate document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

def test_apply_lineage_colors_directly():
    """Test applying lineage colors directly to a mini template document."""
    
    print("\n\nTesting Direct Lineage Color Application")
    print("=" * 50)
    
    # Create a simple mini template document
    doc = Document()
    table = doc.add_table(rows=5, cols=4)
    
    # Add test content to first cell
    cell = table.cell(0, 0)
    cell.text = "MIXED_PRODUCT_TYPE_tincture_IS_CLASSIC_false"
    
    print(f"Before applying colors:")
    print(f"  Cell text: '{cell.text}'")
    
    # Check initial background
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.find(qn('w:shd'))
    if shading is not None:
        print(f"  Initial background: {shading.get(qn('w:fill'))}")
    else:
        print(f"  No initial background")
    
    # Apply lineage colors
    apply_lineage_colors(doc)
    
    print(f"\nAfter applying colors:")
    print(f"  Cell text: '{cell.text}'")
    
    # Check final background
    shading = tcPr.find(qn('w:shd'))
    if shading is not None:
        fill_color = shading.get(qn('w:fill'))
        print(f"  Final background: {fill_color}")
        if fill_color == COLORS['MIXED']:
            print(f"  ✅ Correct color applied")
        else:
            print(f"  ❌ Wrong color. Expected: {COLORS['MIXED']}, Got: {fill_color}")
    else:
        print(f"  ❌ No background color applied")

if __name__ == "__main__":
    debug_mini_template_lineage()
    test_apply_lineage_colors_directly() 