#!/usr/bin/env python3
"""
Comprehensive test to check all markers with various text lengths.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_comprehensive_test_data():
    """Create test data with various text lengths for all markers."""
    return [
        {
            # Short text - should get larger font sizes
            'ProductBrand': 'Short',
            'Price': '$5',
            'Lineage': 'Sativa',
            'Ratio_or_THC_CBD': 'THC: 25%',
            'ProductStrain': 'Strain',
            'Description': 'Short desc',
            'WeightUnits': '1g'
        },
        {
            # Medium text - should get medium font sizes
            'ProductBrand': 'Medium Brand Name',
            'Price': '$25.99',
            'Lineage': 'Hybrid Indica',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductStrain': 'Medium Strain Name',
            'Description': 'This is a medium length description',
            'WeightUnits': '3.5g'
        },
        {
            # Long text - should get smaller font sizes
            'ProductBrand': 'Very Long Brand Name That Should Be Much Smaller',
            'Price': '$125.99',
            'Lineage': 'Hybrid Indica Dominant Strain',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2% Ratio: 1:1:1',
            'ProductStrain': 'Very Long Strain Name That Should Be Smaller',
            'Description': 'This is a very long description that should have much smaller font size to fit properly',
            'WeightUnits': '7.0g'
        }
    ]

def test_all_markers():
    """Test all markers with comprehensive data."""
    
    print("Comprehensive Marker Test")
    print("=" * 60)
    
    test_records = create_comprehensive_test_data()
    
    # Test vertical template
    print("\nTesting VERTICAL template:")
    print("-" * 40)
    
    processor = TemplateProcessor('vertical', {}, 1.0)
    result_doc = processor.process_records(test_records)
    
    if result_doc and result_doc.tables:
        table = result_doc.tables[0]
        
        # Check each cell for markers and font sizes
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"\nCell ({row_idx+1},{col_idx+1}):")
                    print(f"  Content: '{cell.text[:100]}...'")
                    
                    # Check each paragraph
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            print(f"  Paragraph {para_idx+1}: '{paragraph.text[:50]}...'")
                            
                            # Check each run
                            for run_idx, run in enumerate(paragraph.runs):
                                if run.text.strip():
                                    font_size = run.font.size
                                    font_name = run.font.name
                                    is_bold = run.font.bold
                                    
                                    print(f"    Run {run_idx+1}: '{run.text[:30]}...'")
                                    print(f"      Font: {font_name}, Size: {font_size}, Bold: {is_bold}")
                                    
                                    if font_size and hasattr(font_size, 'pt'):
                                        pt_value = font_size.pt
                                        print(f"      Font size: {pt_value}pt")
                                        
                                        # Check if the font size makes sense for the text length
                                        text_length = len(run.text)
                                        if text_length <= 10 and pt_value < 12:
                                            print(f"      ⚠️  WARNING: Short text ({text_length} chars) has small font ({pt_value}pt)")
                                        elif text_length >= 30 and pt_value > 16:
                                            print(f"      ⚠️  WARNING: Long text ({text_length} chars) has large font ({pt_value}pt)")
                                        else:
                                            print(f"      ✓ Font size appropriate for text length")

def test_specific_marker_scenarios():
    """Test specific problematic scenarios."""
    
    print("\n\nTesting Specific Marker Scenarios")
    print("=" * 60)
    
    # Test scenarios that might cause issues
    test_scenarios = [
        {
            'name': 'Very Short Brand',
            'data': {'ProductBrand': 'A', 'Price': '$1', 'Lineage': 'S', 'Ratio_or_THC_CBD': '1:1', 'ProductStrain': 'S', 'Description': 'D', 'WeightUnits': '1g'}
        },
        {
            'name': 'Very Long Brand',
            'data': {'ProductBrand': 'Extremely Long Brand Name That Should Be Much Smaller Than Normal', 'Price': '$1', 'Lineage': 'S', 'Ratio_or_THC_CBD': '1:1', 'ProductStrain': 'S', 'Description': 'D', 'WeightUnits': '1g'}
        },
        {
            'name': 'Very Long Price',
            'data': {'ProductBrand': 'Brand', 'Price': 'Call for pricing information', 'Lineage': 'S', 'Ratio_or_THC_CBD': '1:1', 'ProductStrain': 'S', 'Description': 'D', 'WeightUnits': '1g'}
        },
        {
            'name': 'Very Long Lineage',
            'data': {'ProductBrand': 'Brand', 'Price': '$10', 'Lineage': 'Hybrid Indica Dominant with Sativa characteristics', 'Ratio_or_THC_CBD': '1:1', 'ProductStrain': 'S', 'Description': 'D', 'WeightUnits': '1g'}
        },
        {
            'name': 'Very Long Ratio',
            'data': {'ProductBrand': 'Brand', 'Price': '$10', 'Lineage': 'S', 'Ratio_or_THC_CBD': 'THC: 25% CBD: 2% CBG: 1% CBN: 0.5%', 'ProductStrain': 'S', 'Description': 'D', 'WeightUnits': '1g'}
        }
    ]
    
    processor = TemplateProcessor('vertical', {}, 1.0)
    
    for scenario in test_scenarios:
        print(f"\nTesting: {scenario['name']}")
        print("-" * 30)
        
        result_doc = processor.process_records([scenario['data']])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)  # First cell
            
            print(f"Cell content: '{cell.text[:100]}...'")
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        font_size = run.font.size
                        if font_size and hasattr(font_size, 'pt'):
                            pt_value = font_size.pt
                            text_length = len(run.text)
                            print(f"  '{run.text[:30]}...' ({text_length} chars): {pt_value}pt")

def test_template_comparison():
    """Compare font sizes across different template types."""
    
    print("\n\nTemplate Type Comparison")
    print("=" * 60)
    
    test_data = {
        'ProductBrand': 'Medium Brand Name',
        'Price': '$25.99',
        'Lineage': 'Hybrid Indica',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'ProductStrain': 'Medium Strain',
        'Description': 'Medium description text',
        'WeightUnits': '3.5g'
    }
    
    template_types = ['vertical', 'horizontal', 'mini']
    
    for template_type in template_types:
        print(f"\n{template_type.upper()} template:")
        print("-" * 20)
        
        processor = TemplateProcessor(template_type, {}, 1.0)
        result_doc = processor.process_records([test_data])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        font_size = run.font.size
                        if font_size and hasattr(font_size, 'pt'):
                            pt_value = font_size.pt
                            print(f"  '{run.text[:20]}...': {pt_value}pt")

if __name__ == "__main__":
    print("Comprehensive Marker Testing")
    print("=" * 60)
    
    # Test all markers
    test_all_markers()
    
    # Test specific scenarios
    test_specific_marker_scenarios()
    
    # Test template comparison
    test_template_comparison()
    
    print("\n" + "=" * 60)
    print("Test completed. Check the output above for any warnings or issues.") 