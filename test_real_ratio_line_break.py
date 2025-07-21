#!/usr/bin/env python3
"""
Test script to verify that the ratio line break fix works with real application data.
"""

import sys
import os
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.data.excel_processor import ExcelProcessor, get_default_upload_file
from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt

def test_real_ratio_line_breaks():
    """Test the ratio line break fix with real data from the application."""
    
    print("=== Real Ratio Line Break Test ===\n")
    
    # Initialize the Excel processor
    processor = ExcelProcessor()
    
    # Try to load the default file
    default_file = get_default_upload_file()
    if not default_file:
        print("❌ No default file found")
        return
    
    print(f"Loading file: {default_file}")
    
    # Load the file
    if not processor.load_file(default_file):
        print("❌ Failed to load file")
        return
    
    print("✅ File loaded successfully")
    
    # Get some sample records directly from the dataframe
    sample_records = []
    for i in range(min(5, len(processor.df))):
        record = processor.df.iloc[i].to_dict()
        sample_records.append(record)
    
    if not sample_records:
        print("❌ No records found")
        return
    
    print(f"Found {len(sample_records)} sample records")
    print()
    
    # Create a template processor
    template_processor = TemplateProcessor('vertical', {})
    
    # Test each record's ratio formatting
    for i, record in enumerate(sample_records, 1):
        print(f"Record {i}:")
        print(f"  Product: {record.get('ProductName', 'N/A')}")
        print(f"  Type: {record.get('Product Type*', 'N/A')}")
        
        # Get the ratio text
        ratio_text = record.get('Ratio_or_THC_CBD', '')
        print(f"  Original ratio: '{ratio_text}'")
        
        # Create a test document to simulate the processing
        doc = Document()
        paragraph = doc.add_paragraph()
        
        # Add the ratio text to the paragraph
        run = paragraph.add_run(ratio_text)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(12)
        
        # Apply the line break conversion
        template_processor._convert_br_markers_to_line_breaks(paragraph)
        
        # Check the result
        result_text = paragraph.text
        print(f"  Processed ratio: '{result_text}'")
        
        # Count the number of line breaks
        line_count = len(result_text.split('\n'))
        print(f"  Line count: {line_count}")
        
        # Check if there are any trailing empty lines
        lines = result_text.split('\n')
        if lines and not lines[-1].strip():
            print("  ⚠️  WARNING: Trailing empty line detected")
        else:
            print("  ✅ No trailing empty lines")
        
        print()
    
    print("=== Test Complete ===")

if __name__ == "__main__":
    test_real_ratio_line_breaks() 