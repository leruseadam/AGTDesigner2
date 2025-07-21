#!/usr/bin/env python3
"""
Test line break conversion functionality
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from src.core.generation.template_processor import TemplateProcessor

def test_line_break_conversion():
    """Test that \n characters are properly converted to line breaks in documents."""
    
    print("Testing Line Break Conversion")
    print("=" * 50)
    
    # Create a test document
    doc = Document()
    
    # Add a paragraph with \n characters
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("230mg CBD\n50mg THC\n10mg CBG\n10mg CBN")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(12)
    
    print(f"Original text: '{paragraph.text}'")
    has_newlines = '\n' in paragraph.text
    print(f"Contains newlines: {has_newlines}")
    
    # Create template processor to access the conversion function
    def get_font_scheme(template_type, base_size=12):
        return {'ratio': base_size, 'brand': base_size - 2, 'description': base_size - 1}
    
    processor = TemplateProcessor('classic', get_font_scheme('classic'))
    
    # Convert the line breaks
    processor._convert_br_markers_to_line_breaks(paragraph)
    
    print(f"After conversion: '{paragraph.text}'")
    print(f"Number of runs: {len(paragraph.runs)}")
    
    # Check if line breaks were properly converted
    if len(paragraph.runs) > 1:
        print("✅ Line breaks were successfully converted to separate runs")
        for i, run in enumerate(paragraph.runs):
            print(f"  Run {i+1}: '{run.text}'")
    else:
        print("❌ Line breaks were not converted properly")
    
    # Save the test document
    test_file = "test_line_break_output.docx"
    doc.save(test_file)
    print(f"\nTest document saved as: {test_file}")
    
    return True

if __name__ == "__main__":
    test_line_break_conversion() 