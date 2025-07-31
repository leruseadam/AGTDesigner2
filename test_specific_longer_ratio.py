#!/usr/bin/env python3
"""
Simple test to verify longer THC_or_CBD_Ratio values are bolded.
"""

import sys
from pathlib import Path

# Add the src directory to the Python path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor

def test_specific_longer_ratio():
    """Test specific longer ratio values mentioned by user."""
    
    print("Testing specific longer ratio values...")
    
    # Create test data with the specific longer ratio format mentioned
    test_records = [
        {
            'Product Name*': 'Test Product',
            'Product Type*': 'Edible',
            'Lineage': 'HYBRID',
            'Ratio': '50mg THC 50mg CBD',  # The specific format mentioned
            'Product Strain': 'Test Strain',
            'Product Brand': 'Test Brand',
            'Price': '$25.00',
            'Weight Units': '1g',
            'Description': 'Test description'
        }
    ]
    
    # Test mini template
    print("\nTesting mini template...")
    
    try:
        processor = TemplateProcessor('mini', 'default', 1.0)
        result = processor.process_records(test_records)
        
        # Save to file
        with open('test_specific_longer_ratio_mini.docx', 'wb') as f:
            f.write(result)
        
        print("✓ Generated test_specific_longer_ratio_mini.docx")
        
        # Check the content
        from docx import Document
        from io import BytesIO
        
        # Save to temporary file first
        with open('temp_specific_test.docx', 'wb') as f:
            f.write(result)
        
        doc = Document('temp_specific_test.docx')
        
        # Look for ratio content
        ratio_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'mg THC' in run.text or 'mg CBD' in run.text:
                                print(f"    Found ratio content: '{run.text}'")
                                print(f"    Bold: {run.font.bold}")
                                ratio_found = True
                                if not run.font.bold:
                                    print(f"    ❌ Ratio content is NOT bold!")
                                else:
                                    print(f"    ✓ Ratio content is bold")
        
        if not ratio_found:
            print(f"    ⚠ No ratio content found")
            
    except Exception as e:
        print(f"    ❌ Error: {e}")
    
    print("\nTest completed!")

if __name__ == "__main__":
    test_specific_longer_ratio() 