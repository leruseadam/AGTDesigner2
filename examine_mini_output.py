#!/usr/bin/env python3
"""
Script to examine the generated mini template document.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document

def examine_mini_output():
    """Examine the generated mini template document."""
    
    print("Examining Mini Template Output")
    print("=" * 50)
    
    try:
        # Load the generated document
        doc = Document('test_mini_output.docx')
        
        if doc.tables:
            table = doc.tables[0]
            print(f"Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check each cell in the first few rows
            for row_idx in range(min(3, len(table.rows))):
                for col_idx in range(min(3, len(table.columns))):
                    cell = table.cell(row_idx, col_idx)
                    print(f"Cell ({row_idx}, {col_idx}): '{cell.text.strip()}'")
                    
                    # Check if there are any markers still present
                    if 'LINEAGE_START' in cell.text or 'PRODUCTBRAND_CENTER_START' in cell.text:
                        print(f"  ❌ MARKERS STILL PRESENT!")
                    else:
                        print(f"  ✅ No markers found")
        else:
            print("❌ No tables found in document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    examine_mini_output() 