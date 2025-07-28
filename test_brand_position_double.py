#!/usr/bin/env python3
"""
Test brand positioning in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def test_brand_position_double():
    """Test brand positioning in double template."""
    
    print("Brand Position Double Template Test")
    print("=" * 40)
    
    # Create test records with clear brand names
    test_records = [
        {
            'Description': 'Test Product Description',
            'Price': '$15',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 20% CBD: 1%',
            'ProductBrand': 'BRAND SHOULD BE HERE',
            'ProductStrain': 'TEST STRAIN SHOULD BE HERE',
            'ProductType': 'flower',
            'JointRatio': '3.5g'
        }
    ]
    
    # Initialize template processor
    processor = TemplateProcessor('double', 'standard', scale_factor=1.0)
    
    # Process the records
    print("Processing records...")
    result = processor.process_records(test_records)
    
    if result is None:
        print("❌ Failed to process records")
        return
    
    # Save the result
    output_path = "test_brand_position_output.docx"
    result.save(output_path)
    print(f"✅ Generated document: {output_path}")
    
    # Analyze the content structure
    print("\nContent Structure Analysis:")
    print("=" * 40)
    
    doc = Document(output_path)
    
    # Check all tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"\nCell ({row_idx},{col_idx}):")
                    
                    # Split by lines and analyze each paragraph
                    lines = cell_text.split('\n')
                    for i, line in enumerate(lines):
                        if line.strip():
                            print(f"  Line {i+1}: '{line.strip()}'")
                            
                            # Check if this line contains the brand
                            if 'BRAND SHOULD BE HERE' in line:
                                print(f"    ✓ FOUND BRAND at line {i+1}")
                            elif 'Test Product Description' in line:
                                print(f"    ✓ FOUND DESCRIPTION at line {i+1}")
                            elif '$15' in line:
                                print(f"    ✓ FOUND PRICE at line {i+1}")
                            elif 'HYBRID' in line:
                                print(f"    ✓ FOUND LINEAGE at line {i+1}")
                            elif 'TEST STRAIN SHOULD BE HERE' in line:
                                print(f"    ✓ FOUND PRODUCT STRAIN at line {i+1}")
                            elif 'THC: 20%' in line:
                                print(f"    ✓ FOUND RATIO at line {i+1}")
                            elif '3.5g' in line:
                                print(f"    ✓ FOUND JOINT RATIO at line {i+1}")
    
    print("\nExpected vs Actual Position:")
    print("=" * 40)
    print("Expected order (based on template):")
    print("  1. DescAndWeight (Description)")
    print("  2. Price")
    print("  3. Lineage")
    print("  4. ProductStrain")
    print("  5. Ratio_or_THC_CBD")
    print("  6. ProductBrand (BRAND SHOULD BE HERE)")
    print("  7. JointRatio")
    
    print("\nActual order (from generated document):")
    # Re-analyze to show actual order
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    lines = cell_text.split('\n')
                    print("  Actual content order:")
                    for i, line in enumerate(lines):
                        if line.strip():
                            print(f"    {i+1}. '{line.strip()}'")

if __name__ == "__main__":
    test_brand_position_double() 