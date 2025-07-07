#!/usr/bin/env python3
"""
Debug script to test actual template processing and font sizing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
import logging

# Set up logging to see what's happening
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def create_test_data():
    """Create test data for processing."""
    return [
        {
            'ProductBrand': 'Test Brand Name',
            'Price': '$25.99',
            'Lineage': 'Sativa',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductStrain': 'Test Strain',
            'Description': 'This is a test description',
            'WeightUnits': '3.5g'
        },
        {
            'ProductBrand': 'Very Long Brand Name That Should Be Smaller',
            'Price': '$125.99',
            'Lineage': 'Hybrid Indica Dominant',
            'Ratio_or_THC_CBD': '1:1:1',
            'ProductStrain': 'Another Test Strain',
            'Description': 'This is a very long description that should have smaller font size',
            'WeightUnits': '1g'
        }
    ]

def test_template_processing():
    """Test the actual template processing."""
    
    print("Testing Template Processing")
    print("=" * 50)
    
    # Create test data
    test_records = create_test_data()
    
    # Test different template types
    template_types = ['vertical', 'horizontal', 'mini']
    
    for template_type in template_types:
        print(f"\nTesting template type: {template_type}")
        print("-" * 30)
        
        try:
            # Create template processor
            processor = TemplateProcessor(template_type, {}, 1.0)
            
            # Process the records
            result_doc = processor.process_records(test_records)
            
            if result_doc:
                print(f"✓ Successfully processed {len(test_records)} records")
                
                # Check if the document has content
                if result_doc.tables:
                    print(f"✓ Document has {len(result_doc.tables)} tables")
                    
                    # Check each table for content
                    for i, table in enumerate(result_doc.tables):
                        print(f"  Table {i+1}: {len(table.rows)} rows, {len(table.columns)} columns")
                        
                        # Check each cell for content and font sizes
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                if cell.text.strip():
                                    print(f"    Cell ({row_idx+1},{col_idx+1}): '{cell.text[:50]}...'")
                                    
                                    # Check font sizes in paragraphs
                                    for para_idx, paragraph in enumerate(cell.paragraphs):
                                        if paragraph.text.strip():
                                            print(f"      Paragraph {para_idx+1}: '{paragraph.text[:30]}...'")
                                            
                                            # Check font sizes in runs
                                            for run_idx, run in enumerate(paragraph.runs):
                                                if run.text.strip():
                                                    font_size = run.font.size
                                                    font_name = run.font.name
                                                    is_bold = run.font.bold
                                                    
                                                    print(f"        Run {run_idx+1}: '{run.text[:20]}...'")
                                                    print(f"          Font: {font_name}, Size: {font_size}, Bold: {is_bold}")
                                                    
                                                    # Check if font size is reasonable
                                                    if font_size:
                                                        if hasattr(font_size, 'pt'):
                                                            pt_value = font_size.pt
                                                            if 6 <= pt_value <= 50:
                                                                print(f"          ✓ Font size {pt_value}pt is reasonable")
                                                            else:
                                                                print(f"          ✗ Font size {pt_value}pt is outside reasonable range")
                                                        else:
                                                            print(f"          ✗ Font size {font_size} is not a Pt object")
                                                    else:
                                                        print(f"          ✗ No font size set")
                else:
                    print("✗ Document has no tables")
            else:
                print("✗ Failed to process records")
                
        except Exception as e:
            print(f"✗ Error processing {template_type} template: {e}")
            import traceback
            traceback.print_exc()

def test_marker_processing():
    """Test the marker processing specifically."""
    
    print("\n\nTesting Marker Processing")
    print("=" * 50)
    
    # Create a simple test document
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add test content with markers
    test_content = "PRODUCTBRAND_CENTER_STARTTest BrandPRODUCTBRAND_CENTER_END"
    cell.text = test_content
    
    print(f"Original cell text: '{cell.text}'")
    
    # Test the template processor's marker processing
    processor = TemplateProcessor('vertical', {}, 1.0)
    
    # Process the document
    processor._post_process_and_replace_content(doc)
    
    print(f"After processing: '{cell.text}'")
    
    # Check the formatting
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.text.strip():
                font_size = run.font.size
                font_name = run.font.name
                is_bold = run.font.bold
                
                print(f"Run: '{run.text}'")
                print(f"  Font: {font_name}, Size: {font_size}, Bold: {is_bold}")
                
                if font_size and hasattr(font_size, 'pt'):
                    print(f"  Font size in points: {font_size.pt}")

if __name__ == "__main__":
    print("Template Processing Debug Test")
    print("=" * 50)
    
    # Test marker processing
    test_marker_processing()
    
    # Test full template processing
    test_template_processing() 