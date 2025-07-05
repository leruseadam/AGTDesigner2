#!/usr/bin/env python3
"""
Comprehensive test script to track row height changes from EXACTLY to AT_LEAST
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from docxtpl import DocxTemplate
from io import BytesIO
import logging

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor, get_font_scheme
from src.core.generation.docx_formatting import fix_table_row_heights, enforce_fixed_cell_dimensions

def check_row_heights(doc, stage_name):
    """Check and print row height rules for all tables in the document"""
    print(f"\n=== {stage_name} ===")
    
    # Handle both Document and DocxTemplate objects
    if hasattr(doc, 'docx') and doc.docx is not None:
        # It's a DocxTemplate object with loaded document
        tables = doc.docx.tables
    elif hasattr(doc, 'tables'):
        # It's a Document object
        tables = doc.tables
    else:
        print("  No tables found or document not properly loaded")
        return
    
    for i, table in enumerate(tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            height = row.height
            height_rule = row.height_rule
            status = "✓ EXACTLY" if height_rule == WD_ROW_HEIGHT_RULE.EXACTLY else f"✗ {height_rule}"
            print(f"  Row {j}: height={height}, rule={height_rule} - {status}")

def create_test_template():
    """Create a simple test template with EXACTLY row heights"""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    
    # Set row heights to EXACTLY
    for row in table.rows:
        row.height = Inches(2.25)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Add some template variables
    for i in range(3):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = f"{{{{Label{i*3+j+1}.ProductBrand}}}}"
    
    # Save template
    template_path = "test_template.docx"
    doc.save(template_path)
    return template_path

def test_template_processing():
    """Test the entire template processing pipeline"""
    
    print("=== ROW HEIGHT TRACKING TEST ===")
    
    # Step 1: Create test template
    template_path = create_test_template()
    print(f"Created test template: {template_path}")
    
    # Step 2: Load template and check initial state
    doc = Document(template_path)
    check_row_heights(doc, "1. Initial template state")
    
    # Step 3: Test DocxTemplate loading
    with open(template_path, 'rb') as f:
        template_buffer = BytesIO(f.read())
    
    docx_template = DocxTemplate(template_buffer)
    # DocxTemplate doesn't load the document until render() is called
    print("\n=== 2. After DocxTemplate loading ===")
    print("DocxTemplate loaded (document not yet parsed)")
    
    # Step 4: Test template rendering
    context = {}
    for i in range(9):
        context[f'Label{i+1}'] = {'ProductBrand': f'Test Brand {i+1}'}
    
    docx_template.render(context)
    check_row_heights(docx_template, "3. After DocxTemplate.render()")
    
    # Step 5: Test save and reload
    buffer = BytesIO()
    docx_template.save(buffer)
    buffer.seek(0)
    rendered_doc = Document(buffer)
    check_row_heights(rendered_doc, "4. After save and reload")
    
    # Step 6: Test TemplateProcessor
    font_scheme = get_font_scheme('vertical')
    processor = TemplateProcessor('vertical', font_scheme, 1.0)
    
    # Check the expanded template buffer
    if hasattr(processor, '_expanded_template_buffer'):
        processor._expanded_template_buffer.seek(0)
        expanded_doc = Document(processor._expanded_template_buffer)
        check_row_heights(expanded_doc, "5. After TemplateProcessor expansion")
    
    # Step 7: Test with actual records
    test_records = [
        {
            'ProductBrand': 'Test Brand 1',
            'Description': 'Test Description 1',
            'Price': '$10.00',
            'Lineage': 'Test Lineage 1',
            'Ratio_or_THC_CBD': 'THC: 20%',
            'WeightUnits': '1g',
            'DOH': 'YES',
            'ProductType': 'flower'
        }
    ] * 9  # Create 9 identical records
    
    try:
        final_doc = processor.process_records(test_records)
        check_row_heights(final_doc, "6. After TemplateProcessor.process_records()")
    except Exception as e:
        print(f"Error in process_records: {e}")
        import traceback
        traceback.print_exc()
    
    # Step 8: Test fix_table_row_heights function
    if final_doc:
        fix_table_row_heights(final_doc, 'vertical')
        check_row_heights(final_doc, "7. After fix_table_row_heights()")
    
    # Step 9: Test enforce_fixed_cell_dimensions function
    if final_doc:
        for table in final_doc.tables:
            enforce_fixed_cell_dimensions(table)
        check_row_heights(final_doc, "8. After enforce_fixed_cell_dimensions()")
    
    # Step 10: Final save and check
    if final_doc:
        output_path = "test_final_output.docx"
        final_doc.save(output_path)
        
        final_check_doc = Document(output_path)
        check_row_heights(final_check_doc, "9. After final save and reload")
        
        # Clean up
        if os.path.exists(output_path):
            os.remove(output_path)
    
    # Clean up
    if os.path.exists(template_path):
        os.remove(template_path)
    
    print("\n=== TEST COMPLETE ===")

def test_docxtpl_behavior():
    """Test DocxTemplate behavior specifically"""
    print("\n=== DOCXTEMPLATE BEHAVIOR TEST ===")
    
    # Create a simple document
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    
    # Set row height to EXACTLY
    row = table.rows[0]
    row.height = Inches(2.25)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Add template variable
    table.cell(0, 0).text = "{{test_var}}"
    
    check_row_heights(doc, "1. Initial document")
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Load with DocxTemplate
    docx_template = DocxTemplate(buffer)
    print("\n=== 2. After DocxTemplate loading ===")
    print("DocxTemplate loaded (document not yet parsed)")
    
    # Render
    docx_template.render({'test_var': 'Test Value'})
    check_row_heights(docx_template, "3. After DocxTemplate.render()")
    
    # Save and reload
    output_buffer = BytesIO()
    docx_template.save(output_buffer)
    output_buffer.seek(0)
    
    final_doc = Document(output_buffer)
    check_row_heights(final_doc, "4. After save and reload")

if __name__ == "__main__":
    # Set up logging
    logging.basicConfig(level=logging.INFO)
    
    # Run tests
    test_docxtpl_behavior()
    test_template_processing() 