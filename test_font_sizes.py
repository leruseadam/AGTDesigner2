#!/usr/bin/env python3
"""
Script to examine font sizes in the output document
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt

def examine_font_sizes(docx_path):
    """Examine font sizes in a DOCX document"""
    try:
        doc = Document(docx_path)
        print(f"\n=== Font Size Analysis for {docx_path} ===")
        
        # Check paragraphs outside tables
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"\nParagraph {i}: '{para.text}'")
                for j, run in enumerate(para.runs):
                    if run.text.strip():
                        font_size = run.font.size
                        font_name = run.font.name
                        print(f"  Run {j}: '{run.text}' - Font: {font_name}, Size: {font_size}")
        
        # Check tables
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"\n  Cell [{row_idx},{col_idx}]: '{cell_text}'")
                        for para_idx, para in enumerate(cell.paragraphs):
                            if para.text.strip():
                                print(f"    Paragraph {para_idx}: '{para.text}'")
                                for run_idx, run in enumerate(para.runs):
                                    if run.text.strip():
                                        font_size = run.font.size
                                        font_name = run.font.name
                                        print(f"      Run {run_idx}: '{run.text}' - Font: {font_name}, Size: {font_size}")
                                        
                                        # Check if this is vendor text
                                        if 'Test Vendor Company' in run.text:
                                            print(f"        🔍 THIS IS THE VENDOR TEXT!")
                                            if font_size:
                                                size_pt = font_size.pt if hasattr(font_size, 'pt') else font_size
                                                print(f"        📏 Vendor text font size: {size_pt}pt")
                                                if size_pt < 6:
                                                    print(f"        ⚠️  WARNING: Font size {size_pt}pt is very small!")
                                                elif size_pt < 8:
                                                    print(f"        ⚠️  WARNING: Font size {size_pt}pt is small!")
                                                else:
                                                    print(f"        ✅ Font size {size_pt}pt should be visible")
                                            else:
                                                print(f"        ⚠️  No font size set!")
        
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")

def main():
    # Check the output from our test
    output_path = "test_real_vendor_output.docx"
    if os.path.exists(output_path):
        examine_font_sizes(output_path)
    else:
        print(f"Output file {output_path} not found")

if __name__ == "__main__":
    main() 