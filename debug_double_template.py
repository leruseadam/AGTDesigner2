#!/usr/bin/env python3
"""
Debug script to understand why ProductBrand placeholder is not being added to double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.oxml.ns import qn
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def debug_double_template():
    """Debug the double template expansion to understand the issue."""
    
    print("Debugging Double Template Expansion")
    print("=" * 40)
    
    try:
        # Create template processor
        double_processor = TemplateProcessor('double', {}, 1.0)
        
        # Get the original template
        template_path = double_processor._get_template_path()
        doc = Document(template_path)
        
        print(f"\n1. Original template path: {template_path}")
        print(f"   Template has {len(doc.tables)} tables")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"   Table has {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check the first cell content
            first_cell = table.cell(0, 0)
            print(f"\n2. First cell content:")
            print(f"   Cell text: '{first_cell.text}'")
            
            # Check individual text elements
            print(f"\n3. Individual text elements in first cell:")
            for i, run in enumerate(first_cell.paragraphs[0].runs):
                print(f"   Run {i}: '{run.text}'")
            
            # Check for specific placeholders
            cell_text = first_cell.text
            print(f"\n4. Placeholder analysis:")
            print(f"   Contains 'Label1.Lineage': {'Label1.Lineage' in cell_text}")
            print(f"   Contains 'Label1.ProductStrain': {'Label1.ProductStrain' in cell_text}")
            print(f"   Contains 'Label1.ProductBrand': {'Label1.ProductBrand' in cell_text}")
            
            # Check the source cell content that gets copied
            print(f"\n5. Source cell content for expansion:")
            src_tc = first_cell._tc
            cell_text_parts = []
            for t in src_tc.iter(qn('w:t')):
                if t.text:
                    cell_text_parts.append(t.text)
                    print(f"   Text element: '{t.text}'")
            
            full_cell_text = ''.join(cell_text_parts)
            print(f"   Combined text: '{full_cell_text}'")
            
            # Test the placeholder addition logic
            print(f"\n6. Testing placeholder addition logic:")
            if '{{Label1.ProductBrand}}' not in full_cell_text and 'ProductBrand' not in full_cell_text:
                print("   ✓ ProductBrand placeholder is missing (as expected)")
                
                # Check if Lineage placeholder exists
                if '{{Label1.Lineage}}' in full_cell_text:
                    print("   ✓ Lineage placeholder found")
                    
                    # Try to add ProductBrand placeholder
                    print("   Attempting to add ProductBrand placeholder...")
                    
                    # Find the Lineage text element
                    lineage_element = None
                    for t in src_tc.iter(qn('w:t')):
                        if t.text and '{{Label1.Lineage}}' in t.text:
                            lineage_element = t
                            print(f"   Found Lineage element: '{t.text}'")
                            break
                    
                    if lineage_element:
                        print("   ✓ Found Lineage element, would add ProductBrand after it")
                    else:
                        print("   ❌ Could not find Lineage element")
                else:
                    print("   ❌ Lineage placeholder not found")
            else:
                print("   ❌ ProductBrand placeholder already exists")
        
        print(f"\n7. Testing template expansion...")
        double_processor._expand_template_if_needed()
        
        # Check the expanded template
        expanded_buffer = double_processor._expanded_template_buffer
        if hasattr(expanded_buffer, 'seek'):
            expanded_buffer.seek(0)
        
        expanded_doc = Document(expanded_buffer)
        expanded_table = expanded_doc.tables[0]
        
        print(f"   Expanded table: {len(expanded_table.rows)} rows x {len(expanded_table.columns)} columns")
        
        # Check first few cells for ProductBrand placeholders
        print(f"\n8. Checking expanded cells for ProductBrand placeholders:")
        for r in range(min(2, len(expanded_table.rows))):
            for c in range(min(2, len(expanded_table.columns))):
                cell = expanded_table.cell(r, c)
                cell_text = cell.text
                has_product_brand = 'ProductBrand' in cell_text
                print(f"   Cell ({r},{c}): '{cell_text[:50]}...' - ProductBrand: {has_product_brand}")
        
        return True
        
    except Exception as e:
        print(f"\n❌ DEBUG FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = debug_double_template()
    sys.exit(0 if success else 1) 