#!/usr/bin/env python3
"""
Test script to verify the double template vertical gutter implementation.
This script tests that the 4x3 double template now has a vertical gutter down the middle.
"""

import os
import sys
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def test_double_template_gutter():
    """Test that the double template has a vertical gutter."""
    print("🔍 Testing Double Template Vertical Gutter")
    print("=" * 50)
    
    try:
        # Import the template processor
        from src.core.generation.template_processor import TemplateProcessor
        from src.core.generation.docx_formatting import get_font_scheme
        
        # Create a template processor for double template
        font_scheme = get_font_scheme('double')
        processor = TemplateProcessor('double', font_scheme)
        
        print("✅ Template processor created successfully")
        
        # Force re-expansion to ensure we get the latest gutter implementation
        processor.force_re_expand_template()
        print("✅ Template re-expanded with gutter implementation")
        
        # Create a test document to verify the layout
        from docx import Document
        from io import BytesIO
        
        # Get the expanded template buffer
        template_buffer = processor._expanded_template_buffer
        template_buffer.seek(0)
        
        # Load the template
        doc = Document(template_buffer)
        
        if not doc.tables:
            print("❌ No tables found in template")
            return False
            
        table = doc.tables[0]
        print(f"✅ Found table with {len(table.rows)} rows and {len(table.columns)} columns")
        
        # Verify the table structure
        expected_rows = 3
        expected_cols = 4
        
        if len(table.rows) != expected_rows:
            print(f"❌ Expected {expected_rows} rows, got {len(table.rows)}")
            return False
            
        if len(table.columns) != expected_cols:
            print(f"❌ Expected {expected_cols} columns, got {len(table.columns)}")
            return False
            
        print("✅ Table dimensions are correct (3 rows x 4 columns)")
        
        # Check column widths
        print("\n📏 Column Width Analysis:")
        for i, col in enumerate(table.columns):
            # Get the width from the grid
            grid_cols = table._element.findall('.//w:gridCol', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if i < len(grid_cols):
                width_twips = grid_cols[i].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                width_inches = int(width_twips) / 1440 if width_twips else 0
                
                if i in [1, 3]:  # Gutter columns
                    expected_width = 0.5
                    column_type = "Gutter"
                else:  # Label columns
                    expected_width = 1.75
                    column_type = "Label"
                    
                print(f"   Column {i+1}: {width_inches:.2f}\" ({column_type}) - Expected: {expected_width}\"")
                
                if abs(width_inches - expected_width) > 0.1:
                    print(f"   ⚠️  Width mismatch for column {i+1}")
                else:
                    print(f"   ✅ Width correct for column {i+1}")
        
        # Check cell content
        print("\n📋 Cell Content Analysis:")
        label_count = 0
        gutter_count = 0
        
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                cell = table.cell(r, c)
                cell_text = cell.text.strip()
                
                if c in [1, 3]:  # Gutter columns
                    if not cell_text:
                        print(f"   ✅ Gutter cell ({r+1},{c+1}): Empty (correct)")
                        gutter_count += 1
                    else:
                        print(f"   ⚠️  Gutter cell ({r+1},{c+1}): Has content '{cell_text[:20]}...'")
                else:  # Label columns
                    if cell_text and 'Label' in cell_text:
                        print(f"   ✅ Label cell ({r+1},{c+1}): Contains label placeholder")
                        label_count += 1
                    else:
                        print(f"   ⚠️  Label cell ({r+1},{c+1}): Missing or invalid content")
        
        print(f"\n📊 Summary:")
        print(f"   Label cells: {label_count}/6 (should be 6)")
        print(f"   Gutter cells: {gutter_count}/6 (should be 6)")
        
        if label_count == 6 and gutter_count == 6:
            print("   ✅ All cells correctly configured")
        else:
            print("   ❌ Cell configuration issues detected")
            return False
        
        # Test the layout visually
        print("\n🎨 Layout Structure:")
        print("   ┌─────────┬─────┬─────────┬─────┐")
        print("   │ Label1  │     │ Label2  │     │")
        print("   │         │     │         │     │")
        print("   ├─────────┼─────┼─────────┼─────┤")
        print("   │ Label3  │     │ Label4  │     │")
        print("   │         │     │         │     │")
        print("   ├─────────┼─────┼─────────┼─────┤")
        print("   │ Label5  │     │ Label6  │     │")
        print("   │         │     │         │     │")
        print("   └─────────┴─────┴─────────┴─────┘")
        print("   │ 1.75\"  │0.5\" │ 1.75\"  │0.5\" │")
        
        print("\n✅ Double template vertical gutter test completed successfully!")
        return True
        
    except Exception as e:
        print(f"❌ Error testing double template gutter: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function."""
    print("🚀 Double Template Vertical Gutter Test")
    print("This test verifies that the double template has a vertical gutter down the middle")
    print()
    
    # Run the test
    success = test_double_template_gutter()
    
    if success:
        print("\n🎉 Test passed! The double template now has a vertical gutter.")
        print("The layout provides better visual separation between label groups.")
    else:
        print("\n❌ Test failed! Please check the implementation.")
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 