#!/usr/bin/env python3
"""
Debug script to identify why tables are missing in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def check_template_file():
    """Check if the double template file exists and has tables."""
    print("Checking Double Template File")
    print("=" * 40)
    
    template_path = "src/core/generation/templates/double.docx"
    
    if not os.path.exists(template_path):
        print(f"❌ Template file not found: {template_path}")
        return False
    
    print(f"✓ Template file exists: {template_path}")
    
    try:
        doc = Document(template_path)
        print(f"✓ Template document loaded successfully")
        print(f"  Tables: {len(doc.tables)}")
        print(f"  Paragraphs: {len(doc.paragraphs)}")
        
        if doc.tables:
            for i, table in enumerate(doc.tables):
                print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
                if table.rows and table.columns:
                    cell = table.cell(0, 0)
                    print(f"    First cell content: '{cell.text[:100]}...'")
        else:
            print("  ❌ No tables found in template file")
            return False
            
        return True
        
    except Exception as e:
        print(f"❌ Error loading template: {e}")
        return False

def test_template_expansion():
    """Test template expansion process."""
    print("\nTesting Template Expansion")
    print("=" * 40)
    
    try:
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Check if template was expanded
        if hasattr(processor, '_expanded_template_buffer') and processor._expanded_template_buffer:
            print("✓ Template expansion buffer exists")
            
            # Load the expanded template
            expanded_doc = Document(processor._expanded_template_buffer)
            print(f"✓ Expanded template loaded")
            print(f"  Tables: {len(expanded_doc.tables)}")
            
            if expanded_doc.tables:
                for i, table in enumerate(expanded_doc.tables):
                    print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
                    
                    # Check each cell for placeholders
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if 'Label' in cell_text:
                                print(f"    Cell ({row_idx},{col_idx}): Contains placeholders ✓")
                            else:
                                print(f"    Cell ({row_idx},{col_idx}): '{cell_text[:50]}...'")
            else:
                print("  ❌ No tables in expanded template")
                return False
        else:
            print("❌ Template expansion buffer is empty")
            return False
            
        return True
        
    except Exception as e:
        print(f"❌ Error in template expansion: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_document_generation():
    """Test full document generation process."""
    print("\nTesting Document Generation")
    print("=" * 40)
    
    # Create test records
    test_records = [
        {
            'ProductName': 'HUSTLER\'S AMBITION Lemon Jealousy Wax',
            'Description': 'HUSTLER\'S AMBITION Lemon Jealousy Wax - 1g',
            'ProductBrand': 'HUSTLER\'S AMBITION',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 75.52% CBD: 0.11%',
            'ProductStrain': 'Lemon Jealousy',
            'DOH': 'DOH'
        },
        {
            'ProductName': 'HUSTLER\'S AMBITION Memory Loss Wax',
            'Description': 'HUSTLER\'S AMBITION Memory Loss Wax - 1g',
            'ProductBrand': 'HUSTLER\'S AMBITION',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 73.42% CBD: 0.13%',
            'ProductStrain': 'Memory Loss',
            'DOH': 'DOH'
        },
        {
            'ProductName': 'SUPER MEGA BUSSIN\' Blizzard Wizard Sugar Crumble',
            'Description': 'SUPER MEGA BUSSIN\' Blizzard Wizard Sugar Crumble - 1g',
            'ProductBrand': 'SUPER MEGA BUSSIN\'',
            'Price': '$10',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 77.47% CBD: 0.0%',
            'ProductStrain': 'Blizzard Wizard',
            'DOH': 'DOH'
        },
        {
            'ProductName': 'SUPER MEGA BUSSIN\' Pati Zero Sugar',
            'Description': 'SUPER MEGA BUSSIN\' Pati Zero Sugar - 1g',
            'ProductBrand': 'SUPER MEGA BUSSIN\'',
            'Price': '$10',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 75.0% CBD: 0.1%',
            'ProductStrain': 'Pati Zero Sugar',
            'DOH': 'DOH'
        }
    ]
    
    try:
        processor = TemplateProcessor('double', {}, 1.0)
        
        print(f"Processing {len(test_records)} records...")
        result_doc = processor.process_records(test_records)
        
        if result_doc:
            print("✓ Document generated successfully")
            print(f"  Tables: {len(result_doc.tables)}")
            print(f"  Paragraphs: {len(result_doc.paragraphs)}")
            
            if result_doc.tables:
                for i, table in enumerate(result_doc.tables):
                    print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
                    
                    # Check cell content
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            cell_text = cell.text.strip()
                            if cell_text:
                                print(f"    Cell ({row_idx},{col_idx}): '{cell_text[:100]}...'")
                            else:
                                print(f"    Cell ({row_idx},{col_idx}): Empty")
                
                # Save the document for inspection
                result_doc.save("debug_missing_tables_result.docx")
                print("✓ Result document saved as: debug_missing_tables_result.docx")
            else:
                print("❌ No tables in generated document")
                return False
        else:
            print("❌ Document generation failed - returned None")
            return False
            
        return True
        
    except Exception as e:
        print(f"❌ Error in document generation: {e}")
        import traceback
        traceback.print_exc()
        return False

def check_chunk_processing():
    """Check the chunk processing specifically."""
    print("\nChecking Chunk Processing")
    print("=" * 40)
    
    test_records = [
        {
            'ProductName': 'Test Product',
            'Description': 'Test Description',
            'ProductBrand': 'Test Brand',
            'Price': '$10',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 20% CBD: 2%',
            'ProductStrain': 'Test Strain',
            'DOH': 'DOH'
        }
    ]
    
    try:
        processor = TemplateProcessor('double', {}, 1.0)
        
        print("Testing chunk processing directly...")
        result = processor._process_chunk(test_records)
        
        if result:
            print("✓ Chunk processing successful")
            print(f"  Tables: {len(result.tables)}")
            
            if result.tables:
                table = result.tables[0]
                print(f"  Table: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check cell content
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"    Cell ({row_idx},{col_idx}): '{cell_text[:100]}...'")
                        else:
                            print(f"    Cell ({row_idx},{col_idx}): Empty")
            else:
                print("❌ No tables in chunk result")
                return False
        else:
            print("❌ Chunk processing returned None")
            return False
            
        return True
        
    except Exception as e:
        print(f"❌ Error in chunk processing: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run all diagnostic tests."""
    print("Double Template Missing Tables Diagnostic")
    print("=" * 50)
    
    tests = [
        ("Template File Check", check_template_file),
        ("Template Expansion", test_template_expansion),
        ("Chunk Processing", check_chunk_processing),
        ("Document Generation", test_document_generation)
    ]
    
    results = []
    
    for test_name, test_func in tests:
        print(f"\n{'='*20} {test_name} {'='*20}")
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ Test failed with exception: {e}")
            results.append((test_name, False))
    
    print(f"\n{'='*50}")
    print("DIAGNOSTIC SUMMARY:")
    print("=" * 50)
    
    all_passed = True
    for test_name, result in results:
        status = "✅ PASSED" if result else "❌ FAILED"
        print(f"{test_name}: {status}")
        if not result:
            all_passed = False
    
    if all_passed:
        print("\n🎉 All tests passed! Tables should be working correctly.")
    else:
        print("\n⚠️  Some tests failed. Check the output above for issues.")
    
    return all_passed

if __name__ == "__main__":
    main() 