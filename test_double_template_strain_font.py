#!/usr/bin/env python3
"""
Test script to verify that Product Strain font size is correctly set to 1pt in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.unified_font_sizing import get_font_size_by_marker
from docx.shared import Pt
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_strain_font_size():
    """Test that Product Strain font size is 1pt in double template."""
    
    print("Testing Product Strain font size in double template...")
    
    # Test the font sizing function directly
    test_strain_text = "Test Strain Name"
    font_size = get_font_size_by_marker(test_strain_text, 'PRODUCTSTRAIN', 'double', 1.0)
    
    print(f"Strain text: '{test_strain_text}'")
    print(f"Font size: {font_size.pt}pt")
    
    if font_size.pt == 1.0:
        print("✅ SUCCESS: Product Strain font size is correctly set to 1pt")
    else:
        print(f"❌ FAILURE: Product Strain font size is {font_size.pt}pt, expected 1pt")
    
    # Test with template processor
    print("\nTesting with TemplateProcessor...")
    
    test_data = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'Hybrid',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'ProductStrain': 'Test Strain Name',
        'Description': 'Test description text',
        'WeightUnits': '3.5g'
    }
    
    try:
        processor = TemplateProcessor('double', {}, 1.0)
        result_doc = processor.process_records([test_data])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)
            
            print(f"Document generated successfully")
            print(f"Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check font sizes in the first cell
            strain_font_sizes = []
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip() and run.font.size:
                        if 'Test Strain Name' in run.text:
                            strain_font_sizes.append(f"Strain text: '{run.text}' -> {run.font.size.pt}pt")
            
            if strain_font_sizes:
                print("Found Product Strain text with font sizes:")
                for size_info in strain_font_sizes:
                    print(f"  {size_info}")
                    
                # Check if any strain text has 1pt font
                has_1pt = any("1.0pt" in size_info for size_info in strain_font_sizes)
                if has_1pt:
                    print("✅ SUCCESS: Found Product Strain text with 1pt font size")
                else:
                    print("❌ FAILURE: No Product Strain text found with 1pt font size")
            else:
                print("⚠️  WARNING: No Product Strain text found in the generated document")
                
        else:
            print("❌ FAILURE: Failed to generate document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")

if __name__ == "__main__":
    test_strain_font_size() 