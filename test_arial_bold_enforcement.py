#!/usr/bin/env python3
"""
Test script to verify that all fonts are Arial Bold
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.docx_formatting import enforce_arial_bold_all_text
from docx import Document
from docx.shared import Pt

def test_arial_bold_enforcement():
    """Test that all fonts are Arial Bold"""
    print("🧪 Testing Arial Bold Font Enforcement")
    print("=" * 50)
    
    # Create a test record
    test_data = {
        'ProductName': 'Test Product',
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'Sativa',
        'THC_CBD': 'THC:|BR|CBD:',
        'Ratio_or_THC_CBD': 'THC:|BR|CBD:',
        'Ratio': 'THC:|BR|CBD:',
        'WeightUnits': '1g',
        'ProductStrain': 'Test Strain',
        'DOH': 'Yes'
    }
    
    # Test vertical template
    print("\n1️⃣ Testing Vertical Template...")
    tp_vertical = TemplateProcessor(template_type='vertical', font_scheme='Arial')
    
    # Process the record
    doc = tp_vertical.process_records([test_data])
    
    if not doc:
        print("❌ No document generated")
        return False
    
    # Check fonts before enforcement
    print("\n📋 Checking fonts BEFORE enforcement...")
    non_arial_runs = []
    non_bold_runs = []
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():  # Only check non-empty runs
                            if run.font.name != "Arial":
                                non_arial_runs.append(f"'{run.text}': {run.font.name}")
                            if not run.font.bold:
                                non_bold_runs.append(f"'{run.text}': bold={run.font.bold}")
    
    if non_arial_runs:
        print(f"⚠️  Found {len(non_arial_runs)} runs with non-Arial fonts:")
        for run_info in non_arial_runs[:5]:  # Show first 5
            print(f"   {run_info}")
    else:
        print("✅ All runs are Arial")
    
    if non_bold_runs:
        print(f"⚠️  Found {len(non_bold_runs)} runs that are not bold:")
        for run_info in non_bold_runs[:5]:  # Show first 5
            print(f"   {run_info}")
    else:
        print("✅ All runs are bold")
    
    # Apply enforcement
    print("\n🔧 Applying Arial Bold enforcement...")
    enforce_arial_bold_all_text(doc)
    
    # Check fonts after enforcement
    print("\n📋 Checking fonts AFTER enforcement...")
    non_arial_runs_after = []
    non_bold_runs_after = []
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip():  # Only check non-empty runs
                            if run.font.name != "Arial":
                                non_arial_runs_after.append(f"'{run.text}': {run.font.name}")
                            if not run.font.bold:
                                non_bold_runs_after.append(f"'{run.text}': bold={run.font.bold}")
    
    if non_arial_runs_after:
        print(f"❌ Found {len(non_arial_runs_after)} runs with non-Arial fonts after enforcement:")
        for run_info in non_arial_runs_after[:5]:  # Show first 5
            print(f"   {run_info}")
        return False
    else:
        print("✅ All runs are Arial after enforcement")
    
    if non_bold_runs_after:
        print(f"❌ Found {len(non_bold_runs_after)} runs that are not bold after enforcement:")
        for run_info in non_bold_runs_after[:5]:  # Show first 5
            print(f"   {run_info}")
        return False
    else:
        print("✅ All runs are bold after enforcement")
    
    # Test horizontal template
    print("\n2️⃣ Testing Horizontal Template...")
    tp_horizontal = TemplateProcessor(template_type='horizontal', font_scheme='Arial')
    
    doc_horizontal = tp_horizontal.process_records([test_data])
    
    if doc_horizontal:
        # Apply enforcement
        enforce_arial_bold_all_text(doc_horizontal)
        
        # Check fonts
        non_arial_count = 0
        non_bold_count = 0
        
        for table in doc_horizontal.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                if run.font.name != "Arial":
                                    non_arial_count += 1
                                if not run.font.bold:
                                    non_bold_count += 1
        
        if non_arial_count == 0 and non_bold_count == 0:
            print("✅ Horizontal template: All fonts are Arial Bold")
        else:
            print(f"❌ Horizontal template: {non_arial_count} non-Arial, {non_bold_count} non-bold")
            return False
    
    # Summary
    print("\n" + "=" * 50)
    print("🎉 SUCCESS: All fonts are Arial Bold across all templates")
    return True

if __name__ == "__main__":
    success = test_arial_bold_enforcement()
    sys.exit(0 if success else 1) 