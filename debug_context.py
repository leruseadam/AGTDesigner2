#!/usr/bin/env python3
"""
Debug script to examine context building for templates
"""

import sys
import os
from pathlib import Path

# Add the src directory to the Python path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def debug_context_building():
    """Debug the context building process."""
    print("=== Debugging Context Building ===")
    
    # Create test records
    test_records = [
        {
            'ProductName': 'Test Product 1',
            'ProductBrand': 'Test Brand 1',
            'Price': '$25.99',
            'Description': 'Test Description 1',
            'WeightUnits': '3.5g',
            'Lineage': 'Test Lineage 1',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'DOH': 'YES',
            'ProductStrain': 'Test Strain 1',
            'ProductType': 'flower',
            'Vendor': 'Test Vendor 1'
        }
    ]
    
    try:
        # Create processor for double template
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Get the expanded template buffer
        buffer = processor._expanded_template_buffer
        buffer.seek(0)
        doc = Document(buffer)
        
        # Build context for the first record
        label_context = processor._build_label_context(test_records[0], doc)
        
        print("\n=== Context for Label1 ===")
        for key, value in label_context.items():
            print(f"{key}: {repr(value)}")
        
        # Check what placeholders are in the template
        print("\n=== Placeholders in Template ===")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if '{{Label' in text:
                            print(f"Found placeholder: {text}")
        
        # Create full context
        context = {}
        for i, record in enumerate(test_records):
            label_context = processor._build_label_context(record, doc)
            context[f'Label{i+1}'] = label_context
        
        # Fill remaining labels with empty context
        for i in range(len(test_records), processor.chunk_size):
            context[f'Label{i+1}'] = {}
        
        print(f"\n=== Full Context Keys ===")
        for label_key, label_context in context.items():
            print(f"{label_key}: {list(label_context.keys())}")
        
        # Test DocxTemplate rendering
        print("\n=== Testing DocxTemplate Rendering ===")
        from docxtpl import DocxTemplate
        
        # Reset buffer
        buffer.seek(0)
        docx_template = DocxTemplate(buffer)
        
        print("Context being passed to DocxTemplate:")
        for label_key, label_context in context.items():
            print(f"  {label_key}: {label_context}")
        
        # Try to render
        try:
            docx_template.render(context)
            print("✅ DocxTemplate.render() completed successfully")
            
            # Check the rendered document
            rendered_doc = docx_template.get_docx()
            print("\n=== Rendered Document Content ===")
            for table in rendered_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            if text.strip():
                                print(f"Cell content: {repr(text)}")
        
        except Exception as e:
            print(f"❌ DocxTemplate.render() failed: {e}")
            import traceback
            traceback.print_exc()
        
        return True
        
    except Exception as e:
        print(f"❌ Error during debug: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    debug_context_building() 