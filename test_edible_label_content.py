#!/usr/bin/env python3
"""
Test script to verify that edibles show brand instead of CBD in the label content.
"""

import requests
import json
import time
import zipfile
import io
from docx import Document

def test_edible_label_content():
    """Test that edibles show brand instead of CBD in the label content."""
    base_url = 'http://127.0.0.1:9090'
    
    print("🧪 Testing Edible Label Content")
    print("=" * 50)
    
    # Step 1: Get available tags
    print("\n1. Getting available tags...")
    
    try:
        response = requests.get(f'{base_url}/api/available-tags')
        if response.status_code == 200:
            available_tags = response.json()
            print(f"✅ Got {len(available_tags)} available tags")
        else:
            print(f"❌ Failed to get available tags: {response.status_code}")
            return
    except Exception as e:
        print(f"❌ Error getting available tags: {e}")
        return
    
    # Step 2: Find edible products with CBD lineage
    print("\n2. Finding edible products with CBD lineage...")
    
    edible_cbd_products = []
    for tag in available_tags:
        product_type = tag.get('Product Type*', '').lower()
        lineage = tag.get('Lineage', '').upper()
        if ('edible' in product_type or product_type in ['tincture', 'topical', 'capsule']) and 'CBD' in lineage:
            edible_cbd_products.append(tag)
    
    print(f"✅ Found {len(edible_cbd_products)} edible products with CBD lineage")
    
    if not edible_cbd_products:
        print("❌ No edible products with CBD lineage found to test")
        return
    
    # Step 3: Select some edible products with CBD lineage
    print("\n3. Selecting edible products with CBD lineage for testing...")
    
    # Take first 3 edible products with CBD lineage
    selected_edibles = edible_cbd_products[:3]
    selected_names = [tag['Product Name*'] for tag in selected_edibles]
    
    print(f"✅ Selected {len(selected_edibles)} edible products with CBD lineage:")
    for tag in selected_edibles:
        brand = tag.get('Product Brand', 'Unknown')
        lineage = tag.get('Lineage', 'Unknown')
        print(f"   - {tag['Product Name*']} (Brand: {brand}, Lineage: {lineage})")
    
    # Step 4: Generate labels and check content
    print("\n4. Generating labels to check content...")
    
    try:
        generate_data = {
            'selected_tags': selected_names,
            'template_type': 'horizontal'
        }
        
        response = requests.post(f'{base_url}/api/generate', json=generate_data)
        
        if response.status_code == 200:
            print(f"✅ Generation successful")
            
            # Save the document to examine its content
            docx_content = response.content
            doc = Document(io.BytesIO(docx_content))
            
            # Check for lineage/brand content in the document
            found_brand_instead_of_cbd = False
            found_cbd_instead_of_brand = False
            
            print("\n5. Analyzing document content...")
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.upper()
                        
                        # Check if we find brand names instead of CBD
                        for tag in selected_edibles:
                            brand = tag.get('Product Brand', '').upper()
                            if brand and brand in cell_text and 'CBD' not in cell_text:
                                print(f"✅ Found brand '{brand}' in cell text (no CBD)")
                                found_brand_instead_of_cbd = True
                        
                        # Check if we still find CBD where we shouldn't
                        if 'CBD' in cell_text:
                            # Check if this is in a lineage context
                            if 'LINEAGE' in cell_text or any(brand.upper() in cell_text for brand in [tag.get('Product Brand', '') for tag in selected_edibles]):
                                print(f"⚠️  Found 'CBD' in cell text: '{cell_text[:100]}...'")
                                found_cbd_instead_of_brand = True
            
            # Summary
            print("\n" + "=" * 50)
            print("CONTENT ANALYSIS SUMMARY")
            print("=" * 50)
            
            if found_brand_instead_of_cbd and not found_cbd_instead_of_brand:
                print("✅ SUCCESS: Edibles show brand instead of CBD in label content")
                return True
            elif found_brand_instead_of_cbd and found_cbd_instead_of_brand:
                print("⚠️  PARTIAL: Some edibles show brand, but some still show CBD")
                return False
            else:
                print("❌ FAILED: Edibles still show CBD instead of brand in label content")
                return False
                
        else:
            print(f"❌ Generation failed: {response.status_code}")
            print(f"Response: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Error during generation: {e}")
        return False

def test_non_edible_label_content():
    """Test that non-edibles still show lineage in the label content."""
    base_url = 'http://127.0.0.1:9090'
    
    print("\n🧪 Testing Non-Edible Label Content")
    print("=" * 50)
    
    # Step 1: Get available tags
    print("\n1. Getting available tags...")
    
    try:
        response = requests.get(f'{base_url}/api/available-tags')
        if response.status_code == 200:
            available_tags = response.json()
            print(f"✅ Got {len(available_tags)} available tags")
        else:
            print(f"❌ Failed to get available tags: {response.status_code}")
            return
    except Exception as e:
        print(f"❌ Error getting available tags: {e}")
        return
    
    # Step 2: Find non-edible products
    print("\n2. Finding non-edible products...")
    
    non_edible_products = []
    for tag in available_tags:
        product_type = tag.get('Product Type*', '').lower()
        if 'edible' not in product_type and product_type not in ['tincture', 'topical', 'capsule']:
            non_edible_products.append(tag)
    
    print(f"✅ Found {len(non_edible_products)} non-edible products")
    
    if not non_edible_products:
        print("❌ No non-edible products found to test")
        return
    
    # Step 3: Select some non-edible products
    print("\n3. Selecting non-edible products for testing...")
    
    # Take first 3 non-edible products
    selected_non_edibles = non_edible_products[:3]
    selected_names = [tag['Product Name*'] for tag in selected_non_edibles]
    
    print(f"✅ Selected {len(selected_non_edibles)} non-edible products:")
    for tag in selected_non_edibles:
        brand = tag.get('Product Brand', 'Unknown')
        lineage = tag.get('Lineage', 'Unknown')
        print(f"   - {tag['Product Name*']} (Brand: {brand}, Lineage: {lineage})")
    
    # Step 4: Generate labels and check content
    print("\n4. Generating labels to check content...")
    
    try:
        generate_data = {
            'selected_tags': selected_names,
            'template_type': 'horizontal'
        }
        
        response = requests.post(f'{base_url}/api/generate', json=generate_data)
        
        if response.status_code == 200:
            print(f"✅ Generation successful")
            
            # Save the document to examine its content
            docx_content = response.content
            doc = Document(io.BytesIO(docx_content))
            
            # Check for lineage content in the document
            found_lineage = False
            
            print("\n5. Analyzing document content...")
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.upper()
                        
                        # Check if we find lineage information
                        if any(lineage in cell_text for lineage in ['SATIVA', 'INDICA', 'HYBRID', 'CBD', 'MIXED']):
                            print(f"✅ Found lineage in cell text: '{cell_text[:100]}...'")
                            found_lineage = True
            
            # Summary
            print("\n" + "=" * 50)
            print("CONTENT ANALYSIS SUMMARY")
            print("=" * 50)
            
            if found_lineage:
                print("✅ SUCCESS: Non-edibles still show lineage in label content")
                return True
            else:
                print("❌ FAILED: Non-edibles missing lineage in label content")
                return False
                
        else:
            print(f"❌ Generation failed: {response.status_code}")
            print(f"Response: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Error during generation: {e}")
        return False

if __name__ == "__main__":
    print("Testing Edible vs Non-Edible Label Content")
    print("=" * 60)
    
    # Test edible label content
    edible_result = test_edible_label_content()
    
    # Test non-edible label content
    non_edible_result = test_non_edible_label_content()
    
    # Summary
    print("\n" + "=" * 60)
    print("FINAL SUMMARY")
    print("=" * 60)
    
    if edible_result and non_edible_result:
        print("✅ SUCCESS: Edibles show brand, non-edibles show lineage")
    elif edible_result:
        print("⚠️  PARTIAL: Edibles work, but non-edibles may have issues")
    elif non_edible_result:
        print("⚠️  PARTIAL: Non-edibles work, but edibles may have issues")
    else:
        print("❌ FAILED: Both edible and non-edible label content have issues") 