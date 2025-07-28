#!/usr/bin/env python3
"""
Test script to verify that brand centering is working correctly for the double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def test_double_brand_centering():
    """Test that brand content is properly centered in double template."""
    print("Testing Double Template Brand Centering")
    print("=" * 50)
    
    # Test data for non-classic product type (should be centered)
    test_record = {
        'Description': 'Test Product Description',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand Name',
        'Price': '$25.99',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'DOH',
        'ProductType': 'edible (solid)'  # Non-classic type
    }
    
    # Test double template
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return False
    
    # Check for brand centering
    brand_centered = False
    total_paragraphs = 0
    
    print(f"Document tables: {len(result_doc.tables)}")
    
    for table_idx, table in enumerate(result_doc.tables):
        print(f"Table {table_idx}: {len(table.rows)} rows, {len(table.columns)} columns")
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    total_paragraphs += 1
                    paragraph_text = paragraph.text.strip()
                    
                    # Check if this paragraph contains brand content
                    if ('Test Brand Name' in paragraph_text or 
                        'PRODUCTBRAND_CENTER' in paragraph_text or
                        'PRODUCTBRAND' in paragraph_text):
                        
                        print(f"  Found brand content in table {table_idx}, row {row_idx}, cell {cell_idx}, paragraph {para_idx}")
                        print(f"    Text: '{paragraph_text[:50]}...'")
                        print(f"    Alignment: {paragraph.alignment}")
                        
                        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            brand_centered = True
                            print(f"    ✓ Brand content is CENTERED")
                        else:
                            print(f"    ✗ Brand content is NOT centered (alignment: {paragraph.alignment})")
    
    print(f"\nTotal paragraphs checked: {total_paragraphs}")
    
    if brand_centered:
        print("✓ SUCCESS: Brand content is properly centered in double template")
        return True
    else:
        print("✗ FAILURE: Brand content is not centered in double template")
        return False

def test_classic_type_brand_alignment():
    """Test that classic types are NOT centered in double template."""
    print("\nTesting Classic Type Brand Alignment")
    print("=" * 50)
    
    # Test data for classic product type (should NOT be centered)
    test_record = {
        'Description': 'Test Flower Product',
        'WeightUnits': '3.5g',
        'ProductBrand': 'Test Flower Brand',
        'Price': '$45.99',
        'Lineage': 'SATIVA',
        'THC_CBD': 'THC: 25% CBD: 1%',
        'ProductStrain': 'Test Flower Strain',
        'DOH': 'DOH',
        'ProductType': 'flower'  # Classic type
    }
    
    # Test double template
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return False
    
    # Check for brand alignment
    brand_found = False
    brand_not_centered = False
    
    for table in result_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text.strip()
                    
                    # Check if this paragraph contains brand content
                    if ('Test Flower Brand' in paragraph_text or 
                        'PRODUCTBRAND' in paragraph_text):
                        
                        brand_found = True
                        print(f"Found brand content: '{paragraph_text[:50]}...'")
                        print(f"Alignment: {paragraph.alignment}")
                        
                        # For classic types, we expect it to NOT be centered
                        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                            brand_not_centered = True
                            print("✓ Brand content is NOT centered (as expected for classic type)")
                        else:
                            print("✗ Brand content is centered (unexpected for classic type)")
    
    if brand_found and brand_not_centered:
        print("✓ SUCCESS: Classic type brand content is properly NOT centered")
        return True
    else:
        print("✗ FAILURE: Classic type brand alignment test failed")
        return False

if __name__ == "__main__":
    print("Testing Double Template Brand Centering Fix")
    print("=" * 60)
    
    # Test non-classic type centering
    test1_passed = test_double_brand_centering()
    
    # Test classic type alignment
    test2_passed = test_classic_type_brand_alignment()
    
    print("\n" + "=" * 60)
    if test1_passed and test2_passed:
        print("✓ ALL TESTS PASSED: Brand centering is working correctly")
    else:
        print("✗ SOME TESTS FAILED: Brand centering needs attention")
    
    print("=" * 60) 