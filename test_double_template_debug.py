#!/usr/bin/env python3
"""
Debug test for double template placeholder replacement issue.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.data.excel_processor import ExcelProcessor
from src.core.generation.template_processor import TemplateProcessor

def test_double_template_debug():
    """Debug the double template placeholder replacement issue."""
    print("🔍 Debugging Double Template Placeholder Replacement...")
    
    # Initialize processor
    processor = ExcelProcessor()
    
    # Try to load a default file
    from src.core.data.excel_processor import get_default_upload_file
    default_file = get_default_upload_file()
    
    if not default_file or not os.path.exists(default_file):
        print("❌ No default file found")
        return
    
    print(f"📁 Using default file: {default_file}")
    
    # Load the data
    try:
        success = processor.load_file(default_file)
        if not success:
            print("❌ Failed to load file")
            return
        
        # Get the records from the processor
        records = processor.get_available_tags()
        if not records:
            print("❌ No records found in file")
            return
        
        print(f"✅ Loaded {len(records)} records")
        
        # Get the first few records for testing
        test_records = records[:3]
        
        # Test the template processor
        template_processor = TemplateProcessor('double', 'arial')
        
        print("\n🔍 Testing context building for each record:")
        print("=" * 80)
        
        for i, record in enumerate(test_records):
            print(f"\n📋 Record {i+1}:")
            print(f"  Product Name: {record.get('ProductName', 'N/A')}")
            print(f"  Product Type: {record.get('ProductType', 'N/A')}")
            print(f"  Description: {record.get('Description', 'N/A')}")
            print(f"  Weight Units: {record.get('WeightUnits', 'N/A')}")
            print(f"  Price: {record.get('Price', 'N/A')}")
            print(f"  Lineage: {record.get('Lineage', 'N/A')}")
            print(f"  Product Brand: {record.get('ProductBrand', 'N/A')}")
            print(f"  DOH: {record.get('DOH', 'N/A')}")
            print(f"  Ratio: {record.get('Ratio_or_THC_CBD', 'N/A')}")
            
            # Test context building
            from docx import Document
            test_doc = Document()
            context = template_processor._build_label_context(record, test_doc)
            
            print(f"\n  🔧 Built Context:")
            for key, value in context.items():
                if key in ['DescAndWeight', 'Price', 'Lineage', 'ProductBrand', 'DOH', 'Ratio_or_THC_CBD']:
                    print(f"    {key}: {value}")
        
        print("\n🔍 Testing full template generation:")
        print("=" * 80)
        
        # Test the full process
        result = template_processor._process_chunk(test_records)
        
        if result:
            print("✅ Template generation successful!")
            
            # Save the result to check the content
            output_file = "debug_double_template_output.docx"
            result.save(output_file)
            print(f"📄 Saved output to: {output_file}")
            
            # Check for placeholders in the output
            placeholder_count = 0
            for table in result.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '{{' in paragraph.text and '}}' in paragraph.text:
                                placeholder_count += 1
                                print(f"⚠️  Found placeholder in cell: {paragraph.text[:100]}...")
            
            if placeholder_count == 0:
                print("✅ No placeholders found in output!")
            else:
                print(f"❌ Found {placeholder_count} cells with placeholders")
        else:
            print("❌ Template generation failed")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_double_template_debug() 