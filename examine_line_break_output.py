#!/usr/bin/env python3
"""
Script to examine the generated document and check for line breaks.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document

def examine_line_break_output():
    """Examine the generated document for line breaks."""
    
    print("=== Examining Line Break Output ===")
    
    try:
        # Load the generated document
        doc = Document("test_line_break_output.docx")
        
        print(f"Document has {len(doc.tables)} tables")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx + 1} ({len(table.rows)} rows, {len(table.columns)} columns):")
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"  Cell ({row_idx}, {cell_idx}):")
                        print(f"    Raw text: '{cell_text}'")
                        
                        # Check for line breaks
                        if '\n' in cell_text:
                            print(f"    ✅ Contains line breaks!")
                            lines = cell_text.split('\n')
                            for i, line in enumerate(lines):
                                print(f"      Line {i+1}: '{line}'")
                        else:
                            print(f"    No line breaks found")
                        
                        # Check for |BR| markers (should not be present)
                        if '|BR|' in cell_text:
                            print(f"    ❌ Still contains |BR| markers!")
                        
                        # Check paragraph structure
                        print(f"    Paragraphs: {len(cell.paragraphs)}")
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            print(f"      Paragraph {para_idx + 1}: '{paragraph.text}'")
                            print(f"        Runs: {len(paragraph.runs)}")
                            for run_idx, run in enumerate(paragraph.runs):
                                print(f"          Run {run_idx + 1}: '{run.text}' (font: {run.font.name}, size: {run.font.size})")
                        
                        print()
        
    except FileNotFoundError:
        print("❌ test_line_break_output.docx not found. Run test_real_line_break_generation.py first.")
    except Exception as e:
        print(f"❌ Error examining document: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    examine_line_break_output() 