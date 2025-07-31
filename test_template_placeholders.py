#!/usr/bin/env python3
"""
Script to check what placeholders are in the template files
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from pathlib import Path
import re

def extract_placeholders_from_docx(docx_path):
    """Extract all placeholders from a DOCX file"""
    try:
        doc = Document(docx_path)
        placeholders = set()
        
        # Check all paragraphs
        for para in doc.paragraphs:
            text = para.text
            # Find all {{...}} placeholders
            matches = re.findall(r'\{\{[^}]+\}\}', text)
            placeholders.update(matches)
        
        # Check all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        matches = re.findall(r'\{\{[^}]+\}\}', text)
                        placeholders.update(matches)
        
        return sorted(list(placeholders))
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return []

def main():
    template_dir = Path("src/core/generation/templates")
    
    print("=== Checking Template Placeholders ===")
    
    for template_file in template_dir.glob("*.docx"):
        if template_file.name.startswith("~$"):  # Skip temp files
            continue
            
        print(f"\n--- {template_file.name} ---")
        placeholders = extract_placeholders_from_docx(template_file)
        
        if placeholders:
            for placeholder in placeholders:
                print(f"  {placeholder}")
        else:
            print("  No placeholders found")
        
        # Check specifically for ProductVendor
        productvendor_placeholders = [p for p in placeholders if 'ProductVendor' in p]
        if productvendor_placeholders:
            print(f"  ✅ ProductVendor placeholders found: {productvendor_placeholders}")
        else:
            print("  ❌ No ProductVendor placeholders found")

if __name__ == "__main__":
    main() 