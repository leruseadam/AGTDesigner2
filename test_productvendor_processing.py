#!/usr/bin/env python3
"""
Comprehensive test to debug ProductVendor processing
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from pathlib import Path

def test_productvendor_processing():
    """Test ProductVendor processing with a template that has the placeholder"""
    
    # Create a test record with vendor data
    test_record = {
        'ProductName': 'Test Product',
        'Vendor': 'Test Vendor Company',
        'ProductBrand': 'Test Brand',
        'Price': '$10.99',
        'Description': 'Test Description',
        'Lineage': 'Test Lineage',
        'ProductStrain': 'Test Strain',
        'ProductType': 'flower'
    }
    
    print("=== Testing ProductVendor Processing ===")
    print(f"Test record vendor: '{test_record['Vendor']}'")
    
    # Test with horizontal template (which has ProductVendor placeholder)
    processor = TemplateProcessor('horizontal', {}, 1.0)
    
    # Process a single record
    result = processor.process_records([test_record])
    
    if result:
        print("✅ Processing completed successfully")
        
        # Save the result to examine it
        output_path = "test_productvendor_output.docx"
        result.save(output_path)
        
        print(f"✅ Output saved to {output_path}")
        
        # Examine the output document
        doc = Document(output_path)
        
        # Check all text in the document
        print("\n=== Document Content Analysis ===")
        
        # Check paragraphs outside tables
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"Paragraph {i}: '{para.text}'")
        
        # Check tables
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"  Cell [{row_idx},{col_idx}]: '{cell_text}'")
                        
                        # Check if ProductVendor content is present
                        if 'Test Vendor Company' in cell_text:
                            print(f"    ✅ Found vendor text in cell [{row_idx},{col_idx}]")
                        elif 'PRODUCTVENDOR_START' in cell_text:
                            print(f"    ⚠️  Found ProductVendor markers but no content in cell [{row_idx},{col_idx}]")
                        elif 'ProductVendor' in cell_text:
                            print(f"    ⚠️  Found ProductVendor placeholder in cell [{row_idx},{col_idx}]")
        
        # Check for any ProductVendor markers
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        if 'PRODUCTVENDOR_START' in all_text:
            print("\n✅ ProductVendor markers found in document")
        else:
            print("\n❌ No ProductVendor markers found in document")
            
        if 'Test Vendor Company' in all_text:
            print("✅ Vendor text found in document")
        else:
            print("❌ Vendor text not found in document")
            
        if '{{Label1.ProductVendor}}' in all_text:
            print("⚠️  ProductVendor placeholder still present (not replaced)")
        else:
            print("✅ ProductVendor placeholder was replaced")
            
    else:
        print("❌ Processing failed - no result returned")

def test_context_building():
    """Test the context building specifically"""
    print("\n=== Testing Context Building ===")
    
    test_record = {
        'ProductName': 'Test Product',
        'Vendor': 'Test Vendor Company',
        'ProductBrand': 'Test Brand',
        'Price': '$10.99',
        'Description': 'Test Description',
        'Lineage': 'Test Lineage',
        'ProductStrain': 'Test Strain',
        'ProductType': 'flower'
    }
    
    processor = TemplateProcessor('horizontal', {}, 1.0)
    doc = Document()
    
    context = processor._build_label_context(test_record, doc)
    
    print(f"ProductVendor in context: '{context.get('ProductVendor', 'NOT_FOUND')}'")
    
    if 'ProductVendor' in context:
        vendor_value = context['ProductVendor']
        if 'PRODUCTVENDOR_START' in vendor_value and 'PRODUCTVENDOR_END' in vendor_value:
            print("✅ ProductVendor is properly wrapped with markers")
            # Extract the actual content
            start_marker = "PRODUCTVENDOR_START"
            end_marker = "PRODUCTVENDOR_END"
            start_idx = vendor_value.find(start_marker) + len(start_marker)
            end_idx = vendor_value.find(end_marker)
            content = vendor_value[start_idx:end_idx]
            print(f"✅ Extracted content: '{content}'")
        else:
            print("❌ ProductVendor is not properly wrapped with markers")
    else:
        print("❌ ProductVendor not found in context")

if __name__ == "__main__":
    test_context_building()
    test_productvendor_processing() 