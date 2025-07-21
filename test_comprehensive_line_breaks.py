#!/usr/bin/env python3
"""
Comprehensive test for line break conversion.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.core.generation.template_processor import TemplateProcessor

def test_comprehensive_line_breaks():
    """Comprehensive test for line break conversion."""
    
    print("=== Comprehensive Line Break Test ===")
    
    # Test 1: Direct paragraph manipulation
    print("\n1. Testing direct paragraph manipulation...")
    doc1 = Document()
    table1 = doc1.add_table(rows=1, cols=1)
    cell1 = table1.cell(0, 0)
    
    # Add content with |BR| markers
    test_content = "100mg THC|BR|50mg CBD|BR|5mg CBG"
    cell1.text = test_content
    
    processor = TemplateProcessor('vertical', {}, 1.0)
    paragraph1 = cell1.paragraphs[0]
    
    print(f"Before: '{paragraph1.text}'")
    processor._convert_br_markers_to_line_breaks(paragraph1)
    print(f"After: '{paragraph1.text}'")
    
    # Check runs
    runs = list(paragraph1.runs)
    print(f"Runs: {len(runs)}")
    for i, run in enumerate(runs):
        print(f"  Run {i+1}: '{run.text}'")
    
    # Test 2: Simple document with markers
    print("\n2. Testing simple document with markers...")
    doc2 = Document()
    table2 = doc2.add_table(rows=1, cols=1)
    cell2 = table2.cell(0, 0)
    
    # Add content with markers and |BR|
    cell2.text = "RATIO_START100mg THC|BR|50mg CBDRATIO_END"
    
    print(f"Before processing: '{cell2.text}'")
    processor._post_process_and_replace_content(doc2)
    print(f"After processing: '{cell2.text}'")
    
    # Check the paragraph structure
    for para_idx, paragraph in enumerate(cell2.paragraphs):
        print(f"Paragraph {para_idx + 1}: '{paragraph.text}'")
        for run_idx, run in enumerate(paragraph.runs):
            print(f"  Run {run_idx + 1}: '{run.text}'")
    
    # Test 3: Test with actual edible product data
    print("\n3. Testing with actual edible product data...")
    
    test_records = [
        {
            'Product Name*': 'Test Edible',
            'Product Brand': 'Test Brand',
            'Product Type*': 'edible (solid)',
            'Lineage': 'MIXED',
            'Price': '$25.99',
            'Ratio': '100mg THC 50mg CBD 5mg CBG',
            'Description': 'Test description',
            'Weight Units': '4oz'
        }
    ]
    
    processor2 = TemplateProcessor('vertical', {}, 1.0)
    
    try:
        doc3 = processor2.process_records(test_records)
        
        if doc3:
            print("✅ Document generated successfully!")
            
            # Save for inspection
            doc3.save("test_comprehensive_output.docx")
            print("✅ Saved as test_comprehensive_output.docx")
            
            # Examine all cells
            print("\nExamining all cells in generated document...")
            for table_idx, table in enumerate(doc3.tables):
                print(f"\nTable {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"  Cell ({row_idx}, {cell_idx}): '{cell_text}'")
                            
                            # Check for line breaks
                            if '\n' in cell_text:
                                print(f"    ✅ Contains line breaks!")
                                lines = cell_text.split('\n')
                                for i, line in enumerate(lines):
                                    print(f"      Line {i+1}: '{line}'")
                            
                            # Check for |BR| markers
                            if '|BR|' in cell_text:
                                print(f"    ❌ Still contains |BR| markers!")
        else:
            print("❌ Failed to generate document")
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    test_comprehensive_line_breaks() 