#!/usr/bin/env python3
"""
Debug test to see exactly what content is being generated for edibles.
"""

import requests
import json
import time
import zipfile
import io
from docx import Document

def test_edible_debug_content():
    """Debug test to see exactly what content is being generated for edibles."""
    base_url = 'http://127.0.0.1:9090'
    
    print("🔍 Debug Testing Edible Label Content")
    print("=" * 50)
    
    # Step 1: Get available tags
    print("\n1. Getting available tags...")
    
    try:
        response = requests.get(f'{base_url}/api/available-tags')
        if response.status_code == 200:
            available_tags = response.json()
            print(f"✅ Got {len(available_tags)} available tags")
        else:
            print(f"❌ Failed to get available tags: {response.status_code}")
            return
    except Exception as e:
        print(f"❌ Error getting available tags: {e}")
        return
    
    # Step 2: Find edible products with CBD lineage
    print("\n2. Finding edible products with CBD lineage...")
    
    edible_cbd_products = []
    for tag in available_tags:
        product_type = tag.get('Product Type*', '').lower()
        lineage = tag.get('Lineage', '').upper()
        if ('edible' in product_type or product_type in ['tincture', 'topical', 'capsule']) and 'CBD' in lineage:
            edible_cbd_products.append(tag)
    
    print(f"✅ Found {len(edible_cbd_products)} edible products with CBD lineage")
    
    if not edible_cbd_products:
        print("❌ No edible products with CBD lineage found to test")
        return
    
    # Step 3: Select one edible product for detailed analysis
    print("\n3. Selecting one edible product for detailed analysis...")
    
    selected_edible = edible_cbd_products[0]
    selected_name = selected_edible['Product Name*']
    
    print(f"✅ Selected edible product:")
    print(f"   - Name: {selected_edible['Product Name*']}")
    print(f"   - Brand: {selected_edible.get('Product Brand', 'Unknown')}")
    print(f"   - Lineage: {selected_edible.get('Lineage', 'Unknown')}")
    print(f"   - Product Type: {selected_edible.get('Product Type*', 'Unknown')}")
    print(f"   - Product Strain: {selected_edible.get('Product Strain', 'Unknown')}")
    
    # Step 4: Generate labels and analyze content
    print("\n4. Generating labels to analyze content...")
    
    try:
        generate_data = {
            'selected_tags': [selected_name],
            'template_type': 'horizontal'
        }
        
        response = requests.post(f'{base_url}/api/generate', json=generate_data)
        
        if response.status_code == 200:
            print(f"✅ Generation successful")
            
            # Save the document to examine its content
            docx_content = response.content
            doc = Document(io.BytesIO(docx_content))
            
            print("\n5. Analyzing document content in detail...")
            
            for table_idx, table in enumerate(doc.tables):
                print(f"\n--- Table {table_idx + 1} ---")
                for row_idx, row in enumerate(table.rows):
                    print(f"\nRow {row_idx + 1}:")
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"  Cell {cell_idx + 1}: '{cell_text}'")
                            
                            # Check for specific content
                            if 'CONSTELLATION' in cell_text.upper() or 'GRAVITY' in cell_text.upper():
                                print(f"    ✅ Found brand in cell {cell_idx + 1}")
                            
                            if 'CBD' in cell_text.upper():
                                print(f"    ⚠️  Found 'CBD' in cell {cell_idx + 1}")
                                
                            if 'BLEND' in cell_text.upper():
                                print(f"    ⚠️  Found 'BLEND' in cell {cell_idx + 1}")
                                
                            # Check for lineage-related content
                            if any(lineage in cell_text.upper() for lineage in ['SATIVA', 'INDICA', 'HYBRID', 'MIXED']):
                                print(f"    ⚠️  Found lineage in cell {cell_idx + 1}")
            
            # Summary
            print("\n" + "=" * 50)
            print("DETAILED CONTENT ANALYSIS")
            print("=" * 50)
            
            found_brand = False
            found_cbd = False
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.upper()
                        if 'CONSTELLATION' in cell_text or 'GRAVITY' in cell_text:
                            found_brand = True
                        if 'CBD' in cell_text:
                            found_cbd = True
            
            if found_brand and not found_cbd:
                print("✅ SUCCESS: Found brand, no CBD")
            elif found_brand and found_cbd:
                print("⚠️  PARTIAL: Found brand but also found CBD")
            else:
                print("❌ FAILED: No brand found")
                
        else:
            print(f"❌ Generation failed: {response.status_code}")
            print(f"Response: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Error during generation: {e}")
        return False

if __name__ == "__main__":
    test_edible_debug_content() 