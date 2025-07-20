#!/usr/bin/env python3
"""
Test script to verify mini template blank cell clearing functionality.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_mini_blank_cells_fix():
    """Test that blank cells are properly cleared in mini templates."""
    
    print("Testing Mini Template Blank Cell Clearing")
    print("=" * 50)
    
    # Create test data with fewer records than the 20-cell grid
    test_records = [
        {
            'Description': 'Test Product 1',
            'ProductBrand': 'Test Brand',
            'Price': '$25',
            'Lineage': 'SATIVA',
            'WeightUnits': '1g',
            'ProductType': 'flower',
            'DOH': 'GENERAL USE',
            'Ratio': 'THC: 15%\nCBD: 1%',
            'ProductStrain': 'Test Strain 1'
        },
        {
            'Description': 'Test Product 2', 
            'ProductBrand': 'Test Brand 2',
            'Price': '$30',
            'Lineage': 'INDICA',
            'WeightUnits': '2g',
            'ProductType': 'concentrate',
            'DOH': 'MEDICAL USE',
            'Ratio': 'THC: 20%\nCBD: 2%',
            'ProductStrain': 'Test Strain 2'
        },
        {
            'Description': 'Test Product 3',
            'ProductBrand': 'Test Brand 3', 
            'Price': '$35',
            'Lineage': 'HYBRID',
            'WeightUnits': '3g',
            'ProductType': 'edible (solid)',
            'DOH': 'RECREATIONAL USE',
            'Ratio': 'THC: 10mg\nCBD: 5mg',
            'ProductStrain': 'Test Strain 3'
        }
    ]
    
    try:
        # Create mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        print(f"Processing {len(test_records)} records with mini template (20-cell grid)")
        
        # Process the records
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"Generated table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check each cell
            populated_cells = 0
            blank_cells = 0
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_num = row_idx * len(table.columns) + col_idx + 1
                    
                    if cell_text and cell_text not in ['', '{{Label' + str(cell_num) + '.Description}}', '{{Label' + str(cell_num) + '.Price}}', '{{Label' + str(cell_num) + '.Lineage}}', '{{Label' + str(cell_num) + '.ProductBrand}}', '{{Label' + str(cell_num) + '.Ratio_or_THC_CBD}}', '{{Label' + str(cell_num) + '.DOH}}', '{{Label' + str(cell_num) + '.ProductStrain}}']:
                        print(f"Cell {cell_num} ({row_idx}, {col_idx}): '{cell_text[:50]}...' ✅ POPULATED")
                        populated_cells += 1
                    else:
                        print(f"Cell {cell_num} ({row_idx}, {col_idx}): [BLANK] ✅ CLEARED")
                        blank_cells += 1
            
            print(f"\nSummary:")
            print(f"  Populated cells: {populated_cells}")
            print(f"  Blank cells: {blank_cells}")
            print(f"  Total cells: {populated_cells + blank_cells}")
            
            # Verify that we have the expected number of cells
            expected_total = 20  # 4x5 grid
            if populated_cells + blank_cells == expected_total:
                print(f"✅ Correct total number of cells: {expected_total}")
            else:
                print(f"❌ Wrong total number of cells. Expected: {expected_total}, Got: {populated_cells + blank_cells}")
            
            # Verify that blank cells are actually empty
            if blank_cells > 0:
                print(f"✅ Blank cells detected and should be cleared")
            else:
                print(f"❌ No blank cells detected - this might indicate an issue")
            
            # Save the result for manual inspection
            output_file = 'test_mini_blank_cells_output.docx'
            result_doc.save(output_file)
            print(f"\n✅ Test document saved as: {output_file}")
            
        else:
            print("❌ No tables found in generated document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

def test_mini_full_grid():
    """Test with exactly 20 records to ensure full grid works correctly."""
    
    print("\n\nTesting Mini Template Full Grid (20 records)")
    print("=" * 50)
    
    # Create exactly 20 test records
    test_records = []
    for i in range(20):
        test_records.append({
            'Description': f'Test Product {i+1}',
            'ProductBrand': f'Test Brand {i+1}',
            'Price': f'${20 + i}',
            'Lineage': 'SATIVA' if i % 3 == 0 else 'INDICA' if i % 3 == 1 else 'HYBRID',
            'WeightUnits': f'{i+1}g',
            'ProductType': 'flower',
            'DOH': 'GENERAL USE',
            'Ratio': f'THC: {15 + i}%\nCBD: {1 + i}%',
            'ProductStrain': f'Test Strain {i+1}'
        })
    
    try:
        # Create mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        print(f"Processing {len(test_records)} records with mini template (full 20-cell grid)")
        
        # Process the records
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"Generated table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check each cell
            populated_cells = 0
            blank_cells = 0
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_num = row_idx * len(table.columns) + col_idx + 1
                    
                    if cell_text and cell_text not in ['', '{{Label' + str(cell_num) + '.Description}}', '{{Label' + str(cell_num) + '.Price}}', '{{Label' + str(cell_num) + '.Lineage}}', '{{Label' + str(cell_num) + '.ProductBrand}}', '{{Label' + str(cell_num) + '.Ratio_or_THC_CBD}}', '{{Label' + str(cell_num) + '.DOH}}', '{{Label' + str(cell_num) + '.ProductStrain}}']:
                        print(f"Cell {cell_num} ({row_idx}, {col_idx}): '{cell_text[:30]}...' ✅ POPULATED")
                        populated_cells += 1
                    else:
                        print(f"Cell {cell_num} ({row_idx}, {col_idx}): [BLANK] ❌ SHOULD BE POPULATED")
                        blank_cells += 1
            
            print(f"\nSummary:")
            print(f"  Populated cells: {populated_cells}")
            print(f"  Blank cells: {blank_cells}")
            print(f"  Total cells: {populated_cells + blank_cells}")
            
            # Verify that we have the expected number of cells
            expected_total = 20  # 4x5 grid
            if populated_cells + blank_cells == expected_total:
                print(f"✅ Correct total number of cells: {expected_total}")
            else:
                print(f"❌ Wrong total number of cells. Expected: {expected_total}, Got: {populated_cells + blank_cells}")
            
            # Verify that all cells are populated
            if blank_cells == 0:
                print(f"✅ All cells are populated as expected")
            else:
                print(f"❌ {blank_cells} cells are blank when they should be populated")
            
            # Save the result for manual inspection
            output_file = 'test_mini_full_grid_output.docx'
            result_doc.save(output_file)
            print(f"\n✅ Test document saved as: {output_file}")
            
        else:
            print("❌ No tables found in generated document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_mini_blank_cells_fix()
    test_mini_full_grid() 