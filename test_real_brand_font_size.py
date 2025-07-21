#!/usr/bin/env python3
"""
Test script to verify that the brand font size fix works with real template generation.
"""

import sys
import os
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def test_real_brand_font_size():
    """Test the brand font size fix with real template generation."""
    
    print("=== Real Brand Font Size Test ===\n")
    
    # Test records with different brand name lengths
    test_records = [
        {
            'ProductBrand': 'Short Brand',
            'Price': '$25.99',
            'Lineage': 'Hybrid',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'Description': 'Test description',
            'ProductStrain': 'Test Strain'
        },
        {
            'ProductBrand': 'Very Long Brand Name That Exceeds Twenty Letters',
            'Price': '$35.99',
            'Lineage': 'Indica',
            'Ratio_or_THC_CBD': 'THC: 30% CBD: 1%',
            'Description': 'Another test description',
            'ProductStrain': 'Another Strain'
        },
        {
            'ProductBrand': 'Exactly Twenty Letters Here',
            'Price': '$45.99',
            'Lineage': 'Sativa',
            'Ratio_or_THC_CBD': 'THC: 20% CBD: 3%',
            'Description': 'Third test description',
            'ProductStrain': 'Third Strain'
        }
    ]
    
    print("Testing horizontal template generation:")
    print("-" * 50)
    
    try:
        # Create template processor for horizontal template
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        # Process the test records
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"Generated document with table: {len(table.rows)}x{len(table.columns)}")
            
            # Check font sizes in the first few cells
            for i, record in enumerate(test_records[:3]):
                if i < len(table.rows) and len(table.rows[i].cells) > 0:
                    cell = table.rows[i].cells[0]
                    brand_name = record['ProductBrand']
                    letter_count = len(brand_name)
                    
                    print(f"\nRecord {i+1}: '{brand_name}' ({letter_count} letters)")
                    
                    # Check font sizes in the cell
                    font_sizes = []
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip() and run.font.size:
                                font_sizes.append(f"'{run.text[:30]}...': {run.font.size.pt}pt")
                    
                    if font_sizes:
                        print(f"  Font sizes found:")
                        for size_info in font_sizes[:3]:  # Show first 3
                            print(f"    {size_info}")
                        
                        # Check if any brand text has 14pt font when >20 letters
                        if letter_count > 20:
                            has_14pt = any("14.0pt" in size_info for size_info in font_sizes)
                            print(f"  >20 letters rule applied: {'✅' if has_14pt else '❌'}")
                    else:
                        print(f"  No font sizes found")
        else:
            print("Failed to generate document")
            
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    test_real_brand_font_size() 