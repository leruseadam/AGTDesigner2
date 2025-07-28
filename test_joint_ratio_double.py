#!/usr/bin/env python3
"""
Test joint ratio handling in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt

def test_joint_ratio_double():
    """Test joint ratio handling in double template."""
    
    print("Joint Ratio Double Template Test")
    print("=" * 40)
    
    # Create test records with joint ratio data
    test_records = [
        {
            'Description': 'Amnesia Lemon Pre-Roll',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductBrand': 'Test Brand 1',
            'ProductType': 'pre-roll',
            'JointRatio': '1.0g x 1 Pack'
        },
        {
            'Description': 'Pineapple Cake Pre-Roll',
            'Price': '$27',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 30% CBD: 1%',
            'ProductBrand': 'Test Brand 2',
            'ProductType': 'pre-roll',
            'JointRatio': '1.5g x 2 Pack'
        }
    ]
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Debug: Check what context is being built
        print("Debugging context building...")
        from docx import Document
        temp_doc = Document()
        
        for i, record in enumerate(test_records):
            print(f"\nRecord {i+1}:")
            print(f"  Raw record: {record}")
            context = processor._build_label_context(record, temp_doc)
            print(f"  Built context: {context}")
            
            # Check if JointRatio is in the context
            if 'JointRatio' in context:
                print(f"  ✓ JointRatio found: '{context['JointRatio']}'")
            else:
                print(f"  ✗ JointRatio missing from context")
        
        # Process the test records
        print("\nProcessing test records...")
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            print("✓ Document generated successfully")
            
            # Analyze the document
            print("\nAnalyzing document:")
            print("-" * 25)
            
            table = result_doc.tables[0]
            print(f"Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"\nCell ({row_idx},{col_idx}):")
                        print(f"  Text: '{cell_text}'")
                        
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip():
                                    font_size = run.font.size.pt if run.font.size else "None"
                                    print(f"    Run: '{run.text.strip()}' -> {font_size}pt")
                                    
                                    # Check if this looks like joint ratio content
                                    if any(keyword in run.text.lower() for keyword in ['pack', 'x', 'g']):
                                        print(f"      *** Potential JointRatio content: '{run.text.strip()}' ***")
        else:
            print("✗ Failed to generate document")
            
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_joint_ratio_double() 