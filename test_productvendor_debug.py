#!/usr/bin/env python3
"""
Debug script to test ProductVendor processing
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.formatting.markers import wrap_with_marker, FIELD_MARKERS, MARKER_MAP

def test_productvendor_markers():
    """Test if ProductVendor markers are properly defined"""
    print("=== Testing ProductVendor Markers ===")
    
    # Check if ProductVendor is in FIELD_MARKERS
    if 'PRODUCTVENDOR' in FIELD_MARKERS:
        print(f"✅ PRODUCTVENDOR found in FIELD_MARKERS: {FIELD_MARKERS['PRODUCTVENDOR']}")
    else:
        print("❌ PRODUCTVENDOR NOT found in FIELD_MARKERS")
    
    # Check if ProductVendor is in MARKER_MAP
    if 'ProductVendor' in MARKER_MAP:
        print(f"✅ ProductVendor found in MARKER_MAP: {MARKER_MAP['ProductVendor']}")
    else:
        print("❌ ProductVendor NOT found in MARKER_MAP")
    
    # Test wrapping
    test_vendor = "Test Vendor"
    wrapped = wrap_with_marker(test_vendor, 'PRODUCTVENDOR')
    print(f"✅ Wrapped ProductVendor: {wrapped}")
    
    # Check if it matches the expected format
    expected_start = "PRODUCTVENDOR_START"
    expected_end = "PRODUCTVENDOR_END"
    if expected_start in wrapped and expected_end in wrapped:
        print("✅ Wrapped format is correct")
    else:
        print("❌ Wrapped format is incorrect")

def test_template_processor():
    """Test TemplateProcessor with ProductVendor"""
    print("\n=== Testing TemplateProcessor ===")
    
    # Create a test record
    test_record = {
        'ProductName': 'Test Product',
        'Vendor': 'Test Vendor Company',
        'ProductBrand': 'Test Brand',
        'Price': '$10.99',
        'Description': 'Test Description',
        'Lineage': 'Test Lineage',
        'ProductStrain': 'Test Strain',
        'ProductType': 'flower'
    }
    
    # Create processor
    processor = TemplateProcessor('vertical', {}, 1.0)
    
    # Build context
    from docx import Document
    doc = Document()
    context = processor._build_label_context(test_record, doc)
    
    print(f"✅ Context built successfully")
    print(f"ProductVendor in context: '{context.get('ProductVendor', 'NOT_FOUND')}'")
    
    # Check if ProductVendor is wrapped correctly
    if 'ProductVendor' in context:
        vendor_value = context['ProductVendor']
        if 'PRODUCTVENDOR_START' in vendor_value and 'PRODUCTVENDOR_END' in vendor_value:
            print("✅ ProductVendor is properly wrapped with markers")
        else:
            print("❌ ProductVendor is not properly wrapped with markers")
    else:
        print("❌ ProductVendor not found in context")

def test_simple_template():
    """Test with a simple template"""
    print("\n=== Testing Simple Template ===")
    
    from docx import Document
    from docxtpl import DocxTemplate
    from io import BytesIO
    
    # Create a simple template with ProductVendor placeholder
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "{{Label1.ProductVendor}}"
    
    # Save template to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create DocxTemplate
    template = DocxTemplate(buffer)
    
    # Create context
    context = {
        'Label1': {
            'ProductVendor': wrap_with_marker('Test Vendor', 'PRODUCTVENDOR')
        }
    }
    
    # Render template
    template.render(context)
    
    # Check result
    rendered_text = template.docx.tables[0].cell(0, 0).text
    print(f"Rendered text: '{rendered_text}'")
    
    if 'PRODUCTVENDOR_START' in rendered_text:
        print("✅ ProductVendor placeholder was replaced with markers")
    else:
        print("❌ ProductVendor placeholder was not replaced")

if __name__ == "__main__":
    test_productvendor_markers()
    test_template_processor()
    test_simple_template() 