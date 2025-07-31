#!/usr/bin/env python3
"""
Debug script to understand template processing and DOH placeholder replacement.
"""

import sys
import os
import logging
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from io import BytesIO

def debug_template_processing():
    """Debug template processing and DOH placeholder replacement."""
    print("Debugging Template Processing")
    print("=" * 50)
    
    # Create a test record with DOH image
    test_record = {
        'Description': 'Test Product with DOH',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'classic'  # This should use regular DOH image
    }
    
    # Test double template specifically
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    
    # Check the expanded template
    print("Checking expanded template...")
    expanded_buffer = processor._expand_template_if_needed()
    expanded_doc = Document(expanded_buffer)
    
    print(f"Expanded template has {len(expanded_doc.tables)} tables")
    if expanded_doc.tables:
        table = expanded_doc.tables[0]
        print(f"Expanded table: {len(table.rows)} rows, {len(table.columns)} columns")
        
        # Check first few cells for DOH placeholders
        for i in range(min(2, len(table.rows))):
            for j in range(min(2, len(table.columns))):
                cell = table.cell(i, j)
                cell_text = cell.text.strip()
                print(f"Expanded cell ({i}, {j}): '{cell_text}'")
                
                if 'DOH' in cell_text:
                    print(f"  ✓ DOH placeholder found in expanded cell ({i}, {j})")
                else:
                    print(f"  ✗ DOH placeholder NOT found in expanded cell ({i}, {j})")
    
    # Process the record
    print("\nProcessing record...")
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return
    
    print(f"\nProcessed document has {len(result_doc.tables)} tables")
    
    # Check the processed document
    if result_doc.tables:
        table = result_doc.tables[0]
        print(f"Processed table: {len(table.rows)} rows, {len(table.columns)} columns")
        
        # Check first few cells for DOH images
        for i in range(min(2, len(table.rows))):
            for j in range(min(2, len(table.columns))):
                cell = table.cell(i, j)
                cell_text = cell.text.strip()
                print(f"Processed cell ({i}, {j}): '{cell_text}'")
                
                # Check for images in this cell
                image_count = 0
                for para in cell.paragraphs:
                    for run in para.runs:
                        if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                            image_count += 1
                
                if image_count > 0:
                    print(f"  ✓ Found {image_count} image(s) in processed cell ({i}, {j})")
                else:
                    print(f"  ✗ No images found in processed cell ({i}, {j})")

if __name__ == "__main__":
    debug_template_processing() 