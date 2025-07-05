#!/usr/bin/env python3
"""
Debug script to investigate row height changes in the actual application flow
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from docxtpl import DocxTemplate
from io import BytesIO

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.docx_formatting import fix_table_row_heights, enforce_fixed_cell_dimensions

def check_row_heights(doc, stage_name):
    """Check and print row height rules for all tables in the document"""
    print(f"\n=== {stage_name} ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            height = row.height
            height_rule = row.height_rule
            print(f"  Row {j}: height={height}, rule={height_rule}")

def create_test_template():
    """Create a test template similar to the application templates"""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    
    # Set row heights to EXACTLY
    for row in table.rows:
        row.height = Inches(2.25)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Add template variables like in the real templates
    for i in range(3):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = f"{{{{Label{i*3+j+1}.ProductBrand}}}}\n{{{{Label{i*3+j+1}.Price}}}}"
    
    # Save template
    template_path = "test_app_template.docx"
    doc.save(template_path)
    return template_path

def test_application_flow():
    """Test the actual application flow to see where row heights change"""
    
    # Create test template
    template_path = create_test_template()
    
    try:
        # Step 1: Load template and check initial state
        doc = Document(template_path)
        check_row_heights(doc, "1. After creating template")
        
        # Step 2: Test TemplateProcessor expansion
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        # Check the expanded template buffer
        expanded_doc = Document(processor._expanded_template_buffer)
        check_row_heights(expanded_doc, "2. After template expansion")
        
        # Step 3: Test docxtpl rendering
        template = DocxTemplate(processor._expanded_template_buffer)
        template.init_docx()
        check_row_heights(template.docx, "3. After DocxTemplate init")
        
        # Create context
        context = {}
        for i in range(9):
            context[f'Label{i+1}'] = {
                'ProductBrand': f'Test Brand {i+1}',
                'Price': f'${i+1}.99'
            }
        
        # Render template
        template.render(context)
        check_row_heights(template.docx, "4. After docxtpl render")
        
        # Step 4: Test post-processing
        rendered_doc = template.docx
        processor._post_process_and_replace_content(rendered_doc)
        check_row_heights(rendered_doc, "5. After post-processing")
        
        # Step 5: Test fix_table_row_heights
        fix_table_row_heights(rendered_doc, 'horizontal')
        check_row_heights(rendered_doc, "6. After fix_table_row_heights")
        
        # Step 6: Test enforce_fixed_cell_dimensions
        for table in rendered_doc.tables:
            enforce_fixed_cell_dimensions(table)
        check_row_heights(rendered_doc, "7. After enforce_fixed_cell_dimensions")
        
        # Step 7: Save and check final state
        output_path = "test_app_output.docx"
        rendered_doc.save(output_path)
        
        final_doc = Document(output_path)
        check_row_heights(final_doc, "8. After saving final document")
        
        # Step 8: Check XML
        print("\n=== XML Analysis ===")
        for i, table in enumerate(final_doc.tables):
            print(f"Table {i} XML:")
            table_xml = table._element.xml
            if 'trHeight' in table_xml:
                print("  Found trHeight elements in XML")
                import re
                tr_height_matches = re.findall(r'<w:trHeight[^>]*>', table_xml)
                for match in tr_height_matches:
                    print(f"    {match}")
            else:
                print("  No trHeight elements found in XML")
        
    finally:
        # Clean up
        if os.path.exists(template_path):
            os.remove(template_path)
        if os.path.exists(output_path):
            os.remove(output_path)

if __name__ == "__main__":
    test_application_flow() 