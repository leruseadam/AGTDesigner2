#!/usr/bin/env python3
"""
Detailed debug script to investigate Product Strain font sizing in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.unified_font_sizing import get_font_size_by_marker
from src.core.formatting.markers import wrap_with_marker
from docx.shared import Pt
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def debug_strain_processing():
    """Debug the Product Strain processing in double template."""
    
    print("=== DEBUGGING PRODUCT STRAIN FONT SIZING ===")
    
    # Test data
    test_data = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'Hybrid',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'ProductStrain': 'Test Strain Name',
        'Description': 'Test description text',
        'WeightUnits': '3.5g'
    }
    
    print(f"Test data: {test_data}")
    
    # Test 1: Check font sizing function
    print("\n1. Testing font sizing function...")
    test_strain_text = "Test Strain Name"
    font_size = get_font_size_by_marker(test_strain_text, 'PRODUCTSTRAIN', 'double', 1.0)
    print(f"   Font size for 'PRODUCTSTRAIN' marker: {font_size.pt}pt")
    
    # Test 2: Check marker wrapping
    print("\n2. Testing marker wrapping...")
    wrapped_strain = wrap_with_marker(test_strain_text, 'PRODUCTSTRAIN')
    print(f"   Wrapped strain: {repr(wrapped_strain)}")
    
    # Test 3: Check template processor
    print("\n3. Testing template processor...")
    
    try:
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Check what context is built
        print("   Building label context...")
        context = processor._build_label_context(test_data, None)
        print(f"   ProductStrain in context: {repr(context.get('ProductStrain', 'NOT FOUND'))}")
        
        # Process the document
        print("   Processing document...")
        result_doc = processor.process_records([test_data])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            cell = table.cell(0, 0)
            
            print(f"   Document generated successfully")
            print(f"   Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check all text in the first cell
            print("   All text in first cell:")
            for i, paragraph in enumerate(cell.paragraphs):
                print(f"     Paragraph {i}: {repr(paragraph.text)}")
                for j, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        font_size_pt = run.font.size.pt if run.font.size else "NO SIZE"
                        print(f"       Run {j}: '{run.text}' -> {font_size_pt}pt")
            
            # Look specifically for strain text
            print("   Looking for Product Strain text...")
            found_strain = False
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if 'Test Strain Name' in run.text:
                        found_strain = True
                        font_size_pt = run.font.size.pt if run.font.size else "NO SIZE"
                        print(f"     Found strain text: '{run.text}' -> {font_size_pt}pt")
            
            if not found_strain:
                print("     No Product Strain text found in the document")
                
        else:
            print("   Failed to generate document")
            
    except Exception as e:
        print(f"   ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_strain_processing() 