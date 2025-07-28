#!/usr/bin/env python3
"""
Test brand content in double template to see if it's showing wrong data.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def test_brand_content_double():
    """Test brand content in double template."""
    
    print("Brand Content Double Template Test")
    print("=" * 40)
    
    # Create test records with different scenarios
    test_records = [
        {
            'Description': 'Amnesia Lemon Pre-Roll',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductBrand': 'Test Brand 1',  # Clear brand name
            'ProductStrain': 'Amnesia Lemon',
            'ProductType': 'pre-roll',
            'JointRatio': '1.0g x 1 Pack'
        },
        {
            'Description': 'Blue Dream Flower',
            'Price': '$25',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 22% CBD: 1%',
            'ProductBrand': '',  # Empty brand
            'ProductStrain': 'Blue Dream',
            'ProductType': 'flower',
            'JointRatio': '3.5g'
        },
        {
            'Description': 'CBD Gummies',
            'Price': '$30',
            'Lineage': 'CBD',
            'Ratio_or_THC_CBD': 'CBD: 25mg',
            'ProductBrand': 'Premium Brand',
            'ProductStrain': 'CBD Blend',
            'ProductType': 'edible (solid)',
            'JointRatio': '10 Pack'
        }
    ]
    
    # Initialize template processor
    processor = TemplateProcessor('double', 'standard', scale_factor=1.0)
    
    # Process the records
    print("Processing records...")
    result = processor.process_records(test_records)
    
    if result is None:
        print("❌ Failed to process records")
        return
    
    # Save the result
    output_path = "test_brand_content_output.docx"
    result.save(output_path)
    print(f"✅ Generated document: {output_path}")
    
    # Analyze the content
    print("\nContent Analysis:")
    print("=" * 40)
    
    doc = Document(output_path)
    
    # Check all tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"\nCell ({row_idx},{col_idx}):")
                    
                    # Split by lines and analyze each paragraph
                    lines = cell_text.split('\n')
                    for i, line in enumerate(lines):
                        if line.strip():
                            print(f"  Line {i+1}: '{line.strip()}'")
                            
                            # Check if this line contains brand-related content
                            if 'Test Brand 1' in line:
                                print(f"    ✓ FOUND BRAND 'Test Brand 1' at line {i+1}")
                            elif 'Premium Brand' in line:
                                print(f"    ✓ FOUND BRAND 'Premium Brand' at line {i+1}")
                            elif line.strip() == '':
                                print(f"    ⚠ EMPTY BRAND FIELD at line {i+1}")
                            elif 'Amnesia Lemon Pre-Roll' in line:
                                print(f"    ✓ FOUND DESCRIPTION at line {i+1}")
                            elif 'Blue Dream Flower' in line:
                                print(f"    ✓ FOUND DESCRIPTION at line {i+1}")
                            elif 'CBD Gummies' in line:
                                print(f"    ✓ FOUND DESCRIPTION at line {i+1}")
    
    print("\nBrand Field Analysis:")
    print("=" * 40)
    print("Expected vs Actual Brand Content:")
    
    expected_brands = ['Test Brand 1', '', 'Premium Brand']
    
    for i, expected_brand in enumerate(expected_brands):
        print(f"\nRecord {i+1}:")
        print(f"  Expected Brand: '{expected_brand}'")
        
        # Find the actual brand in the document
        found_brand = None
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        lines = cell_text.split('\n')
                        # Brand should be at line 6 (position 6)
                        if len(lines) >= 6:
                            brand_line = lines[5].strip()  # Line 6 (0-indexed)
                            if brand_line and brand_line not in ['Amnesia Lemon Pre-Roll', 'Blue Dream Flower', 'CBD Gummies']:
                                found_brand = brand_line
                                break
        
        print(f"  Actual Brand: '{found_brand}'")
        
        if expected_brand == found_brand:
            print(f"  ✓ MATCH")
        elif expected_brand == '' and found_brand == '':
            print(f"  ✓ MATCH (both empty)")
        else:
            print(f"  ❌ MISMATCH")

if __name__ == "__main__":
    test_brand_content_double() 