#!/usr/bin/env python3
"""
Test script to verify DOH image centering fix.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from io import BytesIO

def test_doh_centering():
    """Test DOH image centering fix."""
    print("Testing DOH Image Centering Fix")
    print("=" * 50)
    
    # Create a test record with DOH image
    test_record = {
        'Description': 'Test Product with DOH',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'classic'  # This should use regular DOH image
    }
    
    # Test double template specifically
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return False
    
    # Check if DOH images are properly centered
    doh_images_found = 0
    centered_images = 0
    
    for table in result_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Check if this run contains an image
                        if hasattr(run, '_element') and run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
                            doh_images_found += 1
                            print(f"Found DOH image in cell")
                            
                            # Check if paragraph is centered
                            if paragraph.alignment == WD_TABLE_ALIGNMENT.CENTER:
                                centered_images += 1
                                print(f"  ✓ Image paragraph is centered")
                            else:
                                print(f"  ✗ Image paragraph is NOT centered (alignment: {paragraph.alignment})")
                            
                            # Check if cell is vertically centered
                            if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                                print(f"  ✓ Cell is vertically centered")
                            else:
                                print(f"  ✗ Cell is NOT vertically centered (alignment: {cell.vertical_alignment})")
                            
                            # Check paragraph spacing
                            if (paragraph.paragraph_format.space_before == 0 and 
                                paragraph.paragraph_format.space_after == 0):
                                print(f"  ✓ Paragraph spacing is minimal")
                            else:
                                print(f"  ✗ Paragraph spacing is not minimal")
    
    print(f"\nResults:")
    print(f"DOH images found: {doh_images_found}")
    print(f"Centered images: {centered_images}")
    
    if doh_images_found > 0 and centered_images == doh_images_found:
        print("✓ SUCCESS: All DOH images are properly centered")
        return True
    else:
        print("✗ FAILURE: Some DOH images are not properly centered")
        return False

def test_doh_high_cbd():
    """Test DOH image with High CBD product type."""
    print("\nTesting DOH Image with High CBD Product Type")
    print("=" * 50)
    
    # Create a test record with High CBD product type
    test_record = {
        'Description': 'Test High CBD Product',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 5% CBD: 15%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'high cbd classic'  # This should use HighCBD image
    }
    
    # Test double template specifically
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return False
    
    # Check if HighCBD images are properly centered
    high_cbd_images_found = 0
    centered_images = 0
    
    for table in result_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Check if this run contains an image
                        if hasattr(run, '_element') and run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
                            high_cbd_images_found += 1
                            print(f"Found HighCBD image in cell")
                            
                            # Check if paragraph is centered
                            if paragraph.alignment == WD_TABLE_ALIGNMENT.CENTER:
                                centered_images += 1
                                print(f"  ✓ Image paragraph is centered")
                            else:
                                print(f"  ✗ Image paragraph is NOT centered")
    
    print(f"\nResults:")
    print(f"HighCBD images found: {high_cbd_images_found}")
    print(f"Centered images: {centered_images}")
    
    if high_cbd_images_found > 0 and centered_images == high_cbd_images_found:
        print("✓ SUCCESS: All HighCBD images are properly centered")
        return True
    else:
        print("✗ FAILURE: Some HighCBD images are not properly centered")
        return False

if __name__ == "__main__":
    print("DOH Image Centering Fix Test")
    print("=" * 60)
    
    # Test regular DOH image centering
    test1_passed = test_doh_centering()
    
    # Test High CBD DOH image centering
    test2_passed = test_doh_high_cbd()
    
    print("\n" + "=" * 60)
    if test1_passed and test2_passed:
        print("✓ ALL TESTS PASSED: DOH image centering fix is working correctly")
    else:
        print("✗ SOME TESTS FAILED: DOH image centering fix needs more work") 