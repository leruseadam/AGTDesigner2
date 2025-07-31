#!/usr/bin/env python3
"""
Test script to verify that THC/CBD spacing issue has been fixed.
Tests that vertical template THC_CBD values use 1.25 spacing instead of no spacing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_thc_cbd_spacing_fix():
    """Test that THC/CBD spacing is now working correctly."""
    
    print("Testing THC/CBD Spacing Fix")
    print("=" * 40)
    
    # Create test data with THC/CBD content
    test_data = {
        'ProductName': 'Test Product',
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'Sativa',
        'THC_CBD': 'THC: 87.01% CBD: 0.45%',
        'Ratio_or_THC_CBD': 'THC: 87.01% CBD: 0.45%',
        'Ratio': 'THC: 87.01% CBD: 0.45%',
        'WeightUnits': '1g',
        'ProductStrain': 'Test Strain',
        'DOH': 'Yes'
    }
    
    # Test vertical template
    print("\n1️⃣ Testing Vertical Template...")
    tp_vertical = TemplateProcessor(template_type='vertical', font_scheme='Arial')
    
    # Process the record
    doc = tp_vertical.process_records([test_data])
    
    if not doc:
        print("❌ No document generated")
        return False
    
    # Check for THC/CBD content and verify line spacing
    found_thc_cbd = False
    correct_spacing = False
    
    print("🔍 Searching for THC/CBD content in document...")
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'THC:' in paragraph.text and 'CBD:' in paragraph.text:
                        found_thc_cbd = True
                        line_spacing = paragraph.paragraph_format.line_spacing
                        print(f"📋 Found THC/CBD content: '{paragraph.text.strip()}'")
                        print(f"📏 Line spacing: {line_spacing}")
                        
                        if line_spacing == 1.25:
                            correct_spacing = True
                            print("✅ Line spacing is correctly set to 1.25")
                        else:
                            print(f"❌ Line spacing should be 1.25, but is {line_spacing}")
    
    if not found_thc_cbd:
        print("❌ No THC/CBD content found in the document")
        return False
    
    # Test horizontal template for comparison
    print("\n2️⃣ Testing Horizontal Template (should have different spacing)...")
    tp_horizontal = TemplateProcessor(template_type='horizontal', font_scheme='Arial')
    
    doc_horizontal = tp_horizontal.process_records([test_data])
    
    if doc_horizontal:
        for table in doc_horizontal.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'THC:' in paragraph.text and 'CBD:' in paragraph.text:
                            line_spacing_horizontal = paragraph.paragraph_format.line_spacing
                            print(f"📋 Horizontal template THC/CBD line spacing: {line_spacing_horizontal}")
                            
                            if line_spacing_horizontal != 1.25:
                                print("✅ Horizontal template correctly has different spacing")
                            else:
                                print("❌ Horizontal template should not have spacing of 1.25")
    
    # Summary
    print("\n" + "=" * 40)
    if found_thc_cbd and correct_spacing:
        print("🎉 SUCCESS: Vertical template THC/CBD line spacing is correctly set to 1.25")
        print("✅ The 'no spacing' issue has been fixed!")
        return True
    else:
        print("❌ FAILED: Vertical template THC/CBD line spacing is not correct")
        return False

if __name__ == "__main__":
    success = test_thc_cbd_spacing_fix()
    sys.exit(0 if success else 1) 