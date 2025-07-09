#!/usr/bin/env python3
"""
Script to examine the mini template output with 5 records.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document

def examine_mini_5_records():
    """Examine the mini template output with 5 records."""
    
    print("Examining Mini Template Output with 5 Records")
    print("=" * 50)
    
    try:
        # Load the generated document
        doc = Document('test_mini_5_records.docx')
        
        if doc.tables:
            table = doc.tables[0]
            print(f"Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check each cell
            for row_idx in range(len(table.rows)):
                for col_idx in range(len(table.columns)):
                    cell = table.cell(row_idx, col_idx)
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"Cell ({row_idx}, {col_idx}): '{cell_text}'")
                    else:
                        print(f"Cell ({row_idx}, {col_idx}): [empty]")
                        
                    # Check if there are any markers still present
                    if 'LINEAGE_START' in cell.text or 'PRODUCTBRAND_CENTER_START' in cell.text:
                        print(f"  ❌ MARKERS STILL PRESENT!")
        else:
            print("❌ No tables found in document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    examine_mini_5_records() 