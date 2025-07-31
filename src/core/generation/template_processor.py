from copy import deepcopy
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Mm, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docxcompose.composer import Composer
from io import BytesIO
import logging
import os
from pathlib import Path
import re
from typing import Dict, Any, List, Optional
import traceback
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
import time
import pandas as pd

# Local imports
from src.core.utils.common import safe_get
from src.core.generation.docx_formatting import (
    apply_lineage_colors,
    enforce_fixed_cell_dimensions,
    clear_cell_background,
    clear_cell_margins,
    clear_table_cell_padding,
)
from src.core.generation.unified_font_sizing import (
    get_font_size,
    get_font_size_by_marker,
    set_run_font_size,
    is_classic_type,
    get_line_spacing_by_marker
)
from src.core.generation.text_processing import (
    process_doh_image,
    format_ratio_multiline
)
from src.core.formatting.markers import wrap_with_marker, unwrap_marker, is_already_wrapped

# Performance settings
MAX_PROCESSING_TIME_PER_CHUNK = 30  # 30 seconds max per chunk
MAX_TOTAL_PROCESSING_TIME = 300     # 5 minutes max total
CHUNK_SIZE_LIMIT = 50               # Limit chunk size for performance

def get_font_scheme(template_type, base_size=12):
    schemes = {
        'default': {"base_size": base_size, "min_size": 8, "max_length": 25},
        'vertical': {"base_size": base_size, "min_size": 8, "max_length": 25},
        'mini': {"base_size": base_size - 2, "min_size": 6, "max_length": 15},
        'horizontal': {"base_size": base_size + 1, "min_size": 7, "max_length": 20},
        'double': {"base_size": base_size - 1, "min_size": 8, "max_length": 30}
    }
    return {
        field: {**schemes.get(template_type, schemes['default'])}
        for field in ["Description", "ProductBrand", "Price", "Lineage", "DOH", "Ratio_or_THC_CBD", "Ratio"]
    }

class TemplateProcessor:
    def __init__(self, template_type, font_scheme, scale_factor=1.0):
        self.template_type = template_type
        self.font_scheme = font_scheme
        self.scale_factor = scale_factor
        self.logger = logging.getLogger(__name__)
        self._template_path = self._get_template_path()
        self._expanded_template_buffer = self._expand_template_if_needed()
        
        # Set chunk size based on template type with performance limits
        if self.template_type == 'mini':
            self.chunk_size = min(20, CHUNK_SIZE_LIMIT)  # Fixed: 4x5 grid = 20 labels per page
        elif self.template_type == 'double':
            self.chunk_size = min(12, CHUNK_SIZE_LIMIT)  # Fixed: 4x3 grid = 12 labels per page
        else:
            # For standard templates (horizontal, vertical), use 3x3 grid = 9 labels per page
            self.chunk_size = min(9, CHUNK_SIZE_LIMIT)  # Fixed: 3x3 grid = 9 labels per page
        
        self.logger.info(f"Template type: {self.template_type}, Chunk size: {self.chunk_size}")
        
        # Performance tracking
        self.start_time = time.time()
        self.chunk_count = 0

    def _get_template_path(self):
        """Get the template path based on template type."""
        try:
            base_path = Path(__file__).resolve().parent / "templates"
            template_name = f"{self.template_type}.docx"
            template_path = base_path / template_name
            
            if not template_path.exists():
                self.logger.error(f"Template not found: {template_path}")
                raise FileNotFoundError(f"Template not found: {template_path}")
            
            return template_path
        except Exception as e:
            self.logger.error(f"Error getting template path: {e}")
            raise

    def _expand_template_if_needed(self, force_expand=False):
        """Expand template if needed and return buffer."""
        try:
            with open(self._template_path, 'rb') as f:
                buffer = BytesIO(f.read())
            
            # Check if template needs expansion
            doc = Document(buffer)
            text = doc.element.body.xml
            matches = re.findall(r'Label(\d+)\.', text)
            
            # Check if we have all required labels (9 for 3x3, 20 for 4x5, 12 for 4x3)
            required_labels = 9 if self.template_type not in ['mini', 'double'] else (20 if self.template_type == 'mini' else 12)
            unique_labels = set(matches)
            
            if len(unique_labels) < required_labels or force_expand:
                self.logger.info("Template needs expansion")
                if self.template_type == 'mini':
                    return self._expand_template_to_4x5_fixed_scaled()
                elif self.template_type == 'double':
                    return self._expand_template_to_4x3_fixed_double()
                else:
                    return self._expand_template_to_3x3_fixed()
            
            return buffer
        except Exception as e:
            self.logger.error(f"Error expanding template: {e}")
            raise

    def force_re_expand_template(self):
        """Force re-expansion of template."""
        self._expanded_template_buffer = self._expand_template_if_needed(force_expand=True)

    def _expand_template_to_4x5_fixed_scaled(self):
        """Expand template to 4x5 grid for mini templates."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 4, 5
        col_width_twips = str(int(1.75 * 1440))  # 1.75 inches per column for equal width
        row_height_pts = Pt(2.0 * 72)  # 2.0 inches per row for equal height
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D3D3D3')
        tblPr.insert(0, shd)
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        borders = OxmlElement('w:tblBorders')
        for side in ('insideH','insideV'):
            b = OxmlElement(f"w:{side}")
            b.set(qn('w:val'), "single")
            b.set(qn('w:sz'), "4")
            b.set(qn('w:color'), "D3D3D3")
            b.set(qn('w:space'), "0")
            borders.append(b)
        tblPr.append(borders)
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                cell = tbl.cell(r,c)
                cell._tc.clear_content()
                tc = deepcopy(src_tc)
                for t in tc.iter(qn('w:t')):
                    if t.text and 'Label1' in t.text:
                        t.text = t.text.replace('Label1', f'Label{cnt}')
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
                cnt += 1
        from docx.oxml.shared import OxmlElement as OE
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _expand_template_to_4x3_fixed_double(self):
        """Expand template to 4x3 grid for double templates (4 columns, 3 rows)."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 4, 3  # 4 columns, 3 rows for 12 labels total
        
        # Equal width columns: 1.125 inches each for a total of 4.5 inches
        col_width_twips = str(int(1.125 * 1440))  # 1.125 inches per column
        row_height_pts = Pt(2.5 * 72)  # 2.5 inches per row for equal height
        cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D3D3D3')
        tblPr.insert(0, shd)
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        borders = OxmlElement('w:tblBorders')
        for side in ('insideH','insideV'):
            b = OxmlElement(f"w:{side}")
            b.set(qn('w:val'), "single")
            b.set(qn('w:sz'), "4")
            b.set(qn('w:color'), "D3D3D3")
            b.set(qn('w:space'), "0")
            borders.append(b)
        tblPr.append(borders)
        
        # Process all cells normally (no gutters)
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                cell = tbl.cell(r,c)
                cell._tc.clear_content()
                tc = deepcopy(src_tc)
                # Check if ProductBrand placeholder is missing and add it
                cell_text = ''
                for t in tc.iter(qn('w:t')):
                    if t.text:
                        cell_text += t.text
                        if 'Label1' in t.text:
                            t.text = t.text.replace('Label1', f'Label{cnt}')
                
                # If ProductBrand placeholder is missing, add it
                if '{{Label1.ProductBrand}}' not in cell_text and 'ProductBrand' not in cell_text:
                    # Find the position after the Lineage placeholder
                    text_elements = list(tc.iter(qn('w:t')))
                    lineage_end_index = -1
                    
                    # Find where the Lineage placeholder ends
                    for i, t in enumerate(text_elements):
                        if t.text and 'Lineage' in t.text:
                            # Found the Lineage text element, look for the closing }}
                            for j in range(i, len(text_elements)):
                                if text_elements[j].text and '}}' in text_elements[j].text:
                                    lineage_end_index = j
                                    break
                            break
                    
                    if lineage_end_index >= 0:
                        # Insert ProductBrand placeholder after the Lineage placeholder
                        new_text = OxmlElement('w:t')
                        new_text.text = f'\n{{{{Label{cnt}.ProductBrand}}}}'
                        
                        # Insert after the lineage end element
                        lineage_end_element = text_elements[lineage_end_index]
                        lineage_end_element.getparent().insert(
                            lineage_end_element.getparent().index(lineage_end_element) + 1, 
                            new_text
                        )
                
                # Add DOH placeholder if it's missing
                self.logger.debug(f"Cell {cnt} - cell_text: '{cell_text}'")
                self.logger.debug(f"Cell {cnt} - checking for DOH: '{{Label1.DOH}}' not in '{cell_text}' and 'DOH' not in '{cell_text}'")
                if '{{Label1.DOH}}' not in cell_text and 'DOH' not in cell_text:
                    self.logger.debug(f"Adding DOH placeholder to cell {cnt}")
                    # Find the position after the ProductStrain placeholder
                    text_elements = list(tc.iter(qn('w:t')))
                    strain_end_index = -1
                    
                    # Find where the ProductStrain placeholder ends
                    for i, t in enumerate(text_elements):
                        if t.text and 'ProductStrain' in t.text:
                            # Found the ProductStrain text element, look for the closing }}
                            for j in range(i, len(text_elements)):
                                if text_elements[j].text and '}}' in text_elements[j].text:
                                    strain_end_index = j
                                    break
                            break
                    
                    if strain_end_index >= 0:
                        self.logger.debug(f"Found ProductStrain end at index {strain_end_index}")
                        # Insert DOH placeholder after the ProductStrain placeholder
                        new_text = OxmlElement('w:t')
                        new_text.text = f'\n{{{{Label{cnt}.DOH}}}}'
                        
                        # Insert after the strain end element
                        strain_end_element = text_elements[strain_end_index]
                        strain_end_element.getparent().insert(
                            strain_end_element.getparent().index(strain_end_element) + 1, 
                            new_text
                        )
                        self.logger.debug(f"Inserted DOH placeholder: {new_text.text}")
                    else:
                        self.logger.warning(f"Could not find ProductStrain end position for cell {cnt}")
                else:
                    self.logger.debug(f"DOH placeholder already exists in cell {cnt}")
                
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
                cnt += 1
                
        from docx.oxml.shared import OxmlElement as OE
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _expand_template_to_3x3_fixed(self):
        """Expand template to 3x3 grid for standard templates."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        from copy import deepcopy

        num_cols, num_rows = 3, 3
        
        # Set dimensions based on template type - use constants for consistency
        from src.core.constants import CELL_DIMENSIONS
        
        cell_dims = CELL_DIMENSIONS.get(self.template_type, {'width': 2.4, 'height': 2.4})
        col_width_twips = str(int(cell_dims['width'] * 1440))  # Use width from constants
        row_height_pts = Pt(cell_dims['height'] * 72)  # Use height from constants
        # Use minimal spacing for vertical template to ensure all 9 labels fit
        if self.template_type == 'vertical':
            cut_line_twips = int(0.0001 * 1440)  # Minimal spacing for vertical
        else:
            cut_line_twips = int(0.001 * 1440)

        template_path = self._get_template_path()
        doc = Document(template_path)
        if not doc.tables:
            raise RuntimeError("Template must contain at least one table.")
        old = doc.tables[0]
        src_tc = deepcopy(old.cell(0,0)._tc)
        old._element.getparent().remove(old._element)

        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D3D3D3')
        tblPr.insert(0, shd)
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)
        grid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        borders = OxmlElement('w:tblBorders')
        for side in ('insideH','insideV'):
            b = OxmlElement(f"w:{side}")
            b.set(qn('w:val'), "single")
            b.set(qn('w:sz'), "4")
            b.set(qn('w:color'), "D3D3D3")
            b.set(qn('w:space'), "0")
            borders.append(b)
        tblPr.append(borders)
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                cell = tbl.cell(r,c)
                cell._tc.clear_content()
                tc = deepcopy(src_tc)
                
                # Check if ProductBrand placeholder is missing and add it
                cell_text = ''
                for t in tc.iter(qn('w:t')):
                    if t.text:
                        cell_text += t.text
                        if 'Label1' in t.text:
                            t.text = t.text.replace('Label1', f'Label{cnt}')
                
                # If ProductBrand placeholder is missing, add it
                if '{{Label1.ProductBrand}}' not in cell_text and 'ProductBrand' not in cell_text:
                    # Find the position after the Lineage placeholder
                    text_elements = list(tc.iter(qn('w:t')))
                    lineage_end_index = -1
                    
                    # Find where the Lineage placeholder ends
                    for i, t in enumerate(text_elements):
                        if t.text and 'Lineage' in t.text:
                            # Found the Lineage text element, look for the closing }}
                            for j in range(i, len(text_elements)):
                                if text_elements[j].text and '}}' in text_elements[j].text:
                                    lineage_end_index = j
                                    break
                            break
                    
                    if lineage_end_index >= 0:
                        # Insert ProductBrand placeholder after the Lineage placeholder
                        new_text = OxmlElement('w:t')
                        new_text.text = f'\n{{{{Label{cnt}.ProductBrand}}}}'
                        
                        # Insert after the lineage end element
                        lineage_end_element = text_elements[lineage_end_index]
                        lineage_end_element.getparent().insert(
                            lineage_end_element.getparent().index(lineage_end_element) + 1, 
                            new_text
                        )
                
                # Add DOH placeholder if it's missing
                self.logger.debug(f"Cell {cnt} - cell_text: '{cell_text}'")
                self.logger.debug(f"Cell {cnt} - checking for DOH: '{{Label1.DOH}}' not in '{cell_text}' and 'DOH' not in '{cell_text}'")
                if '{{Label1.DOH}}' not in cell_text and 'DOH' not in cell_text:
                    self.logger.debug(f"Adding DOH placeholder to cell {cnt}")
                    # Find the position after the ProductStrain placeholder
                    text_elements = list(tc.iter(qn('w:t')))
                    strain_end_index = -1
                    
                    # Find where the ProductStrain placeholder ends
                    for i, t in enumerate(text_elements):
                        if t.text and 'ProductStrain' in t.text:
                            # Found the ProductStrain text element, look for the closing }}
                            for j in range(i, len(text_elements)):
                                if text_elements[j].text and '}}' in text_elements[j].text:
                                    strain_end_index = j
                                    break
                            break
                    
                    if strain_end_index >= 0:
                        self.logger.debug(f"Found ProductStrain end at index {strain_end_index}")
                        # Insert DOH placeholder after the ProductStrain placeholder
                        new_text = OxmlElement('w:t')
                        new_text.text = f'\n{{{{Label{cnt}.DOH}}}}'
                        
                        # Insert after the strain end element
                        strain_end_element = text_elements[strain_end_index]
                        strain_end_element.getparent().insert(
                            strain_end_element.getparent().index(strain_end_element) + 1, 
                            new_text
                        )
                        self.logger.debug(f"Inserted DOH placeholder: {new_text.text}")
                    else:
                        self.logger.warning(f"Could not find ProductStrain end position for cell {cnt}")
                else:
                    self.logger.debug(f"DOH placeholder already exists in cell {cnt}")
                
                for el in tc.xpath('./*'):
                    cell._tc.append(deepcopy(el))
                cnt += 1
        from docx.oxml.shared import OxmlElement as OE
        tblPr2 = tbl._element.find(qn('w:tblPr'))
        spacing = OxmlElement('w:tblCellSpacing')
        spacing.set(qn('w:w'), str(cut_line_twips))
        spacing.set(qn('w:type'), 'dxa')
        tblPr2.append(spacing)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def process_records(self, records):
        """Process records with performance monitoring and timeout protection."""
        try:
            self.start_time = time.time()
            self.chunk_count = 0
            
            # Debug: Log the overall order of records
            overall_order = [record.get('ProductName', 'Unknown') for record in records]
            self.logger.info(f"Processing {len(records)} records in overall order: {overall_order}")
            
            # Deduplicate records by ProductName to prevent multiple outputs
            seen_products = set()
            unique_records = []
            for record in records:
                product_name = record.get('ProductName', 'Unknown')
                if product_name not in seen_products:
                    seen_products.add(product_name)
                    unique_records.append(record)
                else:
                    self.logger.warning(f"Skipping duplicate product: {product_name}")
            
            if len(unique_records) != len(records):
                self.logger.info(f"Deduplicated records: {len(records)} -> {len(unique_records)}")
                records = unique_records
            
            # Limit total number of records for performance
            if len(records) > 200:
                self.logger.warning(f"Limiting records from {len(records)} to 200 for performance")
                records = records[:200]
            
            documents = []
            for i in range(0, len(records), self.chunk_size):
                # Check total processing time
                if time.time() - self.start_time > MAX_TOTAL_PROCESSING_TIME:
                    self.logger.warning(f"Total processing time limit reached ({MAX_TOTAL_PROCESSING_TIME}s), stopping")
                    break
                
                chunk = records[i:i + self.chunk_size]
                self.chunk_count += 1
                
                self.logger.info(f"Processing chunk {self.chunk_count} ({len(chunk)} records)")
                result = self._process_chunk(chunk)
                if result: 
                    documents.append(result)
            
            if not documents: 
                return None
            if len(documents) == 1: 
                return documents[0]
            
            # Combine documents
            self.logger.info(f"Combining {len(documents)} documents")
            composer = Composer(documents[0])
            for doc in documents[1:]:
                composer.append(doc)
            
            final_doc_buffer = BytesIO()
            composer.save(final_doc_buffer)
            final_doc_buffer.seek(0)
            
            total_time = time.time() - self.start_time
            self.logger.info(f"Template processing completed in {total_time:.2f}s for {len(records)} records")
            
            return Document(final_doc_buffer)
        except Exception as e:
            self.logger.error(f"Error processing records: {e}")
            return None

    def _process_chunk(self, chunk):
        """Process a chunk of records with timeout protection."""
        chunk_start_time = time.time()
        
        try:
            if hasattr(self._expanded_template_buffer, 'seek'):
                self._expanded_template_buffer.seek(0)
            
            doc = DocxTemplate(self._expanded_template_buffer)
            
            # Debug: Log the order of records in this chunk
            chunk_order = [record.get('ProductName', 'Unknown') for record in chunk]
            self.logger.info(f"Processing chunk with {len(chunk)} records in order: {chunk_order}")
            
            # Build context for each record in the chunk
            context = {}
            for i, record in enumerate(chunk):
                # Set current record for brand centering logic
                self.current_record = record
                label_context = self._build_label_context(record, doc)
                context[f'Label{i+1}'] = label_context
                # Debug logging to check field values and order
                product_name = record.get('ProductName', 'Unknown')
                self.logger.debug(f"Label{i+1} -> {product_name} - ProductBrand: '{label_context.get('ProductBrand', 'NOT_FOUND')}', Price: '{label_context.get('Price', 'NOT_FOUND')}', THC: '{label_context.get('THC', 'NOT_FOUND')}', CBD: '{label_context.get('CBD', 'NOT_FOUND')}'")
            for i in range(len(chunk), self.chunk_size):
                context[f'Label{i+1}'] = {}

            # DOH images are already created in _build_label_context, no need for redundant creation here

            doc.render(context)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            rendered_doc = Document(buffer)
            
            # Check timeout before post-processing
            if time.time() - chunk_start_time > MAX_PROCESSING_TIME_PER_CHUNK:
                self.logger.warning(f"Chunk processing timeout reached ({MAX_PROCESSING_TIME_PER_CHUNK}s), skipping post-processing")
                return rendered_doc
            
            # Post-process the document to apply dynamic font sizing first
            self._post_process_and_replace_content(rendered_doc)
            
            # Check timeout before lineage colors
            if time.time() - chunk_start_time > MAX_PROCESSING_TIME_PER_CHUNK:
                self.logger.warning(f"Chunk processing timeout reached ({MAX_PROCESSING_TIME_PER_CHUNK}s), skipping lineage colors")
                return rendered_doc
            
            # Apply lineage colors last to ensure they are not overwritten
            apply_lineage_colors(rendered_doc)
            
            # Final enforcement of fixed cell dimensions to prevent any expansion
            for table in rendered_doc.tables:
                enforce_fixed_cell_dimensions(table)
            
            # CRITICAL: For horizontal, vertical, and double templates, explicitly override cell widths after DocxTemplate rendering
            if self.template_type in ['horizontal', 'vertical', 'double']:
                from src.core.constants import CELL_DIMENSIONS
                individual_cell_width = CELL_DIMENSIONS[self.template_type]['width']
                fixed_col_width = str(int(individual_cell_width * 1440))  # Use individual cell width directly
                
                for table in rendered_doc.tables:
                    # Override each cell width
                    for row in table.rows:
                        for cell in row.cells:
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcW = tcPr.find(qn('w:tcW'))
                            if tcW is not None:
                                tcW.getparent().remove(tcW)
                            
                            # Create new width property with correct value
                            tcW = OxmlElement('w:tcW')
                            tcW.set(qn('w:w'), fixed_col_width)
                            tcW.set(qn('w:type'), 'dxa')
                            tcPr.append(tcW)
            
            # Ensure proper table centering and document setup
            self._ensure_proper_centering(rendered_doc)

            # FINAL ENFORCEMENT: For vertical and double templates, force appropriate line spacing for all paragraphs in any cell containing THC_CBD marker
            if self.template_type in ['vertical', 'double']:
                for table in rendered_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # Check for THC_CBD marker in cell text or runs
                            cell_text = cell.text.lower()
                            has_thc_cbd = 'thc_cbd' in cell_text or 'thc: cbd:' in cell_text
                            # Also check for marker remnants in runs
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if 'THC_CBD' in run.text or 'THC_CBD' in para.text:
                                        has_thc_cbd = True
                            if has_thc_cbd:
                                for para in cell.paragraphs:
                                    # Use unified font sizing system for consistent THC_CBD line spacing
                                    line_spacing = get_line_spacing_by_marker('THC_CBD', self.template_type)
                                    if line_spacing:
                                        para.paragraph_format.line_spacing = line_spacing
                                        pPr = para._element.get_or_add_pPr()
                                        spacing = pPr.find(qn('w:spacing'))
                                        if spacing is None:
                                            spacing = OxmlElement('w:spacing')
                                            pPr.append(spacing)
                                        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
                                        spacing.set(qn('w:lineRule'), 'auto')
            
            chunk_time = time.time() - chunk_start_time
            self.logger.debug(f"Chunk processed in {chunk_time:.2f}s")
            
            # FINAL MARKER CLEANUP: Remove any lingering *_START and *_END markers AFTER font sizing has been applied
            # This cleanup should only remove markers that weren't processed by the font sizing system
            import re
            marker_pattern = re.compile(r'\b\w+_(START|END)\b')
            prefix_pattern = re.compile(r'^(?:[A-Z0-9_]+_)+')
            # Clean in tables
            for table in rendered_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                # Only clean if the run still contains markers (indicating they weren't processed)
                                if marker_pattern.search(run.text):
                                    run.text = marker_pattern.sub('', run.text)
                                if prefix_pattern.search(run.text):
                                    run.text = prefix_pattern.sub('', run.text)
            # Clean in paragraphs outside tables
            for para in rendered_doc.paragraphs:
                for run in para.runs:
                    # Only clean if the run still contains markers (indicating they weren't processed)
                    if marker_pattern.search(run.text):
                        run.text = marker_pattern.sub('', run.text)
                    if prefix_pattern.search(run.text):
                        run.text = prefix_pattern.sub('', run.text)
            
            return rendered_doc
        except Exception as e:
            self.logger.error(f"Error in _process_chunk: {e}\n{traceback.format_exc()}")
            raise

    def _build_label_context(self, record, doc):
        """Ultra-optimized label context building for maximum performance."""
        # Fast dictionary copy
        label_context = dict(record)

        # Fast value cleaning - only process non-empty values
        for key, value in label_context.items():
            if value is not None:
                label_context[key] = str(value).strip()
            else:
                label_context[key] = ""

        # Define product type sets for use throughout the method
        classic_types = {"flower", "pre-roll", "infused pre-roll", "concentrate", "solventless concentrate", "vape cartridge", "rso/co2 tankers"}
        edible_types = {"edible (solid)", "edible (liquid)", "high cbd edible liquid", "tincture", "topical", "capsule"}

        # Fast Description and WeightUnits combination
        desc = label_context.get('Description', '') or ''
        weight = (label_context.get('WeightUnits', '') or '').replace('\u202F', '')
        
        # Ultra-fast string operations
        if desc.endswith('- '):
            desc = desc[:-2]
        if weight.startswith('- '):
            weight = weight[2:]
        
        # Template-specific handling
        if self.template_type == 'mini':
            label_context['DescAndWeight'] = wrap_with_marker(desc, 'DESC')
        else:
            if desc and weight:
                label_context['DescAndWeight'] = wrap_with_marker(f"{desc} -\u00A0{weight}", 'DESC')
            else:
                label_context['DescAndWeight'] = wrap_with_marker(desc or weight, 'DESC')

        # Fast DOH image processing - only if needed
        if label_context.get('DOH'):
            doh_value = label_context.get('DOH', '')
            product_type = (label_context.get('ProductType') or 
                          label_context.get('Product Type*') or 
                          record.get('ProductType') or 
                          record.get('Product Type*') or '')

            image_path = process_doh_image(doh_value, product_type)
            if image_path:
                # Fast width selection
                width_map = {'mini': 9, 'double': 9, 'vertical': 12, 'horizontal': 12}
                image_width = Mm(width_map.get(self.template_type, 12))
                label_context['DOH'] = InlineImage(doc, image_path, width=image_width)
                # Ensure DOH image takes priority - clear any other DOH-related content
                label_context['DOH_TEXT'] = ''  # Clear any text content
            else:
                label_context['DOH'] = ''
                label_context['DOH_TEXT'] = ''
        else:
            label_context['DOH'] = ''
            label_context['DOH_TEXT'] = ''
        
        # Fast ratio processing
        ratio_val = label_context.get('Ratio_or_THC_CBD') or label_context.get('Ratio', '')
        if ratio_val:
            cleaned_ratio = ratio_val.lstrip('- ')
            product_type = (label_context.get('ProductType', '').lower() or 
                          label_context.get('Product Type*', '').lower())
            
            # Fast product type checking with sets
            is_classic = product_type in classic_types
            is_edible = product_type in edible_types
            
            if is_classic and 'mg' in cleaned_ratio.lower():
                cleaned_ratio = format_ratio_multiline(cleaned_ratio)
            elif is_edible and 'mg' in cleaned_ratio.lower():
                cleaned_ratio = format_ratio_multiline(cleaned_ratio)
            elif is_classic:
                cleaned_ratio = self.format_classic_ratio(cleaned_ratio, record)
            
            # Fast marker wrapping
            content = cleaned_ratio.replace('|BR|', '\n')
            # Force line breaks for vertical and double templates
            if self.template_type in ['vertical', 'double'] and content.strip().startswith('THC:') and 'CBD:' in content:
                content = content.replace('THC: CBD:', 'THC:\nCBD:').replace('THC:  CBD:', 'THC:\nCBD:')
                # For vertical template, format with right-aligned percentages
                if self.template_type == 'vertical':
                    content = self.format_thc_cbd_vertical_alignment(content)
            
            marker = 'THC_CBD' if is_classic else 'RATIO'
            label_context['Ratio_or_THC_CBD'] = wrap_with_marker(content, marker)
            
            # Also add separate THC_CBD field for template processing
            if is_classic:
                label_context['THC_CBD'] = wrap_with_marker(content, 'THC_CBD')
            else:
                label_context['THC_CBD'] = ''
        else:
            label_context['Ratio_or_THC_CBD'] = ''
            label_context['THC_CBD'] = ''

        # Fast brand handling - include brand for all product types
        product_brand = (record.get('ProductBrand') or 
                        record.get('Product Brand') or 
                        record.get('product_brand') or 
                        record.get('productbrand') or '')
        
        if product_brand:
            label_context['ProductBrand'] = wrap_with_marker(unwrap_marker(product_brand, 'PRODUCTBRAND_CENTER'), 'PRODUCTBRAND_CENTER')
            label_context['ProductBrand_Center'] = wrap_with_marker(unwrap_marker(product_brand, 'PRODUCTBRAND_CENTER'), 'PRODUCTBRAND_CENTER')
        else:
            label_context['ProductBrand'] = ''
            label_context['ProductBrand_Center'] = ''

        # Fast other field processing
        if label_context.get('Price'):
            label_context['Price'] = wrap_with_marker(unwrap_marker(label_context['Price'], 'PRICE'), 'PRICE')
        
        if label_context.get('Lineage'):
            product_type = (label_context.get('ProductType', '').lower() or 
                          label_context.get('Product Type*', '').lower())
            
            if product_type in edible_types:
                lineage_value = ''
            else:
                lineage_value = label_context['Lineage']
                if self.template_type in {"horizontal", "vertical", "double"} and lineage_value and product_type in classic_types:
                    lineage_value = '\u2022  ' + lineage_value
            
            label_context['Lineage'] = wrap_with_marker(unwrap_marker(lineage_value, 'LINEAGE'), 'LINEAGE')

        # Fast wrapping for remaining fields
        if label_context.get('DescAndWeight'):
            label_context['DescAndWeight'] = wrap_with_marker(unwrap_marker(label_context['DescAndWeight'], 'DESC'), 'DESC')
        
        if 'ProductType' not in label_context:
            label_context['ProductType'] = record.get('ProductType', '')
        
        # Fast strain handling
        product_strain = record.get('ProductStrain') or record.get('Product Strain', '')
        if product_strain:
            label_context['ProductStrain'] = wrap_with_marker(unwrap_marker(product_strain, 'PRODUCTSTRAIN'), 'PRODUCTSTRAIN')
        else:
            label_context['ProductStrain'] = ''

        # Fast joint ratio handling
        if label_context.get('JointRatio'):
            val = label_context['JointRatio']
            # Fix: Handle NaN values in JointRatio
            if pd.isna(val) or str(val).lower() == 'nan':
                val = ''
            marker = 'JOINT_RATIO'
            if is_already_wrapped(val, marker):
                val = unwrap_marker(val, marker)
            formatted_val = self.format_joint_ratio_pack(val)
            label_context['JointRatio'] = wrap_with_marker(formatted_val, marker)

        # Fast description processing
        if label_context.get('Description'):
            label_context['Description'] = self.fix_hyphen_spacing(label_context['Description'])

        # Fast line break processing
        product_type = (label_context.get('ProductType', '').lower() or 
                       label_context.get('Product Type*', '').lower())
        
        if product_type not in classic_types and label_context.get('DescAndWeight'):
            desc_weight = label_context['DescAndWeight']
            if desc_weight.endswith(' - '):
                desc_weight = desc_weight[:-3] + '\n- '
            elif desc_weight.endswith(' -'):
                desc_weight = desc_weight[:-2] + '\n- '
            desc_weight = desc_weight.replace(' - ', '\n- ')
            label_context['DescAndWeight'] = desc_weight
        
        # Fast pre-roll processing
        if product_type in {"pre-roll", "infused pre-roll"} and label_context.get('DescAndWeight'):
            desc_weight = label_context['DescAndWeight']
            desc_weight = desc_weight.replace(' - ', '\n- ')
            label_context['DescAndWeight'] = desc_weight

        # Fast weight and ratio formatting
        for key, marker in [('WeightUnits', 'WEIGHTUNITS'), ('Ratio', 'RATIO')]:
            if label_context.get(key):
                val = label_context[key]
                formatted_val = self.format_with_soft_hyphen(val)
                label_context[key] = wrap_with_marker(unwrap_marker(formatted_val, marker), marker)
        
        # Fast vendor handling - only include for classic product types
        product_type = (label_context.get('ProductType', '').lower() or 
                       label_context.get('Product Type*', '').lower())
        
        if product_type in classic_types:
            # Include vendor for classic product types
            product_vendor = record.get('Vendor') or record.get('Vendor/Supplier*', '')
            # Handle NaN values in vendor data
            if pd.isna(product_vendor) or str(product_vendor).lower() == 'nan':
                product_vendor = ''
            label_context['ProductVendor'] = wrap_with_marker(product_vendor, 'PRODUCTVENDOR')
        else:
            # Remove vendor for non-classic product types
            label_context['ProductVendor'] = ''
        
        return label_context

    def _post_process_and_replace_content(self, doc):
        """
        Ultra-optimized post-processing for maximum performance.
        """
        # Performance optimization: Skip expensive processing for large documents
        if len(doc.tables) > 10:
            self.logger.warning(f"Skipping expensive post-processing for large document with {len(doc.tables)} tables")
            return doc
        
        # Clean up DOH cells before processing to ensure proper image positioning
        try:
            self._clean_doh_cells_before_processing(doc)
        except Exception as e:
            self.logger.warning(f"DOH cell cleanup failed: {e}")
        
        # Fast mini template processing
        if self.template_type == 'mini':
            try:
                self._add_weight_units_markers(doc)
                self._add_brand_markers(doc)
                self._clear_blank_cells_in_mini_template(doc)
            except Exception as e:
                self.logger.warning(f"Mini template processing failed: {e}")

        # Fast font sizing (with timeout protection)
        try:
            self._post_process_template_specific(doc)
        except Exception as e:
            self.logger.warning(f"Font sizing failed: {e}")

        # Fast BR marker conversion - only process if needed
        try:
            br_found = False
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if '|BR|' in paragraph.text:
                                self._convert_br_markers_to_line_breaks(paragraph)
                                br_found = True
            
            # Only process paragraphs outside tables if BR markers were found
            if br_found:
                for paragraph in doc.paragraphs:
                    if '|BR|' in paragraph.text:
                        self._convert_br_markers_to_line_breaks(paragraph)
        except Exception as e:
            self.logger.warning(f"BR marker conversion failed: {e}")
        
        # Fast ratio spacing fix
        try:
            self._fix_ratio_paragraph_spacing(doc)
        except Exception as e:
            self.logger.warning(f"Ratio spacing failed: {e}")

        # Fast Arial Bold enforcement
        try:
            from src.core.generation.docx_formatting import enforce_arial_bold_all_text, enforce_ratio_formatting
            enforce_arial_bold_all_text(doc)
            enforce_ratio_formatting(doc)
        except Exception as e:
            self.logger.warning(f"Arial bold failed: {e}")

        # Fast DOH image centering
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Fast check for image-only cells
                        if len(cell.paragraphs) > 0 and all(len(paragraph.runs) == 1 and not paragraph.text.strip() for paragraph in cell.paragraphs):
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Fast inner table centering
                        for inner_table in cell.tables:
                            inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        # Explicit DOH image centering - check for InlineImage objects
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Check if this run contains an InlineImage (DOH image)
                                if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    # Also center the cell content
                                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                    
            # Additional comprehensive DOH centering pass
            self._ensure_doh_image_centering(doc)
        except Exception as e:
            self.logger.warning(f"DOH centering failed: {e}")
            
        return doc

    def _ensure_doh_image_centering(self, doc):
        """
        Ensure DOH images are properly centered in all cells.
        This method specifically looks for InlineImage objects and centers them.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        image_run = None
                        
                        # First pass: identify if cell has DOH image
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                                    has_doh_image = True
                                    image_run = run
                                    break
                            if has_doh_image:
                                break
                        
                        if has_doh_image:
                            # Clear the entire cell content first
                            cell._tc.clear_content()
                            
                            # Create a new paragraph for the image
                            paragraph = cell.add_paragraph()
                            
                            # Add the image run to the new paragraph
                            if image_run:
                                # Copy the image element to the new paragraph
                                new_run = paragraph.add_run()
                                new_run._element.append(image_run._element)
                                
                                # Ensure the image has proper text content
                                if not new_run.text:
                                    new_run.text = '\u00A0'  # Non-breaking space
                            
                            # Set perfect centering
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Ensure no extra spacing at XML level
                            pPr = paragraph._element.get_or_add_pPr()
                            spacing = pPr.find(qn('w:spacing'))
                            if spacing is None:
                                spacing = OxmlElement('w:spacing')
                                pPr.append(spacing)
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), '240')  # 1.0 line spacing
                            spacing.set(qn('w:lineRule'), 'auto')
                            
                            self.logger.debug("Perfectly centered DOH image in cell")
                                
        except Exception as e:
            self.logger.warning(f"Error in DOH image centering: {e}")

    def _clear_blank_cells_in_mini_template(self, doc):
        """
        Clear blank cells in mini templates when they run out of values.
        This removes empty cells that don't have any meaningful content.
        """
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if cell is essentially empty
                        cell_text = cell.text.strip()
                        
                        # Consider a cell blank if it has no text or only contains template placeholders
                        is_blank = (
                            not cell_text or 
                            cell_text == '' or
                            # Check for empty template placeholders like {{LabelX.Description}}
                            (cell_text.startswith('{{Label') and cell_text.endswith('}}') and 
                             any(field in cell_text for field in ['.Description}}', '.Price}}', '.Lineage}}', '.ProductBrand}}', '.Ratio_or_THC_CBD}}', '.DOH}}', '.ProductStrain}}']))
                        )
                        
                        if is_blank:
                            # Clear the cell content
                            cell._tc.clear_content()
                            
                            # Add a single empty paragraph to maintain cell structure
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set cell background to white/transparent to ensure it's visually clean
                            from src.core.generation.docx_formatting import clear_cell_background
                            clear_cell_background(cell)
                            
                            self.logger.debug(f"Cleared blank cell in mini template")
                            
        except Exception as e:
            self.logger.error(f"Error clearing blank cells in mini template: {e}")
            # Don't raise the exception - this is a cleanup operation that shouldn't break the main process

    def _post_process_template_specific(self, doc):
        """
        Apply template-type-specific font sizing to all markers in the document.
        Uses the original font-sizing functions based on template type.
        """
        # Define marker processing for all template types (including double)
        markers = [
            'DESC', 'PRODUCTBRAND', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 
            'THC_CBD', 'THC_CBD_LABEL', 'RATIO', 'WEIGHTUNITS', 'PRODUCTSTRAIN', 'DOH', 'PRODUCTVENDOR'
        ]
        
        # Process all markers in a single pass to avoid conflicts
        self._recursive_autosize_template_specific_multi(doc, markers)
        
        # Apply vertical template specific optimizations for minimal spacing
        if self.template_type in ['vertical', 'double']:
            self._optimize_vertical_template_spacing(doc)



    def _optimize_vertical_template_spacing(self, doc):
        """
        Apply minimal spacing optimizations specifically for vertical and double templates
        to ensure all labels fit on one page.
        """
        try:
            from docx.shared import Pt
            
            def optimize_paragraph_spacing(paragraph):
                """Set minimal spacing for all paragraphs in vertical and double templates."""
                # Set absolute minimum spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
                # Check if this paragraph contains THC_CBD content and preserve its line spacing
                paragraph_text = paragraph.text
                if 'THC:' in paragraph_text and 'CBD:' in paragraph_text:
                    # Use unified font sizing system for THC_CBD content
                    from src.core.generation.unified_font_sizing import get_line_spacing_by_marker
                    line_spacing = get_line_spacing_by_marker('THC_CBD', self.template_type)
                    if line_spacing:
                        paragraph.paragraph_format.line_spacing = line_spacing
                        # Set at XML level for maximum compatibility
                        pPr = paragraph._element.get_or_add_pPr()
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:before'), '0')
                        spacing.set(qn('w:after'), '0')
                        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
                        spacing.set(qn('w:lineRule'), 'auto')
                        return  # Skip the default 1.0 spacing for THC_CBD content
                
                # Default spacing for non-THC_CBD content
                paragraph.paragraph_format.line_spacing = 1.0
                
                # Set at XML level for maximum compatibility
                pPr = paragraph._element.get_or_add_pPr()
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), '240')  # 1.0 line spacing
                spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            optimize_paragraph_spacing(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                optimize_paragraph_spacing(paragraph)
            
            self.logger.debug("Applied vertical/double template spacing optimizations")
            
        except Exception as e:
            self.logger.error(f"Error optimizing vertical/double template spacing: {e}")
            # Don't raise the exception - this is an optimization that shouldn't break the main process

    def _recursive_autosize_template_specific(self, element, marker_name):
        """
        Recursively find and replace markers in paragraphs and tables using template-specific font sizing.
        """
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                self._process_paragraph_for_marker_template_specific(p, marker_name)

        if hasattr(element, 'tables'):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._recursive_autosize_template_specific(cell, marker_name)

    def _recursive_autosize_template_specific_multi(self, element, markers):
        """
        Recursively find and replace all markers in paragraphs and tables using template-specific font sizing.
        Processes all markers in a single pass to avoid conflicts.
        """
        if hasattr(element, 'paragraphs'):
            for p in element.paragraphs:
                self._process_paragraph_for_markers_template_specific(p, markers)

        if hasattr(element, 'tables'):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self._recursive_autosize_template_specific_multi(cell, markers)

    def _process_paragraph_for_markers_template_specific(self, paragraph, markers):
        """
        Process a single paragraph for multiple markers using template-specific font sizing.
        Handles all markers in a single pass to avoid conflicts.
        """
        full_text = "".join(run.text for run in paragraph.runs)
        
        # First, check if this is a combined lineage/vendor paragraph
        if self._detect_and_process_combined_lineage_vendor(paragraph):
            return
        
        # Check if any markers are present
        found_markers = []
        for marker_name in markers:
            start_marker = f'{marker_name}_START'
            end_marker = f'{marker_name}_END'
            if start_marker in full_text and end_marker in full_text:
                found_markers.append(marker_name)
        
        if found_markers:
            # Process all markers and build the final content
            final_content = full_text
            processed_content = {}
            
            for marker_name in found_markers:
                start_marker = f'{marker_name}_START'
                end_marker = f'{marker_name}_END'
                
                # Extract content for this marker
                start_idx = final_content.find(start_marker)
                end_idx = final_content.find(end_marker) + len(end_marker)
                
                if start_idx != -1 and end_idx != -1:
                    marker_start = final_content.find(start_marker) + len(start_marker)
                    marker_end = final_content.find(end_marker)
                    content = final_content[marker_start:marker_end]
                    
                    # Get font size for this marker
                    font_size = self._get_template_specific_font_size(content, marker_name)
                    processed_content[marker_name] = {
                        'content': content,
                        'font_size': font_size,
                        'start_pos': start_idx,
                        'end_pos': end_idx
                    }
            
            # Clear paragraph and rebuild with all processed content
            paragraph.clear()
            
            # Sort markers by position in text
            sorted_markers = sorted(processed_content.items(), key=lambda x: x[1]['start_pos'])
            
            current_pos = 0
            for marker_name, marker_data in sorted_markers:
                # Add any text before this marker
                if marker_data['start_pos'] > current_pos:
                    text_before = full_text[current_pos:marker_data['start_pos']]
                    # Preserve line breaks and whitespace, but skip if completely empty
                    if text_before or text_before.strip():
                        run = paragraph.add_run(text_before)
                        run.font.name = "Arial"
                        run.font.bold = True
                        run.font.size = Pt(12)  # Default size for non-marker text
                # Add the processed marker content (use the potentially modified content)
                display_content = marker_data.get('display_content', marker_data['content'])
                # --- BULLETPROOF: Only one run for the entire marker content, preserving line breaks ---
                run = paragraph.add_run()
                run.font.name = "Arial"
                
                # Special handling for PRODUCTVENDOR - don't make it bold
                if marker_name == 'PRODUCTVENDOR':
                    run.font.bold = False
                else:
                    run.font.bold = True
                
                run.font.size = marker_data['font_size']
                set_run_font_size(run, marker_data['font_size'])
                lines = display_content.splitlines()
                for i, line in enumerate(lines):
                    if i > 0:
                        run.add_break()
                    run.add_text(line)
                current_pos = marker_data['end_pos']
            
            # Add any remaining text
            if current_pos < len(full_text):
                text_after = full_text[current_pos:]
                # Preserve line breaks and whitespace, but skip if completely empty
                if text_after or text_after.strip():
                    run = paragraph.add_run(text_after)
                    run.font.name = "Arial"
                    run.font.bold = True
                    run.font.size = Pt(12)  # Default size for non-marker text
            
            # Convert |BR| markers to actual line breaks after marker processing
            self._convert_br_markers_to_line_breaks(paragraph)
            
            # Apply special formatting for specific markers
            for marker_name, marker_data in processed_content.items():
                # Special handling for ProductBrand markers in Double template
                if ('PRODUCTBRAND' in marker_name) and self.template_type == 'double':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, 'double', self.scale_factor, product_type))
                    continue
                if marker_name == 'DOH':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    continue
                if marker_name == 'RATIO':
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], 'RATIO', self.template_type, self.scale_factor, product_type))
                        # Ensure ratio values are bold
                        run.font.bold = True
                    continue
                if marker_name == 'LINEAGE':
                    content = marker_data['content']
                    product_type = None
                    if hasattr(self, 'current_product_type'):
                        product_type = self.current_product_type
                    elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                        product_type = self.label_context['ProductType']
                    
                    # Use unified LINEAGE font sizing for all templates including double
                    for run in paragraph.runs:
                        font_size = get_font_size_by_marker(content, 'LINEAGE', self.template_type, self.scale_factor, product_type)
                        set_run_font_size(run, font_size)
                    
                    # Handle alignment based on content type
                    classic_lineages = [
                        "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                        "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                    ]
                    if content.upper() in classic_lineages and content.upper() != "PARAPHERNALIA":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if self.template_type in {"horizontal", "double", "vertical"}:
                            paragraph.paragraph_format.left_indent = Inches(0.15)
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if self.template_type in {"horizontal", "double", "vertical"}:
                            paragraph.paragraph_format.left_indent = Inches(0.15)
                    continue
                # Always center ProductBrand and ProductBrand_Center markers
                if marker_name in ('PRODUCTBRAND', 'PRODUCTBRAND_CENTER') or 'PRODUCTBRAND' in marker_name:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        # Get product type for font sizing
                        product_type = None
                        if hasattr(self, 'current_product_type'):
                            product_type = self.current_product_type
                        elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                            product_type = self.label_context['ProductType']
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor, product_type))
                    continue
                # Special handling for ProductStrain marker - always use 1pt font
                if marker_name in ('PRODUCTSTRAIN', 'STRAIN'):
                    for run in paragraph.runs:
                        # Only apply 1pt font to runs that contain strain content
                        if marker_data['content'] in run.text:
                            set_run_font_size(run, get_font_size_by_marker(marker_data['content'], 'PRODUCTSTRAIN', self.template_type, self.scale_factor))
                    continue
                # Special handling for ProductVendor marker
                if marker_name == 'PRODUCTVENDOR' or marker_name == 'VENDOR':
                    for run in paragraph.runs:
                        set_run_font_size(run, get_font_size_by_marker(marker_data['content'], marker_name, self.template_type, self.scale_factor))
                        # Ensure Arial Bold for vendor text
                        run.font.name = "Arial"
                        run.font.bold = True
                        # Remove white color setting - vendor should be visible
                    continue
            
            self.logger.debug(f"Applied multi-marker processing for: {list(processed_content.keys())}")
        try:
            pass
        except Exception as e:
            self.logger.error(f"Error processing multi-marker template: {e}")
            # Fallback: remove all markers and use default size
            for run in paragraph.runs:
                for marker_name in markers:
                    start_marker = f'{marker_name}_START'
                    end_marker = f'{marker_name}_END'
                    run.text = run.text.replace(start_marker, "").replace(end_marker, "")
                # Use appropriate default size based on template type
                if self.template_type == 'mini':
                    default_size = Pt(8 * self.scale_factor)
                elif self.template_type == 'vertical':
                    default_size = Pt(10 * self.scale_factor)
                else:  # horizontal
                    default_size = Pt(12 * self.scale_factor)
                run.font.size = default_size
        finally:
            # Always check for |BR| markers regardless of success/failure
            self._convert_br_markers_to_line_breaks(paragraph)

    def _process_paragraph_for_marker_template_specific(self, paragraph, marker_name):
        """
        Process a single paragraph for a specific marker using template-type-specific font sizing.
        """
        start_marker = f'{marker_name}_START'
        end_marker = f'{marker_name}_END'
        
        full_text = "".join(run.text for run in paragraph.runs)
        
        if start_marker in full_text and end_marker in full_text:
            try:
                # Extract content
                start_idx = full_text.find(start_marker) + len(start_marker)
                end_idx = full_text.find(end_marker)
                content = full_text[start_idx:end_idx]
                
                # For THC_CBD markers, calculate font size before any splitting to ensure consistency
                if marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL'] and ('\n' in content or '|BR|' in content):
                    # Calculate font size based on the original unsplit content to ensure consistency
                    original_content = content.replace('\n', ' ').replace('|BR|', ' ')
                    font_size = self._get_template_specific_font_size(original_content, marker_name)
                    import logging
                    logging.debug(f"[FONT_DEBUG] Processing marker '{marker_name}' with original content '{original_content}' -> font_size: {font_size}")
                    
                    # Clear and recreate with single run approach
                    paragraph.clear()
                    
                    # Create a single run with the entire content
                    run = paragraph.add_run()
                    run.font.name = "Arial"
                    run.font.bold = True
                    run.font.size = font_size
                    set_run_font_size(run, font_size)
                    
                    # Add the content with line breaks as text
                    run.add_text(content)
                    
                    # Convert line breaks to actual line breaks, passing the font size
                    self._convert_br_markers_to_line_breaks(paragraph, font_size)
                else:
                    # Use template-type-specific font sizing based on original functions
                    font_size = self._get_template_specific_font_size(content, marker_name)
                    import logging
                    logging.debug(f"[FONT_DEBUG] Processing marker '{marker_name}' with content '{content}' -> font_size: {font_size}")
                    
                    # Clear paragraph and re-add content with template-optimized formatting
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.font.name = "Arial"
                    # Special handling for PRODUCTVENDOR - don't make it bold
                    if marker_name == 'PRODUCTVENDOR':
                        run.font.bold = False
                    else:
                        run.font.bold = True
                    run.font.size = font_size
                    
                    # Apply template-specific font size setting
                    set_run_font_size(run, font_size)
                    
                    # Add the content to the run
                    run.add_text(content)
                    
                    # Convert |BR| markers to actual line breaks for other markers
                    self._convert_br_markers_to_line_breaks(paragraph, font_size)
                
                # Handle special formatting for specific markers
                if marker_name in ['PRODUCTBRAND', 'PRODUCTBRAND_CENTER']:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Also ensure all runs in this paragraph are properly sized
                    for run in paragraph.runs:
                        set_run_font_size(run, font_size)
                elif marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                    # Ensure THC_CBD and RATIO values are bold
                    for run in paragraph.runs:
                        run.font.bold = True
                    
                    # For vertical template, apply line spacing from unified font sizing
                    line_spacing = get_line_spacing_by_marker(marker_name, self.template_type)
                    if line_spacing:
                        paragraph.paragraph_format.line_spacing = line_spacing
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Set at XML level for maximum compatibility
                        pPr = paragraph._element.get_or_add_pPr()
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
                        spacing.set(qn('w:lineRule'), 'auto')
                    
                    # For vertical template THC_CBD content, ensure left alignment for proper spacing
                    if self.template_type == 'vertical' and marker_name == 'THC_CBD':
                        # Set paragraph alignment to left for proper spacing behavior
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        self.logger.debug(f"Set left alignment for vertical template THC_CBD content")
                    # Note: Line spacing is now handled by unified font sizing system
                    # The get_line_spacing_by_marker function already applies 1.25 spacing for vertical template THC_CBD
                    # Line spacing for THC: CBD: content across all templates (legacy logic)
                    elif content == 'THC: CBD:':
                        # Use unified font sizing system for consistent spacing
                        legacy_line_spacing = get_line_spacing_by_marker('THC_CBD', self.template_type)
                        paragraph.paragraph_format.line_spacing = legacy_line_spacing
                        
                        if self.template_type == 'vertical':
                            # Add left upper alignment for vertical template
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Set vertical alignment to top for the cell containing this paragraph
                        if paragraph._element.getparent().tag.endswith('tc'):  # Check if in table cell
                            cell = paragraph._element.getparent()
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    # For all other Ratio content in horizontal template, use unified font sizing system
                    elif self.template_type == 'horizontal' and marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                        # Use unified font sizing system for consistent spacing
                        line_spacing = get_line_spacing_by_marker(marker_name, self.template_type)
                        if line_spacing:
                            paragraph.paragraph_format.line_spacing = line_spacing
                        # Set vertical alignment to top for the cell containing this paragraph
                        if paragraph._element.getparent().tag.endswith('tc'):  # Check if in table cell
                            cell = paragraph._element.getparent()
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    # For all other THC/CBD content in other templates, set vertical alignment to top
                    elif marker_name in ['THC_CBD', 'RATIO', 'THC_CBD_LABEL']:
                        # Set vertical alignment to top for the cell containing this paragraph
                        if paragraph._element.getparent().tag.endswith('tc'):  # Check if in table cell
                            cell = paragraph._element.getparent()
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                
                # Center alignment for brand names
                if 'PRODUCTBRAND' in marker_name:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Center alignment for DOH (Date of Harvest)
                if marker_name == 'DOH':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Special handling for lineage markers
                if marker_name == 'LINEAGE':
                    # Extract product type information from the content
                    if '_PRODUCT_TYPE_' in content and '_IS_CLASSIC_' in content:
                        parts = content.split('_PRODUCT_TYPE_')
                        if len(parts) == 2:
                            actual_lineage = parts[0]
                            type_info = parts[1]
                            type_parts = type_info.split('_IS_CLASSIC_')
                            if len(type_parts) == 2:
                                product_type = type_parts[0]
                                is_classic_raw = type_parts[1]
                                # Remove LINEAGE_END marker if present
                                if is_classic_raw.endswith('LINEAGE_END'):
                                    is_classic_raw = is_classic_raw[:-len('LINEAGE_END')]
                                is_classic = is_classic_raw.lower() == 'true'
                                
                                # Center if it's NOT a classic type
                                if not is_classic:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    # For Classic Types, left-justify the text
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                # Update the content to only show the actual lineage (remove any markers)
                                if actual_lineage.startswith('LINEAGE_START'):
                                    actual_lineage = actual_lineage[len('LINEAGE_START'):]
                                content = actual_lineage
                        else:
                            # Fallback: use the old logic for backward compatibility
                            classic_lineages = [
                                "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                                "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                            ]
                            # Only center if the content is NOT a classic lineage (meaning it's likely a brand name)
                            if content.upper() not in classic_lineages:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                # For Classic Types, left-justify the text
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                self.logger.debug(f"Applied template-specific font sizing: {font_size.pt}pt for {marker_name} marker")

            except Exception as e:
                self.logger.error(f"Error processing template-specific marker {marker_name}: {e}")
                # Fallback: remove markers and use default size based on template type
                for run in paragraph.runs:
                    run.text = run.text.replace(start_marker, "").replace(end_marker, "")
                    # Use appropriate default size based on template type
                    if self.template_type == 'mini':
                        default_size = Pt(8 * self.scale_factor)
                    elif self.template_type == 'vertical':
                        default_size = Pt(10 * self.scale_factor)
                    else:  # horizontal
                        default_size = Pt(12 * self.scale_factor)
                    run.font.size = default_size
        elif start_marker in full_text or end_marker in full_text:
            # Log partial markers for debugging
            self.logger.debug(f"Found partial {marker_name} marker in text: '{full_text[:100]}...'")

    def _convert_br_markers_to_line_breaks(self, paragraph, font_size=None):
        """
        Convert |BR| markers and \n characters in paragraph text to actual line breaks.
        This splits the text at |BR| markers or \n characters and creates separate runs for each part.
        """
        try:
            # Get all text from the paragraph and store existing font sizes
            full_text = "".join(run.text for run in paragraph.runs)
            
            # Store existing font sizes for each run
            existing_sizes = []
            for run in paragraph.runs:
                if run.text.strip():
                    existing_sizes.append(run.font.size)
            
            # If we have existing sizes, use the first one for all runs to ensure consistency
            # Or use the passed font_size parameter if provided
            consistent_font_size = None
            if font_size is not None:
                consistent_font_size = font_size
            elif existing_sizes:
                consistent_font_size = existing_sizes[0]
            
            # Check if there are any |BR| markers or \n characters
            if '|BR|' not in full_text and '\n' not in full_text:
                return
            
            # First split by |BR| markers, then by \n characters
            if '|BR|' in full_text:
                parts = full_text.split('|BR|')
            else:
                parts = full_text.split('\n')
            
            # Clear the paragraph
            paragraph.clear()
            
            # Set tight paragraph spacing to prevent excessive gaps
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            # Only set line spacing if it's not already set (to preserve custom line spacing)
            if paragraph.paragraph_format.line_spacing is None:
                paragraph.paragraph_format.line_spacing = 1.0
            # Preserve existing line spacing if it's already set (for THC_CBD markers)
            existing_line_spacing = paragraph.paragraph_format.line_spacing
            
            # Add each part as a separate run, with line breaks between them
            size_index = 0
            for i, part in enumerate(parts):
                if part.strip():  # Only add non-empty parts
                    run = paragraph.add_run(part.strip())
                    run.font.name = "Arial"
                    
                    # Check if this paragraph contains ratio content and should be bold
                    # This ensures multi-line ratio content stays bold
                    if any(pattern in full_text for pattern in [
                        'mg THC', 'mg CBD', 'mg CBC', 'mg CBG', 'mg CBN',
                        'THC:', 'CBD:', 'CBC:', 'CBG:', 'CBN:',
                        '1:1', '2:1', '3:1', '1:1:1', '2:1:1',
                        'RATIO_START', 'THC_CBD_START'
                    ]):
                        run.font.bold = True
                    
                    # Use consistent font size for all runs
                    if consistent_font_size:
                        run.font.size = consistent_font_size
                    else:
                        # Use a default size only if no existing size is available
                        run.font.size = Pt(12)
                    
                    # Add a line break after this part only if the next part is not empty
                    if i < len(parts) - 1 and parts[i + 1].strip():
                        # Use add_break() with WD_BREAK.LINE to create proper line breaks within the same paragraph
                        run.add_break(WD_BREAK.LINE)
            
            # Restore the original line spacing if it was set
            if 'existing_line_spacing' in locals() and existing_line_spacing != 1.0:
                paragraph.paragraph_format.line_spacing = existing_line_spacing
                # Also set at XML level for maximum compatibility
                pPr = paragraph._element.get_or_add_pPr()
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                spacing.set(qn('w:line'), str(int(existing_line_spacing * 240)))
                spacing.set(qn('w:lineRule'), 'auto')
            
            self.logger.debug(f"Converted {len(parts)-1} |BR| markers to line breaks")
            
        except Exception as e:
            self.logger.error(f"Error converting BR markers to line breaks: {e}")
            # Fallback: just remove the BR markers
            for run in paragraph.runs:
                run.text = run.text.replace('|BR|', ' ')

    def _fix_ratio_paragraph_spacing(self, doc):
        """
        Fix paragraph spacing for ratio content to prevent excessive gaps between lines.
        This ensures tight spacing for multi-line ratio content.
        """
        try:
            # Define patterns that indicate ratio content
            ratio_patterns = [
                'mg THC', 'mg CBD', 'mg CBG', 'mg CBN', 'mg CBC',
                'THC:', 'CBD:', 'CBG:', 'CBN:', 'CBC:',
                '1:1', '2:1', '3:1', '1:1:1', '2:1:1'
            ]
            
            def process_paragraph(paragraph):
                # Check if this paragraph contains ratio content
                text = paragraph.text.lower()
                if any(pattern.lower() in text for pattern in ratio_patterns):
                    # Check if this is THC_CBD content and use unified font sizing system
                    if 'thc:' in text and 'cbd:' in text:
                        # Use unified font sizing system for THC_CBD content
                        from src.core.generation.unified_font_sizing import get_line_spacing_by_marker
                        line_spacing = get_line_spacing_by_marker('THC_CBD', self.template_type)
                        if line_spacing:
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = line_spacing
                            # Set at XML level for maximum compatibility
                            pPr = paragraph._element.get_or_add_pPr()
                            spacing = pPr.find(qn('w:spacing'))
                            if spacing is None:
                                spacing = OxmlElement('w:spacing')
                                pPr.append(spacing)
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), str(int(line_spacing * 240)))
                            spacing.set(qn('w:lineRule'), 'auto')
                            return  # Skip the default 1.0 spacing for THC_CBD content
                    
                    # Set tight spacing for other ratio content (not THC_CBD)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    
                    # Also set tight spacing for any child paragraphs (in case of nested content)
                    for child_para in paragraph._element.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if hasattr(child_para, 'pPr') and child_para.pPr is not None:
                            # Set spacing properties at XML level for maximum compatibility
                            spacing = child_para.pPr.find('.//w:spacing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            if spacing is None:
                                spacing = OxmlElement('w:spacing')
                                child_para.pPr.append(spacing)
                            
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), '240')  # 1.0 line spacing (240 twips)
                            spacing.set(qn('w:lineRule'), 'auto')
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            process_paragraph(paragraph)
            
            # Process all paragraphs outside tables
            for paragraph in doc.paragraphs:
                process_paragraph(paragraph)
            
            self.logger.debug("Fixed paragraph spacing for ratio content")
            
        except Exception as e:
            self.logger.error(f"Error fixing ratio paragraph spacing: {e}")
            # Don't raise the exception - this is a formatting enhancement that shouldn't break the main process

    def _ensure_proper_centering(self, doc):
        """
        Ensure tables are properly centered in the document with correct margins and spacing.
        """
        try:
            # Set document margins to ensure proper centering
            for section in doc.sections:
                # Use smaller margins for vertical template to fit all 9 labels
                if self.template_type == 'vertical':
                    section.left_margin = Inches(0.25)
                    section.right_margin = Inches(0.25)
                    section.top_margin = Inches(0.25)
                    section.bottom_margin = Inches(0.25)
                else:
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
            
            # Remove any extra paragraphs that might affect centering
            paragraphs_to_remove = []
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip() and not paragraph.runs:
                    paragraphs_to_remove.append(paragraph)
            
            for paragraph in paragraphs_to_remove:
                paragraph._element.getparent().remove(paragraph._element)
            
            # Ensure all tables are properly centered
            for table in doc.tables:
                # Set table alignment to center
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Ensure table properties are set correctly
                tblPr = table._element.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                
                # Set table to fixed layout
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
                
                # Ensure table is not auto-fit
                table.autofit = False
                if hasattr(table, 'allow_autofit'):
                    table.allow_autofit = False
                

                
                # Calculate and set proper table width for perfect centering
                from src.core.constants import CELL_DIMENSIONS, GRID_LAYOUTS
                
                # Get individual cell dimensions and grid layout
                cell_dims = CELL_DIMENSIONS.get(self.template_type, {'width': 2.4, 'height': 2.4})
                grid_layout = GRID_LAYOUTS.get(self.template_type, {'rows': 3, 'cols': 3})
                
                # Calculate total table width: individual cell width * number of columns
                individual_cell_width = cell_dims['width']
                num_columns = grid_layout['cols']
                total_table_width = individual_cell_width * num_columns
                
                # Set table width to ensure proper centering
                table.width = Inches(total_table_width)
                
                # Also set the table width property in XML to ensure it's properly applied
                tblPr = table._element.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._element.insert(0, tblPr)
                
                # Set table width property
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is None:
                    tblW = OxmlElement('w:tblW')
                    tblPr.append(tblW)
                tblW.set(qn('w:w'), str(int(total_table_width * 1440)))  # Convert to twips
                tblW.set(qn('w:type'), 'dxa')
                
                # Skip width setting for horizontal, mini, and vertical templates since they should already be correct from template expansion
                if self.template_type != 'horizontal' and self.template_type != 'mini' and self.template_type != 'vertical':
                    # Check if table has columns before trying to modify grid
                    if len(table.columns) > 0:
                        tblGrid = table._element.find(qn('w:tblGrid'))
                        if tblGrid is not None:
                            # Remove existing grid and recreate with proper widths
                            tblGrid.getparent().remove(tblGrid)
                        
                        # Create new grid with proper column widths
                        tblGrid = OxmlElement('w:tblGrid')
                        # Use individual cell width directly from CELL_DIMENSIONS
                        col_width = cell_dims['width']
                        
                        for _ in range(len(table.columns)):
                            gridCol = OxmlElement('w:gridCol')
                            gridCol.set(qn('w:w'), str(int(col_width * 1440)))  # Convert to twips
                            tblGrid.append(gridCol)
                        
                        # Insert the grid at the beginning of the table element
                        table._element.insert(0, tblGrid)
                    
                    # Also ensure each cell has the correct width property
                    for row in table.rows:
                        for cell in row.cells:
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcW = tcPr.find(qn('w:tcW'))
                            if tcW is None:
                                tcW = OxmlElement('w:tcW')
                                tcPr.append(tcW)
                            tcW.set(qn('w:w'), str(int(col_width * 1440)))
                            tcW.set(qn('w:type'), 'dxa')
            
            self.logger.debug("Ensured proper table centering and document setup")
            
        except Exception as e:
            self.logger.error(f"Error ensuring proper centering: {e}")

    def _add_weight_units_markers(self, doc):
        """
        Add RATIO markers around weight units content for mini templates with classic types.
        This allows the post-processing to find and apply the correct font sizing.
        """
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Look for weight units content in individual runs
                            for run in paragraph.runs:
                                run_text = run.text
                                # Check if this run contains weight units content (ends with 'g' or 'mg', or contains specific patterns)
                                # More specific check to avoid marking brand names that contain 'g'
                                is_weight_unit = (
                                    run_text.strip().endswith('g') or 
                                    run_text.strip().endswith('mg') or
                                    re.match(r'^\d+\.?\d*\s*g$', run_text.strip()) or  # "1g", "1.5g"
                                    re.match(r'^\d+\.?\d*\s*mg$', run_text.strip()) or  # "100mg", "50.5mg"
                                    re.match(r'^\d+\.?\d*\s*g\s*x\s*\d+', run_text.strip()) or  # "1g x 2"
                                    re.match(r'^\d+\.?\d*\s*mg\s*x\s*\d+', run_text.strip())  # "100mg x 2"
                                )
                                
                                if is_weight_unit and 'RATIO_START' not in run_text:
                                    # This is likely weight units content that needs markers
                                    # Replace the run text with marked content
                                    run.text = f"RATIO_START{run_text}RATIO_END"
                                    run.font.name = "Arial"
                                    run.font.bold = True
                                    run.font.size = Pt(12)  # Default size, will be adjusted by post-processing
                                    
                                    self.logger.debug(f"Added RATIO markers around weight units: {run_text}")
            
        except Exception as e:
            self.logger.error(f"Error adding weight units markers: {e}")

    def _add_brand_markers(self, doc):
        """
        Add PRODUCTBRAND_CENTER markers around brand content for mini templates.
        This allows the post-processing to find and apply the correct font sizing.
        """
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Look for brand content in individual runs
                            for run in paragraph.runs:
                                run_text = run.text
                                # Check if this run contains brand content (not empty and not already marked)
                                # Only add markers to text that looks like brand names (not empty, not marked, not placeholders)
                                # IMPORTANT: Don't add brand markers if content is already wrapped in RATIO markers
                                if (run_text.strip() and 
                                    'PRODUCTBRAND_CENTER_START' not in run_text and 
                                    'RATIO_START' not in run_text and  # Don't mark content already in RATIO markers
                                    'RATIO_END' not in run_text and    # Don't mark content already in RATIO markers
                                    '{{' not in run_text and 
                                    '}}' not in run_text and
                                    len(run_text.strip()) > 0 and
                                    # Only mark content that looks like brand names (not numbers, not empty)
                                    not run_text.strip().isdigit() and
                                    not run_text.strip().startswith('$') and
                                    not run_text.strip().endswith('g') and
                                    not run_text.strip().endswith('mg')):
                                    # This is likely brand content that needs markers
                                    # Replace the run text with marked content
                                    run.text = f"PRODUCTBRAND_CENTER_START{run_text}PRODUCTBRAND_CENTER_END"
                                    run.font.name = "Arial"
                                    run.font.bold = True
                                    run.font.size = Pt(12)  # Default size, will be adjusted by post-processing
                                    
                                    self.logger.debug(f"Added PRODUCTBRAND_CENTER markers around brand: {run_text}")
            
        except Exception as e:
            self.logger.error(f"Error adding brand markers: {e}")



    def _get_template_specific_font_size(self, content, marker_name):
        """
        Get font size using the unified font sizing system.
        """
        # Use unified font sizing with appropriate complexity type
        complexity_type = 'mini' if self.template_type == 'mini' else 'standard'
        return get_font_size_by_marker(content, marker_name, self.template_type, self.scale_factor)

    def fix_hyphen_spacing(self, text):
        """Replace regular hyphens with non-breaking hyphens to prevent line breaks, 
        but add line breaks before hanging hyphens.
        Used for general text formatting to prevent unwanted line breaks at hyphens."""
        if not text:
            return text
        
        # First, normalize hyphen spacing to ensure consistent format
        text = re.sub(r'\s*-\s*', ' - ', text)
        
        # Check for hanging hyphens (hyphen at the end of a line or followed by a space and then end)
        # Pattern: space + hyphen + space + end of string, or space + hyphen + end of string
        if re.search(r' - $', text) or re.search(r' - \s*$', text):
            # Add line break before the hanging hyphen
            text = re.sub(r' - (\s*)$', r'\n- \1', text)
        
        return text

    def format_with_soft_hyphen(self, text):
        """Format text with soft hyphen + nonbreaking space + value pattern.
        Used for specific formatting where you want a soft hyphen followed by nonbreaking space."""
        if not text:
            return text
        # Replace any leading hyphens/spaces with a single soft hyphen + nonbreaking space
        text = re.sub(r'^[\s\-]+', '\u00AD\u00A0', text)
        # If it didn't start with hyphen/space, prepend
        if not text.startswith('\u00AD\u00A0'):
            text = f'\u00AD\u00A0{text}'
        return text

    def format_classic_ratio(self, text, record=None):
        """
        Format ratio for classic types. Handles various input formats and converts them to the standard display format.
        """
        if not text:
            return text
        
        # Clean the text and normalize
        text = text.strip()
        
        # Handle the default "THC:|BR|CBD:" format from excel processor
        if text == "THC:|BR|CBD:":
            # If we have record data, try to get actual THC and CBD values from columns
            if record:
                # Get THC value from AI column (Total THC)
                thc_value = record.get('AI', '')
                if thc_value:
                    thc_value = str(thc_value).strip()
                    # If Total THC is 0 or lower than THCA, use AJ column instead
                    if thc_value == '0' or thc_value == '0.0' or thc_value == '':
                        thc_value = record.get('AJ', '')
                        if thc_value:
                            thc_value = str(thc_value).strip()
                
                # Get CBD value from AK column
                cbd_value = record.get('AK', '')
                if cbd_value:
                    cbd_value = str(cbd_value).strip()
                
                # Clean up values (remove 'nan', empty strings, etc.)
                if thc_value in ['nan', 'NaN', '']:
                    thc_value = ''
                if cbd_value in ['nan', 'NaN', '']:
                    cbd_value = ''
                
                # Format with actual values if available
                if thc_value and cbd_value:
                    return f"THC: {thc_value}% CBD: {cbd_value}%"
                elif thc_value:
                    return f"THC: {thc_value}%"
                elif cbd_value:
                    return f"CBD: {cbd_value}%"
            
            # Fallback to default format if no record data or no values
            return "THC: CBD:"
        
        # If the text already contains THC/CBD format, return as-is
        if 'THC:' in text and 'CBD:' in text:
            return text
        
        # If the text contains mg values, return as-is (let text_processing handle it)
        if 'mg' in text.lower():
            return text
        
        # If the text contains simple ratios (like 1:1:1), format with spaces
        if ':' in text and any(c.isdigit() for c in text):
            # Add spaces around colons for better readability
            # Handle 3-part ratios first to avoid conflicts
            text = re.sub(r'(\d+):(\d+):(\d+)', r'\1: \2: \3', text)
            # Then handle 2-part ratios
            text = re.sub(r'(\d+):(\d+)', r'\1: \2', text)
            return text
        
        # Common patterns for THC/CBD ratios
        thc_patterns = [
            r'THC[:\s]*([0-9.]+)%?',
            r'([0-9.]+)%?\s*THC',
            r'([0-9.]+)\s*THC'
        ]
        
        cbd_patterns = [
            r'CBD[:\s]*([0-9.]+)%?',
            r'([0-9.]+)%?\s*CBD',
            r'([0-9.]+)\s*CBD'
        ]
        
        thc_value = None
        cbd_value = None
        
        # Extract THC value
        for pattern in thc_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                thc_value = match.group(1)
                break
        
        # Extract CBD value
        for pattern in cbd_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                cbd_value = match.group(1)
                break
        
        # If we found both values, format them
        if thc_value and cbd_value:
            # Keep on same line without line breaks
            return f"THC: {thc_value}% CBD: {cbd_value}%"
        elif thc_value:
            return f"THC: {thc_value}%"
        elif cbd_value:
            return f"CBD: {cbd_value}%"
        else:
            # If no clear THC/CBD pattern found, return the original text
            return text

    def format_joint_ratio_pack(self, text):
        """
        Format JointRatio as: [amount]g x [count] Pack
        Handles various input formats and normalizes them to standard format.
        For single units, shows just the weight (e.g., "1g" instead of "1g x 1 Pack").
        """
        if not text:
            return text
            
        # Convert to string and clean up
        text = str(text).strip()
        
        # Remove any leading/trailing spaces and hyphens
        text = re.sub(r'^[\s\-]+', '', text)
        text = re.sub(r'[\s\-]+$', '', text)
        
        # Handle various input patterns
        patterns = [
            # Standard format: "1g x 2 Pack"
            r"([0-9.]+)g\s*x\s*([0-9]+)\s*pack",
            # Compact format: "1gx2Pack"
            r"([0-9.]+)g\s*x?\s*([0-9]+)pack",
            # With spaces: "1g x 2 pack"
            r"([0-9.]+)g\s*x\s*([0-9]+)\s*pack",
            # Just weight: "1g"
            r"([0-9.]+)g",
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                amount = match.group(1).strip()
                # Try to get count, default to 1 if not found
                try:
                    count = match.group(2).strip()
                    if count and count.isdigit():
                        count_int = int(count)
                        if count_int == 1:
                            # For single units, just show the weight
                            formatted = f"{amount}g"
                        else:
                            # For multiple units, show the full pack format
                            formatted = f"{amount}g x {count} Pack"
                    else:
                        # Only amount found (like "1g") - show just the weight
                        formatted = f"{amount}g"
                except IndexError:
                    # Only amount found (like "1g") - show just the weight
                    formatted = f"{amount}g"
                return formatted
        
        # If no pattern matches, return the original text
        return text

    def format_thc_cbd_vertical_alignment(self, text):
        """
        Format THC_CBD content for vertical templates with right-aligned percentages.
        Keeps labels (THC:, CBD:) left-aligned but right-aligns the percentage numbers.
        Uses spaces instead of tabs to prevent percentage value breaking.
        """
        if not text:
            return text
        
        # Split into lines
        lines = text.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line contains both THC and CBD (same line)
            if 'THC:' in line and 'CBD:' in line and '%' in line:
                # Split the line into THC and CBD parts
                if 'CBD:' in line:
                    # Find where CBD starts
                    cbd_start = line.find('CBD:')
                    thc_part = line[:cbd_start].strip()
                    cbd_part = line[cbd_start:].strip()
                    
                    # Check if there are other cannabinoids after CBD
                    remaining_content = ''
                    if 'CBC:' in cbd_part:
                        cbc_start = cbd_part.find('CBC:')
                        cbd_part_only = cbd_part[:cbc_start].strip()
                        remaining_content = cbd_part[cbc_start:].strip()
                        cbd_part = cbd_part_only
                    
                    # Format THC part - use spaces for alignment
                    thc_match = re.search(r'(THC:\s*)([0-9.]+)%', thc_part)
                    if thc_match:
                        thc_label = thc_match.group(1).rstrip()  # Remove trailing spaces
                        thc_percentage = thc_match.group(2)
                        # Add 3 spaces before the percentage value
                        formatted_thc = f"{thc_label}   {thc_percentage}%"
                    else:
                        formatted_thc = thc_part
                    
                    # Format CBD part - use spaces for alignment
                    cbd_match = re.search(r'(CBD:\s*)([0-9.]+)%', cbd_part)
                    if cbd_match:
                        cbd_label = cbd_match.group(1).rstrip()  # Remove trailing spaces
                        cbd_percentage = cbd_match.group(2)
                        # Add 3 spaces before the percentage value
                        formatted_cbd = f"{cbd_label}   {cbd_percentage}%"
                    else:
                        formatted_cbd = cbd_part
                    
                    # Combine with line breaks
                    if remaining_content:
                        formatted_line = f"{formatted_thc}\n{formatted_cbd}\n{remaining_content}"
                    else:
                        formatted_line = f"{formatted_thc}\n{formatted_cbd}"
                    formatted_lines.append(formatted_line)
                else:
                    formatted_lines.append(line)
            # Check if this line contains only THC with percentage
            elif 'THC:' in line and '%' in line and 'CBD:' not in line:
                match = re.search(r'(THC:\s*)([0-9.]+)%', line)
                if match:
                    label = match.group(1).rstrip()  # Remove trailing spaces
                    percentage = match.group(2)
                    # Use 3 spaces to align percentage to the right
                    formatted_line = f"{label}   {percentage}%"
                    formatted_lines.append(formatted_line)
                else:
                    formatted_lines.append(line)
            # Check if this line contains only CBD with percentage
            elif 'CBD:' in line and '%' in line and 'THC:' not in line:
                match = re.search(r'(CBD:\s*)([0-9.]+)%', line)
                if match:
                    label = match.group(1).rstrip()  # Remove trailing spaces
                    percentage = match.group(2)
                    # Use 3 spaces to align percentage to the right
                    formatted_line = f"{label}   {percentage}%"
                    formatted_lines.append(formatted_line)
                else:
                    formatted_lines.append(line)
            else:
                # Keep other lines as-is (like CBC: 1%)
                formatted_lines.append(line)
        
        return '\n'.join(formatted_lines)

    def _process_combined_lineage_vendor(self, paragraph, lineage_content, vendor_content):
        """
        Process combined lineage and vendor text with different font sizes.
        This handles the case where lineage and product vendor are on the same line.
        Lineage is left-aligned, vendor is right-aligned.
        IMPORTANT: Product Vendor should never be split up - if Lineage is too long, it should break to new line.
        SPECIAL RULE: For Vertical template, if Lineage is "Hybrid/Indica" or "Hybrid/Sativa", automatically put ProductVendor on next line.
        """
        try:
            # Clear the paragraph content
            paragraph.clear()
            
            # SPECIAL RULE: For Vertical template, automatically force vendor to next line for specific lineages
            if (self.template_type == 'vertical' and 
                lineage_content and 
                lineage_content.strip().upper() in ['HYBRID/INDICA', 'HYBRID/SATIVA'] and
                vendor_content and vendor_content.strip()):
                
                self.logger.debug(f"Vertical template: Forcing vendor to next line for lineage '{lineage_content}'")
                self._process_lineage_vendor_two_lines(paragraph, lineage_content, vendor_content)
                return
            
            # Check if we need to split to multiple lines due to content length
            # Calculate approximate character limits based on template type
            if self.template_type == 'mini':
                max_chars_per_line = 25
            elif self.template_type == 'vertical':
                max_chars_per_line = 35
            else:  # horizontal, double
                max_chars_per_line = 45
            
            # Check if combined content would be too long for one line
            combined_length = len(lineage_content or '') + len(vendor_content or '')
            
            if combined_length > max_chars_per_line and vendor_content and vendor_content.strip():
                # Split to two lines: lineage on first line, vendor on second line
                self._process_lineage_vendor_two_lines(paragraph, lineage_content, vendor_content)
                return
            
            # Original single-line processing
            # Set paragraph to justified alignment to allow for right-aligned vendor
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add lineage with larger font size (left-aligned)
            if lineage_content and lineage_content.strip():
                lineage_run = paragraph.add_run(lineage_content.strip())
                lineage_run.font.name = "Arial"
                lineage_run.font.bold = True
                
                # Get lineage font size
                product_type = None
                if hasattr(self, 'current_product_type'):
                    product_type = self.current_product_type
                elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                    product_type = self.label_context['ProductType']
                
                lineage_font_size = get_font_size_by_marker(lineage_content, 'LINEAGE', self.template_type, self.scale_factor, product_type)
                set_run_font_size(lineage_run, lineage_font_size)
            
            # Add tab character to push vendor to the right (only if vendor content exists)
            if lineage_content and vendor_content:
                tab_run = paragraph.add_run("\t")
                tab_run.font.name = "Arial"
                tab_run.font.bold = True
                # Use lineage font size for tab to maintain alignment
                set_run_font_size(tab_run, lineage_font_size)
            
            # Add vendor with smaller font size (right-aligned)
            if vendor_content and vendor_content.strip():
                vendor_run = paragraph.add_run(vendor_content.strip())
                vendor_run.font.name = "Arial"
                vendor_run.font.bold = False
                vendor_run.font.italic = True  # Make vendor text italic
                
                # Set vendor color to light gray (#CCCCCC)
                from docx.shared import RGBColor
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Ensure the color is applied by setting it explicitly
                vendor_run.font.color.theme_color = None  # Clear any theme color
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Get vendor font size (smaller than lineage)
                vendor_font_size = get_font_size_by_marker(vendor_content, 'PRODUCTVENDOR', self.template_type, self.scale_factor)
                set_run_font_size(vendor_run, vendor_font_size)
            
            # Set tab stops to position vendor on the right (only if vendor content exists)
            if vendor_content:
                # Clear existing tab stops
                paragraph.paragraph_format.tab_stops.clear_all()
                # Add right-aligned tab stop at the right margin - positioned further right for full justification
                if self.template_type == 'mini':
                    tab_position = Inches(1.7)  # Increased for mini template
                elif self.template_type == 'vertical':
                    tab_position = Inches(2.3)  # Increased for vertical template
                else:  # horizontal, double
                    tab_position = Inches(3.2)  # Further increased for horizontal/double templates
                
                paragraph.paragraph_format.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT)
                
                # Alternative: Use multiple tab stops for more aggressive right positioning
                # This creates additional tab stops to ensure the vendor text reaches the right edge
                if self.template_type in ['horizontal', 'double']:
                    # Add an additional tab stop even further right as backup
                    backup_tab_position = Inches(3.5)
                    paragraph.paragraph_format.tab_stops.add_tab_stop(backup_tab_position, WD_TAB_ALIGNMENT.RIGHT)
            else:
                # For non-classic products without vendor, use left alignment
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Handle left indentation based on lineage content type
            if lineage_content:
                classic_lineages = [
                    "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                    "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                ]
                if lineage_content.upper() in classic_lineages and lineage_content.upper() != "PARAPHERNALIA":
                    if self.template_type in {"horizontal", "double", "vertical"}:
                        paragraph.paragraph_format.left_indent = Inches(0.15)
            
            self.logger.debug(f"Processed combined lineage/vendor with right-aligned vendor: lineage='{lineage_content}', vendor='{vendor_content}'")
            
        except Exception as e:
            self.logger.error(f"Error processing combined lineage/vendor: {e}")
            # Fallback: use default processing
            paragraph.clear()
            combined_text = f"{lineage_content or ''}  {vendor_content or ''}".strip()
            if combined_text:
                run = paragraph.add_run(combined_text)

    def _process_lineage_vendor_two_lines(self, paragraph, lineage_content, vendor_content):
        """
        Process lineage and vendor on two separate lines to prevent vendor splitting.
        Lineage goes on the first line, vendor goes on the second line.
        """
        try:
            # Clear the paragraph content
            paragraph.clear()
            
            # Set paragraph to left alignment for two-line layout
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add lineage on first line with larger font size
            if lineage_content and lineage_content.strip():
                lineage_run = paragraph.add_run(lineage_content.strip())
                lineage_run.font.name = "Arial"
                lineage_run.font.bold = True
                
                # Get lineage font size
                product_type = None
                if hasattr(self, 'current_product_type'):
                    product_type = self.current_product_type
                elif hasattr(self, 'label_context') and 'ProductType' in self.label_context:
                    product_type = self.label_context['ProductType']
                
                lineage_font_size = get_font_size_by_marker(lineage_content, 'LINEAGE', self.template_type, self.scale_factor, product_type)
                set_run_font_size(lineage_run, lineage_font_size)
            
            # Add line break
            if lineage_content and vendor_content:
                paragraph.add_run("\n")
            
            # Add vendor on second line with smaller font size
            if vendor_content and vendor_content.strip():
                vendor_run = paragraph.add_run(vendor_content.strip())
                vendor_run.font.name = "Arial"
                vendor_run.font.bold = False
                vendor_run.font.italic = True  # Make vendor text italic
                
                # Set vendor color to light gray (#CCCCCC)
                from docx.shared import RGBColor
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Ensure the color is applied by setting it explicitly
                vendor_run.font.color.theme_color = None  # Clear any theme color
                vendor_run.font.color.rgb = RGBColor(204, 204, 204)  # #CCCCCC
                
                # Get vendor font size (smaller than lineage)
                vendor_font_size = get_font_size_by_marker(vendor_content, 'PRODUCTVENDOR', self.template_type, self.scale_factor)
                set_run_font_size(vendor_run, vendor_font_size)
            
            # Handle left indentation based on lineage content type
            if lineage_content:
                classic_lineages = [
                    "SATIVA", "INDICA", "HYBRID", "HYBRID/SATIVA", "HYBRID/INDICA", 
                    "CBD", "MIXED", "PARAPHERNALIA", "PARA"
                ]
                if lineage_content.upper() in classic_lineages and lineage_content.upper() != "PARAPHERNALIA":
                    if self.template_type in {"horizontal", "double", "vertical"}:
                        paragraph.paragraph_format.left_indent = Inches(0.15)
            
            self.logger.debug(f"Processed lineage/vendor on two lines: lineage='{lineage_content}', vendor='{vendor_content}'")
            
        except Exception as e:
            self.logger.error(f"Error processing lineage/vendor on two lines: {e}")
            # Fallback: use single line processing
            self._process_combined_lineage_vendor(paragraph, lineage_content, vendor_content)

    def _detect_and_process_combined_lineage_vendor(self, paragraph):
        """
        Detect if paragraph contains combined lineage and vendor markers and process them separately.
        Remove vendor for non-classic product types.
        """
        # Check if this paragraph has already been processed for combined lineage/vendor
        if hasattr(paragraph, '_combined_lineage_vendor_processed'):
            return True
        
        full_text = "".join(run.text for run in paragraph.runs)
        
        # Check if both lineage and vendor markers are present
        lineage_start = "LINEAGE_START"
        lineage_end = "LINEAGE_END"
        vendor_start = "PRODUCTVENDOR_START"
        vendor_end = "PRODUCTVENDOR_END"
        
        if (lineage_start in full_text and lineage_end in full_text and 
            vendor_start in full_text and vendor_end in full_text):
            
            try:
                # Extract lineage content
                lineage_start_idx = full_text.find(lineage_start) + len(lineage_start)
                lineage_end_idx = full_text.find(lineage_end)
                lineage_content = full_text[lineage_start_idx:lineage_end_idx]
                
                # Extract vendor content
                vendor_start_idx = full_text.find(vendor_start) + len(vendor_start)
                vendor_end_idx = full_text.find(vendor_end)
                vendor_content = full_text[vendor_start_idx:vendor_end_idx]
                
                # Note: Product type filtering is now handled in _build_label_context
                # This method only processes the content that's already been filtered
                
                # Process with different font sizes
                self._process_combined_lineage_vendor(paragraph, lineage_content, vendor_content)
                
                # Mark this paragraph as processed to prevent re-processing
                paragraph._combined_lineage_vendor_processed = True
                
                return True
                
            except Exception as e:
                self.logger.error(f"Error detecting combined lineage/vendor: {e}")
                return False
        
        return False

    def _clean_doh_cells_before_processing(self, doc):
        """
        Clean up DOH cells before processing to ensure no content interferes with image positioning.
        This should be called before DOH images are inserted.
        """
        try:
            from docx.oxml.ns import qn
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains DOH placeholder
                        cell_text = cell.text.strip()
                        if '{{Label' in cell_text and '.DOH}}' in cell_text:
                            # Clear the cell content to prepare for image insertion
                            cell._tc.clear_content()
                            
                            # Add a single empty paragraph to maintain cell structure
                            paragraph = cell.add_paragraph()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Set minimal spacing
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            self.logger.debug("Cleaned DOH cell for image insertion")
                            
        except Exception as e:
            self.logger.warning(f"Error cleaning DOH cells: {e}")

    def _ensure_doh_image_centering(self, doc):
        """
        Ensure DOH images are properly centered in all cells.
        This method specifically looks for InlineImage objects and centers them.
        """
        try:
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
            from docx.oxml.ns import qn
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Check if this cell contains a DOH image
                        has_doh_image = False
                        image_run = None
                        
                        # First pass: identify if cell has DOH image
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                                    has_doh_image = True
                                    image_run = run
                                    break
                            if has_doh_image:
                                break
                        
                        if has_doh_image:
                            # Clear the entire cell content first
                            cell._tc.clear_content()
                            
                            # Create a new paragraph for the image
                            paragraph = cell.add_paragraph()
                            
                            # Add the image run to the new paragraph
                            if image_run:
                                # Copy the image element to the new paragraph
                                new_run = paragraph.add_run()
                                new_run._element.append(image_run._element)
                                
                                # Ensure the image has proper text content
                                if not new_run.text:
                                    new_run.text = '\u00A0'  # Non-breaking space
                            
                            # Set perfect centering
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph.paragraph_format.space_before = Pt(0)
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.line_spacing = 1.0
                            
                            # Set cell vertical alignment to center
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            
                            # Ensure no extra spacing at XML level
                            pPr = paragraph._element.get_or_add_pPr()
                            spacing = pPr.find(qn('w:spacing'))
                            if spacing is None:
                                spacing = OxmlElement('w:spacing')
                                pPr.append(spacing)
                            spacing.set(qn('w:before'), '0')
                            spacing.set(qn('w:after'), '0')
                            spacing.set(qn('w:line'), '240')  # 1.0 line spacing
                            spacing.set(qn('w:lineRule'), 'auto')
                            
                            self.logger.debug("Perfectly centered DOH image in cell")
                                
        except Exception as e:
            self.logger.warning(f"Error in DOH image centering: {e}")



__all__ = ['get_font_scheme', 'TemplateProcessor']