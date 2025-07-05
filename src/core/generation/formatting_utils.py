# formatting_utils.py has been split into modular files for maintainability.
# Please use the following modules instead:
#   - font_sizing.py (font size logic)
#   - docx_formatting.py (docx/table formatting)
#   - text_processing.py (text formatting)
#   - context_builders.py (context/chunk builders)

# Import functions from modular files for backward compatibility:
from .docx_formatting import *
from .text_processing import *
from .context_builders import *
from .font_sizing import set_run_font_size
from .template_processor import get_font_scheme, TemplateProcessor

import warnings
warnings.warn(
    'formatting_utils.py is deprecated. Use the new split modules instead.',
    DeprecationWarning
)

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer
from io import BytesIO
import logging
import re
import os
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
import logging
import re
from html import unescape
from src.core.formatting.markers import wrap_with_marker, FIELD_MARKERS, unwrap_marker
from copy import deepcopy
from docxtpl import DocxTemplate, InlineImage
from docxtpl.inline_image import InlineImage
from typing import Optional, Dict, Any
from src.core.utils import resource_path
import traceback
from pathlib import Path

logger = logging.getLogger(__name__)


# Performance optimization: disable debug logging in production
DEBUG_ENABLED = False

__all__ = [
    'process_chunk',
    'get_font_scheme',
    'TemplateProcessor',
    'build_context',
    'build_label_context',
    'format_price',
    'format_ratio_multiline',
    'format_description_text',
    'format_cannabinoid_content',
    'replace_placeholder_with_markers',
    'process_chunk',
    'remove_marker',
    'process_doh_image',
]

# Constants
WORD_WEIGHT = 2
AGT_ORANGE = "FF8C00"
AGT_GREEN = "006400"
AGT_WHITE = "FFFFFF"

def safe_get(record, key, default=''):
    """Safely get a value from a record, handling None and NaN cases."""
    if isinstance(record, dict):
        value = record.get(key, default)
    else:
        value = getattr(record, key, default)
    if value is None or str(value).strip().lower() == 'nan':
        return default
    return str(value).strip()

def safe_get_text(value):
    """Safely get text from a value, handling None and NaN cases."""
    if value is None or str(value).strip().lower() == 'nan':
        return ''
    if isinstance(value, dict):
        return value.get('text', '')
    return str(value).strip()

def format_description_text(text, max_width=40):
    """Format description text with proper wrapping and line breaks."""
    if not text:
        return ""
    
    # Split text into words
    words = text.split()
    lines = []
    current_line = []
    current_length = 0
    
    for word in words:
        # Check if adding this word would exceed max width
        if current_length + len(word) + 1 > max_width:
            # Start a new line
            lines.append(' '.join(current_line))
            current_line = [word]
            current_length = len(word)
        else:
            # Add to current line
            current_line.append(word)
            current_length += len(word) + 1
    
    # Add the last line
    if current_line:
        lines.append(' '.join(current_line))
    
    return '\n'.join(lines)

def format_ratio_multiline(text):
    """Format ratio text into clean content without markers."""
    if not isinstance(text, str):
        return ""
    
    # First remove all markers completely
    content = text.replace('THC_CBD_START', '').replace('THC_CBD_END', '')
    content = content.replace('RATIO_START', '').replace('RATIO_END', '')
    
    # Clean up whitespace
    content = ' '.join(content.split())
    
    # Handle special cases
    if not content or content in ["", "CBD", "THC", "CBD:", "THC:", "CBD:\n", "THC:\n"]:
        return ""
    
    # Handle THC:/CBD: format
    if 'THC:' in content and 'CBD:' in content:
        # Ensure proper line breaks
        if '\n' not in content:
            content = content.replace('CBD:', '\nCBD:')
        return content.strip()
    
    # Handle mg values consistently
    if 'mg' in content.lower():
        parts = []
        current_part = []
        
        for word in content.split():
            word = word.strip()
            if 'mg' in word.lower():
                # If we have accumulated other words, join them
                if current_part:
                    parts.append(' '.join(current_part))
                    current_part = []
                # Add mg value directly
                parts.append(word)
            else:
                current_part.append(word)
        
        # Add any remaining words
        if current_part:
            parts.append(' '.join(current_part))
        
        # Join all parts with spaces
        content = ' '.join(parts)
    
    # Handle ratio format (e.g. "1:1:1", "1:1")
    if ':' in content and any(c.isdigit() for c in content):
        # Ensure proper spacing around colons
        content = re.sub(r'(\d+):(\d+)', r'\1: \2', content)
        content = re.sub(r'(\d+):(\d+):(\d+)', r'\1: \2: \3', content)
    
    if DEBUG_ENABLED:
        logger.debug(f"Formatted ratio from '{text}' to '{content}'")
    return content.strip()

def replace_placeholder_with_markers(doc, placeholder, marker_value):
    """Replace placeholder text with marked content in a document."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = ''.join(run.text for run in para.runs)
                    if placeholder in full_text:
                        # Extract marker type from marker_value
                        match = re.match(r'(\w+)_START.*?(\1)_END', marker_value)
                        if match:
                            marker_type = match.group(1)
                            # Extract content between markers
                            content = re.search(f'{marker_type}_START(.*?){marker_type}_END', marker_value)
                            if content:
                                # Clear existing runs
                                for run in para.runs:
                                    run.text = ''
                                # Add new run with marker value
                                new_run = para.add_run(marker_value)
                                new_run.font.name = "Arial"
                                new_run.font.bold = True
                                if not isinstance(new_run.font.size, Pt):
                                    new_run.font.size = Pt(12)
                                set_run_font_size(new_run, Pt(12))
    return doc

def format_cannabinoid_content(text):
    """Format cannabinoid content text with proper spacing and line breaks."""
    if not text:
        return ""
        
    # Clean up the text
    text = text.strip()
    
    # Handle THC/CBD ratio format
    if 'THC:' in text and 'CBD:' in text:
        return text  # Already formatted correctly
        
    # Handle mg amounts
    if 'mg' in text.lower():
        parts = text.split()
        formatted_parts = []
        for part in parts:
            if 'mg' in part.lower():
                # Add spacing around mg amounts
                formatted_parts.append(part.strip())
            else:
                formatted_parts.append(part)
        return ' '.join(formatted_parts)
        
    # Handle ratio format (e.g. "1:1:1")
    if ':' in text:
        return text.strip()
        
    return text

def process_chunk(chunk, orientation='vertical', scale_factor=1.0):
    """Process a chunk of text with proper formatting and font sizing."""
    if not chunk:
        return None

    # Create a temporary document to handle the chunk
    doc = Document()
    paragraph = doc.add_paragraph()
    
    # Clean the chunk text
    text = chunk.strip()
    if not text:
        return None
        
    # Determine content type
    is_ratio = any(x in text for x in [
        'RATIO_START', 
        'THC_CBD_START',
        'THC:', 
        'CBD:', 
        'mg THC',
        'mg CBD',
        'mg'
    ])
    
    is_price = 'PRICE_START' in text
    is_lineage = 'LINEAGE_START' in text
    is_description = 'DESC_START' in text
    
    # Apply appropriate formatting
    if is_ratio:
        formatted_text = format_ratio_multiline(text)
    elif is_price:
        formatted_text = format_price(text)
    elif is_lineage:
        formatted_text = text.upper()
    elif is_description:
        formatted_text = format_description_text(text)
    else:
        formatted_text = text
    
    # Add the formatted text to the paragraph
    run = paragraph.add_run(formatted_text)
    run.font.name = "Arial"
    run.font.bold = True
    
    return doc

def format_price(price):
    """Format price text with proper formatting."""
    if not price:
        return ""
    
    # Remove any existing markers
    price = price.replace('PRICE_START', '').replace('PRICE_END', '')
    
    # Clean up the price text
    price = price.strip()
    
    # Handle special cases
    if not price or price.lower() in ['nan', 'none', '']:
        return ""
    
    # Ensure proper formatting for currency
    if price and not price.startswith('$'):
        price = f"${price}"
    
    return price

def remove_marker(value):
    """Remove markers from a value."""
    if not isinstance(value, str):
        return value
    
    # Remove common markers
    markers_to_remove = [
        'PRICE_START', 'PRICE_END',
        'LINEAGE_START', 'LINEAGE_END',
        'DESC_START', 'DESC_END',
        'THC_CBD_START', 'THC_CBD_END',
        'RATIO_START', 'RATIO_END',
        'PRODUCTBRAND_CENTER_START', 'PRODUCTBRAND_CENTER_END',
        'PRODUCTSTRAIN_START', 'PRODUCTSTRAIN_END',
        'WEIGHTUNITS_START', 'WEIGHTUNITS_END'
    ]
    
    for marker in markers_to_remove:
        value = value.replace(marker, '')
    
    return value.strip()

def process_doh_image(doh_value, product_type):
    """Process DOH image based on DOH value and product type."""
    doh_value = str(doh_value).strip().upper()
    product_type = str(product_type).strip().lower()
    
    if doh_value == "YES":
        # Use HighCBD.png if product_type starts with 'high cbd'
        if product_type.startswith('high cbd'):
            return resource_path(os.path.join("templates", "HighCBD.png"))
        else:
            return resource_path(os.path.join("templates", "DOH.png"))
    
    return ""





