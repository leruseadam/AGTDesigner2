from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)

# Define colors for lineage
COLORS = {
    'SATIVA': 'ED4123',
    'INDICA': '9900FF',
    'HYBRID': '009900',
    'HYBRID_INDICA': '9900FF',
    'HYBRID_SATIVA': 'ED4123',
    'CBD': 'F1C232',
    'MIXED': '0021F5',
    'PARA': 'FFC0CB'
}

def apply_lineage_colors(doc):
    """Apply lineage colors to all cells based on keywords in cell text."""
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.upper()
                    color_hex = None
                    
                    # Remove marker wrappers for robust matching
                    for marker in ["LINEAGE_START", "LINEAGE_END", "PRODUCTSTRAIN_START", "PRODUCTSTRAIN_END"]:
                        text = text.replace(marker, "")
                    text = text.strip()
                    
                    # Extract lineage and product type information from embedded format
                    lineage_part = text
                    product_type = ""
                    is_classic = False
                    
                    if "_PRODUCT_TYPE_" in text:
                        parts = text.split("_PRODUCT_TYPE_")
                        if len(parts) >= 2:
                            lineage_part = parts[0]
                            remaining = parts[1]
                            # Extract product type and classic flag
                            if "_IS_CLASSIC_" in remaining:
                                type_parts = remaining.split("_IS_CLASSIC_")
                                if len(type_parts) >= 2:
                                    product_type = type_parts[0]
                                    is_classic = type_parts[1].lower() == "true"
                    
                    # Define classic product types
                    classic_types = ["flower", "pre-roll", "infused pre-roll", "concentrate", "solventless concentrate", "vape cartridge", "rso/co2 tankers"]
                    
                    # Apply lineage coloring logic
                    if "PARAPHERNALIA" in text:
                        color_hex = COLORS['PARA']
                    elif "HYBRID/INDICA" in text or "HYBRID INDICA" in text:
                        color_hex = COLORS['HYBRID_INDICA']
                    elif "HYBRID/SATIVA" in text or "HYBRID SATIVA" in text:
                        color_hex = COLORS['HYBRID_SATIVA']
                    elif "SATIVA" in text:
                        color_hex = COLORS['SATIVA']
                    elif "INDICA" in text:
                        color_hex = COLORS['INDICA']
                    elif "HYBRID" in text:
                        color_hex = COLORS['HYBRID']
                    elif "CBD" in text or "CBD_BLEND" in text:
                        color_hex = COLORS['CBD']
                    elif "MIXED" in text:
                        # For non-classic product types, Mixed should be blue
                        if not is_classic and product_type and product_type not in classic_types:
                            color_hex = COLORS['MIXED']  # Blue for non-classic Mixed
                        else:
                            color_hex = COLORS['MIXED']  # Default blue for Mixed
                    
                    if color_hex:
                        # Set cell background color
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        for old_shd in tcPr.findall(qn('w:shd')):
                            tcPr.remove(old_shd)
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), color_hex)
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        tcPr.append(shd)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.bold = True
                                run.font.name = "Arial"
        logger.debug("Applied lineage colors to document")
        return doc
    except Exception as e:
        logger.error(f"Error applying lineage colors: {str(e)}")
        raise

def fix_table_row_heights(doc, template_type):
    """Fix table row heights based on template type."""
    try:
        row_height = {
            'horizontal': 2.4,
            'vertical': 3.4,
            'mini': 1.75,
            'inventory': 2.0
        }.get(template_type, 2.4)
        for table in doc.tables:
            for row in table.rows:
                row.height = Inches(row_height)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        logger.debug(f"Fixed table row heights for template type: {template_type}")
        return doc
    except Exception as e:
        logger.error(f"Error fixing table row heights: {str(e)}")
        raise

def safe_fix_paragraph_spacing(doc):
    """
    Safely adjust paragraph spacing in document without affecting cell colors.
    """
    try:
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
        logger.debug("Successfully fixed paragraph spacing")
    except Exception as e:
        logger.error(f"Error in safe_fix_paragraph_spacing: {str(e)}")
        raise

def apply_conditional_formatting(doc, conditions=None):
    """
    Apply conditional formatting to document elements based on specified conditions.
    """
    try:
        if not conditions:
            conditions = {
                'PRICE_START': {
                    'bold': True,
                    'color': RGBColor(0, 0, 0),
                    'size': Pt(12)
                },
                'LINEAGE_START': {
                    'bold': True,
                    'color': RGBColor(255, 255, 255),
                    'size': Pt(11)
                }
            }
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for marker, formatting in conditions.items():
                                if marker in run.text:
                                    if formatting.get('bold') is not None:
                                        run.font.bold = formatting['bold']
                                    if formatting.get('color') is not None:
                                        run.font.color.rgb = formatting['color']
                                    if formatting.get('size') is not None:
                                        run.font.size = formatting['size']
        logger.debug("Applied conditional formatting to document")
        return doc
    except Exception as e:
        logger.error(f"Error applying conditional formatting: {str(e)}")
        raise

def set_cell_background(cell, color_hex):
    """Set cell background color with white text."""
    try:
        color_hex = color_hex.upper().strip('#')
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for element in tcPr.findall(qn('w:shd')):
            tcPr.remove(element)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:themeFill'), '0')
        tcPr.append(shd)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.name = "Arial"
        return
    except Exception as e:
        logger.error(f"Error in set_cell_background: {str(e)}")
        raise

def clear_cell_background(cell):
    """Clear cell background color and reset to default."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove any existing shading
        for element in tcPr.findall(qn('w:shd')):
            tcPr.remove(element)
        # Set to clear/transparent background
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'auto')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        tcPr.append(shd)
        # Reset text color to black
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        return
    except Exception as e:
        logger.error(f"Error in clear_cell_background: {str(e)}")
        raise

def clear_cell_margins(cell):
    """Remove cell margins."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'right', 'bottom', 'left']:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        for element in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(element)
        tcPr.append(tcMar)
    except Exception as e:
        logger.error(f"Error in clear_cell_margins: {str(e)}")
        raise

def clear_table_cell_padding(cell):
    """Clear padding from a table cell."""
    try:
        if not cell or not hasattr(cell, '_tc'):
            return
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        for old_mar in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(old_mar)
        tcPr.append(tcMar)
        logger.debug(f"Cleared padding for cell")
    except Exception as e:
        logger.error(f"Error in clear_table_cell_padding: {str(e)}")
        raise

def enforce_ratio_formatting(doc):
    """Enforce Arial Bold and consistent font size for ratio-related content."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re
    
    ratio_patterns = [
        'THC:', 'CBD:', 'CBC:', 'CBG:', 'CBN:',
        '100mg', '500mg', '50mg', '25mg', '10mg', '5mg',
        '1:1:1', '1:1', '2:1', '3:1', '4:1', '5:1',
        'CBC/CBG', 'THC/CBD', 'CBD/THC',
        'RATIO_START', 'THC_CBD_START',
        'mg THC', 'mg CBD', 'mg CBC', 'mg CBG', 'mg CBN',
        # Add patterns for longer ratio values
        'RATIO_END', 'THC_CBD_END',
        # Add patterns for any content containing mg values
        'mg', 'MG',
        # Add patterns for ratio values with spaces
        'THC mg', 'CBD mg', 'CBC mg', 'CBG mg', 'CBN mg',
        # Add patterns for any content that looks like cannabinoid ratios
        'THC/CBD', 'CBD/THC', 'THC/CBC', 'CBC/THC', 'THC/CBG', 'CBG/THC'
    ]
    
    def is_ratio_content(text):
        """Check if text contains ratio-like content."""
        # Check for specific patterns
        if any(pattern in text for pattern in ratio_patterns):
            return True
        
        # Check for RATIO markers
        if 'RATIO_START' in text or 'RATIO_END' in text:
            return True
            
        # Check for THC_CBD markers
        if 'THC_CBD_START' in text or 'THC_CBD_END' in text:
            return True
            
        # Check for cannabinoid patterns with numbers and mg
        cannabinoid_pattern = r'\b(THC|CBD|CBC|CBG|CBN)\s*\d+mg\b'
        if re.search(cannabinoid_pattern, text, re.IGNORECASE):
            return True
            
        # Check for ratio patterns like "X:Y"
        ratio_pattern = r'\b\d+:\d+\b'
        if re.search(ratio_pattern, text):
            return True
            
        # Check for partial ratio content (individual cannabinoid values)
        # This catches cases where "50mg CBC" might be in a separate run
        partial_cannabinoid_pattern = r'\b\d+mg\s+(THC|CBD|CBC|CBG|CBN)\b'
        if re.search(partial_cannabinoid_pattern, text, re.IGNORECASE):
            return True
            
        # Check for any text containing "mg" followed by a cannabinoid
        mg_cannabinoid_pattern = r'mg\s+(THC|CBD|CBC|CBG|CBN)'
        if re.search(mg_cannabinoid_pattern, text, re.IGNORECASE):
            return True
            
        # Check for any text containing a cannabinoid followed by "mg"
        cannabinoid_mg_pattern = r'(THC|CBD|CBC|CBG|CBN)\s+mg'
        if re.search(cannabinoid_mg_pattern, text, re.IGNORECASE):
            return True
            
        return False

    def process_paragraph(paragraph):
        # Use comprehensive ratio content detection
        if is_ratio_content(paragraph.text):
            # Store text content and existing font sizes
            text = paragraph.text
            existing_runs = []
            for run in paragraph.runs:
                existing_runs.append({
                    'text': run.text,
                    'size': run.font.size,
                    'bold': run.font.bold
                })
            
            # Clear paragraph
            paragraph.clear()
            
            # Add new run with proper formatting
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.bold = True
            
            # Set font properties at XML level
            rPr = run._element.get_or_add_rPr()
            
            # Force Arial font
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rFonts.set(qn('w:eastAsia'), 'Arial')
            rFonts.set(qn('w:cs'), 'Arial')
            rPr.append(rFonts)
            
            # Force bold
            b = OxmlElement('w:b')
            b.set(qn('w:val'), '1')
            rPr.append(b)
            
            # Preserve existing font size if available, otherwise use default
            if existing_runs and existing_runs[0]['size']:
                font_size = existing_runs[0]['size']
                run.font.size = font_size
                # Set at XML level too
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(font_size.pt * 2)))  # Word uses half-points
                rPr.append(sz)
            else:
                # Use default size (12pt)
                run.font.size = Pt(12)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(12 * 2)))  # Word uses half-points
                rPr.append(sz)

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
                    # Also process individual runs for longer ratio values
                    # First check if the paragraph contains any ratio content
                    paragraph_has_ratio = is_ratio_content(paragraph.text)
                    
                    for run in paragraph.runs:
                        # If paragraph has ratio content, make all runs bold
                        # OR if individual run has ratio content, make it bold
                        if paragraph_has_ratio or is_ratio_content(run.text):
                            run.font.name = "Arial"
                            run.font.bold = True
                            # Set at XML level for maximum compatibility
                            rPr = run._element.get_or_add_rPr()
                            rFonts = OxmlElement('w:rFonts')
                            rFonts.set(qn('w:ascii'), 'Arial')
                            rFonts.set(qn('w:hAnsi'), 'Arial')
                            rFonts.set(qn('w:eastAsia'), 'Arial')
                            rFonts.set(qn('w:cs'), 'Arial')
                            rPr.append(rFonts)
                            b = OxmlElement('w:b')
                            b.set(qn('w:val'), '1')
                            rPr.append(b)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
        # Also process individual runs for longer ratio values
        # First check if the paragraph contains any ratio content
        paragraph_has_ratio = is_ratio_content(paragraph.text)
        
        for run in paragraph.runs:
            # If paragraph has ratio content, make all runs bold
            # OR if individual run has ratio content, make it bold
            if paragraph_has_ratio or is_ratio_content(run.text):
                run.font.name = "Arial"
                run.font.bold = True
                # Set at XML level for maximum compatibility
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Arial')
                rFonts.set(qn('w:hAnsi'), 'Arial')
                rFonts.set(qn('w:eastAsia'), 'Arial')
                rFonts.set(qn('w:cs'), 'Arial')
                rPr.append(rFonts)
                b = OxmlElement('w:b')
                b.set(qn('w:val'), '1')
                rPr.append(b)

    return doc

def enforce_arial_bold_all_text(doc):
    """Enforce Arial Bold font for all text in the document while preserving font sizes."""
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def process_run(run):
        """Apply Arial Bold formatting to a single run while preserving font size."""
        # Store existing font size and bold state
        existing_size = run.font.size
        existing_bold = run.font.bold
        
        # Check if this run contains vendor markers (only check for actual markers)
        run_text = run.text.strip()
        is_vendor_marker = (
            'PRODUCTVENDOR_START' in run_text or 
            'PRODUCTVENDOR_END' in run_text or
            '{{Label' in run_text and 'ProductVendor}}' in run_text
        )
        
        # Set font properties at Python level
        run.font.name = "Arial"
        
        # Only make vendor markers non-bold, everything else should be bold
        if is_vendor_marker:
            run.font.bold = False
        else:
            # Make everything else bold
            run.font.bold = True
        
        # Restore font size if it existed
        if existing_size:
            run.font.size = existing_size

        # Force Arial at XML level for maximum compatibility
        rPr = run._element.get_or_add_rPr()
        
        # Set font family
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rFonts.set(qn('w:eastAsia'), 'Arial')
        rFonts.set(qn('w:cs'), 'Arial')
        rPr.append(rFonts)
        
        # Force bold only if it should be bold
        if run.font.bold:
            b = OxmlElement('w:b')
            b.set(qn('w:val'), '1')
            rPr.append(b)
        
        # Set font size at XML level if it exists
        if existing_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(existing_size.pt * 2)))  # Word uses half-points
            rPr.append(sz)

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        process_run(run)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            process_run(run)

    return doc

def cleanup_all_price_markers(doc):
    """Remove all price markers while preserving formatting."""
    price_patterns = [
        'PRICE_START',
        'PRICE_END',
        '{{PRICE}}',
        '{{/PRICE}}'
    ]
    
    def process_paragraph(paragraph):
        # Store original text and check if it contains price markers
        text = paragraph.text
        if not any(pattern in text for pattern in price_patterns):
            return
        
        # Remove all price markers
        for pattern in price_patterns:
            text = text.replace(pattern, '')
        
        # Clear paragraph and add cleaned text
        paragraph.clear()
        if text.strip():
            run = paragraph.add_run(text.strip())
            # Apply price formatting
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.size = Pt(14)  # Standard price font size

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    return doc

def remove_extra_spacing(doc):
    """Remove extra spacing between paragraphs and set consistent line spacing while preserving font sizes."""
    from docx.shared import Pt
    
    def process_paragraph(paragraph):
        # Store text content and existing font sizes
        if not paragraph.text.strip():
            return
            
        text = paragraph.text
        
        # Store existing font sizes for each run
        existing_sizes = []
        for run in paragraph.runs:
            if run.text.strip():
                existing_sizes.append(run.font.size)
        
        # Clear and reset paragraph
        paragraph.clear()
        
        # Split text into runs and preserve font sizes
        words = text.split()
        size_index = 0
        
        for i, word in enumerate(words):
            run = paragraph.add_run(word)
            
            # Set consistent font
            run.font.name = "Arial"
            run.font.bold = True
            
            # Restore font size if available
            if size_index < len(existing_sizes) and existing_sizes[size_index]:
                run.font.size = existing_sizes[size_index]
                size_index += 1
            
            # Add space between words (except for last word)
            if i < len(words) - 1:
                run.add_text(' ')
        
        # Set paragraph spacing to minimum to prevent cell expansion
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    return doc

def apply_type_formatting(doc, product_type, template_type='vertical'):
    """Apply specific formatting based on product type."""
    def process_paragraph(paragraph):
        # Skip empty paragraphs
        if not paragraph.text.strip():
            return
            
        # Apply type-specific formatting
        if product_type.lower() == 'paraphernalia':
            # Make paraphernalia text smaller and less prominent
            for run in paragraph.runs:
                if run.font.size:
                    current_size = run.font.size.pt
                    run.font.size = Pt(max(6, current_size * 0.8))  # Reduce size by 20%
                run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        elif product_type.lower() in ['concentrate', 'wax', 'shatter', 'rosin']:
            # Make concentrate text bold and prominent
            for run in paragraph.runs:
                run.font.bold = True
                if run.font.size:
                    current_size = run.font.size.pt
                    run.font.size = Pt(min(24, current_size * 1.1))  # Increase size by 10%

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Process all paragraphs outside tables
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    return doc

def create_3x3_grid(doc, template_type='vertical'):
    """Create a 3x3 grid table for label layout."""
    try:
        # Remove existing tables
        for table in doc.tables:
            table._element.getparent().remove(table._element)
        
        # Remove default paragraph if it exists
        if doc.paragraphs:
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
        
        # Create new table
        table = doc.add_table(rows=3, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set table properties
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        table._element.insert(0, tblPr)
        
        # Set column widths
        col_width = Inches(3.3 / 3)  # 3.5 inches divided by 3 columns
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(3):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(col_width.inches * 1440)))  # Convert to twips
            tblGrid.append(gridCol)
        table._element.insert(0, tblGrid)
        
        # Set row heights
        row_height = Inches(2.25)  
        for row in table.rows:
            row.height = row_height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        # Enforce fixed cell dimensions to prevent any growth
        enforce_fixed_cell_dimensions(table)
        
        logger.debug("Created 3x3 grid table")
        return table
    except Exception as e:
        logger.error(f"Error creating 3x3 grid: {str(e)}")
        raise

def disable_autofit(table):
    """Disable autofit for a table."""
    try:
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        table._element.insert(0, tblPr)
        logger.debug("Disabled autofit for table")
    except Exception as e:
        logger.error(f"Error disabling autofit: {str(e)}")
        raise

def enforce_fixed_cell_dimensions(table):
    """Enforce fixed cell dimensions to prevent any cell growth with text."""
    try:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        
        # Disable autofit
        disable_autofit(table)
        
        # Set table properties to prevent any auto-sizing
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
        
        # Ensure fixed layout
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Set table to not auto-fit
        table.autofit = False
        if hasattr(table, 'allow_autofit'):
            table.allow_autofit = False
        
        # Process each cell to enforce fixed dimensions
        for row in table.rows:
            # Set exact row height rule
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            for cell in row.cells:
                # Set cell vertical alignment to top to prevent content from expanding cell
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                
                # Clear any cell margins that might allow expansion
                clear_cell_margins(cell)
                
                # Process paragraphs in the cell to prevent text overflow
                for paragraph in cell.paragraphs:
                    # Set paragraph spacing to minimum
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    
                    # Ensure text doesn't wrap beyond cell boundaries
                    for run in paragraph.runs:
                        # Set font properties to prevent text expansion
                        if not run.font.size:
                            run.font.size = Pt(12)  # Set default size if none
        
        logger.debug("Enforced fixed cell dimensions for table")
        return table
    except Exception as e:
        logger.error(f"Error enforcing fixed cell dimensions: {str(e)}")
        raise

def fix_table(doc, num_rows=3, num_cols=3, template_type='horizontal'):
    """Fix table with proper cell dimensions based on template type."""
    from src.core.constants import CELL_DIMENSIONS
    
    # Get individual cell dimensions
    cell_dims = CELL_DIMENSIONS.get(template_type, {'width': 2.4, 'height': 2.4})
    cell_width = Inches(cell_dims['width'])
    cell_height = Inches(cell_dims['height'])
    
    # Remove all existing tables
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    
    # Remove default paragraph if it exists
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    # Create new table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table properties
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    table._element.insert(0, tblPr)
    
    # Set column widths
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(cell_width.inches * 1440)))
        tblGrid.append(gridCol)
    table._element.insert(0, tblGrid)
    
    # Set row heights
    for row in table.rows:
        row.height = cell_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Enforce fixed cell dimensions to prevent any growth
    enforce_fixed_cell_dimensions(table)
    
    return table

def rebuild_3x3_grid(doc, template_type='horizontal'):
    """Rebuild 3x3 grid with proper cell dimensions based on template type."""
    from src.core.constants import CELL_DIMENSIONS
    
    # Get individual cell dimensions
    cell_dims = CELL_DIMENSIONS.get(template_type, {'width': 2.4, 'height': 2.4})
    cell_width = Inches(cell_dims['width'])
    cell_height = Inches(cell_dims['height'])
    
    # Remove all existing tables
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    
    # Remove default paragraph if it exists
    if doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    
    # Create new 3x3 table
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table properties
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    table._element.insert(0, tblPr)
    
    # Set column widths
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(3):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(cell_width.inches * 1440)))
        tblGrid.append(gridCol)
    table._element.insert(0, tblGrid)
    
    # Set row heights
    for row in table.rows:
        row.height = cell_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Enforce fixed cell dimensions to prevent any growth
    enforce_fixed_cell_dimensions(table)
    
    return table 

def apply_custom_formatting(doc, template_settings):
    """Apply custom formatting based on template settings."""
    try:
        font_family = template_settings.get('font', 'Arial')
        bold_headers = template_settings.get('boldHeaders', False)
        italic_descriptions = template_settings.get('italicDescriptions', False)
        line_spacing = float(template_settings.get('lineSpacing', '1.0'))
        paragraph_spacing = int(template_settings.get('paragraphSpacing', '0'))
        text_color = template_settings.get('textColor', '#000000')
        background_color = template_settings.get('backgroundColor', '#ffffff')
        header_color = template_settings.get('headerColor', '#333333')
        accent_color = template_settings.get('accentColor', '#007bff')
        auto_resize = template_settings.get('autoResize', True)
        smart_truncation = template_settings.get('smartTruncation', True)
        
        # Apply formatting to all paragraphs in the document
        for paragraph in doc.paragraphs:
            # Set line spacing
            paragraph.paragraph_format.line_spacing = line_spacing
            
            # Set paragraph spacing
            if paragraph_spacing > 0:
                paragraph.paragraph_format.space_after = Pt(paragraph_spacing)
            
            # Apply formatting to runs
            for run in paragraph.runs:
                # Set font family
                run.font.name = font_family
                
                # Set font color
                if text_color != '#000000':
                    run.font.color.rgb = RGBColor.from_string(text_color[1:])  # Remove # from hex
                
                # Apply bold to headers (if enabled)
                if bold_headers and any(keyword in run.text.lower() for keyword in ['brand', 'price', 'lineage', 'thc', 'cbd']):
                    run.font.bold = True
                    if header_color != '#333333':
                        run.font.color.rgb = RGBColor.from_string(header_color[1:])
                
                # Apply italic to descriptions (if enabled)
                if italic_descriptions and len(run.text) > 20:  # Assume long text is description
                    run.font.italic = True
        
        # Apply background color to tables if specified
        if background_color != '#ffffff':
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        set_cell_background(cell, background_color)
        
        # Apply accent color to specific elements
        if accent_color != '#007bff':
            # Apply to price elements
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if any(keyword in run.text.lower() for keyword in ['$', 'price', 'cost']):
                        run.font.color.rgb = RGBColor.from_string(accent_color[1:])
        
        # Apply auto-resize if enabled
        if auto_resize:
            # Adjust font sizes to fit content
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.size and run.font.size.pt > 6:
                        # Reduce font size if text is too long
                        if len(run.text) > 50:
                            run.font.size = Pt(max(6, run.font.size.pt - 2))
        
        # Apply smart truncation if enabled
        if smart_truncation:
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if len(run.text) > 100:
                        # Truncate very long text
                        run.text = run.text[:97] + "..."
        
        logger.debug(f"Applied custom formatting with font: {font_family}, line spacing: {line_spacing}")
        
    except Exception as e:
        logger.error(f"Error applying custom formatting: {str(e)}")
        # Fall back to default formatting
        enforce_arial_bold_all_text(doc) 