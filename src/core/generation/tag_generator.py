import io
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from io import BytesIO
import logging
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
import re
from docxtpl import DocxTemplate, InlineImage
from copy import deepcopy
import cProfile
from itertools import groupby

from src.core.generation.docx_formatting import (
    fix_table_row_heights,
    safe_fix_paragraph_spacing,
    apply_conditional_formatting,
    set_cell_background,
    clear_cell_margins,
    clear_table_cell_padding,
    enforce_fixed_cell_dimensions,
)
from src.core.generation.context_builders import (
    build_context,
)
from src.core.formatting.markers import (
    wrap_with_marker,
    FIELD_MARKERS
)
from src.core.utils.resource_utils import resource_path
from src.core.constants import (
    FONT_SCHEME_HORIZONTAL,
    FONT_SCHEME_VERTICAL,
    FONT_SCHEME_MINI,
    LINEAGE_COLOR_MAP
)

# Performance optimization: disable debug logging in production
DEBUG_ENABLED = False

logger = logging.getLogger(__name__)

PLACEHOLDER_MARKERS = {
    "Description": ("DESC_START", "DESC_END"),
    "WeightUnits": ("WEIGHTUNITS_START", "WEIGHTUNITS_END"),
    "ProductBrand": ("PRODUCTBRAND_START", "PRODUCTBRAND_END"),
    "ProductBrand_Center": ("PRODUCTBRAND_CENTER_START", "PRODUCTBRAND_CENTER_END"),
    "Price": ("PRICE_START", "PRICE_END"),
    "Lineage": ("LINEAGE_START", "LINEAGE_END"),
    "DOH": ("{{Label1.DOH}}", ""),
    "Ratio_or_THC_CBD": ("RATIO_START", "RATIO_END"),
    "THC_CBD": ("THC_CBD_START", "THC_CBD_END"),
    "ProductName": ("PRODUCTNAME_START", "PRODUCTNAME_END"),
    "ProductStrain": ("PRODUCTSTRAIN_START", "PRODUCTSTRAIN_END"),
    "ProductType": ("PRODUCTTYPE_START", "PRODUCTTYPE_END")
}

# Import colors from docx_formatting to avoid duplication
from .docx_formatting import COLORS as LINEAGE_COLORS

def get_template_path(template_type):
    """Return the absolute path for a given template type."""
    # Map template types to filenames
    template_files = {
        'horizontal': 'horizontal.docx',
        'vertical': 'vertical.docx',
        'mini': 'mini.docx',
        'inventory': 'inventory.docx'
    }

    # Validate template type
    if template_type not in template_files:
        raise ValueError(f"Invalid template type: {template_type}")

    # Get template directory path
    base_dir = Path(__file__).parent
    template_dir = base_dir / "templates"

    # Ensure template directory exists
    if not template_dir.exists():
        raise ValueError(f"Template directory not found: {template_dir}")

    # Build full path
    template_path = template_dir / template_files[template_type]

    # Verify template exists
    if not template_path.exists():
        raise ValueError(f"Template file not found: {template_path}")

    return str(template_path)

def chunk_records(records, chunk_size=9):
    """Split the list of records into chunks of a given size."""
    return [records[i:i+chunk_size] for i in range(0, len(records), chunk_size)]

def flatten_tags(records):
    """Extract Description values from records for tag generation."""
    flat_tags = []
    for record in records:
        description = record.get("Description", "")
        if description and isinstance(description, str):
            flat_tags.append(description.strip())
    return flat_tags

def expand_template_to_4x5_fixed_scaled(template_path, scale_factor=1.0):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    from copy import deepcopy

    num_cols, num_rows = 4, 5
    col_width_twips = str(int(1.75 * 1440))  # 1.75 inches per column for equal width
    row_height_pts  = Pt(2.0 * 72)  # 2.0 inches per row for equal height
    cut_line_twips  = int(0.001 * 1440)

    doc = Document(template_path)
    if not doc.tables:
        raise RuntimeError("Template must contain at least one table.")
    old = doc.tables[0]
    src_tc = deepcopy(old.cell(0,0)._tc)
    old._element.getparent().remove(old._element)

    while doc.paragraphs and not doc.paragraphs[0].text.strip():
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.alignment = 1
    tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D3D3D3')
    tblPr.insert(0, shd)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    tbl._element.insert(0, tblPr)
    grid = OxmlElement('w:tblGrid')
    for _ in range(num_cols):
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), col_width_twips)
        grid.append(gc)
    tbl._element.insert(0, grid)
    for row in tbl.rows:
        row.height = row_height_pts
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    borders = OxmlElement('w:tblBorders')
    for side in ('insideH','insideV'):
        b = OxmlElement(f"w:{side}")
        b.set(qn('w:val'), "single"); b.set(qn('w:sz'), "4")
        b.set(qn('w:color'), "D3D3D3"); b.set(qn('w:space'), "0")
        borders.append(b)
    tblPr.append(borders)
    cnt = 1
    for r in range(num_rows):
        for c in range(num_cols):
            cell = tbl.cell(r,c)
            cell._tc.clear_content()
            tc = deepcopy(src_tc)
            for t in tc.iter(qn('w:t')):
                if t.text and 'Label1' in t.text:
                    t.text = t.text.replace('Label1', f'Label{cnt}')
            for el in tc.xpath('./*'):
                cell._tc.append(deepcopy(el))
            cnt += 1
    from docx.oxml.shared import OxmlElement as OE
    tblPr2 = tbl._element.find(qn('w:tblPr'))
    spacing = OxmlElement('w:tblCellSpacing'); spacing.set(qn('w:w'), str(cut_line_twips)); spacing.set(qn('w:type'), 'dxa')
    tblPr2.append(spacing)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def process_chunk(args):
    """Process a chunk of records to generate labels."""
    chunk, base_template, font_scheme, orientation, scale_factor = args
    # Always expand to 4x5 grid for mini
    if orientation == "mini":
        local_template_buffer = expand_template_to_4x5_fixed_scaled(base_template, scale_factor=scale_factor)
        num_labels = 20  # Fixed: 4x5 grid = 20 labels per page
    else:
        local_template_buffer = base_template
        num_labels = 9
    tpl = DocxTemplate(local_template_buffer)
    context = {}
    image_width = Mm(8) if orientation == "mini" else Mm(9 if orientation == 'vertical' else 11)
    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
    if DEBUG_ENABLED:
        logger.debug(f"DOH image path: {doh_image_path}")
    
    for i in range(num_labels):
        label_data = {}
        if i < len(chunk):
            row = chunk[i]
            doh_value = str(row.get("DOH", "")).strip().upper()
            if DEBUG_ENABLED:
                logger.debug(f"Processing DOH value: {doh_value}")
            product_type = str(row.get("Product Type*", "")).strip().lower()
            if DEBUG_ENABLED:
                logger.debug(f"Product type: {product_type}")
            
            if doh_value == "YES":
                # Use HighCBD.png if product_type starts with 'high cbd'
                if product_type.startswith('high cbd'):
                    high_cbd_image_path = resource_path(os.path.join("templates", "HighCBD.png"))
                    if DEBUG_ENABLED:
                        logger.debug(f"Using HighCBD image: {high_cbd_image_path}")
                    label_data["DOH"] = InlineImage(tpl, high_cbd_image_path, width=image_width)
                else:
                    doh_image_path = resource_path(os.path.join("templates", "DOH.png"))
                    if DEBUG_ENABLED:
                        logger.debug(f"Using DOH image: {doh_image_path}")
                    label_data["DOH"] = InlineImage(tpl, doh_image_path, width=image_width)
                if DEBUG_ENABLED:
                    logger.debug(f"Created DOH image with width: {image_width}")
            else:
                label_data["DOH"] = ""
                if DEBUG_ENABLED:
                    logger.debug("Skipping DOH image - value is not 'YES'")
                
            # --- Wrap all fields with markers ---
            price_val = f"{row.get('Price', '')}"
            label_data["Price"] = wrap_with_marker(price_val, "PRIC")  # Changed from "PRICE" to "PRIC" to match MAIN.py
            
            lineage_text   = str(row.get("Lineage", "")).strip()
            product_brand  = str(row.get("Product Brand", "")).strip()
            label_data["ProductBrand"] = wrap_with_marker(product_brand.upper(), "PRODUCTBRAND_CENTER")
            
            # Add other fields to label_data
            label_data["Description"] = wrap_with_marker(str(row.get("Description", "")), "DESC")
            label_data["WeightUnits"] = wrap_with_marker(str(row.get("WeightUnits", "")), "WEIGHTUNITS")
            label_data["Lineage"] = wrap_with_marker(lineage_text.upper(), "LINEAGE")
            label_data["Ratio_or_THC_CBD"] = wrap_with_marker(str(row.get("Ratio", "")), "RATIO")
            label_data["ProductStrain"] = wrap_with_marker(str(row.get("Product Strain", "")), "PRODUCTSTRAIN")
            label_data["JointRatio"] = wrap_with_marker(str(row.get("JointRatio", "")), "JOINT_RATIO")
            
            # Combine Description and WeightUnits (use JointRatio for pre-roll products)
            desc = str(row.get("Description", ""))
            if desc:
                desc = re.sub(r'[-\s]+$', '', desc)
            product_type = str(row.get("Product Type*", "")).strip().lower()
            
            if product_type in {"pre-roll", "infused pre-roll"}:
                weight = str(row.get("JointRatio", ""))
                # Remove any leading hyphens, en dashes, or regular spaces (but preserve non-breaking spaces)
                weight = re.sub(r'^[\s\-–]+', '', weight)
                # Preserve non-breaking spaces by replacing them temporarily, then restoring
                weight = weight.replace('\u00A0', '___NBSP___')
                weight = re.sub(r'^[\s\-–]+', '', weight)
                weight = weight.replace('___NBSP___', '\u00A0')
                # Normalize incomplete 'Pack' endings
                weight = re.sub(r'(Pa|P)$', 'Pack', weight, flags=re.IGNORECASE)
                # Accept both .5g and 0.5g, and add leading zero if missing
                weight = re.sub(r'(?<!\d)\.([0-9]+)g', r'0.\1g', weight)
                # Match all forms of '1gx2Pack', '1g x2Pack', '1g x 2 Pack', etc.
                match = re.match(r"([0-9.]+)g\s*[xX]\s*([0-9]+)\s*Pack", weight, re.IGNORECASE)
                if not match:
                    match = re.match(r"([0-9.]+)g\s*[xX]\s*([0-9]+)", weight, re.IGNORECASE)
                if not match:
                    match = re.match(r"([0-9.]+)g[xX]([0-9]+)Pack", weight, re.IGNORECASE)
                if not match:
                    match = re.match(r"([0-9.]+)g[xX]([0-9]+)", weight, re.IGNORECASE)
                if match:
                    g, n = match.groups()
                    weight = f"{g}g x {n} Pack"
            else:
                weight = str(row.get("WeightUnits", ""))
                weight = re.sub(r'^[\s\-–]+', '', weight)
                
            if desc and weight:
                lines = desc.splitlines()
                if lines:
                    # Always put JointRatio on a new line
                    combined = "\n".join(lines) + f"\n- {weight}"
                else:
                    combined = f"- {weight}"
            else:
                combined = desc or weight
            print(f"DEBUG DescAndWeight (before wrap): '{combined}'")
            label_data["DescAndWeight"] = wrap_with_marker(combined, "DESC")
            print(f"DEBUG DescAndWeight: {repr(label_data['DescAndWeight'])}")
            
            context[f"Label{i+1}"] = label_data
            if DEBUG_ENABLED:
                logger.debug(f"Created label data for Label{i+1}")
        else:
            # Empty label data for unused slots
            context[f"Label{i+1}"] = {
                "Description": "",
                "WeightUnits": "",
                "ProductBrand": "",
                "Price": "",
                "Lineage": "",
                "DOH": "",
                "Ratio_or_THC_CBD": "",
                "ProductStrain": "",
                "DescAndWeight": "",
                "JointRatio": ""
            }
            if DEBUG_ENABLED:
                logger.debug(f"Created empty label data for Label{i+1}")

    # Render template
    if DEBUG_ENABLED:
        logger.debug("Rendering template...")
    tpl.render(context)
    if DEBUG_ENABLED:
        logger.debug("Template rendered successfully")
    
    # Save to buffer
    buffer = BytesIO()
    tpl.save(buffer)
    buffer.seek(0)
    if DEBUG_ENABLED:
        logger.debug("Template saved to buffer")
    
    return buffer.getvalue()

def combine_documents(docs):
    """Combine multiple documents into one using a safer method."""
    if not docs:
        return None
        
    try:
        # Use the first document as the base
        master_doc = Document(BytesIO(docs[0]))
        
        # Add spacing between documents
        master_doc.add_paragraph()
        
        # Append each subsequent document
        for doc_bytes in docs[1:]:
            try:
                src_doc = Document(BytesIO(doc_bytes))
                
                # Copy all content from the source document
                for element in src_doc.element.body:
                    # Create a deep copy to avoid reference issues
                    new_element = deepcopy(element)
                    master_doc.element.body.append(new_element)
                    
                # Add spacing between documents
                master_doc.add_paragraph()
                
            except Exception as e:
                logger.error(f"Error combining document: {e}")
                # Continue with other documents instead of failing completely
                continue

        # Save final document to bytes
        final_buffer = BytesIO()
        master_doc.save(final_buffer)
        final_buffer.seek(0)
        
        # Validate the combined document
        try:
            test_doc = Document(final_buffer)
            # Center all tables in the document
            for table in test_doc.tables:
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            final_buffer.seek(0)
            return final_buffer.getvalue()
        except Exception as validation_error:
            logger.error(f"Combined document validation failed: {validation_error}")
            raise ValueError(f"Combined document is corrupted: {validation_error}")
            
    except Exception as e:
        logger.error(f"Error in combine_documents: {e}")
        raise

def run_full_process_by_group(records, template_path, font_scheme):
    """Process all records using the template, grouped by strain color (lineage)."""
    if not records:
        return None
    # Define canonical lineage order
    lineage_order = list(LINEAGE_COLOR_MAP.keys())
    def get_lineage(rec):
        import logging
        # Check multiple possible field names for lineage
        possible_fields = ['Lineage', 'lineage', 'Product Lineage', 'ProductLineage', 'Strain Type', 'StrainType']
        lin = ''
        for field in possible_fields:
            if field in rec and rec[field]:
                lin = str(rec[field]).strip()
                break
        
        # Normalize the lineage value
        lin = lin.upper().replace('PARA', 'PARAPHERNALIA')
        
        # Check if this is a Classic Type
        product_type = str(rec.get('Product Type*', '') or rec.get('ProductType', '')).strip().lower()
        classic_types = {"flower", "pre-roll", "concentrate", "infused pre-roll", "solventless concentrate", "vape cartridge"}
        
        lineage_result = None
        if product_type in classic_types:
            # For classic types, if lineage is not recognized, check database
            if lin and lin in lineage_order:
                lineage_result = lin
            else:
                # Check database for lineage information
                try:
                    from src.core.data.product_database import ProductDatabase
                    db = ProductDatabase()
                    product_name = rec.get('ProductName', '') or rec.get('Product Name*', '')
                    product_strain = rec.get('ProductStrain', '') or rec.get('Product Strain', '')
                    
                    # Try to find lineage in database
                    db_lineage = None
                    if product_name:
                        db_lineage = db.get_lineage_by_product_name(product_name)
                    if not db_lineage and product_strain:
                        db_lineage = db.get_lineage_by_strain(product_strain)
                    
                    if db_lineage:
                        db_lineage = db_lineage.upper().replace('PARA', 'PARAPHERNALIA')
                        if db_lineage in lineage_order:
                            lineage_result = db_lineage
                        else:
                            lineage_result = 'HYBRID'
                    else:
                        lineage_result = 'HYBRID'
                except Exception as e:
                    lineage_result = 'HYBRID'
            if lineage_result == 'MIXED':
                logging.warning(f"[BUG] Classic Type '{product_type}' with original lineage '{lin}' assigned 'MIXED'. Forcing 'HYBRID'. Record: {rec}")
                lineage_result = 'HYBRID'
        else:
            # For non-classic types, keep existing logic
            if lin and lin in lineage_order:
                lineage_result = lin
            else:
                lineage_result = 'MIXED'
        if lineage_result in ('MIXED', 'HYBRID'):
            logging.info(f"Product Type: '{product_type}', Original Lineage: '{lin}', Final Lineage: '{lineage_result}', Record: {rec}")
        return lineage_result
    # Sort records by lineage order, then by ProductName
    records_sorted = sorted(records, key=lambda r: (lineage_order.index(get_lineage(r)), str(r.get('ProductName', ''))))
    # Group by lineage and chunk within each group
    tag_chunks = []
    for lineage, group in groupby(records_sorted, key=get_lineage):
        group_list = list(group)
        for i in range(0, len(group_list), 9):
            tag_chunks.append(group_list[i:i+9])
    def _inner():
        template_type = Path(template_path).stem
        docs = []
        for tag_chunk in tag_chunks:
            doc = process_chunk(
                tag_chunk,
                template_path,
                font_scheme
            )
            if doc:
                docs.append(doc)
        return combine_documents(docs)
    profile = cProfile.Profile()
    result = profile.runcall(_inner)
    profile.dump_stats('profile_group.prof')
    return result

def run_full_process_by_mini(records, template_type, font_scheme, scale_factor=1.0):
    """Process records with the mini template."""
    if not records:
        return None
    def _inner():
        template_path = get_template_path(template_type)
        if hasattr(template_path, "seek"):
            template_path.seek(0)
            try:
                Document(template_path)
                template_path.seek(0)
            except Exception as e:
                raise ValueError(f"Template buffer is not a valid DOCX: {e}")
        chunks = chunk_records(records, chunk_size=20)  # Fixed: 4x5 grid = 20 labels per page
        docs = []
        for chunk in chunks:
            args = (chunk, template_path, font_scheme, template_type, scale_factor)
            docs.append(process_chunk(args))
        return combine_documents(docs)
    profile = cProfile.Profile()
    result = profile.runcall(_inner)
    profile.dump_stats('profile_mini.prof')
    return result

def generate_multiple_label_tables(records, template_path):
    """
    For each record, render the template and append the resulting table to a new document.
    Returns a BytesIO buffer with the final DOCX.
    """
    try:
        # Sort records by canonical lineage order, then by name
        lineage_order = list(LINEAGE_COLOR_MAP.keys())
        def get_lineage(rec):
            possible_fields = ['Lineage', 'lineage', 'Product Lineage', 'ProductLineage', 'Strain Type', 'StrainType']
            lin = ''
            for field in possible_fields:
                if field in rec and rec[field]:
                    lin = str(rec[field]).strip()
                    break
            lin = lin.upper().replace('PARA', 'PARAPHERNALIA')
            # Check if this is a Classic Type
            product_type = str(rec.get('Product Type*', '') or rec.get('ProductType', '')).strip().lower()
            classic_types = {"flower", "pre-roll", "concentrate", "infused pre-roll", "solventless concentrate", "vape cartridge"}
            if product_type in classic_types:
                # For classic types, if lineage is not recognized, check database
                if lin and lin in lineage_order:
                    return lin
                else:
                    # Check database for lineage information
                    try:
                        from src.core.data.product_database import ProductDatabase
                        db = ProductDatabase()
                        product_name = rec.get('ProductName', '') or rec.get('Product Name*', '')
                        product_strain = rec.get('ProductStrain', '') or rec.get('Product Strain', '')
                        
                        # Try to find lineage in database
                        db_lineage = None
                        if product_name:
                            db_lineage = db.get_lineage_by_product_name(product_name)
                        if not db_lineage and product_strain:
                            db_lineage = db.get_lineage_by_strain(product_strain)
                        
                        if db_lineage:
                            db_lineage = db_lineage.upper().replace('PARA', 'PARAPHERNALIA')
                            if db_lineage in lineage_order:
                                return db_lineage
                        
                        # If still no match, default to HYBRID for classic types
                        return 'HYBRID'
                    except Exception as e:
                        # If database lookup fails, default to HYBRID for classic types
                        return 'HYBRID'
            # For non-classic types, keep existing logic
            return lin if lin in lineage_order else 'MIXED'
        records_sorted = sorted(records, key=lambda r: (lineage_order.index(get_lineage(r)), str(r.get('ProductName', ''))))
        final_doc = Document()
        # Remove default paragraph if it exists
        if final_doc.paragraphs:
            p = final_doc.paragraphs[0]
            p._element.getparent().remove(p._element)
        # Group by lineage and chunk within each group
        for lineage, group in groupby(records_sorted, key=get_lineage):
            group_list = list(group)
            for i in range(0, len(group_list), 9):
                chunk = group_list[i:i+9]
                # Create a single 3x3 table for this chunk
                table = final_doc.add_table(rows=3, cols=3)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.allow_autofit = False
                tblPr = table._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'fixed')
                tblPr.append(tblLayout)
                table._element.insert(0, tblPr)
                col_width = Inches(1.0)
                tblGrid = OxmlElement('w:tblGrid')
                for _ in range(3):
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(qn('w:w'), str(int(col_width.inches * 1440)))
                    tblGrid.append(gridCol)
                table._element.insert(0, tblGrid)
                # Fill the table cells with tag data
                for idx in range(9):
                    r, c = divmod(idx, 3)
                    cell = table.cell(r, c)
                    if idx < len(chunk):
                        record = chunk[idx]
                        try:
                            doc = DocxTemplate(template_path)
                            context = build_context(record, doc)
                            doc.render(context)
                            tmp_stream = BytesIO()
                            doc.save(tmp_stream)
                            tmp_stream.seek(0)
                            rendered_doc = Document(tmp_stream)
                            if rendered_doc.tables:
                                src_table = rendered_doc.tables[0]
                                src_cell = src_table.cell(0, 0)
                                for para in cell.paragraphs:
                                    para._element.getparent().remove(para._element)
                                for para in src_cell.paragraphs:
                                    new_para = cell.add_paragraph()
                                    new_para.alignment = para.alignment
                                    for run in para.runs:
                                        try:
                                            new_run = new_para.add_run(run.text)
                                            new_run.bold = run.bold
                                            new_run.italic = run.italic
                                            new_run.underline = run.underline
                                            if run.font.size:
                                                new_run.font.size = run.font.size
                                            if run.font.name:
                                                new_run.font.name = run.font.name
                                            if run.font.color and run.font.color.rgb:
                                                new_run.font.color.rgb = run.font.color.rgb
                                        except Exception as run_error:
                                            logger.warning(f"Error copying run: {run_error}")
                                            new_para.add_run(run.text)
                                if not cell.paragraphs:
                                    cell.add_paragraph()
                        except Exception as record_error:
                            logger.error(f"Error processing record: {record_error}")
                            cell.text = ''
                    else:
                        cell.text = ''
                    cell.width = col_width
                for row in table.rows:
                    row.height = Inches(1.0)
                    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                
                # Enforce fixed cell dimensions to prevent any growth
                enforce_fixed_cell_dimensions(table, "mini")
                table.autofit = False
                if hasattr(table, 'allow_autofit'):
                    table.allow_autofit = False
                
                final_doc.add_paragraph()
        # Center all tables in the final document
        for table in final_doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Save final document
        output = BytesIO()
        final_doc.save(output)
        output.seek(0)
        
        # Validate the final document
        try:
            test_doc = Document(output)
            output.seek(0)
            return output
        except Exception as validation_error:
            logger.error(f"Final document validation failed: {validation_error}")
            raise ValueError(f"Generated document is corrupted: {validation_error}")
            
    except Exception as e:
        logger.error(f"Error in generate_multiple_label_tables: {e}")
        raise

def set_table_borders(table):
    """Apply consistent border formatting matching main application style."""
    tblPr = table._element.find(qn('w:tblPr'))
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
        
    tblBorders = OxmlElement('w:tblBorders')
    
    # Remove outer borders
    for side in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "nil")
        tblBorders.append(bd)
        
    # Add light gray interior lines
    for side in ("insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn('w:val'), "single")
        bd.set(qn('w:sz'), "4")
        bd.set(qn('w:color'), "D3D3D3")
        bd.set(qn('w:space'), "0")
        tblBorders.append(bd)
        
    tblPr.append(tblBorders)

def debug_markers(text):
    """Debug helper to identify marker issues."""
    markers = ['DESC', 'WEIGHTUNITS', 'PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 'THC_CBD']
    found_markers = []
    
    for marker in markers:
        start = f"{marker}_START"
        end = f"{marker}_END"
        if start in text:
            pos = text.find(start)
            found_markers.append(f"{marker} start at {pos}")
        if end in text:
            pos = text.find(end)
            found_markers.append(f"{marker} end at {pos}")
            
    if found_markers:
        if DEBUG_ENABLED:
            logger.debug(f"Found markers in text: {text}")
        for marker in found_markers:
            if DEBUG_ENABLED:
                logger.debug(f"  {marker}")
    return found_markers

def validate_and_repair_document(doc_bytes):
    """Validate a document and attempt to repair common issues."""
    try:
        # Try to load the document
        doc = Document(BytesIO(doc_bytes))
        
        # Check for common corruption issues
        issues_found = []
        
        # Check if document has content
        if not doc.paragraphs and not doc.tables:
            issues_found.append("Document has no content")
            
        # Check for malformed tables
        for table in doc.tables:
            try:
                # Try to access table properties
                _ = table.rows
                _ = table.columns
            except Exception as e:
                issues_found.append(f"Malformed table: {e}")
                
        # Check for malformed paragraphs
        for para in doc.paragraphs:
            try:
                # Try to access paragraph properties
                _ = para.runs
            except Exception as e:
                issues_found.append(f"Malformed paragraph: {e}")
        
        if issues_found:
            logger.warning(f"Document validation issues found: {issues_found}")
            return False, issues_found
            
        return True, []
        
    except Exception as e:
        logger.error(f"Document validation failed: {e}")
        return False, [f"Document cannot be loaded: {e}"]

def create_safe_document():
    """Create a minimal safe document as a fallback."""
    try:
        doc = Document()
        # Add a simple paragraph to ensure the document is valid
        doc.add_paragraph("Document generated successfully.")
        return doc
    except Exception as e:
        logger.error(f"Error creating safe document: {e}")
        raise



