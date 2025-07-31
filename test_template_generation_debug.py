#!/usr/bin/env python3
"""
Comprehensive test script to debug template generation issues.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.data.excel_processor import ExcelProcessor

def test_template_generation():
    """Test template generation and identify missing information."""
    print("🔍 Testing Template Generation Issues...")
    
    # Initialize processor
    processor = ExcelProcessor()
    
    # Try to load a default file
    from src.core.data.excel_processor import get_default_upload_file
    default_file = get_default_upload_file()
    
    if not default_file or not os.path.exists(default_file):
        print("❌ No default file found for testing")
        return False
    
    print(f"📁 Loading file: {default_file}")
    
    try:
        # Load the file
        processor.load_file(default_file)
        print("✅ File loaded successfully")
        
        # Get available tags
        tags = processor.get_available_tags()
        print(f"📊 Found {len(tags)} tags")
        
        if not tags:
            print("❌ No tags found")
            return False
        
        # Test first few tags for template generation
        test_tags = tags[:5]
        
        print("\n🔍 Testing template generation for first 5 tags:")
        print("=" * 80)
        
        for i, tag in enumerate(test_tags, 1):
            print(f"\n📋 Tag {i}: {tag.get('Product Name*', 'Unknown')}")
            print("-" * 40)
            
            # Check what fields are available
            print("Available fields:")
            for key, value in tag.items():
                if value and str(value).strip():
                    print(f"  ✅ {key}: {value}")
                else:
                    print(f"  ❌ {key}: {value}")
        
        # Test specific template types
        print("\n🔍 Testing different template types:")
        print("=" * 80)
        
        template_types = ['double', 'mini', 'horizontal']
        test_tag = test_tags[0] if test_tags else None
        
        if test_tag:
            for template_type in template_types:
                print(f"\n📄 Testing {template_type.upper()} template:")
                print("-" * 40)
                
                try:
                    # Import TemplateProcessor with correct parameters
                    from src.core.generation.template_processor import TemplateProcessor
                    template_processor = TemplateProcessor(template_type, 'arial')
                    
                    # Build context for this template type
                    context = template_processor._build_label_context(test_tag, 1)
                    
                    # Check if template file exists
                    template_file = f"src/core/generation/templates/{template_type}.docx"
                    if os.path.exists(template_file):
                        print(f"  ✅ Template file exists: {template_file}")
                    else:
                        print(f"  ❌ Template file missing: {template_file}")
                    
                    # Show context fields
                    print("  Context fields:")
                    for key, value in context.items():
                        if value and str(value).strip():
                            print(f"    ✅ {key}: {value}")
                        else:
                            print(f"    ❌ {key}: {value}")
                            
                except Exception as e:
                    print(f"  ❌ Error with {template_type} template: {e}")
                    import traceback
                    traceback.print_exc()
        
        # Test manual placeholder replacement
        print("\n🔍 Testing manual placeholder replacement:")
        print("=" * 80)
        
        if test_tags:
            test_tag = test_tags[0]
            try:
                # Test double template manual replacement
                from src.core.generation.template_processor import TemplateProcessor
                template_processor = TemplateProcessor('double', 'arial')
                label_context = template_processor._build_label_context(test_tag, 1)
                
                # Create the proper context structure that manual replacement expects
                context = {'Label1': label_context}
                
                # Debug: Show the actual context structure
                print("Actual context structure:")
                for label_key, label_data in context.items():
                    print(f"  {label_key}:")
                    for field_key, field_value in label_data.items():
                        print(f"    {field_key}: {field_value}")
                
                # Check what placeholders would be replaced
                placeholders = [
                    'Label1.DescAndWeight',
                    'Label1.DOH',
                    'Label1.Price',
                    'Label1.Ratio_or_THC_CBD',
                    'Label1.ProductBrand',
                    'Label1.ProductName',
                    'Label1.Vendor'
                ]
                
                print("\nPlaceholder replacement test:")
                for placeholder in placeholders:
                    # Parse the placeholder to get label and field
                    if '.' in placeholder:
                        label_key, field_key = placeholder.split('.', 1)
                        if label_key in context and field_key in context[label_key]:
                            value = context[label_key][field_key]
                            if value and str(value).strip():
                                print(f"  ✅ {placeholder}: {value}")
                            else:
                                print(f"  ❌ {placeholder}: {value}")
                        else:
                            print(f"  ❌ {placeholder}: NOT FOUND")
                    else:
                        print(f"  ❌ {placeholder}: INVALID FORMAT")
                        
            except Exception as e:
                print(f"❌ Error testing placeholder replacement: {e}")
                import traceback
                traceback.print_exc()
        
        # Test actual template generation process
        print("\n🔍 Testing actual template generation process:")
        print("=" * 80)
        
        if test_tags:
            test_tag = test_tags[0]
            try:
                from src.core.generation.template_processor import TemplateProcessor
                template_processor = TemplateProcessor('double', 'arial')
                
                # Simulate the actual process_chunk method
                print("Testing process_chunk method...")
                
                # Create a test document
                from docx import Document
                test_doc = Document()
                test_doc.add_paragraph("Test document")
                
                # Process the chunk
                result = template_processor._process_chunk([test_tag])
                
                print(f"Process chunk result type: {type(result)}")
                
                # Check if the result is a Document object (success)
                if result and hasattr(result, 'tables'):
                    print("✅ Template generation successful!")
                    print(f"Document has {len(result.tables)} tables")
                    
                    # Check if there are any tables with content
                    if result.tables:
                        first_table = result.tables[0]
                        print(f"First table has {len(first_table.rows)} rows and {len(first_table.columns)} columns")
                        
                        # Check if there's content in the first cell
                        if first_table.rows and first_table.columns:
                            first_cell = first_table.rows[0].cells[0]
                            if first_cell.text.strip():
                                print("✅ First cell has content!")
                            else:
                                print("⚠️ First cell is empty")
                    else:
                        print("⚠️ No tables found in document")
                else:
                    print("❌ Template generation failed - no document returned")
                    
            except Exception as e:
                print(f"❌ Error in template generation: {e}")
                import traceback
                traceback.print_exc()
        
        return True
        
    except Exception as e:
        print(f"❌ Error during testing: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_template_generation() 