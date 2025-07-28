#!/usr/bin/env python3
"""
Test full pipeline post-processing to identify brand truncation issue.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docxtpl import DocxTemplate
from io import BytesIO

def test_full_pipeline_brand():
    """Test full pipeline post-processing to identify brand truncation issue."""
    
    print("Full Pipeline Brand Test")
    print("=" * 40)
    
    # Create a single test record
    test_record = {
        'Description': 'Pipeline Product',
        'Price': '$15',
        'Lineage': 'HYBRID',
        'Ratio_or_THC_CBD': 'THC: 22%',
        'ProductBrand': 'PIPELINE_BRAND_TEST',
        'ProductStrain': 'Pipeline Strain',
        'ProductType': 'flower',
        'JointRatio': '3.5g'
    }
    
    # Initialize template processor
    processor = TemplateProcessor('double', 'standard', scale_factor=1.0)
    
    # Get the template buffer
    template_buffer = processor._expanded_template_buffer
    template_buffer.seek(0)
    
    # Create DocxTemplate and build context
    doc = DocxTemplate(template_buffer)
    context = {}
    for i in range(1, 13):  # Double template has 12 labels (4x3 grid)
        if i == 1:
            context[f'Label{i}'] = processor._build_label_context(test_record, doc)
        else:
            context[f'Label{i}'] = {}
    
    # Render the template
    doc.render(context)
    
    # Save the raw rendered document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    rendered_doc = Document(buffer)
    
    print("After Jinja2 rendering:")
    print_brand_content(rendered_doc, "Jinja2")
    
    # Apply post-processing in the exact same order as _post_process_and_replace_content
    print("\n" + "="*50)
    print("Applying post-processing in full pipeline order...")
    
    # Step 1: Template-specific font sizing
    print("\nStep 1: _post_process_template_specific...")
    processor._post_process_template_specific(rendered_doc)
    print_brand_content(rendered_doc, "After template-specific font sizing")
    
    # Step 2: Convert BR markers
    print("\nStep 2: _convert_br_markers_to_line_breaks...")
    for table in rendered_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    processor._convert_br_markers_to_line_breaks(paragraph)
    for paragraph in rendered_doc.paragraphs:
        processor._convert_br_markers_to_line_breaks(paragraph)
    print_brand_content(rendered_doc, "After BR marker conversion")
    
    # Step 3: Fix ratio spacing
    print("\nStep 3: _fix_ratio_paragraph_spacing...")
    processor._fix_ratio_paragraph_spacing(rendered_doc)
    print_brand_content(rendered_doc, "After ratio spacing fix")
    
    # Step 4: Enforce Arial Bold
    print("\nStep 4: enforce_arial_bold_all_text...")
    from src.core.generation.docx_formatting import enforce_arial_bold_all_text
    enforce_arial_bold_all_text(rendered_doc)
    print_brand_content(rendered_doc, "After Arial bold enforcement")
    
    # Step 5: Apply brand centering
    print("\nStep 5: _apply_brand_centering_for_double_template...")
    processor.current_record = test_record  # Set current record like in the pipeline
    processor._apply_brand_centering_for_double_template(rendered_doc)
    print_brand_content(rendered_doc, "After brand centering")
    
    # Save the final result
    output_path = "test_full_pipeline_brand_output.docx"
    rendered_doc.save(output_path)
    print(f"\n✅ Final document saved: {output_path}")

def print_brand_content(doc, stage_name):
    """Print brand content from document."""
    print(f"\n{stage_name}:")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    lines = cell_text.split('\n')
                    for i, line in enumerate(lines):
                        if line.strip():
                            if 'PIPELINE_BRAND_TEST' in line:
                                print(f"  Line {i+1}: '{line.strip()}' ✓ FULL BRAND")
                            elif 'PIPELINE_BRAND' in line:
                                print(f"  Line {i+1}: '{line.strip()}' ⚠ PARTIAL BRAND")
                            elif 'BRAND_TEST' in line:
                                print(f"  Line {i+1}: '{line.strip()}' ⚠ PARTIAL BRAND")

if __name__ == "__main__":
    test_full_pipeline_brand() 