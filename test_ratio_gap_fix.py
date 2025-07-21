#!/usr/bin/env python3
"""
Test script to verify that the excessive ratio gap fix is working correctly.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt

def test_ratio_gap_fix():
    """Test that ratio content has proper tight spacing without excessive gaps."""
    
    print("=== Testing Ratio Gap Fix ===")
    
    # Create test data with ratio content that should have tight spacing
    test_records = [
        {
            'Product Name*': 'Test Edible Product',
            'Product Brand': 'Test Brand',
            'Product Type*': 'edible (solid)',
            'Lineage': 'MIXED',
            'Price': '$25.99',
            'Ratio': '100mg THC 50mg CBD 5mg CBG',
            'Description': 'Test description',
            'Weight Units': '4oz'
        },
        {
            'Product Name*': 'Test RSO Product',
            'Product Brand': 'Test Brand',
            'Product Type*': 'rso/co2 tankers',
            'Lineage': 'MIXED',
            'Price': '$45.99',
            'Ratio': '200mg THC 100mg CBD 10mg CBG',
            'Description': 'Test RSO description',
            'Weight Units': '1g'
        }
    ]
    
    # Create template processor
    processor = TemplateProcessor('vertical', {}, 1.0)
    
    try:
        # Process the records
        print("Processing test records...")
        doc = processor.process_records(test_records)
        
        if doc:
            print("✅ Document generated successfully!")
            
            # Check the content for proper spacing
            print("\nChecking paragraph spacing in generated content...")
            
            for table_idx, table in enumerate(doc.tables):
                print(f"\nTable {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"  Cell ({row_idx}, {cell_idx}): '{cell_text}'")
                            
                            # Check paragraph spacing for ratio content
                            for para_idx, paragraph in enumerate(cell.paragraphs):
                                if paragraph.text.strip():
                                    print(f"    Paragraph {para_idx + 1}: '{paragraph.text}'")
                                    
                                    # Check spacing settings
                                    space_before = paragraph.paragraph_format.space_before
                                    space_after = paragraph.paragraph_format.space_after
                                    line_spacing = paragraph.paragraph_format.line_spacing
                                    
                                    print(f"      Space before: {space_before}")
                                    print(f"      Space after: {space_after}")
                                    print(f"      Line spacing: {line_spacing}")
                                    
                                    # Check if this is ratio content
                                    ratio_patterns = ['mg THC', 'mg CBD', 'mg CBG', 'THC:', 'CBD:']
                                    if any(pattern in paragraph.text for pattern in ratio_patterns):
                                        print(f"      ✅ Ratio content detected")
                                        
                                        # Verify tight spacing
                                        if space_before and space_before.pt == 0:
                                            print(f"      ✅ Space before is 0pt")
                                        else:
                                            print(f"      ❌ Space before should be 0pt, got {space_before}")
                                        
                                        if space_after and space_after.pt == 0:
                                            print(f"      ✅ Space after is 0pt")
                                        else:
                                            print(f"      ❌ Space after should be 0pt, got {space_after}")
                                        
                                        if line_spacing and line_spacing == 1.0:
                                            print(f"      ✅ Line spacing is 1.0")
                                        else:
                                            print(f"      ❌ Line spacing should be 1.0, got {line_spacing}")
                                    
                                    # Check for line breaks
                                    if '\n' in paragraph.text:
                                        print(f"      ✅ Contains line breaks")
                                        lines = paragraph.text.split('\n')
                                        for i, line in enumerate(lines):
                                            print(f"        Line {i+1}: '{line}'")
                                    else:
                                        print(f"      No line breaks found")
            
            # Save the document for inspection
            output_path = "test_ratio_gap_fix_output.docx"
            doc.save(output_path)
            print(f"\n✅ Document saved to: {output_path}")
            print("You can open this file in Word to verify the ratio gap fix is working correctly.")
            
        else:
            print("❌ Failed to generate document")
            
    except Exception as e:
        print(f"❌ Error during document generation: {e}")
        import traceback
        traceback.print_exc()

def test_direct_paragraph_spacing():
    """Test direct paragraph spacing fix."""
    
    print("\n=== Testing Direct Paragraph Spacing ===")
    
    # Create a simple document
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add content with |BR| markers
    test_content = "100mg THC|BR|50mg CBD|BR|5mg CBG"
    cell.text = test_content
    
    print(f"Original content: '{cell.text}'")
    
    # Create template processor
    processor = TemplateProcessor('vertical', {}, 1.0)
    
    # Test the line break conversion function directly
    paragraph = cell.paragraphs[0]
    print(f"Before conversion: '{paragraph.text}'")
    
    processor._convert_br_markers_to_line_breaks(paragraph)
    print(f"After conversion: '{paragraph.text}'")
    
    # Check spacing settings
    space_before = paragraph.paragraph_format.space_before
    space_after = paragraph.paragraph_format.space_after
    line_spacing = paragraph.paragraph_format.line_spacing
    
    print(f"Space before: {space_before}")
    print(f"Space after: {space_after}")
    print(f"Line spacing: {line_spacing}")
    
    # Check runs
    runs = list(paragraph.runs)
    print(f"Number of runs: {len(runs)}")
    for i, run in enumerate(runs):
        print(f"  Run {i+1}: '{run.text}'")
    
    # Save for inspection
    doc.save("test_direct_spacing.docx")
    print("✅ Direct spacing test saved to: test_direct_spacing.docx")

if __name__ == "__main__":
    test_ratio_gap_fix()
    test_direct_paragraph_spacing()
    print("\n=== Test Complete ===") 