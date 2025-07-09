#!/usr/bin/env python3
"""
Debug script to check mini template expansion.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def debug_mini_expansion():
    """Debug the mini template expansion process."""
    
    print("Debugging Mini Template Expansion")
    print("=" * 50)
    
    try:
        # Create processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Check the expanded template buffer
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            
            if doc.tables:
                table = doc.tables[0]
                print(f"Expanded template table dimensions: {len(table.rows)}x{len(table.columns)}")
                
                # Check the first cell content
                cell = table.cell(0, 0)
                print(f"First cell text: '{cell.text}'")
                
                # Check for template variables
                if '{{Label1.' in cell.text:
                    print("✅ Template variables found")
                    # Extract all template variables
                    import re
                    variables = re.findall(r'\{\{Label1\.(\w+)\}\}', cell.text)
                    print(f"Template variables: {variables}")
                else:
                    print("❌ No template variables found")
            else:
                print("❌ No tables found in expanded template")
        else:
            print("❌ No expanded template buffer found")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_mini_expansion() 