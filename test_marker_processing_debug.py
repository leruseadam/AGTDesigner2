#!/usr/bin/env python3
"""
Debug script to test marker processing logic for Product Strain.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.unified_font_sizing import get_font_size_by_marker
from src.core.formatting.markers import wrap_with_marker
from docx.shared import Pt
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def test_marker_processing():
    """Test the marker processing logic step by step."""
    
    print("=== TESTING MARKER PROCESSING LOGIC ===")
    
    # Test data
    test_data = {
        'ProductStrain': 'Test Strain Name'
    }
    
    print(f"Test data: {test_data}")
    
    # Test 1: Check marker wrapping
    print("\n1. Testing marker wrapping...")
    wrapped_strain = wrap_with_marker(test_data['ProductStrain'], 'PRODUCTSTRAIN')
    print(f"   Wrapped: {repr(wrapped_strain)}")
    
    # Test 2: Check font sizing
    print("\n2. Testing font sizing...")
    font_size = get_font_size_by_marker(test_data['ProductStrain'], 'PRODUCTSTRAIN', 'double', 1.0)
    print(f"   Font size: {font_size.pt}pt")
    
    # Test 3: Check marker detection
    print("\n3. Testing marker detection...")
    start_marker = 'PRODUCTSTRAIN_START'
    end_marker = 'PRODUCTSTRAIN_END'
    
    if start_marker in wrapped_strain and end_marker in wrapped_strain:
        print(f"   ✅ Markers found in wrapped text")
        start_idx = wrapped_strain.find(start_marker)
        end_idx = wrapped_strain.find(end_marker)
        print(f"   Start marker at: {start_idx}")
        print(f"   End marker at: {end_idx}")
        
        # Extract content
        content_start = start_idx + len(start_marker)
        content_end = end_idx
        content = wrapped_strain[content_start:content_end]
        print(f"   Extracted content: {repr(content)}")
    else:
        print(f"   ❌ Markers not found in wrapped text")
    
    # Test 4: Check template processor marker processing
    print("\n4. Testing template processor marker processing...")
    
    try:
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Build context
        context = processor._build_label_context(test_data, None)
        product_strain = context.get('ProductStrain', '')
        print(f"   Context ProductStrain: {repr(product_strain)}")
        
        # Test the marker processing function directly
        from docx import Document
        from docx.shared import Pt
        
        # Create a test document with the marker
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run(product_strain)
        
        print(f"   Created test document with marker")
        print(f"   Paragraph text: {repr(paragraph.text)}")
        
        # Process the paragraph
        processor._process_paragraph_for_markers_template_specific(paragraph, ['PRODUCTSTRAIN'])
        
        print(f"   After processing:")
        print(f"   Paragraph text: {repr(paragraph.text)}")
        
        for i, run in enumerate(paragraph.runs):
            if run.text.strip():
                font_size_pt = run.font.size.pt if run.font.size else "NO SIZE"
                print(f"   Run {i}: '{run.text}' -> {font_size_pt}pt")
        
    except Exception as e:
        print(f"   ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_marker_processing() 