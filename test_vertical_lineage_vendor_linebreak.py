#!/usr/bin/env python3
"""
Test script to verify Vertical template automatic line breaking for ProductVendor
when Lineage is "Hybrid/Indica" or "Hybrid/Sativa".
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_vertical_lineage_vendor_linebreak():
    """Test the automatic line breaking for ProductVendor in Vertical template."""
    
    print("Testing Vertical Template Lineage/Vendor Line Breaking")
    print("=" * 60)
    
    # Create a test document
    doc = Document()
    
    # Test cases
    test_cases = [
        {
            'name': 'Hybrid/Indica - Should force vendor to next line',
            'lineage': 'HYBRID/INDICA',
            'vendor': 'Test Vendor',
            'template_type': 'vertical',
            'expected_lines': 2
        },
        {
            'name': 'Hybrid/Sativa - Should force vendor to next line',
            'lineage': 'HYBRID/SATIVA',
            'vendor': 'Another Vendor',
            'template_type': 'vertical',
            'expected_lines': 2
        },
        {
            'name': 'Regular Hybrid - Should stay on same line',
            'lineage': 'HYBRID',
            'vendor': 'Regular Vendor',
            'template_type': 'vertical',
            'expected_lines': 1
        },
        {
            'name': 'Sativa - Should stay on same line',
            'lineage': 'SATIVA',
            'vendor': 'Sativa Vendor',
            'template_type': 'vertical',
            'expected_lines': 1
        },
        {
            'name': 'Indica - Should stay on same line',
            'lineage': 'INDICA',
            'vendor': 'Indica Vendor',
            'template_type': 'vertical',
            'expected_lines': 1
        },
        {
            'name': 'Horizontal template Hybrid/Indica - Should stay on same line',
            'lineage': 'HYBRID/INDICA',
            'vendor': 'Horizontal Vendor',
            'template_type': 'horizontal',
            'expected_lines': 1
        },
        {
            'name': 'Mini template Hybrid/Sativa - Should stay on same line',
            'lineage': 'HYBRID/SATIVA',
            'vendor': 'Mini Vendor',
            'template_type': 'mini',
            'expected_lines': 1
        }
    ]
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\nTest {i}: {test_case['name']}")
        print(f"  Lineage: {test_case['lineage']}")
        print(f"  Vendor: {test_case['vendor']}")
        print(f"  Template: {test_case['template_type']}")
        print(f"  Expected lines: {test_case['expected_lines']}")
        
        # Create template processor
        processor = TemplateProcessor(test_case['template_type'], {}, 1.0)
        
        # Create a test paragraph
        paragraph = doc.add_paragraph()
        
        # Process the lineage and vendor
        processor._process_combined_lineage_vendor(paragraph, test_case['lineage'], test_case['vendor'])
        
        # Count the number of lines (by counting \n characters + 1)
        paragraph_text = paragraph.text
        line_count = paragraph_text.count('\n') + 1 if paragraph_text else 0
        
        print(f"  Result: {line_count} lines")
        print(f"  Paragraph text: '{paragraph_text}'")
        
        # Check if the result matches expectation
        if line_count == test_case['expected_lines']:
            print(f"  Status: ✅ PASS")
        else:
            print(f"  Status: ❌ FAIL (expected {test_case['expected_lines']}, got {line_count})")
        
        # Add spacing between tests
        doc.add_paragraph()
    
    print("\n" + "=" * 60)
    print("Testing with actual data from file:")
    
    # Test with actual data
    from src.core.data.excel_processor import ExcelProcessor
    from src.core.data.excel_processor import get_default_upload_file
    
    processor = ExcelProcessor()
    default_file = get_default_upload_file()
    
    if default_file:
        print(f"Found default file: {default_file}")
        
        # Load the file
        success = processor.load_file(default_file)
        if success:
            print(f"✅ File loaded successfully: {len(processor.df)} rows")
            
            # Find products with Hybrid/Indica or Hybrid/Sativa lineage
            hybrid_products = processor.df[
                processor.df['Lineage'].str.upper().isin(['HYBRID/INDICA', 'HYBRID/SATIVA'])
            ].head(5)
            
            print(f"\nFound {len(hybrid_products)} products with Hybrid/Indica or Hybrid/Sativa lineage:")
            
            for idx, row in hybrid_products.iterrows():
                product_name = row.get('Product Name*', f'Product {idx}')
                lineage = row.get('Lineage', '')
                vendor = row.get('Vendor', '')
                
                print(f"\n  Product: {product_name}")
                print(f"  Lineage: {lineage}")
                print(f"  Vendor: {vendor}")
                
                # Test with template processor
                template_processor = TemplateProcessor('vertical', {}, 1.0)
                test_doc = Document()
                test_paragraph = test_doc.add_paragraph()
                
                template_processor._process_combined_lineage_vendor(test_paragraph, lineage, vendor)
                
                line_count = test_paragraph.text.count('\n') + 1 if test_paragraph.text else 0
                print(f"  Result: {line_count} lines")
                print(f"  Text: '{test_paragraph.text}'")
                
                if line_count == 2:
                    print(f"  Status: ✅ Correctly forced to 2 lines")
                else:
                    print(f"  Status: ⚠️  Expected 2 lines, got {line_count}")
        else:
            print("❌ Failed to load file")
    else:
        print("❌ No default file found")
    
    print("\n✅ Vertical template lineage/vendor line breaking test completed!")

if __name__ == "__main__":
    test_vertical_lineage_vendor_linebreak() 