#!/usr/bin/env python3
"""
Debug script to investigate the template expansion process
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from io import BytesIO

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor

def check_table_structure(doc, stage_name):
    """Check and print table structure details"""
    print(f"\n=== {stage_name} ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        print(f"  Rows: {len(table.rows)}")
        print(f"  Columns: {len(table.columns)}")
        
        # Check each row
        for j, row in enumerate(table.rows):
            height = row.height
            height_rule = row.height_rule
            cells = len(row.cells)
            print(f"  Row {j}: height={height}, rule={height_rule}, cells={cells}")
            
            # Check each cell
            for k, cell in enumerate(row.cells):
                paragraphs = len(cell.paragraphs)
                print(f"    Cell {k}: paragraphs={paragraphs}")
                for p_idx, para in enumerate(cell.paragraphs):
                    runs = len(para.runs)
                    text = para.text[:50] + "..." if len(para.text) > 50 else para.text
                    print(f"      Paragraph {p_idx}: runs={runs}, text='{text}'")

def test_template_expansion():
    """Test the template expansion process step by step"""
    
    # Step 1: Load original template
    template_path = "src/core/generation/templates/horizontal.docx"
    doc = Document(template_path)
    check_table_structure(doc, "1. Original template")
    
    # Step 2: Create TemplateProcessor
    processor = TemplateProcessor('horizontal', {}, 1.0)
    
    # Step 3: Check expanded template
    expanded_doc = Document(processor._expanded_template_buffer)
    check_table_structure(expanded_doc, "2. After template expansion")
    
    # Step 4: Check XML structure of expanded template
    print("\n=== XML Analysis of Expanded Template ===")
    for i, table in enumerate(expanded_doc.tables):
        print(f"Table {i} XML:")
        table_xml = table._element.xml
        
        # Count tr elements (rows)
        import re
        tr_matches = re.findall(r'<w:tr[^>]*>', table_xml)
        print(f"  Total <w:tr> elements: {len(tr_matches)}")
        
        # Count trHeight elements
        if 'trHeight' in table_xml:
            tr_height_matches = re.findall(r'<w:trHeight[^>]*>', table_xml)
            print(f"  Total trHeight elements: {len(tr_height_matches)}")
            
            with_hrule = 0
            without_hrule = 0
            for idx, match in enumerate(tr_height_matches):
                if 'w:hRule="exact"' in match:
                    with_hrule += 1
                    print(f"    {idx}: {match} - HAS hRule")
                else:
                    without_hrule += 1
                    print(f"    {idx}: {match} - NO hRule")
            
            print(f"  Summary: {with_hrule} with hRule, {without_hrule} without hRule")
        else:
            print("  No trHeight elements found in XML")
    
    # Step 5: Save expanded template and check
    expanded_path = "test_expanded_template.docx"
    expanded_doc.save(expanded_path)
    
    reloaded_doc = Document(expanded_path)
    check_table_structure(reloaded_doc, "3. After saving expanded template")
    
    # Clean up
    if os.path.exists(expanded_path):
        os.remove(expanded_path)

if __name__ == "__main__":
    test_template_expansion() 