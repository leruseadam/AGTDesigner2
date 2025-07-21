#!/usr/bin/env python3
"""
Test script to verify line break conversion is working correctly.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.core.generation.template_processor import TemplateProcessor

def test_line_break_conversion():
    """Test that |BR| markers are converted to actual line breaks."""
    
    print("=== Testing Line Break Conversion ===")
    
    # Create a test document
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add test content with |BR| markers
    test_content = "100mg THC|BR|50mg CBD|BR|5mg CBG"
    cell.text = test_content
    
    print(f"Original content: '{cell.text}'")
    
    # Create template processor
    processor = TemplateProcessor('vertical', {}, 1.0)
    
    # Test the line break conversion function directly
    paragraph = cell.paragraphs[0]
    print(f"Before conversion: '{paragraph.text}'")
    
    processor._convert_br_markers_to_line_breaks(paragraph)
    print(f"After conversion: '{paragraph.text}'")
    
    # Check if line breaks were created
    runs = list(paragraph.runs)
    print(f"Number of runs: {len(runs)}")
    
    for i, run in enumerate(runs):
        print(f"  Run {i+1}: '{run.text}' (font: {run.font.name}, size: {run.font.size})")
    
    # Test with marker processing
    print("\n=== Testing with Marker Processing ===")
    
    # Create another test document with markers
    doc2 = Document()
    table2 = doc2.add_table(rows=1, cols=1)
    cell2 = table2.cell(0, 0)
    
    # Add content with markers and |BR|
    test_content_with_markers = "RATIO_START100mg THC|BR|50mg CBDRATIO_END"
    cell2.text = test_content_with_markers
    
    print(f"Original content with markers: '{cell2.text}'")
    
    # Process the document
    processor._post_process_and_replace_content(doc2)
    
    print(f"After processing: '{cell2.text}'")
    
    # Check the results
    for para_idx, paragraph in enumerate(cell2.paragraphs):
        print(f"Paragraph {para_idx+1}:")
        for run_idx, run in enumerate(paragraph.runs):
            print(f"  Run {run_idx+1}: '{run.text}' (font: {run.font.name}, size: {run.font.size})")
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    test_line_break_conversion() 