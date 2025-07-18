#!/usr/bin/env python3

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.data.excel_processor import ExcelProcessor
from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.docx_formatting import apply_lineage_colors
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)

def test_description_lineage_processing():
    """Test Description and Lineage processing to see what's happening."""
    
    # Create ExcelProcessor and add a test record
    excel_processor = ExcelProcessor()
    
    # Create a test DataFrame with our sample data
    import pandas as pd
    
    test_data = {
        'Product Name*': ['Hustler\'s Ambition - Wax - Acapulco Gold (S) - 1g'],
        'Product Brand': ['Hustler\'s Ambition'],
        'Price': ['$25.00'],
        'Lineage': ['SATIVA'],
        'Product Type*': ['concentrate'],
        'Weight Units': ['g'],
        'Weight*': ['1'],
        'Ratio': ['THC: 22% CBD: 1%'],
        'DOH': ['Yes']
    }
    
    excel_processor.df = pd.DataFrame(test_data)
    excel_processor.selected_tags = ['Hustler\'s Ambition - Wax - Acapulco Gold (S) - 1g']
    
    print("=== TESTING EXCEL PROCESSOR ===")
    print(f"DataFrame columns: {list(excel_processor.df.columns)}")
    print(f"DataFrame data:")
    print(excel_processor.df.to_string())
    
    # Process the DataFrame to extract Description
    excel_processor.complete_processing()
    
    print(f"\n=== AFTER COMPLETE PROCESSING ===")
    print(f"Description column: {excel_processor.df['Description'].iloc[0]}")
    print(f"Ratio column: {excel_processor.df['Ratio'].iloc[0]}")
    
    # Get selected records
    processed_records = excel_processor.get_selected_records('horizontal')
    
    print(f"\n=== PROCESSED RECORDS ===")
    for i, record in enumerate(processed_records):
        print(f"Record {i}:")
        for key, value in record.items():
            if key in ['Description', 'Lineage', 'DescAndWeight', 'ProductBrand', 'Price', 'ProductStrain']:
                print(f"  {key}: {repr(value)}")
    
    if not processed_records:
        print("No processed records returned!")
        return
    
    # Test template processing
    template_processor = TemplateProcessor('horizontal', {}, 1.0)
    
    # Create a dummy document for testing
    from docx import Document
    doc = Document()
    
    # Build label context
    label_context = template_processor._build_label_context(processed_records[0], doc)
    
    print("\n=== AFTER BUILDING LABEL CONTEXT ===")
    for key, value in label_context.items():
        if key in ['Description', 'Lineage', 'DescAndWeight', 'ProductBrand', 'Price']:
            print(f"{key}: {repr(value)}")
    
    # Test marker processing
    print("\n=== TESTING MARKER PROCESSING ===")
    
    # Create a test paragraph with markers
    test_para = doc.add_paragraph()
    test_para.add_run("Description: ")
    test_para.add_run(label_context.get('Description', ''))
    test_para.add_run(" | Lineage: ")
    test_para.add_run(label_context.get('Lineage', ''))
    
    print(f"Before processing: {test_para.text}")
    
    # Process the paragraph
    template_processor._process_paragraph_for_marker_template_specific(test_para, 'DESC')
    template_processor._process_paragraph_for_marker_template_specific(test_para, 'LINEAGE')
    
    print(f"After processing: {test_para.text}")
    
    # Check if markers are still present
    full_text = "".join(run.text for run in test_para.runs)
    print(f"Full text with runs: {repr(full_text)}")
    
    # Check for marker presence
    desc_markers = ['DESC_START', 'DESC_END']
    lineage_markers = ['LINEAGE_START', 'LINEAGE_END']
    
    for marker in desc_markers:
        if marker in full_text:
            print(f"Found {marker} in text")
        else:
            print(f"Missing {marker} in text")
    
    for marker in lineage_markers:
        if marker in full_text:
            print(f"Found {marker} in text")
        else:
            print(f"Missing {marker} in text")

if __name__ == "__main__":
    test_description_lineage_processing() 