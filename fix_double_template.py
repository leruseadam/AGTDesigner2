#!/usr/bin/env python3
"""
Fix the double template by adding missing Jinja2 variables.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def fix_double_template():
    """Fix the double template by adding missing Jinja2 variables."""
    
    print("Fixing Double Template")
    print("=" * 30)
    
    # Get the template path
    template_path = Path(__file__).parent / "src" / "core" / "generation" / "templates" / "double.docx"
    
    if not template_path.exists():
        print(f"❌ Template not found: {template_path}")
        return False
    
    try:
        # Load the template
        doc = Document(template_path)
        
        if not doc.tables:
            print("❌ No tables found in template")
            return False
        
        table = doc.tables[0]
        if not table.rows or not table.columns:
            print("❌ Empty table found")
            return False
        
        # Get the first cell (this is what gets copied to all other cells)
        first_cell = table.cell(0, 0)
        
        print(f"Current first cell content: '{first_cell.text}'")
        
        # Clear the cell content
        for paragraph in first_cell.paragraphs[:]:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Add all the necessary variables with proper formatting
        # Description at the top
        desc_para = first_cell.add_paragraph()
        desc_para.text = "{{Label1.Description}}"
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Price
        price_para = first_cell.add_paragraph()
        price_para.text = "{{Label1.Price}}"
        price_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lineage (already existed)
        lineage_para = first_cell.add_paragraph()
        lineage_para.text = "{{Label1.Lineage}}"
        lineage_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Product Strain (already existed)
        strain_para = first_cell.add_paragraph()
        strain_para.text = "{{Label1.ProductStrain}}"
        strain_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ratio/THC_CBD
        ratio_para = first_cell.add_paragraph()
        ratio_para.text = "{{Label1.Ratio_or_THC_CBD}}"
        ratio_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Product Brand
        brand_para = first_cell.add_paragraph()
        brand_para.text = "{{Label1.ProductBrand}}"
        brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Joint Ratio
        joint_ratio_para = first_cell.add_paragraph()
        joint_ratio_para.text = "{{Label1.JointRatio}}"
        joint_ratio_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        print(f"Updated first cell content: '{first_cell.text}'")
        
        # Save the template
        doc.save(template_path)
        
        print("✅ Double template fixed successfully!")
        print("Added variables:")
        print("  - {{Label1.Description}}")
        print("  - {{Label1.Price}}")
        print("  - {{Label1.Lineage}} (already existed)")
        print("  - {{Label1.ProductStrain}} (already existed)")
        print("  - {{Label1.Ratio_or_THC_CBD}}")
        print("  - {{Label1.ProductBrand}}")
        print("  - {{Label1.JointRatio}}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error fixing template: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = fix_double_template()
    if success:
        print("\nTemplate has been fixed. You can now test the double template generation.")
    else:
        print("\nFailed to fix template.") 