#!/usr/bin/env python3
"""
Script to examine template structure
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from pathlib import Path

def examine_template_structure(docx_path):
    """Examine the structure of a DOCX template"""
    try:
        doc = Document(docx_path)
        print(f"\n=== {docx_path.name} Structure ===")
        
        # Check paragraphs outside tables
        if doc.paragraphs:
            print("Paragraphs outside tables:")
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    print(f"  Para {i}: '{para.text}'")
        
        # Check tables
        if doc.tables:
            print(f"Tables: {len(doc.tables)}")
            for table_idx, table in enumerate(doc.tables):
                print(f"  Table {table_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            print(f"    Cell [{row_idx},{col_idx}]: '{cell_text}'")
        
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")

def main():
    template_dir = Path("src/core/generation/templates")
    
    # Check vertical and mini templates specifically
    for template_name in ["vertical.docx", "mini.docx"]:
        template_path = template_dir / template_name
        if template_path.exists():
            examine_template_structure(template_path)

if __name__ == "__main__":
    main() 