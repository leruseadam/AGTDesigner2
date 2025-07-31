#!/usr/bin/env python3
"""
Simple test to verify DocxTemplate is working correctly
"""

import sys
import os
from pathlib import Path

# Add the src directory to the Python path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO

def test_simple_docxtpl():
    """Test DocxTemplate with a simple template."""
    print("=== Testing Simple DocxTemplate ===")
    
    # Create a simple template
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "{{Label1.Lineage}} {{Label1.ProductVendor}} {{Label1.ProductStrain}}"
    
    # Save template to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create context
    context = {
        'Label1': {
            'Lineage': 'Test Lineage',
            'ProductVendor': 'Test Vendor',
            'ProductStrain': 'Test Strain'
        }
    }
    
    print(f"Template content: {cell.text}")
    print(f"Context: {context}")
    
    # Test DocxTemplate
    try:
        template = DocxTemplate(buffer)
        template.render(context)
        
        # Get rendered document
        rendered_doc = template.get_docx()
        
        # Check result
        result_text = rendered_doc.tables[0].cell(0, 0).text
        print(f"Rendered content: {repr(result_text)}")
        
        if '{{Label1' in result_text:
            print("❌ Placeholders not replaced")
            return False
        else:
            print("✅ Placeholders replaced successfully")
            return True
            
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    test_simple_docxtpl() 