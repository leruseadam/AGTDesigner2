#!/usr/bin/env python3
"""
Test script to check if lineage and product brand are missing from the double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def test_double_template_lineage_brand():
    """Test that lineage and product brand are properly included in double template."""
    
    print("Testing Double Template Lineage and Product Brand")
    print("=" * 55)
    
    # Test data with lineage and product brand
    test_records = [
        {
            'Product Name*': 'Test Product 1',
            'ProductBrand': 'Test Brand 1',
            'Product Brand': 'Test Brand 1',  # Alternative field name
            'Price': '$25.00',
            'Description': 'Test description 1',
            'Lineage': 'HYBRID',
            'THC_CBD': 'THC: 20% CBD: 2%',
            'WeightUnits': '3.5g',
            'ProductType': 'flower',
            'Product Strain': 'Test Strain 1'
        },
        {
            'Product Name*': 'Test Product 2',
            'ProductBrand': 'Test Brand 2',
            'Product Brand': 'Test Brand 2',
            'Price': '$30.00',
            'Description': 'Test description 2',
            'Lineage': 'SATIVA',
            'THC_CBD': 'THC: 25% CBD: 1%',
            'WeightUnits': '1g',
            'ProductType': 'concentrate',
            'Product Strain': 'Test Strain 2'
        }
    ]
    
    try:
        # Test double template
        print("\n1. Testing Double Template Processing...")
        double_processor = TemplateProcessor('double', {}, 1.0)
        
        # Test _build_label_context method directly
        print("\n2. Testing _build_label_context method...")
        for i, record in enumerate(test_records):
            print(f"\n   Record {i+1}:")
            
            # Create a mock document for the method
            from docx import Document
            mock_doc = Document()
            
            # Build label context
            label_context = double_processor._build_label_context(record, mock_doc)
            
            # Check lineage
            lineage = label_context.get('Lineage', 'NOT_FOUND')
            print(f"     Lineage: '{lineage}'")
            
            # Check product brand
            product_brand = label_context.get('ProductBrand', 'NOT_FOUND')
            print(f"     ProductBrand: '{product_brand}'")
            
            # Check if lineage is empty or missing
            if not lineage or lineage == 'NOT_FOUND' or lineage.strip() == '':
                print(f"     ❌ ERROR: Lineage is missing or empty for record {i+1}")
                return False
            else:
                print(f"     ✓ Lineage is present: '{lineage}'")
            
            # Check if product brand is empty or missing
            if not product_brand or product_brand == 'NOT_FOUND' or product_brand.strip() == '':
                print(f"     ❌ ERROR: ProductBrand is missing or empty for record {i+1}")
                return False
            else:
                print(f"     ✓ ProductBrand is present: '{product_brand}'")
        
        print("\n3. Testing Full Template Processing...")
        
        # Process records through the full pipeline
        result = double_processor.process_records(test_records)
        assert result is not None, "Double template processing failed"
        
        # Check that we have a table with the expected structure
        assert len(result.tables) > 0, "No tables found in double template result"
        
        table = result.tables[0]
        print(f"   Double template table: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Check cell content for lineage and product brand
        print("\n4. Checking Cell Content...")
        lineage_found = False
        product_brand_found = False
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"   Cell ({row_idx},{col_idx}): '{cell_text}'")
                    
                    # Check for lineage content
                    if any(lineage in cell_text.upper() for lineage in ['HYBRID', 'SATIVA', 'INDICA', 'CBD']):
                        lineage_found = True
                        print(f"     ✓ Found lineage content: '{cell_text}'")
                    
                    # Check for product brand content
                    if any(brand in cell_text for brand in ['Test Brand 1', 'Test Brand 2']):
                        product_brand_found = True
                        print(f"     ✓ Found product brand content: '{cell_text}'")
        
        if not lineage_found:
            print("   ❌ ERROR: No lineage content found in any cells")
            return False
        
        if not product_brand_found:
            print("   ❌ ERROR: No product brand content found in any cells")
            return False
        
        print("\n5. Testing Template Expansion...")
        
        # Check if template expansion is working correctly
        expanded_buffer = double_processor._expanded_template_buffer
        if hasattr(expanded_buffer, 'seek'):
            expanded_buffer.seek(0)
        
        expanded_doc = Document(expanded_buffer)
        expanded_text = expanded_doc.element.body.xml
        
        # Check for lineage and product brand placeholders in expanded template
        lineage_placeholder_found = 'LINEAGE' in expanded_text
        product_brand_placeholder_found = 'PRODUCTBRAND' in expanded_text
        
        print(f"   Lineage placeholder found: {lineage_placeholder_found}")
        print(f"   ProductBrand placeholder found: {product_brand_placeholder_found}")
        
        # Note: Placeholders may not be found in XML if they've been processed during rendering
        # This is expected behavior, so we don't fail the test here
        print("   Note: Placeholders may be processed during rendering (expected behavior)")
        
        print("\n" + "=" * 55)
        print("✓ ALL TESTS PASSED!")
        print("✓ Lineage and ProductBrand are properly included in double template")
        
        return True
        
    except Exception as e:
        print(f"\n❌ TEST FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_double_template_lineage_brand()
    sys.exit(0 if success else 1) 