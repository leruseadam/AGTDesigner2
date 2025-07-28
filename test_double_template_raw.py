#!/usr/bin/env python3
"""
Test to check double template immediately after Jinja2 rendering.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
from io import BytesIO

def test_double_template_raw():
    """Test double template immediately after Jinja2 rendering."""
    
    print("Double Template Raw Jinja2 Rendering Test")
    print("=" * 50)
    
    # Create test records
    test_records = [
        {
            'Description': 'Amnesia Lemon Pre-Roll',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductBrand': 'Test Brand 1',
            'ProductType': 'pre-roll'
        }
    ]
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Get the expanded template buffer
        if hasattr(processor._expanded_template_buffer, 'seek'):
            processor._expanded_template_buffer.seek(0)
        
        # Create DocxTemplate and render without post-processing
        from docxtpl import DocxTemplate
        doc = DocxTemplate(processor._expanded_template_buffer)
        
        # Build context
        context = {}
        for i, record in enumerate(test_records):
            processor.current_record = record
            context[f'Label{i+1}'] = processor._build_label_context(record, doc)
        
        # Fill empty slots
        for i in range(len(test_records), processor.chunk_size):
            context[f'Label{i+1}'] = {}
        
        print("Context built:")
        for key, value in context.items():
            if value:
                print(f"  {key}: {value}")
        
        # Render the template
        print("\nRendering template...")
        doc.render(context)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Load the rendered document
        rendered_doc = Document(buffer)
        
        print("✓ Document rendered successfully")
        
        # Analyze the raw rendered document
        print("\nAnalyzing raw rendered document:")
        print("-" * 35)
        
        if rendered_doc.tables:
            table = rendered_doc.tables[0]
            print(f"Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    print(f"\nCell ({row_idx},{col_idx}):")
                    cell_text = cell.text.strip()
                    print(f"  Text: '{cell_text}'")
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                font_size = run.font.size.pt if run.font.size else "None"
                                print(f"    Run: '{run.text.strip()}' -> {font_size}pt")
        else:
            print("No tables found in document")
            
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_double_template_raw() 