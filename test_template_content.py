#!/usr/bin/env python3
"""
Examine the double template content to see what fields are actually used.
"""

from docx import Document
import re

def examine_double_template():
    """Examine the double template content."""
    print("=== DOUBLE TEMPLATE CONTENT EXAMINATION ===")
    
    # Load the double template
    doc = Document('src/core/generation/templates/double.docx')
    
    # Extract all text with more detail
    print("All text in template:")
    print("-" * 50)
    
    # Check paragraphs outside tables
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"Paragraph {i}: '{paragraph.text}'")
    
    # Check table content
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}:")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"  Cell ({row_idx}, {col_idx}): '{cell_text}'")
                    
                    # Check individual paragraphs in the cell
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            print(f"    Paragraph {para_idx}: '{paragraph.text}'")
                            
                            # Check individual runs
                            for run_idx, run in enumerate(paragraph.runs):
                                if run.text.strip():
                                    print(f"      Run {run_idx}: '{run.text}'")
    
    print("\n" + "=" * 50)
    
    # Look for Jinja2 placeholders
    all_text = ""
    for paragraph in doc.paragraphs:
        all_text += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_text += paragraph.text + "\n"
    
    # Find Jinja2 placeholders
    placeholder_pattern = r'\{\{[^}]+\}\}'
    placeholders = re.findall(placeholder_pattern, all_text)
    
    print("Jinja2 placeholders found:")
    for placeholder in placeholders:
        print(f"  {placeholder}")
    
    # Check if there are any other patterns
    print("\nOther patterns found:")
    other_patterns = re.findall(r'[A-Z][a-zA-Z0-9_]*', all_text)
    unique_patterns = list(set(other_patterns))
    for pattern in sorted(unique_patterns):
        if len(pattern) > 2:  # Filter out short patterns
            print(f"  {pattern}")

if __name__ == "__main__":
    examine_double_template() 