#!/usr/bin/env python3
"""
Direct test of THC:CBD line spacing logic for vertical template
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_thc_cbd_line_spacing_direct():
    """Test the THC:CBD line spacing logic directly"""
    print("🧪 Direct Test of THC:CBD Line Spacing Logic")
    print("=" * 50)
    
    # Create a test document
    doc = Document()
    
    # Add a table with cells
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    
    # Add THC:CBD content to a paragraph
    paragraph = cell.paragraphs[0]
    paragraph.text = "THC:\nCBD:"
    
    # Apply the line spacing logic that should be used for vertical template
    if "THC:\nCBD:" in paragraph.text:
        paragraph.paragraph_format.line_spacing = 2.0
        print("✅ Applied line spacing of 2.0 to THC:CBD content")
    
    # Check the line spacing
    line_spacing = paragraph.paragraph_format.line_spacing
    print(f"📏 Line spacing: {line_spacing}")
    
    if line_spacing == 2.0:
        print("🎉 SUCCESS: Line spacing is correctly set to 2.0")
        return True
    else:
        print(f"❌ FAILED: Line spacing should be 2.0, but is {line_spacing}")
        return False

if __name__ == "__main__":
    success = test_thc_cbd_line_spacing_direct()
    sys.exit(0 if success else 1) 