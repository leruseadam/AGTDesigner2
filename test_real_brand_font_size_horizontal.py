#!/usr/bin/env python3
"""
Test script to verify that the brand font size fix works with real template generation for horizontal templates.
"""

from src.core.generation.template_processor import TemplateProcessor
from docx.shared import Pt

def test_real_brand_font_size_horizontal():
    """Test the brand font size fix with real template generation for horizontal templates."""
    
    print("=== Real Brand Font Size Horizontal Test ===\n")
    
    # Test data with long brand names
    test_records = [
        {
            'ProductBrand': 'Fairwinds Manufacturing',
            'Price': '$25.99',
            'Lineage': 'Hybrid',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'Description': 'Test description text',
            'ProductStrain': 'Test Strain',
            'ProductType': 'flower'
        },
        {
            'ProductBrand': 'Very Long Brand Name That Exceeds Limits',
            'Price': '$30.99',
            'Lineage': 'Sativa',
            'Ratio_or_THC_CBD': 'THC: 30% CBD: 1%',
            'Description': 'Another test description',
            'ProductStrain': 'Another Strain',
            'ProductType': 'concentrate'
        },
        {
            'ProductBrand': 'Short Brand',
            'Price': '$15.99',
            'Lineage': 'Indica',
            'Ratio_or_THC_CBD': 'THC: 20% CBD: 3%',
            'Description': 'Short description',
            'ProductStrain': 'Short Strain',
            'ProductType': 'pre-roll'
        }
    ]
    
    try:
        # Create horizontal template processor
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        # Process the records
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"✅ Document generated successfully")
            print(f"   Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check font sizes for brand names
            brand_font_sizes = []
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                # Look for brand-related content
                                if any(brand in run.text for brand in ['Fairwinds Manufacturing', 'Very Long Brand Name', 'Short Brand']):
                                    font_size = run.font.size.pt if run.font.size else 'No size'
                                    brand_font_sizes.append(f"Row {row_idx+1}, Col {col_idx+1}: '{run.text[:30]}...' = {font_size}pt")
            
            if brand_font_sizes:
                print(f"\n📋 Brand font sizes found:")
                for size_info in brand_font_sizes:
                    print(f"   {size_info}")
                
                # Check if any brand names have 14pt font
                has_14pt = any('14' in size_info for size_info in brand_font_sizes)
                if has_14pt:
                    print(f"\n✅ SUCCESS: Found brand names with 14pt font!")
                else:
                    print(f"\n⚠️  WARNING: No brand names found with 14pt font")
            else:
                print(f"\n⚠️  WARNING: No brand content found in the document")
            
            # Save the document for inspection
            output_path = "test_brand_font_size_horizontal_output.docx"
            result_doc.save(output_path)
            print(f"\n💾 Document saved as: {output_path}")
            
        else:
            print(f"❌ Failed to generate document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
    
    print(f"\n🎯 Test completed!")

if __name__ == "__main__":
    test_real_brand_font_size_horizontal() 