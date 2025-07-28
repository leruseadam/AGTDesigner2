#!/usr/bin/env python3
"""
Test lineage and brand display in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def test_lineage_brand_double():
    """Test lineage and brand display in double template."""
    
    print("Lineage/Brand Double Template Test")
    print("=" * 40)
    
    # Create test records with different lineage and brand scenarios
    test_records = [
        {
            'Description': 'Amnesia Lemon Pre-Roll',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductBrand': 'Test Brand 1',
            'ProductType': 'pre-roll',
            'JointRatio': '1.0g x 1 Pack'
        },
        {
            'Description': 'Pineapple Cake Pre-Roll',
            'Price': '$27',
            'Lineage': 'INDICA',
            'Ratio_or_THC_CBD': 'THC: 30% CBD: 1%',
            'ProductBrand': 'Test Brand 2',
            'ProductType': 'pre-roll',
            'JointRatio': '1.5g x 2 Pack'
        },
        {
            'Description': 'Blue Dream Flower',
            'Price': '$45',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 22% CBD: 1%',
            'ProductBrand': 'Premium Brand',
            'ProductType': 'flower',
            'JointRatio': ''
        }
    ]
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Debug: Check what context is being built
        print("Debugging context building...")
        from docx import Document
        temp_doc = Document()
        
        for i, record in enumerate(test_records):
            print(f"\nRecord {i+1}:")
            print(f"  Raw Lineage: '{record.get('Lineage', '')}'")
            print(f"  Raw ProductBrand: '{record.get('ProductBrand', '')}'")
            context = processor._build_label_context(record, temp_doc)
            print(f"  Processed Lineage: '{context.get('Lineage', '')}'")
            print(f"  Processed ProductBrand: '{context.get('ProductBrand', '')}'")
            
            # Check for any unexpected content
            if '•' in context.get('Lineage', ''):
                print(f"  ⚠️  Lineage contains bullet point: '{context.get('Lineage', '')}'")
            if context.get('Lineage', '') != record.get('Lineage', ''):
                print(f"  ⚠️  Lineage was modified from '{record.get('Lineage', '')}' to '{context.get('Lineage', '')}'")
        
        # Process the test records
        print("\nProcessing test records...")
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            print("✓ Document generated successfully")
            
            # Analyze the document
            print("\nAnalyzing document:")
            print("-" * 25)
            
            table = result_doc.tables[0]
            print(f"Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"\nCell ({row_idx},{col_idx}):")
                        print(f"  Full text: '{cell_text}'")
                        
                        # Split by lines to analyze each field
                        lines = cell_text.split('\n')
                        for line_num, line in enumerate(lines):
                            line = line.strip()
                            if line:
                                print(f"    Line {line_num+1}: '{line}'")
                                
                                # Identify what each line likely represents
                                if line.startswith('$'):
                                    print(f"      → Price field")
                                elif line.startswith('•'):
                                    print(f"      → Lineage field (with bullet)")
                                elif any(keyword in line.upper() for keyword in ['THC:', 'CBD:', 'CBC:', 'CBG:', 'CBN:']):
                                    print(f"      → Ratio/THC_CBD field")
                                elif any(keyword in line.lower() for keyword in ['pack', 'x', 'g']):
                                    print(f"      → JointRatio field")
                                elif 'Brand' in line or 'brand' in line:
                                    print(f"      → ProductBrand field")
                                else:
                                    print(f"      → Likely Description field")
        else:
            print("✗ Failed to generate document")
            
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_lineage_brand_double() 