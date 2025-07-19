#!/usr/bin/env python3
"""
Comprehensive test to verify Arial Bold enforcement across all document generation paths
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.tag_generator import generate_multiple_label_tables
from src.core.generation.docx_formatting import enforce_arial_bold_all_text
from docx import Document
from docx.shared import Pt

def test_comprehensive_arial_bold():
    """Test Arial Bold enforcement across all document generation paths"""
    print("🧪 Comprehensive Arial Bold Font Enforcement Test")
    print("=" * 60)
    
    # Create test records
    test_records = [
        {
            'Product Name*': 'Test Product 1',
            'Description': 'Test Description 1',
            'Product Brand': 'Test Brand 1',
            'Price': '$10.00',
            'Lineage': 'HYBRID',
            'Product Type*': 'flower',
            'THC_CBD': 'THC:\nCBD:',
            'Ratio_or_THC_CBD': 'THC:\nCBD:',
            'Ratio': 'THC:\nCBD:',
            'WeightUnits': '1g',
            'Product Strain': 'Test Strain 1',
            'DOH': 'Yes'
        },
        {
            'Product Name*': 'Test Product 2',
            'Description': 'Test Description 2',
            'Product Brand': 'Test Brand 2',
            'Price': '$15.00',
            'Lineage': 'SATIVA',
            'Product Type*': 'concentrate',
            'THC_CBD': 'THC: 25%\nCBD: 2%',
            'Ratio_or_THC_CBD': 'THC: 25%\nCBD: 2%',
            'Ratio': 'THC: 25%\nCBD: 2%',
            'WeightUnits': '0.5g',
            'Product Strain': 'Test Strain 2',
            'DOH': 'No'
        }
    ]
    
    results = []
    
    # Test 1: Template Processor - Vertical Template
    print("\n1️⃣ Testing Template Processor - Vertical Template...")
    try:
        tp_vertical = TemplateProcessor(template_type='vertical', font_scheme='Arial')
        doc_vertical = tp_vertical.process_records(test_records)
        
        if doc_vertical:
            # Check fonts
            non_arial_count = 0
            non_bold_count = 0
            
            for table in doc_vertical.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip():
                                    if run.font.name != "Arial":
                                        non_arial_count += 1
                                    if not run.font.bold:
                                        non_bold_count += 1
            
            if non_arial_count == 0 and non_bold_count == 0:
                print("✅ Vertical Template: All fonts are Arial Bold")
                results.append(("Template Processor - Vertical", True))
            else:
                print(f"❌ Vertical Template: {non_arial_count} non-Arial, {non_bold_count} non-bold")
                results.append(("Template Processor - Vertical", False))
        else:
            print("❌ Vertical Template: No document generated")
            results.append(("Template Processor - Vertical", False))
    except Exception as e:
        print(f"❌ Vertical Template Error: {e}")
        results.append(("Template Processor - Vertical", False))
    
    # Test 2: Template Processor - Horizontal Template
    print("\n2️⃣ Testing Template Processor - Horizontal Template...")
    try:
        tp_horizontal = TemplateProcessor(template_type='horizontal', font_scheme='Arial')
        doc_horizontal = tp_horizontal.process_records(test_records)
        
        if doc_horizontal:
            # Check fonts
            non_arial_count = 0
            non_bold_count = 0
            
            for table in doc_horizontal.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip():
                                    if run.font.name != "Arial":
                                        non_arial_count += 1
                                    if not run.font.bold:
                                        non_bold_count += 1
            
            if non_arial_count == 0 and non_bold_count == 0:
                print("✅ Horizontal Template: All fonts are Arial Bold")
                results.append(("Template Processor - Horizontal", True))
            else:
                print(f"❌ Horizontal Template: {non_arial_count} non-Arial, {non_bold_count} non-bold")
                results.append(("Template Processor - Horizontal", False))
        else:
            print("❌ Horizontal Template: No document generated")
            results.append(("Template Processor - Horizontal", False))
    except Exception as e:
        print(f"❌ Horizontal Template Error: {e}")
        results.append(("Template Processor - Horizontal", False))
    
    # Test 3: Template Processor - Mini Template
    print("\n3️⃣ Testing Template Processor - Mini Template...")
    try:
        tp_mini = TemplateProcessor(template_type='mini', font_scheme='Arial')
        doc_mini = tp_mini.process_records(test_records)
        
        if doc_mini:
            # Check fonts
            non_arial_count = 0
            non_bold_count = 0
            
            for table in doc_mini.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip():
                                    if run.font.name != "Arial":
                                        non_arial_count += 1
                                    if not run.font.bold:
                                        non_bold_count += 1
            
            if non_arial_count == 0 and non_bold_count == 0:
                print("✅ Mini Template: All fonts are Arial Bold")
                results.append(("Template Processor - Mini", True))
            else:
                print(f"❌ Mini Template: {non_arial_count} non-Arial, {non_bold_count} non-bold")
                results.append(("Template Processor - Mini", False))
        else:
            print("❌ Mini Template: No document generated")
            results.append(("Template Processor - Mini", False))
    except Exception as e:
        print(f"❌ Mini Template Error: {e}")
        results.append(("Template Processor - Mini", False))
    
    # Test 4: Tag Generator
    print("\n4️⃣ Testing Tag Generator...")
    try:
        from src.core.generation.tag_generator import get_template_path
        template_path = get_template_path('vertical')
        
        # Generate document using tag generator
        output_buffer = generate_multiple_label_tables(test_records, template_path)
        
        if output_buffer:
            # Load the document from buffer
            output_buffer.seek(0)
            doc_tag_gen = Document(output_buffer)
            
            # Check fonts
            non_arial_count = 0
            non_bold_count = 0
            
            for table in doc_tag_gen.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text.strip():
                                    if run.font.name != "Arial":
                                        non_arial_count += 1
                                    if not run.font.bold:
                                        non_bold_count += 1
            
            if non_arial_count == 0 and non_bold_count == 0:
                print("✅ Tag Generator: All fonts are Arial Bold")
                results.append(("Tag Generator", True))
            else:
                print(f"❌ Tag Generator: {non_arial_count} non-Arial, {non_bold_count} non-bold")
                results.append(("Tag Generator", False))
        else:
            print("❌ Tag Generator: No document generated")
            results.append(("Tag Generator", False))
    except Exception as e:
        print(f"❌ Tag Generator Error: {e}")
        results.append(("Tag Generator", False))
    
    # Test 5: Direct Enforcement Function
    print("\n5️⃣ Testing Direct Enforcement Function...")
    try:
        # Create a document with mixed fonts
        doc_mixed = Document()
        table = doc_mixed.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        
        # Add text with different fonts
        para = cell.paragraphs[0]
        run1 = para.add_run("Arial Text")
        run1.font.name = "Arial"
        run1.font.bold = True
        
        run2 = para.add_run("Times Text")
        run2.font.name = "Times New Roman"
        run2.font.bold = False
        
        # Apply enforcement
        enforce_arial_bold_all_text(doc_mixed)
        
        # Check results
        all_arial = True
        all_bold = True
        
        for table in doc_mixed.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                if run.font.name != "Arial":
                                    all_arial = False
                                if not run.font.bold:
                                    all_bold = False
        
        if all_arial and all_bold:
            print("✅ Direct Enforcement: All fonts converted to Arial Bold")
            results.append(("Direct Enforcement Function", True))
        else:
            print("❌ Direct Enforcement: Some fonts not converted")
            results.append(("Direct Enforcement Function", False))
    except Exception as e:
        print(f"❌ Direct Enforcement Error: {e}")
        results.append(("Direct Enforcement Function", False))
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 TEST RESULTS SUMMARY")
    print("=" * 60)
    
    passed = 0
    total = len(results)
    
    for test_name, success in results:
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{status} {test_name}")
        if success:
            passed += 1
    
    print(f"\n🎯 Overall Result: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 SUCCESS: All document generation paths enforce Arial Bold fonts")
        return True
    else:
        print("❌ FAILED: Some document generation paths do not enforce Arial Bold fonts")
        return False

if __name__ == "__main__":
    success = test_comprehensive_arial_bold()
    sys.exit(0 if success else 1) 