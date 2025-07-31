#!/usr/bin/env python3
"""
Test with real vendor data to see what's happening
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.data.excel_processor import ExcelProcessor
from src.core.generation.template_processor import TemplateProcessor
from pathlib import Path

def test_real_vendor_data():
    """Test with real data to see vendor processing"""
    
    print("=== Testing Real Vendor Data ===")
    
    # Get the Excel processor
    excel_processor = ExcelProcessor()
    
    # Try to load the uploaded file
    uploads_dir = Path("uploads")
    if uploads_dir.exists():
        excel_files = list(uploads_dir.glob("*.xlsx"))
        if excel_files:
            file_path = str(excel_files[0])
            print(f"Loading file: {file_path}")
            
            if excel_processor.load_file(file_path):
                print(f"✅ File loaded successfully: {len(excel_processor.df)} records")
            else:
                print("❌ Failed to load file")
                return
        else:
            print("❌ No Excel files found in uploads directory")
            return
    else:
        print("❌ Uploads directory not found")
        return
    
    # Check vendor column
    vendor_columns = [col for col in excel_processor.df.columns if 'vendor' in col.lower()]
    print(f"Vendor columns found: {vendor_columns}")
    
    # Check first few records for vendor data
    print("\n=== First 5 Records Vendor Data ===")
    for i in range(min(5, len(excel_processor.df))):
        record = excel_processor.df.iloc[i]
        vendor_value = record.get('Vendor', '') or record.get('Vendor/Supplier*', '')
        product_name = record.get('Product Name*', 'Unknown')
        print(f"Record {i+1}: '{product_name}' -> Vendor: '{vendor_value}'")
    
    # Check if any records have vendor data
    has_vendor_data = False
    for col in vendor_columns:
        if col in excel_processor.df.columns:
            non_empty_vendors = excel_processor.df[col].dropna()
            if len(non_empty_vendors) > 0:
                has_vendor_data = True
                print(f"\n✅ Found {len(non_empty_vendors)} records with vendor data in column '{col}'")
                print(f"Sample vendors: {non_empty_vendors.head(3).tolist()}")
                break
    
    if not has_vendor_data:
        print("\n❌ No vendor data found in any vendor columns")
        # Show all column names to help debug
        print(f"All columns: {list(excel_processor.df.columns)}")
        return
    
    # Test processing with real data
    print("\n=== Testing Template Processing with Real Data ===")
    
    # Get first few records with vendor data
    test_records = []
    for i in range(min(3, len(excel_processor.df))):
        record = excel_processor.df.iloc[i].to_dict()
        vendor_value = record.get('Vendor', '') or record.get('Vendor/Supplier*', '')
        if vendor_value and str(vendor_value).lower() not in ['nan', '']:
            test_records.append(record)
            print(f"Adding record {i+1}: '{record.get('Product Name*', 'Unknown')}' with vendor '{vendor_value}'")
    
    if not test_records:
        print("❌ No records with vendor data found for testing")
        return
    
    # Test with horizontal template
    processor = TemplateProcessor('horizontal', {}, 1.0)
    
    # Process the records
    result = processor.process_records(test_records)
    
    if result:
        print("✅ Processing completed successfully")
        
        # Save the result
        output_path = "test_real_vendor_output.docx"
        result.save(output_path)
        print(f"✅ Output saved to {output_path}")
        
        # Examine the output
        from docx import Document
        doc = Document(output_path)
        
        print("\n=== Output Document Analysis ===")
        
        # Check all text in the document
        all_text = ""
        for para in doc.paragraphs:
            all_text += para.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        # Check for vendor data
        vendor_found = False
        for record in test_records:
            vendor_value = record.get('Vendor', '') or record.get('Vendor/Supplier*', '')
            if vendor_value and vendor_value in all_text:
                print(f"✅ Found vendor '{vendor_value}' in output")
                vendor_found = True
        
        if not vendor_found:
            print("❌ No vendor data found in output")
            
            # Check what IS in the output
            print(f"\nOutput text preview: '{all_text[:200]}...'")
            
            # Check for any vendor-related text
            if 'vendor' in all_text.lower():
                print("⚠️  Found 'vendor' in output (case insensitive)")
            if 'PRODUCTVENDOR' in all_text:
                print("⚠️  Found 'PRODUCTVENDOR' markers in output")
            if '{{Label' in all_text:
                print("⚠️  Found template placeholders in output")
    else:
        print("❌ Processing failed - no result returned")

if __name__ == "__main__":
    test_real_vendor_data() 