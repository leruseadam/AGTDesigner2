#!/usr/bin/env python3
"""
Test script to verify that ratio line breaks don't add extra breaks when text runs out.
"""

import sys
import os
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

from src.core.generation.template_processor import TemplateProcessor

def test_ratio_line_break_fix():
    """Test that ratio formatting doesn't add extra line breaks when text runs out."""
    
    print("=== Ratio Line Break Fix Test ===\n")
    
    # Create a test document
    doc = Document()
    paragraph = doc.add_paragraph()
    
    # Test case 1: Text with |BR| at the end (should not add extra line break)
    test_text = "THC:|BR|CBD:"
    print(f"Test 1: '{test_text}'")
    
    # Add the text to the paragraph
    run = paragraph.add_run(test_text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(12)
    
    # Create a template processor to test the fix
    processor = TemplateProcessor('vertical', {})
    
    # Apply the line break conversion
    processor._convert_br_markers_to_line_breaks(paragraph)
    
    # Check the result
    result_text = paragraph.text
    print(f"  Result: '{result_text}'")
    
    # Count the number of line breaks (should be 1, not 2)
    line_count = len(result_text.split('\n'))
    print(f"  Line count: {line_count}")
    
    if line_count == 2:
        print("  ✅ PASS: Correct number of lines (no extra break at end)")
    else:
        print(f"  ❌ FAIL: Expected 2 lines, got {line_count}")
    
    print()
    
    # Test case 2: Text with |BR| in the middle (should add line break)
    doc2 = Document()
    paragraph2 = doc2.add_paragraph()
    
    test_text2 = "THC: 25%|BR|CBD: 15%"
    print(f"Test 2: '{test_text2}'")
    
    # Add the text to the paragraph
    run2 = paragraph2.add_run(test_text2)
    run2.font.name = "Arial"
    run2.font.bold = True
    run2.font.size = Pt(12)
    
    # Apply the line break conversion
    processor._convert_br_markers_to_line_breaks(paragraph2)
    
    # Check the result
    result_text2 = paragraph2.text
    print(f"  Result: '{result_text2}'")
    
    # Count the number of line breaks (should be 2)
    line_count2 = len(result_text2.split('\n'))
    print(f"  Line count: {line_count2}")
    
    if line_count2 == 2:
        print("  ✅ PASS: Correct number of lines (line break in middle)")
    else:
        print(f"  ❌ FAIL: Expected 2 lines, got {line_count2}")
    
    print()
    
    # Test case 3: Text with multiple |BR| markers
    doc3 = Document()
    paragraph3 = doc3.add_paragraph()
    
    test_text3 = "THC: 25%|BR|CBD: 15%|BR|CBG: 5%"
    print(f"Test 3: '{test_text3}'")
    
    # Add the text to the paragraph
    run3 = paragraph3.add_run(test_text3)
    run3.font.name = "Arial"
    run3.font.bold = True
    run3.font.size = Pt(12)
    
    # Apply the line break conversion
    processor._convert_br_markers_to_line_breaks(paragraph3)
    
    # Check the result
    result_text3 = paragraph3.text
    print(f"  Result: '{result_text3}'")
    
    # Count the number of line breaks (should be 3)
    line_count3 = len(result_text3.split('\n'))
    print(f"  Line count: {line_count3}")
    
    if line_count3 == 3:
        print("  ✅ PASS: Correct number of lines (multiple line breaks)")
    else:
        print(f"  ❌ FAIL: Expected 3 lines, got {line_count3}")
    
    print()
    
    # Test case 4: Text with |BR| at the end and empty content
    doc4 = Document()
    paragraph4 = doc4.add_paragraph()
    
    test_text4 = "THC: 25%|BR|"
    print(f"Test 4: '{test_text4}'")
    
    # Add the text to the paragraph
    run4 = paragraph4.add_run(test_text4)
    run4.font.name = "Arial"
    run4.font.bold = True
    run4.font.size = Pt(12)
    
    # Apply the line break conversion
    processor._convert_br_markers_to_line_breaks(paragraph4)
    
    # Check the result
    result_text4 = paragraph4.text
    print(f"  Result: '{result_text4}'")
    
    # Count the number of line breaks (should be 1, not 2)
    line_count4 = len(result_text4.split('\n'))
    print(f"  Line count: {line_count4}")
    
    if line_count4 == 1:
        print("  ✅ PASS: Correct number of lines (no extra break with empty end)")
    else:
        print(f"  ❌ FAIL: Expected 1 line, got {line_count4}")
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    test_ratio_line_break_fix() 