#!/usr/bin/env python3
"""
Check what markers are in the double template.
"""

from docx import Document
import re

def check_template_markers():
    """Check what markers are in the double template."""
    print("=== DOUBLE TEMPLATE MARKERS ===")
    
    # Load the double template
    doc = Document('src/core/generation/templates/double.docx')
    
    # Extract all text
    all_text = ""
    for paragraph in doc.paragraphs:
        all_text += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_text += paragraph.text + "\n"
    
    print("All text in template:")
    print(all_text)
    print()
    
    # Look for markers
    marker_pattern = r'[A-Z_]+_START.*?[A-Z_]+_END'
    markers = re.findall(marker_pattern, all_text, re.DOTALL)
    
    print("Found markers:")
    for marker in markers:
        print(f"  {marker}")
    
    # Also look for any text that might be placeholders
    placeholder_pattern = r'\{[^}]+\}'
    placeholders = re.findall(placeholder_pattern, all_text)
    
    print("\nFound placeholders:")
    for placeholder in placeholders:
        print(f"  {placeholder}")

if __name__ == "__main__":
    check_template_markers() 