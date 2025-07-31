#!/usr/bin/env python3
"""
Debug script to understand DOH image insertion issues.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from io import BytesIO

def debug_doh_image_insertion():
    """Debug DOH image insertion process."""
    print("Debugging DOH Image Insertion")
    print("=" * 50)
    
    # Create a test record with DOH image
    test_record = {
        'Description': 'Test Product with DOH',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'classic'  # This should use regular DOH image
    }
    
    # Test double template specifically
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return
    
    print(f"Document has {len(result_doc.tables)} tables")
    
    # Examine each table and cell
    for table_idx, table in enumerate(result_doc.tables):
        print(f"\nTable {table_idx + 1}:")
        print(f"  Rows: {len(table.rows)}")
        print(f"  Columns: {len(table.columns)}")
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                print(f"    Cell ({row_idx}, {col_idx}):")
                print(f"      Paragraphs: {len(cell.paragraphs)}")
                
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    print(f"        Paragraph {para_idx + 1}:")
                    print(f"          Text: '{paragraph.text}'")
                    print(f"          Runs: {len(paragraph.runs)}")
                    
                    for run_idx, run in enumerate(paragraph.runs):
                        print(f"            Run {run_idx + 1}:")
                        print(f"              Text: '{run.text}'")
                        print(f"              Has element: {hasattr(run, '_element')}")
                        
                        if hasattr(run, '_element'):
                            # Check for drawing elements
                            drawing_elements = run._element.findall(qn('w:drawing'))
                            print(f"              Drawing elements: {len(drawing_elements)}")
                            
                            # Check for other image-related elements
                            blip_elements = run._element.findall('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                            print(f"              Blip elements: {len(blip_elements)}")
                            
                            # Check for any image-related content
                            if drawing_elements or blip_elements:
                                print(f"              ✓ FOUND IMAGE ELEMENT!")
                                
                                # Check alignment
                                if paragraph.alignment == WD_TABLE_ALIGNMENT.CENTER:
                                    print(f"              ✓ Paragraph is centered")
                                else:
                                    print(f"              ✗ Paragraph is NOT centered (alignment: {paragraph.alignment})")
                                
                                # Check cell vertical alignment
                                if cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                                    print(f"              ✓ Cell is vertically centered")
                                else:
                                    print(f"              ✗ Cell is NOT vertically centered (alignment: {cell.vertical_alignment})")

if __name__ == "__main__":
    debug_doh_image_insertion() 