#!/usr/bin/env python3
"""
Debug script to see what the expanded double template contains
"""

import sys
import os
from pathlib import Path

# Add the src directory to the Python path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def debug_expanded_template():
    """Debug what the expanded double template contains."""
    print("=== Debugging Expanded Double Template ===")
    
    try:
        # Create processor for double template
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Get the expanded template buffer
        buffer = processor._expanded_template_buffer
        buffer.seek(0)
        doc = Document(buffer)
        
        print("\n=== Expanded Template Content ===")
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"  Cell ({row_idx},{col_idx}): {repr(cell_text)}")
        
        # Check for all placeholders
        print("\n=== All Placeholders Found ===")
        all_placeholders = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        # Find all placeholders
                        import re
                        placeholders = re.findall(r'\{\{Label\d+\.[^}]+\}\}', text)
                        all_placeholders.update(placeholders)
        
        for placeholder in sorted(all_placeholders):
            print(f"  {placeholder}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error during debug: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    debug_expanded_template() 