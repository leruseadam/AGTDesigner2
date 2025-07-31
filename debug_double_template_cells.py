#!/usr/bin/env python3
"""
Debug script to identify and fix double template table cell corruption issues.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def diagnose_double_template():
    """Diagnose double template issues."""
    print("Double Template Cell Corruption Diagnostic")
    print("=" * 50)
    
    # Create test records
    test_records = [
        {
            'ProductName': 'HUSTLER\'S AMBITION Cheesecake - 14g',
            'Description': 'HUSTLER\'S AMBITION Cheesecake - 14g',
            'ProductBrand': 'HUSTLER\'S AMBITION',
            'Price': '$100',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 24.95% CBD: 0.05%',
            'ProductStrain': 'Cheesecake',
            'DOH': 'DOH'
        },
        {
            'ProductName': 'HUSTLER\'S AMBITION Churros - 14g',
            'Description': 'HUSTLER\'S AMBITION Churros - 14g',
            'ProductBrand': 'HUSTLER\'S AMBITION',
            'Price': '$100',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 19.0% CBD: 0.03%',
            'ProductStrain': 'Churros',
            'DOH': 'DOH'
        },
        {
            'ProductName': 'HUSTLER\'S AMBITION Jelly Donuts - 14g',
            'Description': 'HUSTLER\'S AMBITION Jelly Donuts - 14g',
            'ProductBrand': 'HUSTLER\'S AMBITION',
            'Price': '$100',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 18.57% CBD: 0.05%',
            'ProductStrain': 'Jelly Donuts',
            'DOH': 'DOH'
        },
        {
            'ProductName': 'HUSTLER\'S AMBITION Gelato - 14g',
            'Description': 'HUSTLER\'S AMBITION Gelato - 14g',
            'ProductBrand': 'HUSTLER\'S AMBITION',
            'Price': '$100',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 25.0% CBD: 0%',
            'ProductStrain': 'Gelato',
            'DOH': 'DOH'
        }
    ]
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', {}, 1.0)
        
        print("1. Testing template expansion...")
        
        # Test template expansion
        expanded_buffer = processor._expand_template_if_needed(force_expand=True)
        if expanded_buffer:
            print("✓ Template expansion successful")
            
            # Load the expanded template
            expanded_doc = Document(expanded_buffer)
            print(f"✓ Expanded template loaded: {len(expanded_doc.tables)} tables")
            
            # Check table structure
            for i, table in enumerate(expanded_doc.tables):
                print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check each cell
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        print(f"    Cell ({row_idx},{col_idx}): '{cell_text[:50]}...'")
                        
                        # Check for XML corruption
                        try:
                            cell_xml = cell._tc.xml
                            if 'Label' in cell_xml:
                                print(f"      ✓ Contains Label placeholders")
                            else:
                                print(f"      ⚠ Missing Label placeholders")
                        except Exception as xml_error:
                            print(f"      ❌ XML corruption detected: {xml_error}")
        else:
            print("❌ Template expansion failed")
            return
        
        print("\n2. Testing document generation...")
        
        # Test full document generation
        result_doc = processor.process_records(test_records)
        
        if result_doc:
            print("✓ Document generation successful")
            
            # Save test document
            test_filename = "debug_double_template_test.docx"
            result_doc.save(test_filename)
            print(f"✓ Test document saved as: {test_filename}")
            
            # Analyze the generated document
            print("\n3. Analyzing generated document...")
            
            for i, table in enumerate(result_doc.tables):
                print(f"  Table {i}: {len(table.rows)} rows x {len(table.columns)} columns")
                
                # Check each cell for content and formatting
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        print(f"    Cell ({row_idx},{col_idx}):")
                        print(f"      Content: '{cell_text[:100]}...'")
                        
                        # Check cell formatting
                        try:
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcW = tcPr.find(qn('w:tcW'))
                            if tcW is not None:
                                width = tcW.get(qn('w:w'))
                                print(f"      Width: {width} twips")
                            else:
                                print(f"      Width: Not set")
                        except Exception as format_error:
                            print(f"      ❌ Formatting error: {format_error}")
                        
                        # Check for corruption indicators
                        if not cell_text:
                            print(f"      ⚠ Empty cell")
                        elif 'Label' in cell_text:
                            print(f"      ⚠ Unprocessed placeholder")
                        elif len(cell_text) < 10:
                            print(f"      ⚠ Very short content")
        else:
            print("❌ Document generation failed")
            
    except Exception as e:
        print(f"❌ Error during diagnosis: {e}")
        import traceback
        traceback.print_exc()

def fix_double_template_cells():
    """Fix double template cell issues."""
    print("\nDouble Template Cell Fix")
    print("=" * 30)
    
    # The issue is likely in the _expand_template_to_4x3_fixed_double method
    # Let's create a safer version
    
    def create_safe_double_template():
        """Create a safe double template without XML corruption."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from io import BytesIO
        
        num_cols, num_rows = 4, 3
        col_width_twips = str(int(1.75 * 1440))  # 1.75 inches per column
        row_height_pts = Pt(2.5 * 72)  # 2.5 inches per row
        
        # Create a new document
        doc = Document()
        
        # Remove default paragraph
        if doc.paragraphs:
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)
        
        # Create table
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set up table properties safely
        tblPr = tbl._element.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
        
        # Set layout to fixed
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
        tbl._element.insert(0, tblPr)
        
        # Set up grid
        grid = OxmlElement('w:tblGrid')
        for i in range(num_cols):
            gc = OxmlElement('w:gridCol')
            gc.set(qn('w:w'), col_width_twips)
            grid.append(gc)
        tbl._element.insert(0, grid)
        
        # Set row heights
        for row in tbl.rows:
            row.height = row_height_pts
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        # Set borders
        borders = OxmlElement('w:tblBorders')
        for side in ('insideH', 'insideV'):
            b = OxmlElement(f"w:{side}")
            b.set(qn('w:val'), "single")
            b.set(qn('w:sz'), "4")
            b.set(qn('w:color'), "D3D3D3")
            b.set(qn('w:space'), "0")
            borders.append(b)
        tblPr.append(borders)
        
        # Fill cells with placeholders safely
        cnt = 1
        for r in range(num_rows):
            for c in range(num_cols):
                cell = tbl.cell(r, c)
                
                # Clear cell content
                cell._tc.clear_content()
                
                # Create paragraph with placeholders
                p = cell.paragraphs[0]
                
                # Add placeholders in order
                placeholders = [
                    f'{{{{Label{cnt}.ProductBrand}}}}',
                    f'{{{{Label{cnt}.DescAndWeight}}}}',
                    f'{{{{Label{cnt}.Price}}}}',
                    f'{{{{Label{cnt}.Ratio_or_THC_CBD}}}}',
                    f'{{{{Label{cnt}.DOH}}}}',
                    f'{{{{Label{cnt}.Lineage}}}}',
                    f'{{{{Label{cnt}.ProductStrain}}}}'
                ]
                
                # Add placeholders with line breaks
                for i, placeholder in enumerate(placeholders):
                    if i > 0:
                        p.add_run('\n')
                    p.add_run(placeholder)
                
                cnt += 1
        
        # Save to buffer
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    
    try:
        # Create safe template
        safe_buffer = create_safe_double_template()
        
        # Test the safe template
        safe_doc = Document(safe_buffer)
        print(f"✓ Safe template created: {len(safe_doc.tables)} tables")
        
        # Save safe template for inspection
        safe_filename = "safe_double_template.docx"
        safe_doc.save(safe_filename)
        print(f"✓ Safe template saved as: {safe_filename}")
        
        # Test with actual data
        test_records = [
            {
                'ProductName': 'Test Product',
                'Description': 'Test Description',
                'ProductBrand': 'Test Brand',
                'Price': '$10.00',
                'Lineage': 'HYBRID',
                'Ratio_or_THC_CBD': 'THC: 20% CBD: 2%',
                'ProductStrain': 'Test Strain',
                'DOH': 'DOH'
            }
        ]
        
        # Create processor with safe template
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Override the template buffer with our safe version
        processor._expanded_template_buffer = safe_buffer
        
        # Process records
        result_doc = processor.process_records(test_records)
        
        if result_doc:
            result_filename = "fixed_double_template_result.docx"
            result_doc.save(result_filename)
            print(f"✓ Fixed template result saved as: {result_filename}")
            print("✓ Double template cell fix completed successfully!")
        else:
            print("❌ Fixed template processing failed")
            
    except Exception as e:
        print(f"❌ Error during fix: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    print("Double Template Cell Corruption Diagnostic and Fix")
    print("=" * 60)
    
    # Run diagnosis
    diagnose_double_template()
    
    # Run fix
    fix_double_template_cells()
    
    print("\nDiagnostic and fix completed!")
    print("Check the generated files for results.") 