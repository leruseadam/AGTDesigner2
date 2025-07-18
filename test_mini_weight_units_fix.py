#!/usr/bin/env python3
"""
Test script to verify that mini templates now use weight units instead of THC/CBD for classic types.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.template_processor import get_font_scheme
from src.core.formatting.markers import wrap_with_marker, unwrap_marker
from docx import Document
from docx.shared import Pt

def test_mini_weight_units_for_classic_types():
    """Test that mini templates use weight units for classic types."""
    
    print("Testing Mini Template Weight Units for Classic Types")
    print("=" * 60)
    
    # Test data for classic types
    classic_test_cases = [
        {
            'name': 'Flower Product',
            'ProductType': 'flower',
            'WeightUnits': '3.5g',
            'Ratio': 'THC: 25%\nCBD: 2%',
            'Description': 'Test Flower Strain',
            'ProductBrand': 'Test Brand',
            'Price': '$45',
            'Lineage': 'HYBRID'
        },
        {
            'name': 'Vape Cartridge',
            'ProductType': 'vape cartridge',
            'WeightUnits': '1g',
            'Ratio': 'THC: 85%\nCBD: 1%',
            'Description': 'Test Vape Cartridge',
            'ProductBrand': 'Test Brand',
            'Price': '$35',
            'Lineage': 'SATIVA'
        },
        {
            'name': 'Concentrate',
            'ProductType': 'concentrate',
            'WeightUnits': '0.5g',
            'Ratio': 'THC: 90%\nCBD: 0.5%',
            'Description': 'Test Concentrate',
            'ProductBrand': 'Test Brand',
            'Price': '$40',
            'Lineage': 'INDICA'
        }
    ]
    
    # Test data for non-classic types
    non_classic_test_cases = [
        {
            'name': 'Edible Product',
            'ProductType': 'edible (solid)',
            'WeightUnits': '100mg',
            'Ratio': 'THC: 10mg\nCBD: 5mg',
            'Description': 'Test Edible',
            'ProductBrand': 'Test Brand',
            'Price': '$15',
            'Lineage': 'MIXED'
        }
    ]
    
    # Test mini template processor
    font_scheme = get_font_scheme('mini')
    processor = TemplateProcessor('mini', font_scheme, 1.0)
    
    print("\n1. Testing Classic Types (should use WeightUnits):")
    print("-" * 50)
    
    for case in classic_test_cases:
        print(f"\n  {case['name']}:")
        print(f"    Product Type: {case['ProductType']}")
        print(f"    Weight Units: {case['WeightUnits']}")
        print(f"    Original Ratio: {case['Ratio']}")
        
        # Build label context
        label_context = processor._build_label_context(case, Document())
        
        # Check what was set for Ratio_or_THC_CBD
        ratio_field = label_context.get('Ratio_or_THC_CBD', '')
        print(f"    Final Ratio_or_THC_CBD: {repr(ratio_field)}")
        
        # Check if it contains RATIO marker with weight units content
        if 'RATIO_START' in ratio_field and 'RATIO_END' in ratio_field:
            # Extract the actual content
            start_marker = 'RATIO_START'
            end_marker = 'RATIO_END'
            start_idx = ratio_field.find(start_marker) + len(start_marker)
            end_idx = ratio_field.find(end_marker)
            content = ratio_field[start_idx:end_idx].strip()
            print(f"    Content: {repr(content)}")
            if content == case['WeightUnits']:
                print(f"    ✓ Correctly uses RATIO marker with weight units content")
                print(f"    ✓ Content matches WeightUnits: {case['WeightUnits']}")
            else:
                print(f"    ✗ Content mismatch! Expected: {case['WeightUnits']}, Got: {content}")
        else:
            print(f"    ✗ Does NOT use RATIO marker")
    
    print("\n2. Testing Non-Classic Types (should use original logic):")
    print("-" * 50)
    
    for case in non_classic_test_cases:
        print(f"\n  {case['name']}:")
        print(f"    Product Type: {case['ProductType']}")
        print(f"    Weight Units: {case['WeightUnits']}")
        print(f"    Original Ratio: {case['Ratio']}")
        
        # Build label context
        label_context = processor._build_label_context(case, Document())
        
        # Check what was set for Ratio_or_THC_CBD
        ratio_field = label_context.get('Ratio_or_THC_CBD', '')
        print(f"    Final Ratio_or_THC_CBD: {repr(ratio_field)}")
        
        # Check if it contains RATIO marker (not WEIGHTUNITS)
        if 'RATIO_START' in ratio_field and 'RATIO_END' in ratio_field:
            print(f"    ✓ Correctly uses RATIO marker for non-classic type")
        elif 'WEIGHTUNITS_START' in ratio_field and 'WEIGHTUNITS_END' in ratio_field:
            print(f"    ✗ Incorrectly uses WEIGHTUNITS marker for non-classic type")
        else:
            print(f"    ? Uses different marker or no marker")
    
    print("\n3. Testing Font Sizing for WEIGHTUNITS marker:")
    print("-" * 50)
    
    # Test font sizing for weight units
    test_weights = ['1g', '3.5g', '0.5g', '100mg', '500mg']
    
    for weight in test_weights:
        font_size = processor._get_template_specific_font_size(weight, 'WEIGHTUNITS')
        print(f"    Weight: {weight} -> Font Size: {font_size.pt}pt")
    
    print("\n4. Testing Template Processing:")
    print("-" * 50)
    
    # Test with a classic type
    test_record = classic_test_cases[0]
    print(f"    Processing: {test_record['name']}")
    
    try:
        # Process the record
        doc = Document()
        processed_doc = processor.process_records([test_record])
        
        if processed_doc:
            print(f"    ✓ Successfully processed record")
            
            # Check if the document contains the expected content
            full_text = ""
            for paragraph in processed_doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            if test_record['WeightUnits'] in full_text:
                print(f"    ✓ Weight units found in document: {test_record['WeightUnits']}")
            else:
                print(f"    ✗ Weight units NOT found in document")
                print(f"    Document content: {full_text[:200]}...")
        else:
            print(f"    ✗ Failed to process record")
            
    except Exception as e:
        print(f"    ✗ Error processing record: {e}")
    
    print("\n" + "=" * 60)
    print("Test completed!")

if __name__ == "__main__":
    test_mini_weight_units_for_classic_types() 