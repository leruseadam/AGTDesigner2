#!/usr/bin/env python3
"""
Test script to verify double template placeholder replacement fix
"""

import sys
import os
from pathlib import Path

# Add the src directory to the Python path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def test_double_template_placeholder_replacement():
    """Test that double template placeholders are replaced with actual data."""
    print("Testing double template placeholder replacement...")
    
    # Create test records
    test_records = [
        {
            'ProductName': 'Test Product 1',
            'ProductBrand': 'Test Brand 1',
            'Price': '$25.99',
            'Description': 'Test Description 1',
            'WeightUnits': '3.5g',
            'Lineage': 'Test Lineage 1',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'DOH': 'YES',
            'ProductStrain': 'Test Strain 1',
            'ProductType': 'flower'
        },
        {
            'ProductName': 'Test Product 2',
            'ProductBrand': 'Test Brand 2',
            'Price': '$30.99',
            'Description': 'Test Description 2',
            'WeightUnits': '1g',
            'Lineage': 'Test Lineage 2',
            'Ratio_or_THC_CBD': 'THC: 30% CBD: 1%',
            'DOH': 'YES',
            'ProductStrain': 'Test Strain 2',
            'ProductType': 'concentrate'
        }
    ]
    
    try:
        # Create processor for double template
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Process the records
        result = processor.process_records(test_records)
        
        if result is None:
            print("❌ Failed to process records")
            return False
        
        # Save the result to a temporary file
        output_path = "test_double_template_output.docx"
        result.save(output_path)
        
        # Load the saved document and check for placeholders
        doc = Document(output_path)
        
        # Check if any placeholders remain in the document
        placeholder_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if '{{Label' in text and '}}' in text:
                            print(f"❌ Found placeholder in cell: {text}")
                            placeholder_found = True
        
        if placeholder_found:
            print("❌ Placeholders found in output document")
            return False
        
        # Check if actual data is present (double template only shows Lineage, ProductVendor, ProductStrain)
        data_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if 'Test Lineage' in text or 'Test Vendor' in text or 'Test Strain' in text:
                            data_found = True
                            print(f"✅ Found actual data: {text}")
        
        if not data_found:
            print("❌ No actual data found in output document")
            return False
        
        print("✅ Double template placeholder replacement test completed successfully!")
        print(f"✅ Output saved to: {output_path}")
        
        # Clean up
        if os.path.exists(output_path):
            os.remove(output_path)
        
        return True
        
    except Exception as e:
        print(f"❌ Error during test: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_horizontal_template_comparison():
    """Test that horizontal template works correctly for comparison."""
    print("\nTesting horizontal template for comparison...")
    
    # Create test records
    test_records = [
        {
            'ProductName': 'Test Product 1',
            'ProductBrand': 'Test Brand 1',
            'Price': '$25.99',
            'Description': 'Test Description 1',
            'WeightUnits': '3.5g',
            'Lineage': 'Test Lineage 1',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'DOH': 'YES',
            'ProductStrain': 'Test Strain 1',
            'ProductType': 'flower'
        }
    ]
    
    try:
        # Create processor for horizontal template
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        # Process the records
        result = processor.process_records(test_records)
        
        if result is None:
            print("❌ Failed to process horizontal template records")
            return False
        
        # Save the result to a temporary file
        output_path = "test_horizontal_template_output.docx"
        result.save(output_path)
        
        # Load the saved document and check for placeholders
        doc = Document(output_path)
        
        # Check if any placeholders remain in the document
        placeholder_found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if '{{Label' in text and '}}' in text:
                            print(f"❌ Found placeholder in horizontal template: {text}")
                            placeholder_found = True
        
        if placeholder_found:
            print("❌ Placeholders found in horizontal template output")
            return False
        
        print("✅ Horizontal template placeholder replacement test completed successfully!")
        
        # Clean up
        if os.path.exists(output_path):
            os.remove(output_path)
        
        return True
        
    except Exception as e:
        print(f"❌ Error during horizontal template test: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("=== Double Template Placeholder Replacement Test ===")
    
    # Test horizontal template first
    horizontal_success = test_horizontal_template_comparison()
    
    # Test double template
    double_success = test_double_template_placeholder_replacement()
    
    if horizontal_success and double_success:
        print("\n🎉 All tests passed! Double template now uses the same replacement logic as horizontal template.")
    else:
        print("\n❌ Some tests failed. Please check the output above.")
        sys.exit(1) 