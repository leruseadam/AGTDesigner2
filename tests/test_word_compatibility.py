#!/usr/bin/env python3
"""
Test script to check if row height changes occur when document is opened in Word
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from io import BytesIO

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor

def check_row_heights(doc, stage_name):
    """Check and print row height rules for all tables in the document"""
    print(f"\n=== {stage_name} ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            height = row.height
            height_rule = row.height_rule
            print(f"  Row {j}: height={height}, rule={height_rule}")

def test_word_compatibility():
    """Test if the issue occurs when document is opened in Word"""
    
    # Create a test document using the actual application flow
    processor = TemplateProcessor('horizontal', {}, 1.0)
    
    # Create context
    context = {}
    for i in range(9):
        context[f'Label{i+1}'] = {
            'ProductBrand': f'Test Brand {i+1}',
            'Price': f'${i+1}.99',
            'Description': f'Test Description {i+1}',
            'Lineage': f'Test Lineage {i+1}',
            'Ratio_or_THC_CBD': f'THC: {i+1}%',
            'DOH': 'YES',
            'DescAndWeight': f'Test Description {i+1} - 3.5g'
        }
    
    # Process the document
    doc = processor._process_chunk([context[f'Label{i+1}'] for i in range(9)])
    
    if doc is None:
        print("Error: Failed to process document")
        return
    
    # Check initial state
    check_row_heights(doc, "1. After processing")
    
    # Save the document
    output_path = "test_word_compatibility.docx"
    doc.save(output_path)
    
    # Check state after saving
    final_doc = Document(output_path)
    check_row_heights(final_doc, "2. After saving")
    
    # Check XML
    print("\n=== XML Analysis ===")
    for i, table in enumerate(final_doc.tables):
        print(f"Table {i} XML:")
        table_xml = table._element.xml
        if 'trHeight' in table_xml:
            print("  Found trHeight elements in XML")
            import re
            tr_height_matches = re.findall(r'<w:trHeight[^>]*>', table_xml)
            print(f"  Total trHeight elements: {len(tr_height_matches)}")
            
            # Count elements with and without hRule
            with_hrule = 0
            without_hrule = 0
            for idx, match in enumerate(tr_height_matches):
                if 'w:hRule="exact"' in match:
                    with_hrule += 1
                    print(f"    {idx}: {match} - HAS hRule")
                else:
                    without_hrule += 1
                    print(f"    {idx}: {match} - NO hRule")
            
            print(f"  Summary: {with_hrule} with hRule, {without_hrule} without hRule")
        else:
            print("  No trHeight elements found in XML")
    
    print(f"\nDocument saved as: {output_path}")
    print("Please open this document in Microsoft Word, make a small change (like adding a space),")
    print("save it, and then run the check_word_after_edit.py script to see if row heights changed.")

if __name__ == "__main__":
    test_word_compatibility() 