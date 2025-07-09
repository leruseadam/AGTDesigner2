#!/usr/bin/env python3
"""
Debug script to investigate what's changing table row heights from EXACTLY to AT_LEAST
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from docxtpl import DocxTemplate
from io import BytesIO

def check_row_heights(doc, stage_name):
    """Check and print row height rules for all tables in the document"""
    print(f"\n=== {stage_name} ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            height = row.height
            height_rule = row.height_rule
            print(f"  Row {j}: height={height}, rule={height_rule}")

def create_test_template():
    """Create a simple test template with EXACTLY row heights"""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    
    # Set row heights to EXACTLY
    for row in table.rows:
        row.height = Inches(2.25)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Add some template variables
    for i in range(3):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = f"{{{{Label{i*3+j+1}.ProductBrand}}}}"
    
    # Save template
    template_path = "test_template.docx"
    doc.save(template_path)
    return template_path

def test_docxtpl_rendering():
    """Test what happens to row heights during docxtpl rendering"""
    
    # Create test template
    template_path = create_test_template()
    
    try:
        # Load template and check initial state
        doc = Document(template_path)
        check_row_heights(doc, "After creating template")
        
        # Create DocxTemplate and check state
        template = DocxTemplate(template_path)
        template.init_docx()
        check_row_heights(template.docx, "After DocxTemplate init")
        
        # Create context
        context = {}
        for i in range(9):
            context[f'Label{i+1}'] = {'ProductBrand': f'Test Brand {i+1}'}
        
        # Render template
        template.render(context)
        check_row_heights(template.docx, "After docxtpl render")
        
        # Save rendered document
        output_path = "test_output.docx"
        template.save(output_path)
        
        # Check final state
        final_doc = Document(output_path)
        check_row_heights(final_doc, "After saving rendered document")
        
        # Also check the XML directly
        print("\n=== XML Analysis ===")
        for i, table in enumerate(final_doc.tables):
            print(f"Table {i} XML:")
            table_xml = table._element.xml
            # Look for trHeight elements
            if 'trHeight' in table_xml:
                print("  Found trHeight elements in XML")
                # Extract trHeight elements
                import re
                tr_height_matches = re.findall(r'<w:trHeight[^>]*>', table_xml)
                for match in tr_height_matches:
                    print(f"    {match}")
            else:
                print("  No trHeight elements found in XML")
        
    finally:
        # Clean up
        if os.path.exists(template_path):
            os.remove(template_path)
        if os.path.exists(output_path):
            os.remove(output_path)

if __name__ == "__main__":
    test_docxtpl_rendering() 