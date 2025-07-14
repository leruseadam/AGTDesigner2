#!/usr/bin/env python3
"""
Debug script to isolate the issue with enforce_fixed_cell_dimensions function
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from io import BytesIO

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.docx_formatting import enforce_fixed_cell_dimensions

def check_row_heights(doc, stage_name):
    """Check and print row height rules for all tables in the document"""
    print(f"\n=== {stage_name} ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            height = row.height
            height_rule = row.height_rule
            print(f"  Row {j}: height={height}, rule={height_rule}")

def test_enforce_fixed_cell_dimensions():
    """Test the enforce_fixed_cell_dimensions function in isolation"""
    
    # Create a simple document with a table
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    
    # Set row heights to EXACTLY
    for row in table.rows:
        row.height = Inches(2.25)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    # Add some content to cells
    for i in range(3):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = f"Test content {i*3+j+1}"
    
    # Check initial state
    check_row_heights(doc, "1. Initial state")
    
    # Save and reload to ensure clean state
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    doc = Document(buffer)
    check_row_heights(doc, "2. After save/reload")
    
    # Apply enforce_fixed_cell_dimensions
    for table in doc.tables:
        enforce_fixed_cell_dimensions(table)
    
    check_row_heights(doc, "3. After enforce_fixed_cell_dimensions")
    
    # Save and check final state
    output_path = "test_enforce_output.docx"
    doc.save(output_path)
    
    final_doc = Document(output_path)
    check_row_heights(final_doc, "4. After saving final document")
    
    # Check XML
    print("\n=== XML Analysis ===")
    for i, table in enumerate(final_doc.tables):
        print(f"Table {i} XML:")
        table_xml = table._element.xml
        if 'trHeight' in table_xml:
            print("  Found trHeight elements in XML")
            import re
            tr_height_matches = re.findall(r'<w:trHeight[^>]*>', table_xml)
            print(f"  Total trHeight elements: {len(tr_height_matches)}")
            for idx, match in enumerate(tr_height_matches):
                print(f"    {idx}: {match}")
        else:
            print("  No trHeight elements found in XML")
    
    # Clean up
    if os.path.exists(output_path):
        os.remove(output_path)

if __name__ == "__main__":
    test_enforce_fixed_cell_dimensions() 