#!/usr/bin/env python3
"""
Test to manually add lineage content to mini template cells.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.docx_formatting import apply_lineage_colors, COLORS
from docx import Document
from docx.oxml.ns import qn

def test_mini_lineage_manual():
    """Test manually adding lineage content to mini template cells."""
    
    print("Testing Manual Lineage Addition to Mini Template")
    print("=" * 50)
    
    # Test data
    test_record = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'MIXED',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'Description': 'Test description text',
        'ProductStrain': 'Mixed',
        'ProductType': 'tincture'
    }
    
    try:
        # Create mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Process a single record
        result_doc = processor.process_records([test_record])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"Table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Manually add lineage content to the first cell
            cell = table.cell(0, 0)
            print(f"\nOriginal cell text: '{cell.text}'")
            
            # Add lineage content manually
            lineage_content = "MIXED_PRODUCT_TYPE_tincture_IS_CLASSIC_false"
            cell.text = f"{cell.text}\n{lineage_content}"
            print(f"After adding lineage: '{cell.text}'")
            
            # Apply lineage colors
            apply_lineage_colors(result_doc)
            
            # Check if the cell now has background color
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = tcPr.find(qn('w:shd'))
            
            if shading is not None:
                fill_color = shading.get(qn('w:fill'))
                print(f"Background color: {fill_color}")
                if fill_color == COLORS['MIXED']:
                    print("✅ Correct color applied for MIXED")
                else:
                    print(f"❌ Wrong color. Expected: {COLORS['MIXED']}, Got: {fill_color}")
            else:
                print("❌ No background color applied")
            
            # Test with CBD content
            print(f"\nTesting with CBD content:")
            cell.text = "CBD_PRODUCT_TYPE_tincture_IS_CLASSIC_false"
            print(f"Cell text: '{cell.text}'")
            
            # Clear any existing background
            for old_shd in tcPr.findall(qn('w:shd')):
                tcPr.remove(old_shd)
            
            # Apply lineage colors again
            apply_lineage_colors(result_doc)
            
            # Check the new color
            shading = tcPr.find(qn('w:shd'))
            if shading is not None:
                fill_color = shading.get(qn('w:fill'))
                print(f"Background color: {fill_color}")
                if fill_color == COLORS['CBD']:
                    print("✅ Correct color applied for CBD")
                else:
                    print(f"❌ Wrong color. Expected: {COLORS['CBD']}, Got: {fill_color}")
            else:
                print("❌ No background color applied")
                
        else:
            print("❌ Failed to generate document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

def test_mini_template_content():
    """Test what content is actually in the mini template."""
    
    print("\n\nTesting Mini Template Content")
    print("=" * 50)
    
    try:
        # Create mini template processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Check the expanded template buffer
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            
            if doc.tables:
                table = doc.tables[0]
                print(f"Template table dimensions: {len(table.rows)}x{len(table.columns)}")
                
                # Check the first cell content
                cell = table.cell(0, 0)
                print(f"First cell text: '{cell.text}'")
                
                # Check for template variables
                if '{{Label1.' in cell.text:
                    print("✅ Template variables found")
                    # Extract all template variables
                    import re
                    variables = re.findall(r'\{\{Label1\.(\w+)\}\}', cell.text)
                    print(f"Template variables: {variables}")
                    
                    if 'Lineage' in variables:
                        print("✅ Lineage field found in template")
                    else:
                        print("❌ Lineage field NOT found in template")
                        print("This explains why lineage coloring doesn't work for mini templates")
                else:
                    print("❌ No template variables found")
            else:
                print("❌ No tables found in template")
        else:
            print("❌ No expanded template buffer found")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_mini_template_content()
    test_mini_lineage_manual() 