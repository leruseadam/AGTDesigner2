#!/usr/bin/env python3
"""
Debug script using actual template files to investigate row height changes
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from docxtpl import DocxTemplate
from io import BytesIO

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.core.generation.template_processor import TemplateProcessor

def check_row_heights(doc, stage_name):
    """Check and print row height rules for all tables in the document"""
    print(f"\n=== {stage_name} ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            height = row.height
            height_rule = row.height_rule
            print(f"  Row {j}: height={height}, rule={height_rule}")

def test_real_template():
    """Test with the actual horizontal template"""
    
    # Use the actual template
    template_path = "src/core/generation/templates/horizontal.docx"
    
    try:
        # Step 1: Load original template and check state
        doc = Document(template_path)
        check_row_heights(doc, "1. Original template")
        
        # Step 2: Test TemplateProcessor with real template
        processor = TemplateProcessor('horizontal', {}, 1.0)
        
        # Check the expanded template buffer
        expanded_doc = Document(processor._expanded_template_buffer)
        check_row_heights(expanded_doc, "2. After template expansion")
        
        # Step 3: Test docxtpl rendering
        template = DocxTemplate(processor._expanded_template_buffer)
        template.init_docx()
        check_row_heights(template.docx, "3. After DocxTemplate init")
        
        # Create context
        context = {}
        for i in range(9):
            context[f'Label{i+1}'] = {
                'ProductBrand': f'Test Brand {i+1}',
                'Price': f'${i+1}.99',
                'Description': f'Test Description {i+1}',
                'Lineage': f'Test Lineage {i+1}',
                'Ratio_or_THC_CBD': f'THC: {i+1}%',
                'DOH': 'YES',
                'DescAndWeight': f'Test Description {i+1} - 3.5g'
            }
        
        # Render template
        template.render(context)
        check_row_heights(template.docx, "4. After docxtpl render")
        
        # Step 4: Test post-processing
        rendered_doc = template.docx
        processor._post_process_and_replace_content(rendered_doc)
        check_row_heights(rendered_doc, "5. After post-processing")
        
        # Step 5: Save and check final state
        output_path = "test_real_output.docx"
        rendered_doc.save(output_path)
        
        final_doc = Document(output_path)
        check_row_heights(final_doc, "6. After saving final document")
        
        # Step 6: Check XML
        print("\n=== XML Analysis ===")
        for i, table in enumerate(final_doc.tables):
            print(f"Table {i} XML:")
            table_xml = table._element.xml
            if 'trHeight' in table_xml:
                print("  Found trHeight elements in XML")
                import re
                tr_height_matches = re.findall(r'<w:trHeight[^>]*>', table_xml)
                print(f"  Total trHeight elements: {len(tr_height_matches)}")
                for idx, match in enumerate(tr_height_matches):
                    print(f"    {idx}: {match}")
            else:
                print("  No trHeight elements found in XML")
        
        # Step 7: Check if there are any rows without height rules
        print("\n=== Row Analysis ===")
        for i, table in enumerate(final_doc.tables):
            print(f"Table {i}:")
            for j, row in enumerate(table.rows):
                height = row.height
                height_rule = row.height_rule
                if height_rule != WD_ROW_HEIGHT_RULE.EXACTLY:
                    print(f"  Row {j}: height={height}, rule={height_rule} - NOT EXACTLY!")
                else:
                    print(f"  Row {j}: height={height}, rule={height_rule} - OK")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_real_template() 