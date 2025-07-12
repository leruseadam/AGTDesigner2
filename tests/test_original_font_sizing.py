#!/usr/bin/env python3
"""
Test script to verify that the original font-sizing functions are being used correctly.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.font_sizing import (
    get_thresholded_font_size,
    get_thresholded_font_size_ratio,
    get_thresholded_font_size_brand,
    get_thresholded_font_size_price,
    get_thresholded_font_size_lineage,
    get_thresholded_font_size_description,
    get_thresholded_font_size_strain
)
from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def test_original_font_sizing_functions():
    """Test that the original font-sizing functions work correctly for different template types."""
    
    print("Testing Original Font-Sizing Functions")
    print("=" * 50)
    
    # Test data
    test_content = "Test Brand Name"
    template_types = ['vertical', 'horizontal', 'mini']
    
    print(f"Testing content: '{test_content}'")
    print()
    
    for template_type in template_types:
        print(f"{template_type.upper()} TEMPLATE:")
        print("-" * 20)
        
        # Test different field types
        field_tests = [
            ('Description', get_thresholded_font_size_description, 'description'),
            ('Brand', get_thresholded_font_size_brand, 'brand'),
            ('Price', get_thresholded_font_size_price, 'price'),
            ('Lineage', get_thresholded_font_size_lineage, 'lineage'),
            ('Ratio', get_thresholded_font_size_ratio, 'ratio'),
            ('Strain', get_thresholded_font_size_strain, 'strain'),
            ('General', get_thresholded_font_size, 'default')
        ]
        
        for field_name, func, field_type in field_tests:
            try:
                if field_type == 'default':
                    font_size = func(test_content, template_type, 1.0, field_type)
                else:
                    font_size = func(test_content, template_type, 1.0)
                print(f"  {field_name}: {font_size.pt}pt")
            except Exception as e:
                print(f"  {field_name}: ERROR - {e}")
        
        print()

def test_template_processor_font_sizing():
    """Test that the template processor uses the correct font-sizing functions."""
    
    print("\nTesting Template Processor Font Sizing")
    print("=" * 50)
    
    # Test data
    test_record = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'Hybrid',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'Description': 'Test description text',
        'ProductStrain': 'Test Strain'
    }
    
    template_types = ['vertical', 'horizontal', 'mini']
    
    for template_type in template_types:
        print(f"\n{template_type.upper()} TEMPLATE:")
        print("-" * 20)
        
        try:
            # Create template processor
            processor = TemplateProcessor(template_type, {}, 1.0)
            
            # Process a single record
            result_doc = processor.process_records([test_record])
            
            if result_doc and result_doc.tables:
                table = result_doc.tables[0]
                cell = table.cell(0, 0)
                
                print(f"  Document generated successfully")
                print(f"  Table dimensions: {len(table.rows)}x{len(table.columns)}")
                
                # Check font sizes in the first cell
                font_sizes = []
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text.strip() and run.font.size:
                            font_sizes.append(f"{run.text[:20]}...: {run.font.size.pt}pt")
                
                if font_sizes:
                    print(f"  Font sizes found:")
                    for size_info in font_sizes[:5]:  # Show first 5
                        print(f"    {size_info}")
                else:
                    print(f"  No font sizes found")
            else:
                print(f"  Failed to generate document")
                
        except Exception as e:
            print(f"  ERROR: {e}")

def test_font_size_comparison():
    """Compare font sizes between different template types for the same content."""
    
    print("\nFont Size Comparison Across Template Types")
    print("=" * 50)
    
    test_content = "Medium length brand name"
    field_type = 'brand'
    
    print(f"Content: '{test_content}'")
    print(f"Field type: {field_type}")
    print()
    
    template_types = ['vertical', 'horizontal', 'mini']
    
    print(f"{'Template':<12} {'Font Size':<10} {'Scale Factor':<12}")
    print("-" * 35)
    
    for template_type in template_types:
        for scale_factor in [0.8, 1.0, 1.2]:
            try:
                font_size = get_thresholded_font_size_brand(test_content, template_type, scale_factor)
                print(f"{template_type:<12} {font_size.pt:<10.1f} {scale_factor:<12.1f}")
            except Exception as e:
                print(f"{template_type:<12} ERROR     {scale_factor:<12.1f}")

if __name__ == "__main__":
    test_original_font_sizing_functions()
    test_template_processor_font_sizing()
    test_font_size_comparison() 