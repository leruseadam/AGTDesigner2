#!/usr/bin/env python3
"""
Detailed debug script to trace lineage coloring in mini template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.docx_formatting import apply_lineage_colors, COLORS
from docx import Document
from docx.oxml.ns import qn

def debug_mini_coloring_step_by_step():
    """Debug the mini template coloring step by step."""
    
    print("Detailed Mini Template Coloring Debug")
    print("=" * 50)
    
    # Test data
    test_record = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'MIXED',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'Description': 'Test description text',
        'ProductStrain': 'Mixed',
        'ProductType': 'tincture'
    }
    
    try:
        # Step 1: Create processor
        print("Step 1: Creating template processor...")
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Step 2: Check template content
        print("\nStep 2: Checking template content...")
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            if doc.tables:
                cell = doc.tables[0].cell(0, 0)
                print(f"Template cell text: '{cell.text}'")
        
        # Step 3: Process record
        print("\nStep 3: Processing record...")
        result_doc = processor.process_records([test_record])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            print(f"Result table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Step 4: Check each cell before coloring
            print("\nStep 4: Checking cells before applying colors...")
            for row_idx in range(min(2, len(table.rows))):  # Check first 2 rows
                for col_idx in range(min(2, len(table.columns))):  # Check first 2 columns
                    cell = table.cell(row_idx, col_idx)
                    print(f"Cell ({row_idx}, {col_idx}): '{cell.text}'")
                    
                    # Check if cell has any background color
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = tcPr.find(qn('w:shd'))
                    if shading is not None:
                        print(f"  Has background: {shading.get(qn('w:fill'))}")
                    else:
                        print(f"  No background color")
            
            # Step 5: Apply lineage colors
            print("\nStep 5: Applying lineage colors...")
            apply_lineage_colors(result_doc)
            
            # Step 6: Check cells after coloring
            print("\nStep 6: Checking cells after applying colors...")
            for row_idx in range(min(2, len(table.rows))):
                for col_idx in range(min(2, len(table.columns))):
                    cell = table.cell(row_idx, col_idx)
                    print(f"Cell ({row_idx}, {col_idx}): '{cell.text}'")
                    
                    # Check if cell has background color
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = tcPr.find(qn('w:shd'))
                    if shading is not None:
                        fill_color = shading.get(qn('w:fill'))
                        print(f"  Background color: {fill_color}")
                        
                        # Check if it matches expected lineage color
                        if 'MIXED' in cell.text:
                            expected = COLORS['MIXED']
                            if fill_color == expected:
                                print(f"  ✅ Correct MIXED color")
                            else:
                                print(f"  ❌ Wrong color for MIXED. Expected: {expected}, Got: {fill_color}")
                        elif 'CBD' in cell.text:
                            expected = COLORS['CBD']
                            if fill_color == expected:
                                print(f"  ✅ Correct CBD color")
                            else:
                                print(f"  ❌ Wrong color for CBD. Expected: {expected}, Got: {fill_color}")
                    else:
                        print(f"  ❌ No background color")
                        
                        # If no color and cell has lineage content, debug why
                        if 'MIXED' in cell.text or 'CBD' in cell.text:
                            print(f"  🔍 Cell has lineage content but no color - investigating...")
                            # Check if the content is in the right format
                            if '_PRODUCT_TYPE_' in cell.text:
                                print(f"  ✅ Content has product type info")
                            else:
                                print(f"  ❌ Content missing product type info")
        else:
            print("❌ Failed to generate document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

def test_apply_lineage_colors_manual():
    """Test applying lineage colors manually to verify the function works."""
    
    print("\n\nTesting Manual Lineage Color Application")
    print("=" * 50)
    
    # Create a simple document with lineage content
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Test different lineage content formats
    test_cases = [
        "MIXED_PRODUCT_TYPE_tincture_IS_CLASSIC_false",
        "CBD_PRODUCT_TYPE_tincture_IS_CLASSIC_false",
        "MIXED",
        "CBD"
    ]
    
    for i, content in enumerate(test_cases):
        print(f"\nTest case {i+1}: '{content}'")
        
        # Set cell content
        cell.text = content
        
        # Apply colors
        apply_lineage_colors(doc)
        
        # Check result
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.find(qn('w:shd'))
        
        if shading is not None:
            fill_color = shading.get(qn('w:fill'))
            print(f"  Background color: {fill_color}")
            
            if 'MIXED' in content:
                expected = COLORS['MIXED']
                if fill_color == expected:
                    print(f"  ✅ Correct MIXED color")
                else:
                    print(f"  ❌ Wrong color for MIXED. Expected: {expected}, Got: {fill_color}")
            elif 'CBD' in content:
                expected = COLORS['CBD']
                if fill_color == expected:
                    print(f"  ✅ Correct CBD color")
                else:
                    print(f"  ❌ Wrong color for CBD. Expected: {expected}, Got: {fill_color}")
        else:
            print(f"  ❌ No background color applied")
        
        # Clear background for next test
        for old_shd in tcPr.findall(qn('w:shd')):
            tcPr.remove(old_shd)

if __name__ == "__main__":
    debug_mini_coloring_step_by_step()
    test_apply_lineage_colors_manual() 