#!/usr/bin/env python3
"""
Debug script to check if markers are being found and processed in templates.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def check_template_content():
    """Check what content is in the template files."""
    
    print("Checking Template Content")
    print("=" * 50)
    
    template_types = ['vertical', 'horizontal', 'mini']
    
    for template_type in template_types:
        print(f"\n{template_type.upper()} template:")
        print("-" * 30)
        
        try:
            processor = TemplateProcessor(template_type, {}, 1.0)
            
            # Check the expanded template buffer
            if hasattr(processor, '_expanded_template_buffer'):
                processor._expanded_template_buffer.seek(0)
                doc = Document(processor._expanded_template_buffer)
                
                print(f"Template has {len(doc.tables)} tables")
                
                if doc.tables:
                    table = doc.tables[0]
                    print(f"First table: {len(table.rows)} rows, {len(table.columns)} columns")
                    
                    # Check each cell for content
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            if cell.text.strip():
                                print(f"  Cell ({row_idx+1},{col_idx+1}): '{cell.text[:100]}...'")
                                
                                # Check for markers
                                markers_found = []
                                for marker in ['PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 'THC_CBD', 'PRODUCTSTRAIN', 'DESC']:
                                    if f'{marker}_START' in cell.text:
                                        markers_found.append(marker)
                                
                                if markers_found:
                                    print(f"    Found markers: {markers_found}")
                                else:
                                    print(f"    No markers found")
        except Exception as e:
            print(f"Error checking {template_type} template: {e}")

def test_marker_processing_step_by_step():
    """Test marker processing step by step."""
    
    print("\n\nTesting Marker Processing Step by Step")
    print("=" * 50)
    
    # Create test data with all markers
    test_data = [
        {
            'ProductBrand': 'Short Brand',
            'Price': '$25.99',
            'Lineage': 'Sativa',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductStrain': 'Test Strain',
            'Description': 'Short description',
            'WeightUnits': '3.5g'
        }
    ]
    
    processor = TemplateProcessor('vertical', {}, 1.0)
    
    # Step 1: Build context
    print("Step 1: Building context...")
    context = processor._build_label_context(test_data[0], None)
    for key, value in context.items():
        if value and any(marker in value for marker in ['_START', '_END']):
            print(f"  {key}: '{value[:50]}...'")
    
    # Step 2: Process chunk
    print("\nStep 2: Processing chunk...")
    result_doc = processor._process_chunk(test_data)
    
    if result_doc and result_doc.tables:
        table = result_doc.tables[0]
        print(f"Result document has {len(table.rows)} rows, {len(table.columns)} columns")
        
        # Check each cell
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"  Cell ({row_idx+1},{col_idx+1}): '{cell.text[:100]}...'")
                    
                    # Check for markers
                    markers_found = []
                    for marker in ['PRODUCTBRAND_CENTER', 'PRICE', 'LINEAGE', 'THC_CBD', 'PRODUCTSTRAIN', 'DESC']:
                        if f'{marker}_START' in cell.text:
                            markers_found.append(marker)
                    
                    if markers_found:
                        print(f"    Found markers: {markers_found}")
                    else:
                        print(f"    No markers found")
                    
                    # Check font sizes
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if run.text.strip():
                                font_size = run.font.size
                                pt_value = font_size.pt if font_size else "None"
                                print(f"    Paragraph {para_idx+1}, Run {run_idx+1}: '{run.text[:30]}...' - {pt_value}pt")

def test_simple_marker_processing():
    """Test simple marker processing with a basic document."""
    
    print("\n\nTesting Simple Marker Processing")
    print("=" * 50)
    
    # Create a simple document with markers
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add content with markers
    cell.text = "PRODUCTBRAND_CENTER_STARTShort BrandPRODUCTBRAND_CENTER_END\nPRICE_START$25.99PRICE_END\nLINEAGE_STARTSativaLINEAGE_END"
    
    print(f"Original cell text: '{cell.text}'")
    
    # Test the marker processing
    processor = TemplateProcessor('vertical', {}, 1.0)
    processor._post_process_and_replace_content(doc)
    
    print(f"After processing: '{cell.text}'")
    
    # Check the results
    for para_idx, paragraph in enumerate(cell.paragraphs):
        for run_idx, run in enumerate(paragraph.runs):
            if run.text.strip():
                font_size = run.font.size
                pt_value = font_size.pt if font_size else "None"
                print(f"  Paragraph {para_idx+1}, Run {run_idx+1}: '{run.text}' - {pt_value}pt")

if __name__ == "__main__":
    print("Template Marker Debug Test")
    print("=" * 50)
    
    # Check template content
    check_template_content()
    
    # Test step by step processing
    test_marker_processing_step_by_step()
    
    # Test simple marker processing
    test_simple_marker_processing()
    
    print("\n" + "=" * 50)
    print("Debug test completed.") 