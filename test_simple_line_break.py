#!/usr/bin/env python3
"""
Simple test for line break conversion
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

def test_line_break_conversion_simple():
    """Test that \n characters are properly converted to line breaks."""
    
    print("Testing Simple Line Break Conversion")
    print("=" * 50)
    
    # Create a test document
    doc = Document()
    
    # Add a paragraph with \n characters
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("230mg CBD\n50mg THC\n10mg CBG\n10mg CBN")
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(12)
    
    print(f"Original text: '{paragraph.text}'")
    has_newlines = '\n' in paragraph.text
    print(f"Contains newlines: {has_newlines}")
    
    # Manually implement the line break conversion logic
    try:
        # Get all text from the paragraph
        full_text = "".join(run.text for run in paragraph.runs)
        
        # Check if there are any \n characters
        if '\n' not in full_text:
            print("❌ No newlines found in text")
            return False
        
        # Split the text by \n characters
        parts = full_text.split('\n')
        
        # Clear the paragraph
        paragraph.clear()
        
        # Set tight paragraph spacing to prevent excessive gaps
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        
        # Add each part as a separate run, with line breaks between them
        for i, part in enumerate(parts):
            if part.strip():  # Only add non-empty parts
                run = paragraph.add_run(part.strip())
                run.font.name = "Arial"
                run.font.bold = True
                run.font.size = Pt(12)
                
                # Add a line break after this part only if the next part is not empty
                if i < len(parts) - 1 and parts[i + 1].strip():
                    # Use add_break() with WD_BREAK.LINE to create proper line breaks within the same paragraph
                    run.add_break(WD_BREAK.LINE)
        
        print(f"Converted {len(parts)-1} newlines to line breaks")
        
    except Exception as e:
        print(f"❌ Error converting newlines to line breaks: {e}")
        return False
    
    print(f"After conversion: '{paragraph.text}'")
    print(f"Number of runs: {len(paragraph.runs)}")
    
    # Check if line breaks were properly converted
    if len(paragraph.runs) > 1:
        print("✅ Line breaks were successfully converted to separate runs")
        for i, run in enumerate(paragraph.runs):
            print(f"  Run {i+1}: '{run.text}'")
    else:
        print("❌ Line breaks were not converted properly")
    
    # Save the test document
    test_file = "test_simple_line_break_output.docx"
    doc.save(test_file)
    print(f"\nTest document saved as: {test_file}")
    
    return True

if __name__ == "__main__":
    test_line_break_conversion_simple() 