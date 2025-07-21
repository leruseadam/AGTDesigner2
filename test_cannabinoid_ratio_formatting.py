#!/usr/bin/env python3
"""
Test script to verify that cannabinoid ratios are properly formatted with line breaks.
"""

import sys
import os
from pathlib import Path

# Add the src directory to the path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

from src.core.generation.template_processor import TemplateProcessor

def test_cannabinoid_ratio_formatting():
    """Test that cannabinoid ratios are properly formatted with line breaks."""
    
    print("=== Cannabinoid Ratio Formatting Test ===\n")
    
    # Test case: The exact content from the image
    test_text = "230mg CBD|BR|50mg THC|BR|10mg CBG|BR|10mg CBN"
    print(f"Test: '{test_text}'")
    print("Expected: Each cannabinoid on its own line")
    print()
    
    # Create a test document
    doc = Document()
    paragraph = doc.add_paragraph()
    
    # Add the text to the paragraph
    run = paragraph.add_run(test_text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(12)
    
    # Create a template processor to test the fix
    processor = TemplateProcessor('vertical', {})
    
    # Apply the line break conversion
    processor._convert_br_markers_to_line_breaks(paragraph)
    
    # Check the result
    result_text = paragraph.text
    print(f"Result:")
    print(f"'{result_text}'")
    
    # Count the number of line breaks (should be 4)
    line_count = len(result_text.split('\n'))
    print(f"Line count: {line_count}")
    
    # Show each line separately
    lines = result_text.split('\n')
    print("\nIndividual lines:")
    for i, line in enumerate(lines, 1):
        print(f"  Line {i}: '{line}'")
    
    if line_count == 4:
        print("\n✅ PASS: Correct number of lines (4 cannabinoids on separate lines)")
    else:
        print(f"\n❌ FAIL: Expected 4 lines, got {line_count}")
    
    print()
    
    # Test case 2: Different cannabinoid content
    doc2 = Document()
    paragraph2 = doc2.add_paragraph()
    
    test_text2 = "100mg THC|BR|100mg CBD|BR|50mg CBG"
    print(f"Test 2: '{test_text2}'")
    
    # Add the text to the paragraph
    run2 = paragraph2.add_run(test_text2)
    run2.font.name = "Arial"
    run2.font.bold = True
    run2.font.size = Pt(12)
    
    # Apply the line break conversion
    processor._convert_br_markers_to_line_breaks(paragraph2)
    
    # Check the result
    result_text2 = paragraph2.text
    print(f"Result: '{result_text2}'")
    
    # Count the number of line breaks (should be 3)
    line_count2 = len(result_text2.split('\n'))
    print(f"Line count: {line_count2}")
    
    if line_count2 == 3:
        print("✅ PASS: Correct number of lines (3 cannabinoids on separate lines)")
    else:
        print(f"❌ FAIL: Expected 3 lines, got {line_count2}")
    
    print()
    
    # Test case 3: Single cannabinoid (should not add extra line break)
    doc3 = Document()
    paragraph3 = doc3.add_paragraph()
    
    test_text3 = "500mg CBD"
    print(f"Test 3: '{test_text3}'")
    
    # Add the text to the paragraph
    run3 = paragraph3.add_run(test_text3)
    run3.font.name = "Arial"
    run3.font.bold = True
    run3.font.size = Pt(12)
    
    # Apply the line break conversion
    processor._convert_br_markers_to_line_breaks(paragraph3)
    
    # Check the result
    result_text3 = paragraph3.text
    print(f"Result: '{result_text3}'")
    
    # Count the number of line breaks (should be 1)
    line_count3 = len(result_text3.split('\n'))
    print(f"Line count: {line_count3}")
    
    if line_count3 == 1:
        print("✅ PASS: Correct number of lines (single cannabinoid, no extra breaks)")
    else:
        print(f"❌ FAIL: Expected 1 line, got {line_count3}")
    
    print("\n=== Test Complete ===")

if __name__ == "__main__":
    test_cannabinoid_ratio_formatting() 