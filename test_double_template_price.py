#!/usr/bin/env python3
"""
Test script to debug double template price processing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.generation.unified_font_sizing import get_font_size_by_marker
from docx import Document
from docx.shared import Pt

def test_double_template_price_processing():
    """Test double template price processing."""
    
    print("Testing Double Template Price Processing")
    print("=" * 50)
    
    # Create a test record
    test_record = {
        'Description': 'Test Product',
        'Price': '$12',
        'Lineage': 'SATIVA',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'ProductBrand': 'Test Brand',
        'ProductType': 'pre-roll'
    }
    
    # Create template processor
    processor = TemplateProcessor('double', {}, 1.0)
    
    # Test marker type detection
    print("\nTesting marker type detection:")
    print("-" * 30)
    
    test_prices = ["$12", "$27", "$4", "$110", "$75", "$7"]
    
    for price in test_prices:
        field_type = processor._determine_field_type_from_content(price, 'pre-roll')
        print(f"Price '{price}' -> Field type: {field_type}")
        
        if field_type == 'price':
            from src.core.generation.unified_font_sizing import get_font_size
            font_size = get_font_size(price, field_type, 'double', 1.0, 'standard')
            print(f"  Font size: {font_size.pt}pt")
        else:
            print(f"  ERROR: Not detected as price!")
    
    # Test actual template processing
    print("\nTesting actual template processing:")
    print("-" * 35)
    
    try:
        # Process the test record
        result_doc = processor.process_records([test_record])
        
        if result_doc and result_doc.tables:
            table = result_doc.tables[0]
            
            # Check all cells for price content
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip() and any(char in run.text for char in '$€£¥'):
                                print(f"Found price text: '{run.text}' in cell ({row_idx}, {col_idx})")
                                if run.font.size:
                                    print(f"  Font size: {run.font.size.pt}pt")
                                else:
                                    print(f"  Font size: None")
        else:
            print("No document or tables generated")
            
    except Exception as e:
        print(f"Error processing template: {e}")
        import traceback
        traceback.print_exc()

def test_price_detection_patterns():
    """Test price detection patterns."""
    
    print("\n\nTesting Price Detection Patterns")
    print("=" * 40)
    
    import re
    
    test_cases = [
        "$12",
        "$27", 
        "$4",
        "$110",
        "$75",
        "$7",
        "12",
        "27.99",
        "$12.99",
        "€15",
        "£20"
    ]
    
    for text in test_cases:
        # Test the exact pattern from the code
        has_currency = any(char in text for char in '$€£¥')
        is_number = bool(re.match(r'^\d+\.?\d*$', text.strip()))
        
        print(f"'{text}':")
        print(f"  Has currency: {has_currency}")
        print(f"  Is number: {is_number}")
        print(f"  Would be PRICE: {has_currency or is_number}")

if __name__ == "__main__":
    test_double_template_price_processing()
    test_price_detection_patterns() 