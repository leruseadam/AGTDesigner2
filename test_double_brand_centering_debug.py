#!/usr/bin/env python3
"""
Comprehensive test to debug why brand centering isn't working in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO

def debug_brand_centering():
    """Debug why brand centering isn't working."""
    print("Debugging Double Template Brand Centering")
    print("=" * 60)
    
    # Test data with actual brand names from the screenshot
    test_records = [
        {
            'Description': '2:1 CBN Blueberry Lemonade Shot -0.07oz',
            'WeightUnits': '0.07oz',
            'ProductBrand': 'JOURNEYMAN',
            'Price': '$18',
            'Lineage': '100mg THC 50mg',
            'THC_CBD': 'THC: 20% CBD: 2%',
            'ProductStrain': 'Test Strain',
            'DOH': 'DOH',
            'ProductType': 'edible (liquid)'
        },
        {
            'Description': '20:1 CBD Ratio Tincture -0.25oz',
            'WeightUnits': '0.25oz',
            'ProductBrand': 'FAIRWINDS',
            'Price': '$38',
            'Lineage': '190mg CBD 10mg',
            'THC_CBD': 'THC: 20% CBD: 2%',
            'ProductStrain': 'Test Strain',
            'DOH': 'DOH',
            'ProductType': 'tincture'
        },
        {
            'Description': '2:1:1 Agave Lime MAX Wildside Shot -2oz',
            'WeightUnits': '2oz',
            'ProductBrand': 'GREEN',
            'Price': '$18',
            'Lineage': '100mg THC 50mg',
            'THC_CBD': 'THC: 20% CBD: 2%',
            'ProductStrain': 'Test Strain',
            'DOH': 'DOH',
            'ProductType': 'edible (liquid)'
        }
    ]
    
    # Process the records
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records(test_records)
    
    if not result_doc:
        print("ERROR: Failed to process test records")
        return
    
    print(f"Document processed successfully")
    print(f"Tables found: {len(result_doc.tables)}")
    
    # Examine each table and cell
    for table_idx, table in enumerate(result_doc.tables):
        print(f"\nTable {table_idx + 1}:")
        print(f"  Table alignment: {table.alignment}")
        print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                print(f"  Cell [{row_idx},{col_idx}]:")
                print(f"    Cell text: '{cell.text.strip()}'")
                
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    print(f"    Paragraph {para_idx + 1}:")
                    print(f"      Text: '{paragraph.text.strip()}'")
                    print(f"      Alignment: {paragraph.alignment}")
                    print(f"      Runs: {len(paragraph.runs)}")
                    
                    for run_idx, run in enumerate(paragraph.runs):
                        print(f"        Run {run_idx + 1}: '{run.text}' (font size: {run.font.size})")
                        
                        # Check if this looks like brand content
                        if any(brand in run.text.upper() for brand in ['JOURNEYMAN', 'FAIRWINDS', 'GREEN', 'CQ', 'HOT SHOTZ']):
                            print(f"        *** BRAND CONTENT DETECTED: {run.text} ***")
                            print(f"        Current alignment: {paragraph.alignment}")
                            
                            # Force center this paragraph
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print(f"        *** FORCED CENTERING APPLIED ***")
    
    # Save the document for inspection
    output_path = "debug_brand_centering_output.docx"
    result_doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    
    # Now let's manually apply centering to all brand content
    print("\n" + "=" * 60)
    print("MANUALLY APPLYING BRAND CENTERING")
    print("=" * 60)
    
    for table in result_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text.strip()
                    
                    # Check if this paragraph contains brand content
                    brand_names = ['JOURNEYMAN', 'FAIRWINDS', 'GREEN', 'CQ', 'HOT SHOTZ', 'CERES']
                    if any(brand in paragraph_text.upper() for brand in brand_names):
                        print(f"Found brand content: '{paragraph_text}'")
                        print(f"  Before centering: {paragraph.alignment}")
                        
                        # Force center
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        print(f"  After centering: {paragraph.alignment}")
                        print(f"  ✓ Centered")
    
    # Save the manually corrected document
    corrected_path = "debug_brand_centering_corrected.docx"
    result_doc.save(corrected_path)
    print(f"\nCorrected document saved to: {corrected_path}")
    
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print("1. Check the original output document to see the current state")
    print("2. Check the corrected document to see the manually applied centering")
    print("3. Compare the two to understand what needs to be fixed in the code")

if __name__ == "__main__":
    debug_brand_centering() 