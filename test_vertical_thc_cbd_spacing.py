#!/usr/bin/env python3
"""
Test script to verify THC:CBD line spacing is set to 2.0 for vertical template
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_vertical_thc_cbd_spacing():
    """Test that THC:CBD line spacing is set to 2.0 for vertical template"""
    print("🧪 Testing Vertical Template THC:CBD Line Spacing")
    print("=" * 50)
    
    # Create a test record with THC:CBD content
    test_record = {
        'Product Name*': 'Test Product',
        'Description': 'Test Description',
        'Product Brand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'HYBRID',
        'Product Type*': 'flower',  # Add product type to trigger classic type logic
        'THC_CBD': 'THC:\nCBD:',
        'Ratio_or_THC_CBD': 'THC:\nCBD:',
        'Ratio': 'THC:\nCBD:',
        'WeightUnits': '1g',
        'Product Strain': 'Test Strain',
        'DOH': 'Yes'
    }
    
    # Test vertical template
    print("\n1️⃣ Testing Vertical Template...")
    tp_vertical = TemplateProcessor(template_type='vertical', font_scheme='Arial')
    
    # Process the record
    doc = tp_vertical.process_records([test_record])
    
    if not doc:
        print("❌ No document generated")
        return False
    
    # Check for THC:CBD content and verify line spacing
    found_thc_cbd = False
    correct_spacing = False
    
    print("🔍 Searching for THC:CBD content in document...")
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(f"📄 Paragraph text: '{paragraph.text.strip()}'")
                    if 'THC:' in paragraph.text and 'CBD:' in paragraph.text:
                        found_thc_cbd = True
                        line_spacing = paragraph.paragraph_format.line_spacing
                        print(f"📋 Found THC:CBD content: '{paragraph.text.strip()}'")
                        print(f"📏 Line spacing: {line_spacing}")
                        
                        if line_spacing == 2.0:
                            correct_spacing = True
                            print("✅ Line spacing is correctly set to 2.0")
                        else:
                            print(f"❌ Line spacing should be 2.0, but is {line_spacing}")
    
    if not found_thc_cbd:
        print("❌ No THC:CBD content found in the document")
        return False
    
    # Test horizontal template for comparison
    print("\n2️⃣ Testing Horizontal Template (should have different spacing)...")
    tp_horizontal = TemplateProcessor(template_type='horizontal', font_scheme='Arial')
    
    doc_horizontal = tp_horizontal.process_records([test_record])
    
    if doc_horizontal:
        
        for table in doc_horizontal.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'THC:' in paragraph.text and 'CBD:' in paragraph.text:
                            line_spacing_horizontal = paragraph.paragraph_format.line_spacing
                            print(f"📋 Horizontal template THC:CBD line spacing: {line_spacing_horizontal}")
                            
                            if line_spacing_horizontal != 2.0:
                                print("✅ Horizontal template correctly has different spacing")
                            else:
                                print("❌ Horizontal template should not have spacing of 2.0")
    
    # Summary
    print("\n" + "=" * 50)
    if found_thc_cbd and correct_spacing:
        print("🎉 SUCCESS: Vertical template THC:CBD line spacing is correctly set to 2.0")
        return True
    else:
        print("❌ FAILED: Vertical template THC:CBD line spacing is not correct")
        return False

if __name__ == "__main__":
    success = test_vertical_thc_cbd_spacing()
    sys.exit(0 if success else 1) 