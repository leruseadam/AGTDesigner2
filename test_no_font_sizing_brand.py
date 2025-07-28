#!/usr/bin/env python3
"""
Test brand processing without font sizing to see if that's where the truncation happens.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docxtpl import DocxTemplate
from io import BytesIO

def test_no_font_sizing_brand():
    """Test brand processing without font sizing."""
    
    print("No Font Sizing Brand Test")
    print("=" * 40)
    
    # Create a single test record
    test_record = {
        'Description': 'No Font Product',
        'Price': '$15',
        'Lineage': 'HYBRID',
        'Ratio_or_THC_CBD': 'THC: 22%',
        'ProductBrand': 'NO_FONT_BRAND_TEST',
        'ProductStrain': 'No Font Strain',
        'ProductType': 'flower',
        'JointRatio': '3.5g'
    }
    
    # Initialize template processor
    processor = TemplateProcessor('double', 'standard', scale_factor=1.0)
    
    # Get the template buffer
    template_buffer = processor._expanded_template_buffer
    template_buffer.seek(0)
    
    # Create DocxTemplate and build context
    doc = DocxTemplate(template_buffer)
    context = {}
    for i in range(1, 13):  # Double template has 12 labels (4x3 grid)
        if i == 1:
            context[f'Label{i}'] = processor._build_label_context(test_record, doc)
        else:
            context[f'Label{i}'] = {}
    
    # Render the template
    doc.render(context)
    
    # Save the raw rendered document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    rendered_doc = Document(buffer)
    
    print("After Jinja2 rendering:")
    print_brand_content(rendered_doc, "Jinja2")
    
    # Apply only brand centering (skip font sizing)
    print("\n" + "="*50)
    print("Testing only brand centering (no font sizing)...")
    processor.current_record = test_record
    processor._apply_brand_centering_for_double_template(rendered_doc)
    print_brand_content(rendered_doc, "After brand centering only")
    
    # Apply only BR marker conversion
    print("\n" + "="*50)
    print("Testing BR marker conversion...")
    for table in rendered_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    processor._convert_br_markers_to_line_breaks(paragraph)
    print_brand_content(rendered_doc, "After BR marker conversion")
    
    # Apply only ratio spacing fix
    print("\n" + "="*50)
    print("Testing ratio spacing fix...")
    processor._fix_ratio_paragraph_spacing(rendered_doc)
    print_brand_content(rendered_doc, "After ratio spacing fix")
    
    # Apply only Arial bold enforcement
    print("\n" + "="*50)
    print("Testing Arial bold enforcement...")
    from src.core.generation.docx_formatting import enforce_arial_bold_all_text
    enforce_arial_bold_all_text(rendered_doc)
    print_brand_content(rendered_doc, "After Arial bold enforcement")
    
    # Save the final result
    output_path = "test_no_font_sizing_brand_output.docx"
    rendered_doc.save(output_path)
    print(f"\n✅ Final document saved: {output_path}")

def print_brand_content(doc, stage_name):
    """Print brand content from document."""
    print(f"\n{stage_name}:")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    lines = cell_text.split('\n')
                    for i, line in enumerate(lines):
                        if line.strip():
                            if 'NO_FONT_BRAND_TEST' in line:
                                print(f"  Line {i+1}: '{line.strip()}' ✓ FULL BRAND")
                            elif 'NO_FONT_BRAND' in line:
                                print(f"  Line {i+1}: '{line.strip()}' ⚠ PARTIAL BRAND")
                            elif 'BRAND_TEST' in line:
                                print(f"  Line {i+1}: '{line.strip()}' ⚠ PARTIAL BRAND")

if __name__ == "__main__":
    test_no_font_sizing_brand() 