#!/usr/bin/env python3
"""
Comprehensive test for double template font sizing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from docx.shared import Pt

def test_double_template_comprehensive():
    """Test double template with actual document generation."""
    
    print("Comprehensive Double Template Font Sizing Test")
    print("=" * 55)
    
    # Create test records with various price formats
    test_records = [
        {
            'Description': 'Amnesia Lemon Pre-Roll',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductBrand': 'Test Brand 1',
            'ProductType': 'pre-roll'
        },
        {
            'Description': 'Pineapple Cake Distillate Cartridge - 1g',
            'Price': '$27',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 30% CBD: 1%',
            'ProductBrand': 'Test Brand 2',
            'ProductType': 'vape cartridge'
        },
        {
            'Description': 'Fugi Fruit Pre-Roll',
            'Price': '$4',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 20% CBD: 3%',
            'ProductBrand': 'Test Brand 3',
            'ProductType': 'pre-roll'
        },
        {
            'Description': 'Fugi Fruit Pre-Roll',
            'Price': '$110',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 35% CBD: 0%',
            'ProductBrand': 'Test Brand 4',
            'ProductType': 'pre-roll'
        }
    ]
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Debug: Check what context is being built
        print("Debugging context building...")
        from docx import Document
        temp_doc = Document()
        
        for i, record in enumerate(test_records[:2]):  # Just check first 2 records
            print(f"\nRecord {i+1}:")
            print(f"  Raw record: {record}")
            context = processor._build_label_context(record, temp_doc)
            print(f"  Built context: {context}")
        
        # Process the test records
        print("\nProcessing test records...")
        result_doc = processor.process_records(test_records)
        
        if result_doc and result_doc.tables:
            print("✓ Document generated successfully")
            
            # Analyze font sizes in the document
            print("\nAnalyzing font sizes:")
            print("-" * 25)
            
            table = result_doc.tables[0]
            price_font_sizes = []
            other_font_sizes = []
            
            print(f"Table dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    print(f"\nCell ({row_idx},{col_idx}) content:")
                    cell_text = cell.text.strip()
                    print(f"  Full text: '{cell_text}'")
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text.strip():
                                font_size = run.font.size.pt if run.font.size else "None"
                                run_text = run.text.strip()
                                
                                print(f"    Run: '{run_text}' -> {font_size}pt")
                                
                                # Check if this looks like a price
                                if any(char in run_text for char in '$€£¥') or (run_text.replace('.', '').replace(',', '').isdigit() and len(run_text) <= 10):
                                    price_font_sizes.append(f"Price '{run_text}': {font_size}pt (cell {row_idx},{col_idx})")
                                else:
                                    other_font_sizes.append(f"'{run_text[:20]}...': {font_size}pt (cell {row_idx},{col_idx})")
            
            # Report price font sizes
            print("\nPrice Font Sizes:")
            for price_info in price_font_sizes:
                print(f"  {price_info}")
            
            # Report other font sizes (first 10)
            print(f"\nOther Font Sizes (first 10):")
            for other_info in other_font_sizes[:10]:
                print(f"  {other_info}")
            
            # Check if prices have correct font sizes
            correct_prices = 0
            total_prices = len(price_font_sizes)
            
            for price_info in price_font_sizes:
                if "20.0pt" in price_info or "20pt" in price_info:
                    correct_prices += 1
            
            print(f"\nPrice Font Size Summary:")
            print(f"  Total prices found: {total_prices}")
            print(f"  Prices with correct size (20pt): {correct_prices}")
            print(f"  Success rate: {correct_prices/total_prices*100:.1f}%" if total_prices > 0 else "  No prices found")
            
            if correct_prices == total_prices and total_prices > 0:
                print("✓ All prices have correct font size!")
            elif total_prices > 0:
                print("✗ Some prices have incorrect font size")
            else:
                print("⚠ No prices found in document")
                
        else:
            print("✗ Failed to generate document")
            
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_double_template_comprehensive() 