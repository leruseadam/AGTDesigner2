#!/usr/bin/env python3
"""
Test to check what variables are in the double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
import re

def test_template_variables():
    """Test to see what variables are in the double template."""
    
    print("Double Template Variables Test")
    print("=" * 40)
    
    try:
        # Create template processor
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Get the template path
        template_path = processor._get_template_path()
        print(f"Template path: {template_path}")
        
        # Load the original template
        doc = Document(template_path)
        
        print("\nAnalyzing template content:")
        print("-" * 30)
        
        # Check for Jinja2 variables in the document
        jinja2_variables = []
        
        # Check paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"Paragraph {i}: '{paragraph.text}'")
                # Find Jinja2 variables
                variables = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
                for var in variables:
                    jinja2_variables.append(var)
        
        # Check tables
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"  Cell ({row_idx},{col_idx}): '{cell_text}'")
                        # Find Jinja2 variables
                        variables = re.findall(r'\{\{([^}]+)\}\}', cell_text)
                        for var in variables:
                            jinja2_variables.append(var)
        
        # Report found variables
        print(f"\nFound Jinja2 variables:")
        print("-" * 25)
        if jinja2_variables:
            unique_vars = list(set(jinja2_variables))
            for var in sorted(unique_vars):
                print(f"  {var}")
        else:
            print("  No Jinja2 variables found!")
            
        # Check if we have the expected variables
        expected_vars = ['Label1.Price', 'Label1.Description', 'Label1.Lineage', 'Label1.Ratio_or_THC_CBD']
        print(f"\nChecking for expected variables:")
        print("-" * 35)
        for expected in expected_vars:
            found = any(expected in var for var in jinja2_variables)
            status = "✓" if found else "✗"
            print(f"  {status} {expected}")
            
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_template_variables() 