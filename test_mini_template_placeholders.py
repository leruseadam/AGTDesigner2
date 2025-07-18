#!/usr/bin/env python3
"""
Test script to check what placeholders are in the expanded mini template and add missing ones.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from src.core.generation.template_processor import get_font_scheme, TemplateProcessor
from docx.oxml.ns import qn
import re

def check_and_fix_mini_template():
    """Check what placeholders are in the mini template and add missing ones."""
    
    print("Checking and Fixing Mini Template Placeholders")
    print("=" * 60)
    
    # Get the template path
    processor = TemplateProcessor('mini', get_font_scheme('mini'), 1.0)
    template_path = processor._get_template_path()
    
    print(f"Template path: {template_path}")
    
    try:
        # Load the original template
        doc = Document(template_path)
        
        print(f"\nOriginal template has {len(doc.tables)} tables")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"Original table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check the first cell for placeholders
            if table.rows and table.rows[0].cells:
                first_cell = table.rows[0].cells[0]
                cell_text = first_cell.text
                print(f"\nOriginal first cell text: {repr(cell_text)}")
                
                # Look for common placeholders
                expected_placeholders = [
                    '{{Label1.ProductBrand}}',
                    '{{Label1.Description}}',
                    '{{Label1.Price}}',
                    '{{Label1.Lineage}}',
                    '{{Label1.Ratio_or_THC_CBD}}',
                    '{{Label1.DescAndWeight}}',
                    '{{Label1.DOH}}'
                ]
                
                print(f"\nExpected placeholders:")
                for placeholder in expected_placeholders:
                    if placeholder in cell_text:
                        print(f"  ✓ {placeholder}")
                    else:
                        print(f"  ✗ {placeholder}")
                
                # Add missing placeholders
                missing_placeholders = []
                for placeholder in expected_placeholders:
                    if placeholder not in cell_text:
                        missing_placeholders.append(placeholder)
                
                if missing_placeholders:
                    print(f"\nAdding missing placeholders: {missing_placeholders}")
                    
                    # Add missing placeholders to the cell
                    for placeholder in missing_placeholders:
                        if first_cell.text.strip():
                            first_cell.text += f"\n{placeholder}"
                        else:
                            first_cell.text = placeholder
                    
                    # Save the updated template
                    doc.save(template_path)
                    print(f"✓ Updated template saved to: {template_path}")
                    
                    # Verify the update
                    updated_doc = Document(template_path)
                    updated_cell = updated_doc.tables[0].rows[0].cells[0]
                    print(f"\nUpdated first cell text: {repr(updated_cell.text)}")
                    
                    # Check again
                    print(f"\nPlaceholders after update:")
                    for placeholder in expected_placeholders:
                        if placeholder in updated_cell.text:
                            print(f"  ✓ {placeholder}")
                        else:
                            print(f"  ✗ {placeholder}")
                else:
                    print(f"\n✓ All expected placeholders are present")
        
        # Now test the expanded template
        print(f"\n" + "=" * 60)
        print("Testing Expanded Template")
        print("=" * 60)
        
        # Create a new processor to get the expanded template
        processor2 = TemplateProcessor('mini', get_font_scheme('mini'), 1.0)
        
        if hasattr(processor2, '_expanded_template_buffer'):
            processor2._expanded_template_buffer.seek(0)
            expanded_doc = Document(processor2._expanded_template_buffer)
            
            print(f"Expanded template has {len(expanded_doc.tables)} tables")
            
            if expanded_doc.tables:
                expanded_table = expanded_doc.tables[0]
                print(f"Expanded table: {len(expanded_table.rows)} rows x {len(expanded_table.columns)} columns")
                
                # Check a few cells for placeholders
                for i in range(min(3, len(expanded_table.rows))):
                    for j in range(min(3, len(expanded_table.columns))):
                        cell = expanded_table.cell(i, j)
                        cell_text = cell.text
                        label_num = i * len(expanded_table.columns) + j + 1
                        print(f"\nCell ({i},{j}) - Label{label_num}: {repr(cell_text)}")
                        
                        # Check for expected placeholders
                        expected_placeholders = [
                            f'{{{{Label{label_num}.ProductBrand}}}}',
                            f'{{{{Label{label_num}.Description}}}}',
                            f'{{{{Label{label_num}.Price}}}}',
                            f'{{{{Label{label_num}.Lineage}}}}',
                            f'{{{{Label{label_num}.Ratio_or_THC_CBD}}}}',
                            f'{{{{Label{label_num}.DescAndWeight}}}}',
                            f'{{{{Label{label_num}.DOH}}}}'
                        ]
                        
                        found_placeholders = []
                        for placeholder in expected_placeholders:
                            if placeholder in cell_text:
                                found_placeholders.append(placeholder)
                        
                        print(f"  Found placeholders: {found_placeholders}")
                
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_and_fix_mini_template() 