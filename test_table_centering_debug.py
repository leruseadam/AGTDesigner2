#!/usr/bin/env python3
"""
Test script to debug table centering issues.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO

def debug_table_centering():
    """Debug table centering issues."""
    print("Debugging Table Centering")
    print("=" * 50)
    
    # Create a test record
    test_record = {
        'Description': 'Test Product',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'DOH'
    }
    
    # Test double template specifically
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return
    
    # Examine document structure
    print(f"Document sections: {len(result_doc.sections)}")
    for i, section in enumerate(result_doc.sections):
        print(f"Section {i}:")
        print(f"  Left margin: {section.left_margin}")
        print(f"  Right margin: {section.right_margin}")
        print(f"  Page width: {section.page_width}")
        print(f"  Available width: {section.page_width - section.left_margin - section.right_margin}")
    
    # Examine tables
    print(f"\nTables found: {len(result_doc.tables)}")
    for i, table in enumerate(result_doc.tables):
        print(f"\nTable {i}:")
        print(f"  Alignment: {table.alignment}")
        print(f"  Columns: {len(table.columns)}")
        print(f"  Rows: {len(table.rows)}")
        
        # Check table properties
        tblPr = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
        if tblPr is not None:
            print(f"  Table properties found")
            
            # Check for table alignment property
            tblAlign = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblAlign')
            if tblAlign is not None:
                print(f"  Table alignment property: {tblAlign.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')}")
            else:
                print(f"  No table alignment property found")
        
        # Check table grid
        tblGrid = table._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblGrid')
        if tblGrid is not None:
            print(f"  Table grid found with {len(tblGrid)} columns")
            for j, gridCol in enumerate(tblGrid):
                width = gridCol.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                print(f"    Column {j}: width = {width} twips ({int(width)/1440:.2f} inches)")
        
        # Check first cell properties
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0]
            tcPr = first_cell._tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            if tcPr is not None:
                tcW = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
                if tcW is not None:
                    width = tcW.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                    print(f"  First cell width: {width} twips ({int(width)/1440:.2f} inches)")
    
    # Check paragraphs outside tables
    print(f"\nParagraphs outside tables: {len(result_doc.paragraphs)}")
    for i, paragraph in enumerate(result_doc.paragraphs):
        if paragraph.text.strip():
            print(f"  Paragraph {i}: '{paragraph.text[:50]}...'")
            print(f"    Alignment: {paragraph.alignment}")
    
    # Save for manual inspection
    output_path = "debug_table_centering_output.docx"
    result_doc.save(output_path)
    print(f"\nDebug output saved to: {output_path}")

if __name__ == "__main__":
    debug_table_centering() 