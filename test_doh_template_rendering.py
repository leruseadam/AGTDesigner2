#!/usr/bin/env python3
"""
Simple test to verify DOH template rendering.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn

def test_doh_template_rendering():
    """Test if DocxTemplate can render DOH images."""
    print("Testing DOH Template Rendering")
    print("=" * 50)
    
    # Create a simple template with DOH placeholder
    template_content = """
    {{Label1.DOH}}
    """
    
    # Create a simple document with the template content
    doc = Document()
    doc.add_paragraph(template_content)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create DocxTemplate
    tpl = DocxTemplate(buffer)
    
    # Create context with DOH image
    context = {
        'Label1': {
            'DOH': 'TEST_DOH_IMAGE'  # Use a string first to test placeholder replacement
        }
    }
    
    # Render template
    tpl.render(context)
    
    # Save rendered document
    output_buffer = BytesIO()
    tpl.save(output_buffer)
    output_buffer.seek(0)
    
    # Check if placeholder was replaced
    rendered_doc = Document(output_buffer)
    print(f"Rendered document paragraphs: {len(rendered_doc.paragraphs)}")
    
    for i, para in enumerate(rendered_doc.paragraphs):
        print(f"Paragraph {i}: '{para.text}'")
    
    # Now test with actual image
    print("\nTesting with actual image...")
    
    # Create a simple image placeholder
    image_path = "src/core/generation/templates/DOH.png"
    if os.path.exists(image_path):
        print(f"Image path exists: {image_path}")
        
        # Create InlineImage
        image_width = Mm(12)
        doh_image = InlineImage(tpl, image_path, width=image_width)
        
        # Create context with image
        context_with_image = {
            'Label1': {
                'DOH': doh_image
            }
        }
        
        # Render template
        tpl.render(context_with_image)
        
        # Save rendered document
        output_buffer2 = BytesIO()
        tpl.save(output_buffer2)
        output_buffer2.seek(0)
        
        # Check if image was inserted
        rendered_doc2 = Document(output_buffer2)
        print(f"Rendered document with image paragraphs: {len(rendered_doc2.paragraphs)}")
        
        # Check for images in the document
        image_count = 0
        for para in rendered_doc2.paragraphs:
            for run in para.runs:
                if hasattr(run, '_element') and run._element.find(qn('w:drawing')) is not None:
                    image_count += 1
                    print(f"Found image in paragraph: '{para.text}'")
        
        print(f"Total images found: {image_count}")
        
        if image_count > 0:
            print("✓ SUCCESS: DOH image was properly rendered")
            return True
        else:
            print("✗ FAILURE: No DOH image found in rendered document")
            return False
    else:
        print(f"✗ FAILURE: Image path does not exist: {image_path}")
        return False

if __name__ == "__main__":
    test_doh_template_rendering() 