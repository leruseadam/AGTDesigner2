#!/usr/bin/env python3
"""
Debug script to understand template expansion and DOH placeholder addition.
"""

import sys
import os
import logging
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)

from src.core.generation.template_processor import TemplateProcessor
from src.core.constants import FONT_SCHEME_DOUBLE
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from io import BytesIO

def debug_template_expansion():
    """Debug template expansion process."""
    print("Debugging Template Expansion")
    print("=" * 50)
    
    # Create a test record with DOH image
    test_record = {
        'Description': 'Test Product with DOH',
        'WeightUnits': '1g',
        'ProductBrand': 'Test Brand',
        'Price': '$10.00',
        'Lineage': 'Test Lineage',
        'THC_CBD': 'THC: 20% CBD: 2%',
        'ProductStrain': 'Test Strain',
        'DOH': 'YES',  # This should trigger DOH image
        'Product Type*': 'classic'  # This should use regular DOH image
    }
    
    # Test double template specifically
    processor = TemplateProcessor('double', FONT_SCHEME_DOUBLE)
    
    # Check the original template
    template_path = processor._get_template_path()
    print(f"Template path: {template_path}")
    
    original_doc = Document(template_path)
    print(f"Original template has {len(original_doc.tables)} tables")
    
    if original_doc.tables:
        original_cell = original_doc.tables[0].cell(0, 0)
        print(f"Original cell text: '{original_cell.text}'")
        
        # Check for DOH placeholder in original template
        if 'DOH' in original_cell.text:
            print("✓ DOH placeholder found in original template")
        else:
            print("✗ DOH placeholder NOT found in original template")
    
    # Process the record
    result_doc = processor.process_records([test_record])
    
    if not result_doc:
        print("ERROR: Failed to process test record")
        return
    
    print(f"\nProcessed document has {len(result_doc.tables)} tables")
    
    # Check the first few cells for DOH placeholders
    if result_doc.tables:
        table = result_doc.tables[0]
        print(f"Processed table: {len(table.rows)} rows, {len(table.columns)} columns")
        
        # Check first few cells
        for i in range(min(3, len(table.rows))):
            for j in range(min(3, len(table.columns))):
                cell = table.cell(i, j)
                cell_text = cell.text.strip()
                print(f"Cell ({i}, {j}): '{cell_text}'")
                
                if 'DOH' in cell_text:
                    print(f"  ✓ DOH placeholder found in cell ({i}, {j})")
                else:
                    print(f"  ✗ DOH placeholder NOT found in cell ({i}, {j})")

if __name__ == "__main__":
    debug_template_expansion() 