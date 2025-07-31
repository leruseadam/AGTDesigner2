#!/usr/bin/env python3
"""
Detailed test script to examine multi-line THC/CBD ratio value formatting.
"""

import os
import sys
from pathlib import Path

# Add the src directory to the Python path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def examine_document_content(doc, template_type):
    """Examine the actual content of the document to see what's happening."""
    
    print(f"\n=== Examining {template_type} template document ===")
    
    # Check all tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    if paragraph.text.strip():
                        print(f"  Cell[{row_idx},{cell_idx}] Para[{para_idx}]: '{paragraph.text}'")
                        
                        # Check each run in the paragraph
                        for run_idx, run in enumerate(paragraph.runs):
                            if run.text.strip():
                                bold_status = "BOLD" if run.font.bold else "NOT BOLD"
                                font_size = f"{run.font.size.pt}pt" if run.font.size else "No size"
                                print(f"    Run[{run_idx}]: '{run.text}' - {bold_status} - {font_size}")
                                
                                # Check for ratio patterns
                                if any(pattern in run.text for pattern in [
                                    'THC:', 'CBD:', 'CBC:', 'CBG:', 'CBN:',
                                    'mg THC', 'mg CBD', 'mg CBC', 'mg CBG', 'mg CBN',
                                    '1:1', '2:1', '3:1', '1:1:1', '2:1:1'
                                ]):
                                    print(f"      *** RATIO CONTENT DETECTED: '{run.text}' - {bold_status} ***")
                                    
                                    # Check if this is multi-line content
                                    if '\n' in run.text or '|BR|' in run.text:
                                        print(f"      *** MULTI-LINE RATIO: '{run.text[:50]}...' - {bold_status} ***")
    
    # Check paragraphs outside tables
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"\nParagraph {para_idx + 1}: '{paragraph.text}'")
            for run_idx, run in enumerate(paragraph.runs):
                if run.text.strip():
                    bold_status = "BOLD" if run.font.bold else "NOT BOLD"
                    font_size = f"{run.font.size.pt}pt" if run.font.size else "No size"
                    print(f"  Run[{run_idx}]: '{run.text}' - {bold_status} - {font_size}")

def test_multiline_ratio_bold_detailed():
    """Detailed test of multi-line ratio value bold formatting."""
    
    print("Testing multi-line ratio value bold formatting (detailed)...")
    
    # Create test data with multi-line ratio formats
    test_records = [
        {
            'Product Name*': 'Multi-Line Test Product 1',
            'Product Type*': 'Edible',
            'Lineage': 'HYBRID',
            'Product Brand': 'CONSTELLATION CANNABIS',
            'Vendor/Supplier*': 'JOURNEYMAN',
            'Weight*': '1.7',
            'Weight Unit* (grams/gm or ounces/oz)': 'oz',
            'Price* (Tier Name for Bulk)': '$18',
            'Ratio_or_THC_CBD': 'THC: 25%\nCBD: 2%',  # Use the correct field name
            'Product Strain': 'Test Strain'
        },
        {
            'Product Name*': 'Multi-Line Test Product 2',
            'Product Type*': 'Edible',
            'Lineage': 'HYBRID',
            'Product Brand': 'CONSTELLATION CANNABIS',
            'Vendor/Supplier*': 'JOURNEYMAN',
            'Weight*': '0.07',
            'Weight Unit* (grams/gm or ounces/oz)': 'oz',
            'Price* (Tier Name for Bulk)': '$15',
            'Ratio_or_THC_CBD': '100mg THC\n50mg CBD\n25mg CBC',  # Use the correct field name
            'Product Strain': 'Test Strain 2'
        }
    ]
    
    # Test just the vertical template for detailed examination
    template_type = 'vertical'
    
    try:
        # Create template processor
        processor = TemplateProcessor(template_type, {}, scale_factor=1.0)
        
        # Process the test records
        doc = processor.process_records(test_records)
        
        if doc:
            # Save the document
            output_file = f'test_multiline_ratio_bold_detailed_{template_type}.docx'
            doc.save(output_file)
            print(f"✅ Generated {output_file}")
            
            # Examine the document content
            examine_document_content(doc, template_type)
            
        else:
            print(f"❌ Failed to generate document for {template_type} template")
            
    except Exception as e:
        print(f"❌ Error testing {template_type} template: {e}")
        import traceback
        traceback.print_exc()
    
    print("\nDetailed test completed.")

if __name__ == "__main__":
    test_multiline_ratio_bold_detailed() 