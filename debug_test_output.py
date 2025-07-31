#!/usr/bin/env python3
"""
Debug script to see what the actual output contains
"""

import sys
import os
from pathlib import Path

# Add the src directory to the Python path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document
from io import BytesIO

def debug_test_output():
    """Debug what the actual output contains."""
    print("=== Debugging Test Output ===")
    
    # Create test records
    test_records = [
        {
            'ProductName': 'Test Product 1',
            'ProductBrand': 'Test Brand 1',
            'Price': '$25.99',
            'Description': 'Test Description 1',
            'WeightUnits': '3.5g',
            'Lineage': 'Test Lineage 1',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'DOH': 'YES',
            'ProductStrain': 'Test Strain 1',
            'ProductType': 'flower',
            'Vendor': 'Test Vendor 1'
        },
        {
            'ProductName': 'Test Product 2',
            'ProductBrand': 'Test Brand 2',
            'Price': '$30.99',
            'Description': 'Test Description 2',
            'WeightUnits': '1g',
            'Lineage': 'Test Lineage 2',
            'Ratio_or_THC_CBD': 'THC: 30% CBD: 1%',
            'DOH': 'YES',
            'ProductStrain': 'Test Strain 2',
            'ProductType': 'concentrate',
            'Vendor': 'Test Vendor 2'
        }
    ]
    
    try:
        # Create processor for double template
        processor = TemplateProcessor('double', {}, 1.0)
        
        # Process the records
        result = processor.process_records(test_records)
        
        if result is None:
            print("❌ Failed to process records")
            return False
        
        # Save the result to a temporary file
        output_path = "debug_test_output.docx"
        result.save(output_path)
        
        # Load the saved document and check content
        doc = Document(output_path)
        
        print("\n=== Actual Output Content ===")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if text.strip():
                            print(f"Cell content: {repr(text)}")
        
        # Check for specific data patterns
        print("\n=== Data Pattern Check ===")
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text += paragraph.text + " "
        
        print(f"All text: {repr(all_text)}")
        
        # Check for expected patterns
        patterns = ['Test Brand', 'Test Product', '$25.99', '$30.99', 'Test Lineage', 'Test Strain', 'Test Vendor']
        for pattern in patterns:
            if pattern in all_text:
                print(f"✅ Found pattern: {pattern}")
            else:
                print(f"❌ Missing pattern: {pattern}")
        
        # Clean up
        if os.path.exists(output_path):
            os.remove(output_path)
        
        return True
        
    except Exception as e:
        print(f"❌ Error during debug: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    debug_test_output() 