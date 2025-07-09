#!/usr/bin/env python3
"""
Detailed debug script to trace mini template processing.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def debug_mini_detailed():
    """Detailed debugging of mini template processing."""
    
    print("Detailed Mini Template Debug")
    print("=" * 50)
    
    # Test data - using the actual data from the API
    test_record = {
        'ProductName': 'Core Reactor Quartz Banger',
        'ProductBrand': 'Paraphernalia',
        'Price': '$25.99',
        'Lineage': 'HYBRID',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'Description': 'Core Reactor Quartz Banger',
        'ProductStrain': '',
        'ProductType': 'Paraphernalia',
        'Vendor': 'One Stop Wholesale'
    }
    
    try:
        # Step 1: Create processor
        print("Step 1: Creating template processor...")
        processor = TemplateProcessor('mini', {}, 1.0)
        print(f"  Template type: {processor.template_type}")
        print(f"  Chunk size: {processor.chunk_size}")
        
        # Step 2: Check expanded template
        print("\nStep 2: Checking expanded template...")
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            if doc.tables:
                table = doc.tables[0]
                print(f"  Expanded template table: {len(table.rows)}x{len(table.columns)}")
                cell = table.cell(0, 0)
                print(f"  First cell text: '{cell.text}'")
            else:
                print("  ❌ No tables in expanded template")
        else:
            print("  ❌ No expanded template buffer")
        
        # Step 3: Process records
        print("\nStep 3: Processing records...")
        result_doc = processor.process_records([test_record])
        
        if result_doc:
            print("  ✅ Document created successfully")
            if result_doc.tables:
                table = result_doc.tables[0]
                print(f"  Result table: {len(table.rows)}x{len(table.columns)}")
                
                # Check first few cells
                for row_idx in range(min(3, len(table.rows))):
                    for col_idx in range(min(3, len(table.columns))):
                        cell = table.cell(row_idx, col_idx)
                        print(f"  Cell ({row_idx}, {col_idx}): '{cell.text.strip()}'")
            else:
                print("  ❌ No tables in result document")
        else:
            print("  ❌ Failed to create document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_mini_detailed() 