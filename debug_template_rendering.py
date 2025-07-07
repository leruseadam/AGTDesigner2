#!/usr/bin/env python3
"""
Debug script to check template rendering for lineage field.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def debug_template_rendering():
    """Debug template rendering to see if lineage field is being processed."""
    
    print("Debugging Template Rendering")
    print("=" * 50)
    
    # Test data
    test_record = {
        'ProductBrand': 'Test Brand',
        'Price': '$25.99',
        'Lineage': 'MIXED',
        'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
        'Description': 'Test description text',
        'ProductStrain': 'Mixed',
        'ProductType': 'tincture'
    }
    
    try:
        # Create processor
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Check the expanded template before rendering
        print("Step 1: Check expanded template before rendering...")
        if hasattr(processor, '_expanded_template_buffer'):
            processor._expanded_template_buffer.seek(0)
            doc = Document(processor._expanded_template_buffer)
            if doc.tables:
                cell = doc.tables[0].cell(0, 0)
                print(f"Template cell before rendering: '{cell.text}'")
        
        # Build context
        print("\nStep 2: Build label context...")
        from docxtpl import DocxTemplate
        doc_template = DocxTemplate(processor._expanded_template_buffer)
        
        # Build context for all 20 labels (mini template has 20 labels)
        context = {}
        for i in range(1, 21):  # Label1 to Label20
            if i == 1:
                context[f'Label{i}'] = processor._build_label_context(test_record, doc_template)
            else:
                context[f'Label{i}'] = {}  # Empty context for unused labels
        
        print(f"Context keys: {list(context.keys())}")
        print(f"Label1 context: {context['Label1']}")
        
        # Render template
        print("\nStep 3: Render template...")
        doc_template.render(context)
        
        # Save and reload to check what was rendered
        from io import BytesIO
        buffer = BytesIO()
        doc_template.save(buffer)
        buffer.seek(0)
        rendered_doc = Document(buffer)
        
        print("\nStep 4: Check rendered document...")
        if rendered_doc.tables:
            table = rendered_doc.tables[0]
            print(f"Rendered table dimensions: {len(table.rows)}x{len(table.columns)}")
            
            # Check first few cells
            for row_idx in range(min(2, len(table.rows))):
                for col_idx in range(min(2, len(table.columns))):
                    cell = table.cell(row_idx, col_idx)
                    print(f"Cell ({row_idx}, {col_idx}): '{cell.text}'")
                    
                    # Check paragraph content
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        print(f"  Paragraph {para_idx}: '{paragraph.text}'")
                        for run_idx, run in enumerate(paragraph.runs):
                            print(f"    Run {run_idx}: '{run.text}'")
        else:
            print("❌ No tables found in rendered document")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

def test_simple_template():
    """Test with a simple template to see if the issue is with the template itself."""
    
    print("\n\nTesting Simple Template")
    print("=" * 50)
    
    try:
        # Create a simple document with just lineage
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "{{Label1.Lineage}}"
        
        # Save to buffer
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create template processor with this simple template
        processor = TemplateProcessor('mini', {}, 1.0)
        
        # Override the template buffer with our simple template
        processor._expanded_template_buffer = buffer
        
        # Test record
        test_record = {
            'Lineage': 'MIXED'
        }
        
        # Process
        result_doc = processor.process_records([test_record])
        
        if result_doc and result_doc.tables:
            cell = result_doc.tables[0].cell(0, 0)
            print(f"Simple template result: '{cell.text}'")
        else:
            print("❌ Failed to process simple template")
            
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_template_rendering()
    test_simple_template() 