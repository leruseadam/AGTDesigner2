#!/usr/bin/env python3
"""
Test brand display in double template.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.core.generation.template_processor import TemplateProcessor
from docx import Document

def test_brand_double():
    """Test brand display in double template."""
    
    print("Brand Double Template Test")
    print("=" * 40)
    
    # Create test records with different brand scenarios
    test_records = [
        {
            'Description': 'Amnesia Lemon Pre-Roll',
            'Price': '$12',
            'Lineage': 'SATIVA',
            'Ratio_or_THC_CBD': 'THC: 25% CBD: 2%',
            'ProductBrand': 'Test Brand 1',
            'ProductType': 'pre-roll',
            'JointRatio': '1.0g x 1 Pack'
        },
        {
            'Description': 'Blue Dream Flower',
            'Price': '$25',
            'Lineage': 'HYBRID',
            'Ratio_or_THC_CBD': 'THC: 22% CBD: 1%',
            'ProductBrand': 'Premium Brand',
            'ProductType': 'flower',
            'JointRatio': '1.5g x 2 Pack'
        }
    ]
    
    # Initialize template processor
    processor = TemplateProcessor('double', 'standard', scale_factor=1.0)
    
    # Process the records
    print("Processing records...")
    result = processor.process_records(test_records)
    
    if result is None:
        print("❌ Failed to process records")
        return
    
    # Save the result
    output_path = "test_brand_double_output.docx"
    result.save(output_path)
    print(f"✅ Generated document: {output_path}")
    
    # Analyze the content
    print("\nAnalyzing content...")
    
    # Check the context that was built
    print("\nContext analysis:")
    for i, record in enumerate(test_records):
        print(f"\nRecord {i+1}:")
        print(f"  Description: '{record.get('Description', 'N/A')}'")
        print(f"  ProductBrand: '{record.get('ProductBrand', 'N/A')}'")
        print(f"  Price: '{record.get('Price', 'N/A')}'")
        print(f"  Lineage: '{record.get('Lineage', 'N/A')}'")
    
    # Analyze the generated document
    print("\nGenerated document analysis:")
    doc = Document(output_path)
    
    # Check all tables
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f"  Cell ({row_idx},{col_idx}): '{cell_text}'")
                    
                    # Check font sizes for each run
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if paragraph.text.strip():
                            print(f"    Paragraph {para_idx + 1}: '{paragraph.text.strip()}'")
                            for run_idx, run in enumerate(paragraph.runs):
                                if run.text.strip():
                                    font_size = run.font.size
                                    font_name = run.font.name
                                    print(f"      Run {run_idx + 1}: '{run.text.strip()}' - Font: {font_name}, Size: {font_size}")

if __name__ == "__main__":
    test_brand_double() 